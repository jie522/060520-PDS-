import csv
import io
import json
import os
import sys
import re
import sqlite3
import time
import threading
import uuid
import webbrowser
import urllib.parse
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date
from pathlib import Path
from flask import Flask, render_template, request, jsonify, redirect
import requests
from requests_ntlm import HttpNtlmAuth
from urllib.parse import quote
try:
    import openpyxl
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
    _DOCXTPL_OK = True
except ImportError:
    _DOCXTPL_OK = False

# ── PyInstaller 路徑處理 ────────────────────────────────────────
# 執行為 EXE 時：_MEIPASS 為解壓目錄（含 templates）；
# EXE 同目錄為使用者可編輯的 config.py / pdm_search.db
if getattr(sys, 'frozen', False):
    _BASE = sys._MEIPASS          # 唯讀資源（templates 等）
    _APP_DIR = os.path.dirname(sys.executable)  # EXE 所在目錄
    sys.path.insert(0, _APP_DIR)  # 讓 import config 找到 EXE 旁的 config.py
else:
    _BASE = os.path.dirname(os.path.abspath(__file__))
    _APP_DIR = _BASE
    # python311._pth 存在時，嵌入式 Python 不自動加入腳本目錄
    # 需手動插入，讓 import config / templates 等都能正常運作
    if _APP_DIR not in sys.path:
        sys.path.insert(0, _APP_DIR)
    _PARENT = os.path.dirname(_APP_DIR)
    if os.path.exists(os.path.join(_PARENT, 'config.py')) and _PARENT not in sys.path:
        sys.path.insert(0, _PARENT)

import config

# 自主巡檢表路徑
INSPECTION_ROOT = r"\\192.168.1.99\共用區\品保加工共用平台\加工課自主巡檢表(空白)"

# PDM 圖面索引資料庫（可在 config.py 覆寫）
PDM_DB_PATH = getattr(config, 'PDM_DB_PATH',
    r"C:\Users\user\AppData\Local\PDMSearch\pdm_search.db")

# Fallback：若設定路徑不存在，嘗試同目錄下的 pdm_search.db（打包時隨包）
if not os.path.exists(PDM_DB_PATH):
    _bundled_db = os.path.join(_APP_DIR, 'pdm_search.db')
    if os.path.exists(_bundled_db):
        PDM_DB_PATH = _bundled_db

# ── 技術資料清單資料庫（zume-n.com 圖號↔URL 對照表）──────────────────────
ZUME_DB_PATH = os.path.join(_APP_DIR, 'zume_drawings.db')

# ZUMEN 圖面欄位 → CSV 標題關鍵字對照（依關鍵字動態偵測欄位，缺的留空）
_ZUME_EXTRA_COLS = ['line', 'prod_group', 'category', 'vendor']
_ZUME_COL_KEYWORDS = {
    'line':       ['生產線別', '線別'],
    'prod_group': ['生產群組', '群組'],
    'category':   ['分類'],
    'vendor':     ['廠商', '客戶'],
}

def _ensure_zume_columns(con):
    """相容遷移：drawings 表若缺新欄位則補上（line/prod_group/category/vendor）"""
    existing = {r[1] for r in con.execute('PRAGMA table_info(drawings)')}
    for col in _ZUME_EXTRA_COLS:
        if col not in existing:
            con.execute(f'ALTER TABLE drawings ADD COLUMN {col} TEXT')

def _init_zume_db():
    con = sqlite3.connect(ZUME_DB_PATH)
    con.execute('''
        CREATE TABLE IF NOT EXISTS drawings (
            part_no   TEXT PRIMARY KEY,
            part_name TEXT,
            url       TEXT NOT NULL
        )
    ''')
    con.execute('''
        CREATE TABLE IF NOT EXISTS import_log (
            filename  TEXT PRIMARY KEY,
            imported_at TEXT,
            count     INTEGER
        )
    ''')
    _ensure_zume_columns(con)
    con.commit(); con.close()

def _zume_header_indices(headers):
    """依標題關鍵字回傳各欄位的索引 dict（找不到的為 None）"""
    idx = {
        'part_no':   next((i for i, h in enumerate(headers) if '圖號' in h), None),
        'url':       next((i for i, h in enumerate(headers) if h.strip().upper() == 'URL'), None),
        'part_name': next((i for i, h in enumerate(headers) if '品名' in h), 1),
    }
    for col, kws in _ZUME_COL_KEYWORDS.items():
        idx[col] = next((i for i, h in enumerate(headers) if any(k in h for k in kws)), None)
    return idx

def _zume_row_to_rec(row, idx):
    """依欄位索引把一列轉成 dict；圖號/URL 缺或無效則回傳 None"""
    def cell(col):
        i = idx.get(col)
        return str(row[i]).strip() if (i is not None and i < len(row) and row[i] is not None) else ''
    no  = cell('part_no')
    url = cell('url')
    if not no or not url.startswith('http'):
        return None
    return {
        'part_no': no, 'part_name': cell('part_name'), 'url': url,
        'line': cell('line'), 'prod_group': cell('prod_group'),
        'category': cell('category'), 'vendor': cell('vendor'),
    }

def _insert_zume_recs(con, recs):
    """把解析後的 dict 清單寫入 drawings 表"""
    con.executemany(
        'INSERT OR REPLACE INTO drawings'
        '(part_no, part_name, url, line, prod_group, category, vendor) '
        'VALUES (:part_no, :part_name, :url, :line, :prod_group, :category, :vendor)',
        recs)

def _parse_zume_csv(filepath):
    """解析 zume-n_data_list_*.csv，回傳 [dict, ...]"""
    recs = []
    try:
        with open(filepath, encoding='utf-8-sig', newline='') as f:
            reader  = csv.reader(f)
            headers = [h.strip() for h in next(reader)]
            idx     = _zume_header_indices(headers)
            if idx['part_no'] is None or idx['url'] is None:
                return []
            for row in reader:
                rec = _zume_row_to_rec(row, idx)
                if rec:
                    recs.append(rec)
    except Exception:
        pass
    return recs

def _zume_csv_search_dirs():
    """回傳要掃描 zume-n_data_list_*.csv 的候選資料夾（去重、僅保留存在者）"""
    home = os.path.expanduser('~')
    cands = [
        os.path.join(home, 'Downloads'),
        os.path.join(home, '下載'),
        os.path.join(home, 'Desktop'),
        os.path.join(home, 'OneDrive', 'Downloads'),
        os.path.join(home, 'OneDrive', '下載'),
        os.path.join(home, 'OneDrive', 'Desktop'),
    ]
    seen, dirs = set(), []
    for d in cands:
        key = os.path.normcase(os.path.abspath(d))
        if key not in seen and os.path.isdir(d):
            seen.add(key); dirs.append(d)
    return dirs


def _find_zume_csvs():
    """跨多個常見資料夾搜尋 zume-n_data_list_*.csv，依修改時間新→舊排序"""
    import glob as _glob
    files = []
    for d in _zume_csv_search_dirs():
        files.extend(_glob.glob(os.path.join(d, 'zume-n_data_list_*.csv')))
    return sorted(files, key=os.path.getmtime, reverse=True)


def _auto_import_zume_csv():
    """啟動時自動掃描下載/桌面等資料夾，匯入最新的 zume-n_data_list_*.csv"""
    files = _find_zume_csvs()
    if not files:
        return
    latest = files[0]
    fname  = os.path.basename(latest)
    con = sqlite3.connect(ZUME_DB_PATH)
    already = con.execute('SELECT count FROM import_log WHERE filename=?', (fname,)).fetchone()
    con.close()
    if already:
        return  # 已匯入過，略過
    recs = _parse_zume_csv(latest)
    if not recs:
        return
    con = sqlite3.connect(ZUME_DB_PATH)
    _insert_zume_recs(con, recs)
    con.execute('INSERT OR REPLACE INTO import_log(filename,imported_at,count) VALUES(?,datetime("now"),?)',
                (fname, len(recs)))
    con.commit(); con.close()

_init_zume_db()
_auto_import_zume_csv()   # 啟動時自動掃描匯入

app = Flask(__name__,
            template_folder=os.path.join(_BASE, 'templates'),
            static_folder=os.path.join(_BASE, 'static'))
app.config['TEMPLATES_AUTO_RELOAD'] = True

@app.after_request
def set_no_cache(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    # 動態注入版本號到所有 HTML 頁面（繞過模板快取）
    if response.content_type and 'text/html' in response.content_type:
        try:
            html = response.get_data(as_text=True)
            if '</head>' in html and APP_VERSION not in html:
                inject = f'<script>document.title="製令查詢 {APP_VERSION}";</script>\n</head>'
                html = html.replace('</head>', inject)
                response.set_data(html)
        except Exception:
            pass
    return response

# ── 全域錯誤處理（確保所有錯誤都回傳 JSON，而非 HTML）────────
@app.errorhandler(400)
def bad_request(e):
    return jsonify({'success': False, 'error': f'請求錯誤：{str(e)}'}), 400

@app.errorhandler(404)
def not_found(e):
    return jsonify({'success': False, 'error': f'路由不存在：{str(e)}'}), 404

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'success': False, 'error': f'伺服器錯誤：{str(e)}'}), 500

@app.errorhandler(Exception)
def unhandled_exception(e):
    return jsonify({'success': False, 'error': f'未預期的錯誤：{str(e)}'}), 500

# ── TTL 快取（避免每次搜尋都重新呼叫 SSRS）──────────────────
CACHE_TTL = 120  # 快取有效期 120 秒

_cache = {}
_cache_lock = threading.Lock()


def cache_get(key):
    """取得快取資料，過期回傳 None"""
    with _cache_lock:
        entry = _cache.get(key)
        if entry and (time.time() - entry['ts']) < CACHE_TTL:
            return entry['data']
    return None


def cache_set(key, data):
    """設定快取資料"""
    with _cache_lock:
        _cache[key] = {'data': data, 'ts': time.time()}


def cache_clear(*keys):
    """清除指定的快取 key（用於「重新整理」強制重抓最新資料）"""
    with _cache_lock:
        for k in keys:
            _cache.pop(k, None)


def match_and_not(query, text):
    """AND/NOT 模糊搜尋語法
    空格分隔 = AND（全部都要符合）
    -前綴 = NOT（排除）
    例: 'FTB 夾管座'   → 包含 FTB 且包含 夾管座
        'FTB -夾管座'  → 包含 FTB 但不包含 夾管座
    """
    if not query:
        return True
    text_upper = text.upper()
    for token in query.strip().split():
        if token.startswith('-') and len(token) > 1:
            if token[1:].upper() in text_upper:
                return False
        else:
            if token.upper() not in text_upper:
                return False
    return True


def fetch_all_unfinished():
    """取得所有未完工製令（帶 TTL 快取）"""
    cached = cache_get('unfinished')
    if cached is not None:
        return cached
    params = {
        '發放情況': '已發放',
        'SFT完工碼': '尚未',
        '加工單位': '*',
    }
    csv_text = fetch_report_csv(config.REPORT_PATHS['unfinished'], params)
    data = parse_csv(csv_text)
    cache_set('unfinished', data)
    return data


_ssrs_session = None
_ssrs_session_lock = threading.Lock()


def get_ssrs_session():
    """建立/重用 SSRS NTLM 認證連線（連線池，避免每次建立新連線）"""
    global _ssrs_session
    with _ssrs_session_lock:
        if _ssrs_session is None:
            _ssrs_session = requests.Session()
            _ssrs_session.auth = HttpNtlmAuth(config.SSRS_USERNAME, config.SSRS_PASSWORD)
        return _ssrs_session


def fetch_report_csv(report_path, params=None):
    """從 SSRS 取得報表 CSV 資料"""
    session = get_ssrs_session()
    url = f'{config.SSRS_BASE_URL}?{quote(report_path)}&rs:Command=Render&rs:Format=CSV'
    if params:
        param_str = '&'.join(f'{quote(k)}={quote(str(v))}' for k, v in params.items())
        url += f'&{param_str}'
    resp = session.get(url, timeout=30)
    resp.raise_for_status()
    return resp.content.decode('utf-8-sig')


def fetch_google_sheet_csv(sheet_id, sheet_name=None, gid=None):
    """從 Google 試算表（共用設定：知道連結的人均可檢視）取得指定分頁的 CSV 資料
    可用分頁名稱（sheet_name）或分頁 gid 指定要讀取的分頁"""
    if gid is not None:
        url = f'https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&gid={gid}'
    else:
        url = f'https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={quote(sheet_name)}'
    resp = requests.get(url, timeout=10)
    resp.raise_for_status()
    text = resp.content.decode('utf-8-sig')
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return []
    header = rows[0]
    return [dict(zip(header, row)) for row in rows[1:] if any(cell.strip() for cell in row)]


def fetch_employee_name_map():
    """讀取 Google 試算表『員工登錄系統』，回傳 {工號ID: 姓名} 對照表（含快取）"""
    cache_key = 'employee_name_map'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached

    rows = fetch_google_sheet_csv(config.EMPLOYEE_SHEET_ID, gid=config.EMPLOYEE_SHEET_GID)
    name_map = {}
    for r in rows:
        emp_id = (r.get('工號ID') or '').strip()
        emp_name = (r.get('姓名') or '').strip()
        if emp_id and emp_name:
            name_map[emp_id] = emp_name

    cache_set(cache_key, name_map)
    return name_map


def fetch_employee_roster():
    """讀取 Google 試算表『員工登錄系統』(1.1員工登錄系統分頁)，回傳完整人員名冊（含快取）"""
    cache_key = 'employee_roster'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached

    rows = fetch_google_sheet_csv(config.EMPLOYEE_SHEET_ID, gid=config.EMPLOYEE_SHEET_GID)
    roster = []
    for r in rows:
        emp_id = (r.get('工號ID') or '').strip()
        name   = (r.get('姓名') or '').strip()
        if not emp_id or not name:
            continue
        roster.append({
            'emp_id':      emp_id,
            'name':        name,
            'group':       (r.get('組別') or '').strip(),
            'title':       (r.get('職務') or '').strip(),
            'status':      (r.get('狀態') or '').strip(),
            'nationality': (r.get('國籍') or '').strip(),
            'hire_date':   (r.get('到職日') or '').strip(),
            'leave_date':  (r.get('離職日') or '').strip(),
            'remark':      (r.get('備註') or '').strip(),
        })

    cache_set(cache_key, roster)
    return roster


def fetch_efficiency_all_groups():
    """透過 SSRS SOAP Execution Service 取得加工部總效率(全部)報表 CSV"""
    import base64 as _b64
    import re as _re
    session = get_ssrs_session()
    soap_url = f'http://192.168.1.212/ReportServer/ReportExecution2005.asmx'

    def _soap(body, action):
        return session.post(soap_url, data=body.encode('utf-8'),
            headers={'Content-Type': 'text/xml; charset=utf-8',
                     'SOAPAction': f'"http://schemas.microsoft.com/sqlserver/2005/06/30/reporting/reportingservices/{action}"'},
            timeout=30)

    NS = 'xmlns:rs="http://schemas.microsoft.com/sqlserver/2005/06/30/reporting/reportingservices"'
    ENV_OPEN = f'<?xml version="1.0" encoding="utf-8"?><soap:Envelope xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/" {NS}>'

    # Step 1: LoadReport
    r1 = _soap(f'{ENV_OPEN}<soap:Body><rs:LoadReport><rs:Report>/加工課/加工部總效率</rs:Report></rs:LoadReport></soap:Body></soap:Envelope>', 'LoadReport')
    r1.raise_for_status()
    eid = _re.search(r'<ExecutionID>([^<]+)</ExecutionID>', r1.text)
    if not eid:
        raise RuntimeError('LoadReport: 無法取得 ExecutionID')
    exec_id = eid.group(1)

    # Step 2: SetExecutionParameters ORG=全部
    r2 = _soap(f'{ENV_OPEN}<soap:Header><rs:ExecutionHeader><rs:ExecutionID>{exec_id}</rs:ExecutionID></rs:ExecutionHeader></soap:Header><soap:Body><rs:SetExecutionParameters><rs:Parameters><rs:ParameterValue><rs:Name>ORG</rs:Name><rs:Value>全部</rs:Value></rs:ParameterValue></rs:Parameters><rs:ParameterLanguage>zh-TW</rs:ParameterLanguage></rs:SetExecutionParameters></soap:Body></soap:Envelope>', 'SetExecutionParameters')
    r2.raise_for_status()

    # Step 3: Render as CSV
    r3 = _soap(f'{ENV_OPEN}<soap:Header><rs:ExecutionHeader><rs:ExecutionID>{exec_id}</rs:ExecutionID></rs:ExecutionHeader></soap:Header><soap:Body><rs:Render><rs:Format>CSV</rs:Format><rs:DeviceInfo></rs:DeviceInfo></rs:Render></soap:Body></soap:Envelope>', 'Render')
    r3.raise_for_status()
    b64 = _re.search(r'<Result>([^<]+)</Result>', r3.text)
    if not b64:
        raise RuntimeError('Render: 無法取得 CSV 結果')
    csv_bytes = _b64.b64decode(b64.group(1))
    return csv_bytes.decode('utf-8-sig', errors='replace')


def parse_csv(csv_text):
    """解析 CSV 為 dict list"""
    reader = csv.reader(io.StringIO(csv_text))
    rows = list(reader)
    if not rows:
        return []
    header = rows[0]
    results = []
    for row in rows[1:]:
        if len(row) < len(header):
            continue
        results.append(dict(zip(header, row)))
    return results


UNIT_OPTIONS = {
    '000-1': '加工課',
    '000-3': '製三課',
    '000-2': '成品部',
}
KNOWN_UNITS = set(UNIT_OPTIONS.keys())


def search_orders(order_id='', product_name='', unit='000-1'):
    """從未完工製令報表搜尋，支援製令號碼、品名模糊搜尋和生產線篩選
    unit: '000-1'/'000-2'/'000-3'/'other'/'*'
    """
    all_records = fetch_all_unfinished()
    filtered = []
    for r in all_records:
        if order_id and order_id.upper() not in r.get('單別', '').upper():
            continue
        if product_name and not match_and_not(product_name, r.get('品名', '')):
            continue
        # 生產線篩選
        if unit and unit != '*':
            rec_unit = r.get('加工單位', '').strip()
            if unit == 'other':
                if rec_unit in KNOWN_UNITS:
                    continue
            else:
                if rec_unit != unit:
                    continue
        filtered.append(r)
    return filtered


def fetch_std_time(order_type, order_num):
    """從加工課-生產日報表取得標準工時和製程名稱（帶 TTL 快取）"""
    cache_key = f'std_{order_type}-{order_num}'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached
    try:
        params = {'p1': order_type, 'p2': order_num}
        csv_text = fetch_report_csv(config.REPORT_PATHS['daily'], params)
        reader = csv.reader(io.StringIO(csv_text))
        rows = list(reader)
        result = {}
        for row in rows[1:]:
            if len(row) < 6:
                continue
            std_text = row[1]       # "標準工時：44 PCS/H"
            proc_code_text = row[3] # "製程代號：N19"
            proc_name_text = row[5] # "製程名稱：鑽+攻+倒"
            m_std = re.search(r'標準工時：(.+)', std_text)
            m_proc = re.search(r'製程代號：(.+)', proc_code_text)
            m_name = re.search(r'製程名稱：(.+)', proc_name_text)
            if m_proc:
                code = m_proc.group(1).strip()
                result[code] = {
                    'std_time': m_std.group(1).strip() if m_std else '',
                    'process_name': m_name.group(1).strip() if m_name else '',
                }
        cache_set(cache_key, result)
        return result
    except Exception:
        return {}


def fetch_efficiency_data():
    """從加工部總效率報表取得 SFT 預計開工日/完工日/標工/即時生產量（帶 TTL 快取）
    回傳 dict: { '製令號碼': { '製程代號': {...} } }
    """
    cached = cache_get('efficiency')
    if cached is not None:
        return cached
    try:
        csv_text = fetch_report_csv(config.REPORT_PATHS['efficiency'], {})
        records = parse_csv(csv_text)
        result = {}
        for r in records:
            oid = r.get('製令', '').strip()
            proc = r.get('製程', '').strip()
            if not oid:
                continue
            if oid not in result:
                result[oid] = {}
            result[oid][proc] = {
                'SFT預計開工日': r.get('SFT預計開工日', '').strip(),
                'SFT預計完工日': r.get('SFT預計完工日', '').strip(),
                '標工_秒': r.get('ERP變動人時3', '').strip(),
                '生產機台': r.get('生產機台', '').strip(),
                '即時生產量': r.get('現在產量', '').strip(),
            }
        cache_set('efficiency', result)
        return result
    except Exception:
        return {}


# ── ServCloud 設備稼動 ─────────────────────────────────────────────
_servcloud_session = None
_servcloud_session_lock = threading.Lock()
SERVCLOUD_BASE = getattr(config, 'SERVCLOUD_BASE', 'http://192.168.1.69:58080/ServCloud')
SERVCLOUD_USER = getattr(config, 'SERVCLOUD_USER', 'adminstd')
SERVCLOUD_PASS = getattr(config, 'SERVCLOUD_PASS', 'adminstd')

_STATUS_MAP = {
    '0':  ('離線', '#bdbdbd'),
    '11': ('運轉', '#43a047'),
    '12': ('待機', '#fb8c00'),
    '13': ('警報', '#e53935'),
    'B':  ('離線', '#bdbdbd'),
}

def _sc_login(s):
    s.post(f'{SERVCLOUD_BASE}/api/user/login',
           data={'id': SERVCLOUD_USER, 'password': SERVCLOUD_PASS}, timeout=10)

def get_servcloud_session():
    global _servcloud_session
    with _servcloud_session_lock:
        if _servcloud_session is None:
            s = requests.Session()
            _sc_login(s)
            _servcloud_session = s
        return _servcloud_session

def _sc_post(path, payload, retry=True):
    """POST to ServCloud; auto re-login on session expiry (type==2)."""
    session = get_servcloud_session()
    resp = session.post(f'{SERVCLOUD_BASE}{path}', json=payload, timeout=30)
    data = resp.json()
    if data.get('type') == 2 and retry:
        global _servcloud_session
        with _servcloud_session_lock:
            _sc_login(session)
            _servcloud_session = session
        return _sc_post(path, payload, retry=False)
    return data

def _sc_parse(raw):
    """ServCloud data 欄位可能是 JSON 字串或已解析的 list/dict，統一回傳 list。"""
    if isinstance(raw, list):
        return raw
    if isinstance(raw, str):
        return json.loads(raw)
    return []

def servcloud_get_machines():
    data = _sc_post('/api/getdata/db',
                    {'table': 'm_device', 'columns': ['device_id', 'device_name']})
    machines = _sc_parse(data.get('data', []))
    return [m for m in machines if m.get('device_name') and m['device_name'] != m['device_id']]

def servcloud_get_history(machine_ids, date_str):
    """date_str: YYYYMMDD。最多 10 台/次，自動分批。"""
    all_recs = []
    for i in range(0, len(machine_ids), 10):
        batch = machine_ids[i:i+10]
        data = _sc_post('/api/hippo/simple', {
            'space': 'machine_status_history',
            'index': {'machine_id': batch},
            'indexRange': {'key': 'date', 'start': date_str, 'end': date_str},
            'columns': ['machine_id', 'status', 'start_time', 'end_time', 'duration']
        })
        if data.get('type') == 0 and data.get('data'):
            all_recs.extend(_sc_parse(data['data']))
    return all_recs


APP_VERSION = 'V20260626d'

@app.route('/ver')
def ver_check():
    """版本確認（純文字，繞過模板和快取）"""
    return f'<h1>{APP_VERSION}</h1><p>template_folder={app.template_folder}</p>'

@app.route('/')
def index():
    return render_template('index.html', app_version=APP_VERSION)


@app.route('/drawing')
def drawing_page():
    return render_template('drawing.html')


@app.route('/production')
def production_page():
    return render_template('production.html', app_version=APP_VERSION)


@app.route('/api/production', methods=['GET'])
def production_data():
    """生產中製令 API：用 SOAP 取得加工部總效率(全部)"""
    try:
        csv_text = fetch_efficiency_all_groups()
        records = parse_csv(csv_text)
        # 回傳原始欄位名稱供前端自動對應
        raw_columns = list(records[0].keys()) if records else []
        return jsonify({'success': True, 'data': records, 'count': len(records), 'columns': raw_columns})
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500


@app.route('/api/production/debug')
def production_debug():
    """除錯：顯示 SSRS CSV 原始欄位名稱"""
    try:
        csv_text = fetch_report_csv(config.REPORT_PATHS['efficiency'], {})
        reader = csv.reader(io.StringIO(csv_text))
        rows = list(reader)
        if not rows:
            return '<h2>CSV 空白</h2>'
        header = rows[0]
        html = '<h2>SSRS CSV 欄位名稱</h2><ol>'
        for i, col in enumerate(header):
            html += f'<li><b>{col}</b> (repr: {repr(col)})</li>'
        html += '</ol>'
        if len(rows) > 1:
            html += '<h3>第一筆資料</h3><table border="1" cellpadding="4"><tr><th>欄位</th><th>值</th></tr>'
            for col, val in zip(header, rows[1]):
                html += f'<tr><td>{col}</td><td>{val}</td></tr>'
            html += '</table>'
        return html
    except Exception as e:
        import traceback
        return f'<pre>Error: {traceback.format_exc()}</pre>', 500


@app.route('/api/query', methods=['GET'])
def query():
    """製令查詢 API：支援製令號碼和品名模糊搜尋"""
    order_id = request.args.get('order_id', '').strip()
    product_name_q = request.args.get('product_name', '').strip()
    unit = request.args.get('unit', '000-1').strip()

    if not order_id and not product_name_q:
        return jsonify({'success': False, 'error': '請輸入製令號碼或品名'}), 400

    try:
        # ── 第 1 步：取未完工製令 + 效率報表（並行） ──
        with ThreadPoolExecutor(max_workers=2) as pool:
            fut_orders = pool.submit(search_orders, order_id, product_name_q, unit)
            fut_eff = pool.submit(fetch_efficiency_data)
            records = fut_orders.result()
            eff_data = fut_eff.result()

        if not records:
            return jsonify({'success': False, 'error': '查無符合條件的資料'})

        # ── 第 2 步：日報表標工查詢（並行，限最多 20 筆） ──
        std_time_cache = {}
        unique_oids = list({r.get('單別', '') for r in records})
        oids_to_fetch = []
        for oid in unique_oids[:20]:
            oparts = oid.split('-', 1)
            if len(oparts) == 2:
                # 先檢查快取，只查沒快取的
                cache_key = f'std_{oparts[0]}-{oparts[1]}'
                cached = cache_get(cache_key)
                if cached is not None:
                    std_time_cache[oid] = cached
                else:
                    oids_to_fetch.append((oid, oparts[0], oparts[1]))

        if oids_to_fetch:
            with ThreadPoolExecutor(max_workers=min(8, len(oids_to_fetch))) as pool:
                futures = {
                    pool.submit(fetch_std_time, otype, onum): oid
                    for oid, otype, onum in oids_to_fetch
                }
                for fut in as_completed(futures):
                    oid = futures[fut]
                    try:
                        std_time_cache[oid] = fut.result()
                    except Exception:
                        std_time_cache[oid] = {}

        # ── 第 3 步：組合結果 ──
        data = []
        for r in records:
            rid = r.get('單別', '')
            proc_code = r.get('製程代號', '').strip()

            daily_info = std_time_cache.get(rid, {}).get(proc_code, {})
            eff_info = eff_data.get(rid, {}).get(proc_code, {})

            product_name = r.get('品名', '').rstrip('|').strip()

            std_time = daily_info.get('std_time', '')
            if not std_time and eff_info.get('標工_秒'):
                std_time = eff_info['標工_秒'] + ' 秒'
            # 標工：移除 "/H"，若只有 "- PCS" 等無數字則顯示空白
            if std_time:
                std_time = std_time.replace('/H', '').strip()
                if not re.search(r'\d', std_time):
                    std_time = ''

            realtime_qty = eff_info.get('即時生產量', '') or r.get('即時生產量', '')

            data.append({
                '製令': rid,
                '品號': r.get('品號', ''),
                '品名': product_name,
                '加工順序': r.get('加工順序', ''),
                '製程代號': proc_code,
                '製程名稱': daily_info.get('process_name', ''),
                '標工': std_time,
                '預計生產量': r.get('預計生產數', ''),
                '預計開工日': eff_info.get('SFT預計開工日', '') or r.get('製程預計開工日', ''),
                # SFT效率報表只有正在生產的製程；未開工製程用製程預計開工日作為估算完工日（has_sft_finish=False 表示估算值）
                '預計完工日': eff_info.get('SFT預計完工日', '') or r.get('製程預計開工日', ''),
                'has_sft_finish': bool(eff_info.get('SFT預計完工日', '')),
                '即時生產量': realtime_qty,
                '生產機台': eff_info.get('生產機台', ''),
            })

        return jsonify({'success': True, 'data': data})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/detail', methods=['GET'])
def detail():
    """製令詳細資訊 API：用料、製程、移轉單"""
    order_id = request.args.get('order_id', '').strip()
    if not order_id:
        return jsonify({'success': False, 'error': '請輸入製令號碼'}), 400

    try:
        params = {'製令': order_id}
        csv_text = fetch_report_csv(config.REPORT_PATHS['all'], params)

        # 報表回傳三段式 CSV（空行分隔）：材料、製程、移轉單
        csv_text = csv_text.replace('\r\n', '\n').replace('\r', '\n')
        sections = [s.strip() for s in csv_text.split('\n\n') if s.strip()]
        result = {'materials': [], 'processes': [], 'transfers': []}

        for section in sections:
            reader = csv.reader(io.StringIO(section))
            rows = list(reader)
            if len(rows) < 2:
                continue
            header = rows[0]
            data = []
            for row in rows[1:]:
                if len(row) >= len(header) and any(cell.strip() for cell in row):
                    data.append(dict(zip(header, row)))

            if '材料品號' in header:
                result['materials'] = data
            elif '完工否' in header:
                result['processes'] = data
            elif '移轉單' in header:
                result['transfers'] = data

        return jsonify({'success': True, 'data': result})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── 列印暫存（pywebview 無法用 window.open，改用 URL 方式）──
_print_cache = {}
_print_cache_lock = threading.Lock()


def _render_print_html(records):
    """將列印資料渲染成 HTML"""
    from datetime import datetime
    today = datetime.now().strftime('%Y/%m/%d %H:%M')

    std_time_cache = {}
    for r in records:
        order_id = r.get('order_id', '')
        if order_id and order_id not in std_time_cache:
            parts = order_id.split('-')
            if len(parts) >= 2:
                std_time_cache[order_id] = fetch_std_time(parts[0], parts[1])

    items = []
    for r in records:
        order_id = r.get('order_id', '')
        process_seq = r.get('process_seq', '')
        process_code = r.get('process_code', '')
        process_name = r.get('process_name', '')
        unit = r.get('unit', '').strip()
        product_name = r.get('product_name', '')
        qty = r.get('qty', '')

        proc_info = std_time_cache.get(order_id, {}).get(process_code, {})
        std_time = proc_info.get('std_time', '') if isinstance(proc_info, dict) else ''
        actual_process_name = proc_info.get('process_name', '') if isinstance(proc_info, dict) else ''
        if not actual_process_name:
            actual_process_name = process_name

        barcode = f"{order_id};{process_seq};{process_code};{unit}"

        items.append({
            'order_id': f"{order_id}-{process_seq}",
            'product_name': product_name,
            'process_code': process_code,
            'process_name': actual_process_name,
            'qty': qty,
            'std_time': std_time,
            'barcode': barcode,
            'machine': r.get('machine', ''),
        })

    return render_template('print_report.html',
                           items=items,
                           items_json=json.dumps(items, ensure_ascii=False),
                           today=today)


@app.route('/print', methods=['POST'])
def print_report():
    """生產日報表列印頁面 — POST 暫存 HTML，回傳 key 供 GET 取用"""
    try:
        records = request.get_json() or []
    except Exception:
        return jsonify({'success': False, 'error': '資料格式錯誤'}), 400

    html = _render_print_html(records)
    key = str(uuid.uuid4())
    with _print_cache_lock:
        _print_cache[key] = html
    return jsonify({'success': True, 'key': key})


@app.route('/print/view/<key>', methods=['GET'])
def print_view(key):
    """取得暫存的列印 HTML"""
    with _print_cache_lock:
        html = _print_cache.get(key)
        # 列出目前快取的 keys 用於除錯
        cached_keys = list(_print_cache.keys())
    if html:
        return html
    return f'列印資料不存在。要求 key={key}，目前快取 keys={cached_keys}', 404


@app.route('/api/inspection', methods=['GET'])
def find_inspection():
    """根據品號查找自主巡檢表"""
    product_no = request.args.get('product_no', '').strip()
    if not product_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        root = Path(INSPECTION_ROOT)
        if not root.exists():
            return jsonify({'success': False, 'error': '無法存取巡檢表網路資料夾'}), 500

        found = []
        for f in root.rglob('*.xls*'):
            fname = f.stem
            if product_no.upper() in fname.upper():
                found.append({
                    'file_path': str(f),
                    'file_name': f.name,
                    'dir': str(f.parent),
                    'need_inspect': '需送檢' in fname,
                })

        if found:
            return jsonify({'success': True, 'files': found, 'count': len(found)})
        else:
            return jsonify({'success': False, 'error': '無對應自主檢查表', 'files': []})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/inspection/open', methods=['POST'])
def open_inspection():
    """開啟自主巡檢表，並將製令、預計產量寫入 Excel 指定欄位
    - 單張製令：G6 寫製令號碼（垂直置中），J6+K6 寫預計產量
    - 多張製令：G6 寫所有製令號碼（換行，垂直置中），J6/K6 不寫
    """
    import time as _time
    data = request.get_json() or {}
    file_path = data.get('file_path', '').strip()
    order_ids = data.get('order_ids', [])       # 製令號碼列表
    qty = data.get('qty', '').strip()            # 預計產量（僅單張時寫入）

    if not file_path:
        return jsonify({'success': False, 'error': '請提供檔案路徑'}), 400

    try:
        import pythoncom
        import win32com.client

        # 使用 STA threading model，避免 RPC_E_CALL_REJECTED 錯誤
        pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
        try:
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = True
            wb = excel.Workbooks.Open(file_path)

            # ── COM retry 機制：Excel 開啟後可能忙碌，需重試 ──
            max_retries = 3
            ws = None
            for attempt in range(max_retries):
                try:
                    _time.sleep(2)  # 等待 Excel 載入完成
                    ws = wb.ActiveSheet

                    # xlVAlignCenter = -4108
                    if order_ids:
                        cell_g6 = ws.Range("G6")
                        cell_g6.Value = '\n'.join(order_ids)
                        cell_g6.WrapText = True
                        cell_g6.VerticalAlignment = -4108   # 垂直置中
                        cell_g6.ShrinkToFit = False
                        # 自動調整文字大小以符合儲存格（避免跨到下一排）
                        if len(order_ids) > 1:
                            cell_g6.ShrinkToFit = True

                        # 同時填入 J6 和 K6 預計產量
                        if qty:
                            for cell_name in ["J6", "K6"]:
                                cell = ws.Range(cell_name)
                                cell.Value = qty
                                cell.VerticalAlignment = -4108

                    break  # 成功，跳出重試迴圈
                except pythoncom.com_error as ce:
                    if attempt < max_retries - 1:
                        _time.sleep(2)  # 等待後重試
                    else:
                        raise  # 最後一次仍失敗，拋出錯誤

            return jsonify({'success': True, 'message': f'已開啟並寫入 {os.path.basename(file_path)}'})
        finally:
            pythoncom.CoUninitialize()
    except ImportError:
        # pywin32 未安裝（如 Embeddable Python 環境）→ 直接用 os.startfile 開啟
        try:
            os.startfile(file_path)
            return jsonify({'success': True,
                            'message': f'已開啟 {os.path.basename(file_path)}（此電腦無法自動填入製令號碼，請手動輸入）'})
        except Exception as e2:
            return jsonify({'success': False, 'error': f'無法開啟檔案：{e2}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def get_pdm_db():
    """取得 PDMSearch SQLite 連線（自動嘗試多個路徑）"""
    candidates = [
        PDM_DB_PATH,                                                          # config.py 設定值
        os.path.join(_APP_DIR, 'pdm_search.db'),                              # 隨包附帶 (_app/)
        os.path.join(os.path.dirname(_APP_DIR), 'pdm_search.db'),             # 上一層 (製令查詢/)
        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'PDMSearch', 'pdm_search.db'),  # 使用者 AppData
    ]
    for path in candidates:
        if path and os.path.isfile(path):
            conn = sqlite3.connect(path)
            conn.row_factory = sqlite3.Row
            return conn
    return None


@app.route('/api/debug/paths')
def debug_paths():
    """診斷用：顯示 PDM 路徑解析結果"""
    candidates = [
        ('config.py 設定值', PDM_DB_PATH),
        ('_APP_DIR/pdm_search.db', os.path.join(_APP_DIR, 'pdm_search.db')),
        ('上一層/pdm_search.db', os.path.join(os.path.dirname(_APP_DIR), 'pdm_search.db')),
        ('LOCALAPPDATA', os.path.join(os.environ.get('LOCALAPPDATA', ''), 'PDMSearch', 'pdm_search.db')),
        ('sys.executable 同層', os.path.join(os.path.dirname(sys.executable), 'pdm_search.db')),
        ('__file__ 目錄', os.path.dirname(os.path.abspath(__file__))),
    ]
    result = {
        '_APP_DIR': _APP_DIR,
        '_BASE': _BASE,
        'PDM_DB_PATH': PDM_DB_PATH,
        'sys.executable': sys.executable,
        '__file__': os.path.abspath(__file__),
        'cwd': os.getcwd(),
        'candidates': []
    }
    for label, path in candidates:
        result['candidates'].append({
            'label': label,
            'path': path,
            'exists': os.path.isfile(path) if not label.endswith('目錄') else os.path.isdir(path),
        })
    # 列出 _APP_DIR 裡有什麼檔案
    try:
        result['_APP_DIR_files'] = os.listdir(_APP_DIR)
    except Exception as e:
        result['_APP_DIR_files'] = str(e)
    return jsonify(result)


@app.route('/api/drawing', methods=['GET'])
def find_drawing():
    """根據品號從 PDMSearch 索引資料庫查找圖面"""
    product_no = request.args.get('product_no', '').strip()
    if not product_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        conn = get_pdm_db()
        if not conn:
            return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

        cur = conn.cursor()
        cur.execute(
            """SELECT file_path, pin_hao, tu_mian_pin_ming, ji_xing, xing_hao, modified_at
               FROM drawing_index WHERE UPPER(pin_hao) = UPPER(?)""",
            (product_no,)
        )
        rows = cur.fetchall()
        conn.close()

        if rows:
            files = _rows_to_drawing_list(rows)
            return jsonify({'success': True, 'files': files, 'count': len(files)})
        else:
            return jsonify({'success': False, 'error': f'PDM 索引中找不到品號 {product_no} 的圖面', 'files': []})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/drawing/search', methods=['GET'])
def search_drawing():
    """圖面模糊搜尋 API：支援品號或品名模糊搜尋（AND/NOT 語法）"""
    q = request.args.get('q', '').strip()
    if not q:
        return jsonify({'success': False, 'error': '請輸入品號或品名'}), 400

    try:
        conn = get_pdm_db()
        if not conn:
            return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

        cur = conn.cursor()
        # 取出所有的關鍵字（排除 NOT 語法用的）做 SQL LIKE 初篩
        tokens = q.strip().split()
        positive = [t for t in tokens if not t.startswith('-')]
        conditions = []
        params = []
        for t in positive[:3]:  # 最多取前3個正向關鍵字做 SQL 篩選
            conditions.append(
                "(UPPER(pin_hao) LIKE ? OR UPPER(tu_mian_pin_ming) LIKE ? OR UPPER(file_path) LIKE ?)"
            )
            like = f'%{t.upper()}%'
            params.extend([like, like, like])

        where = ' AND '.join(conditions) if conditions else '1=1'
        cur.execute(
            f"""SELECT file_path, pin_hao, tu_mian_pin_ming, ji_xing, xing_hao, modified_at
                FROM drawing_index WHERE {where} LIMIT 200""",
            params
        )
        rows = cur.fetchall()
        conn.close()

        # 再用 match_and_not 做精確的 AND/NOT 過濾
        files = []
        for row in rows:
            combined = f"{row['pin_hao'] or ''} {row['tu_mian_pin_ming'] or ''} {row['file_path'] or ''}"
            if match_and_not(q, combined):
                files.append(_row_to_drawing_dict(row))

        if files:
            return jsonify({'success': True, 'files': files, 'count': len(files)})
        else:
            return jsonify({'success': False, 'error': '查無符合條件的圖面', 'files': []})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _row_to_drawing_dict(row):
    """將 SQLite Row 轉為圖面 dict"""
    fp = row['file_path']
    return {
        'file_path': fp,
        'file_name': os.path.basename(fp),
        'dir': os.path.dirname(fp),
        'pin_hao': row['pin_hao'] or '',
        'product_name': row['tu_mian_pin_ming'] or '',
        'ji_xing': row['ji_xing'] or '',
        'xing_hao': row['xing_hao'] or '',
        'modified_at': row['modified_at'] or '',
    }


def _rows_to_drawing_list(rows):
    """將多筆 SQLite Row 轉為圖面 dict list"""
    return [_row_to_drawing_dict(row) for row in rows]


def pdm_get_file(file_path):
    """用 PowerShell 呼叫 PDM COM API 取出檔案（Python COM 的 GetFileCopy 有型別問題，改用 PowerShell）"""
    import subprocess

    # PowerShell 腳本：透過 PDM COM 取出檔案到本機快取
    ps_script = f'''
$ErrorActionPreference = "Stop"
try {{
    $vault = New-Object -ComObject "ConisioLib.EdmVault"
    $vault.LoginAuto("MAXCLAW", 0)
    $path = "{file_path.replace(chr(92), chr(92)+chr(92))}"
    $dir = [System.IO.Path]::GetDirectoryName($path)
    $fname = [System.IO.Path]::GetFileName($path)
    $folder = $vault.GetFolderFromPath($dir)
    if ($null -eq $folder) {{
        Write-Error "PDM 找不到資料夾: $dir"
        exit 1
    }}
    $fileObj = $folder.GetFile($fname)
    if ($null -eq $fileObj) {{
        Write-Error "PDM 找不到檔案: $fname"
        exit 1
    }}
    $fileObj.GetFileCopy(0)
    Write-Output "OK"
}} catch {{
    Write-Error $_.Exception.Message
    exit 1
}}
'''
    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-Command', ps_script],
            capture_output=True, text=True, timeout=30
        )

        if result.returncode == 0 and 'OK' in result.stdout:
            return True, 'OK'
        else:
            err = result.stderr.strip()
            if not err:
                err = result.stdout.strip() or '未知錯誤'
            return False, err
    except subprocess.TimeoutExpired:
        return False, 'PDM 取出逾時'
    except Exception as e:
        return False, str(e)


@app.route('/api/drawing/open', methods=['POST'])
def open_drawing():
    """開啟 PDM 圖面檔案"""
    data = request.get_json() or {}
    file_path = data.get('file_path', '').strip()

    if not file_path:
        return jsonify({'success': False, 'error': '請提供圖面路徑'}), 400

    try:
        if os.path.exists(file_path):
            # 檔案已在本機（或可直接存取的網路路徑）
            # 嘗試 PDM 取最新版，但失敗不阻擋開啟
            pdm_get_file(file_path)
            os.startfile(file_path)
            return jsonify({'success': True, 'message': f'已開啟 {os.path.basename(file_path)}'})
        else:
            # 本機沒有此檔案，嘗試透過 PDM COM 取出
            ok, msg = pdm_get_file(file_path)
            if ok:
                os.startfile(file_path)
                return jsonify({'success': True, 'message': f'已開啟 {os.path.basename(file_path)}'})
            else:
                # PDM 取出失敗（可能未安裝 PDM）→ 仍嘗試直接開啟
                try:
                    os.startfile(file_path)
                    return jsonify({'success': True, 'message': f'已開啟 {os.path.basename(file_path)}'})
                except Exception:
                    pdm_hint = '（此電腦未安裝 SolidWorks PDM，無法自動取出圖面）' if 'ConisioLib' in msg or 'EdmVault' in msg or '無法建立' in msg else ''
                    return jsonify({'success': False,
                                    'error': f'無法開啟圖面：{msg}{pdm_hint}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── PDM 索引重建 ───────────────────────────────────────────────────────────────

import threading as _threading
import subprocess as _subprocess
import re as _re

_reindex_state = {
    'running':    False,
    'phase':      'idle',   # idle | scanning | indexing | done | error
    'scanned':    0,
    'total':      0,
    'indexed':    0,
    'message':    '',
    'error':      '',
    'last_run':   None,
    'last_count': 0,
}
_reindex_lock = _threading.Lock()

# 持久化狀態檔：重啟後仍可讀取最後一次 reindex 結果
_REINDEX_STATE_FILE = os.path.join(_APP_DIR, 'last_reindex.json')


def _save_last_reindex_state(ts: str, cnt: int):
    """reindex 完成後將時間與筆數寫入 JSON（重啟後仍保留）"""
    try:
        with open(_REINDEX_STATE_FILE, 'w', encoding='utf-8') as f:
            json.dump({'last_run': ts, 'last_count': cnt}, f)
    except Exception:
        pass


def _load_last_reindex_state():
    """從 JSON 讀取最後一次 reindex 結果（比 DB indexed_at 更可靠）"""
    try:
        with open(_REINDEX_STATE_FILE, encoding='utf-8') as f:
            d = json.load(f)
            return d.get('last_run'), int(d.get('last_count', 0))
    except Exception:
        return None, 0


def _run_reindex(update_only: bool):
    """背景執行索引重建，解析 stdout 更新進度狀態"""
    global _reindex_state
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_pdm_index.py')
    cmd = [sys.executable, script] + (['--update'] if update_only else [])

    with _reindex_lock:
        _reindex_state.update(running=True, phase='scanning', scanned=0,
                              total=0, indexed=0, message='啟動中...', error='')

    try:
        proc = _subprocess.Popen(
            cmd, stdout=_subprocess.PIPE, stderr=_subprocess.STDOUT,
            text=True, encoding='utf-8', errors='replace'
        )
        for line in proc.stdout:
            line = line.rstrip()
            if not line:
                continue
            with _reindex_lock:
                _reindex_state['message'] = line

            # 解析掃描進度：「掃描中... 1,200 個檔案」
            m = _re.search(r'掃描中.*?(\d[\d,]*)\s*個', line)
            if m:
                with _reindex_lock:
                    _reindex_state['phase']   = 'scanning'
                    _reindex_state['scanned'] = int(m.group(1).replace(',', ''))
                continue

            # 解析找到總數：「找到 X 個 .SLDDRW」
            m = _re.search(r'找到\s*(\d[\d,]*)\s*個', line)
            if m:
                with _reindex_lock:
                    _reindex_state['total'] = int(m.group(1).replace(',', ''))
                continue

            # 解析寫入進度：「 33%  1000/3000  新增:X」
            m = _re.search(r'(\d+)%.*?(\d[\d,]*)/(\d[\d,]*)', line)
            if m:
                with _reindex_lock:
                    _reindex_state['phase']   = 'indexing'
                    _reindex_state['indexed'] = int(m.group(2).replace(',', ''))
                    _reindex_state['total']   = int(m.group(3).replace(',', ''))
                continue

            # 完成行：「完成！新增:X  更新:X  ...」
            m = _re.search(r'完成.*新增:(\d[\d,]*)', line)
            if m:
                with _reindex_lock:
                    _reindex_state['indexed'] = int(m.group(1).replace(',', ''))

        proc.wait()

        if proc.returncode == 0:
            # 取得資料庫最新筆數
            try:
                conn_tmp = get_pdm_db()
                if conn_tmp:
                    cnt = conn_tmp.execute('SELECT COUNT(*) FROM drawing_index').fetchone()[0]
                    conn_tmp.close()
                else:
                    cnt = _reindex_state['indexed']
            except Exception:
                cnt = _reindex_state['indexed']

            ts_done = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            with _reindex_lock:
                _reindex_state.update(
                    running=False, phase='done', last_count=cnt,
                    last_run=ts_done,
                    message=f'更新完成，共 {cnt:,} 筆圖面資料'
                )
            # 持久化到 JSON，重啟後仍可讀取
            _save_last_reindex_state(ts_done, cnt)
        else:
            with _reindex_lock:
                _reindex_state.update(running=False, phase='error',
                                      error='重建失敗，請查看伺服器日誌')
    except Exception as exc:
        with _reindex_lock:
            _reindex_state.update(running=False, phase='error', error=str(exc))


@app.route('/api/drawing/reindex', methods=['POST'])
def drawing_reindex():
    """啟動 PDM 圖面索引重建（背景執行）"""
    with _reindex_lock:
        if _reindex_state['running']:
            return jsonify({'success': False, 'error': '索引重建已在執行中，請稍候'}), 409

    data        = request.get_json() or {}
    update_only = data.get('update_only', True)   # 預設增量更新

    t = _threading.Thread(target=_run_reindex, args=(update_only,), daemon=True)
    t.start()
    return jsonify({'success': True, 'message': '索引重建已啟動'})


def _pdm_last_indexed():
    """從 drawing_index 讀取最後一次 indexed_at（持久，不受重啟影響）"""
    try:
        c = get_pdm_db()
        if not c:
            return None, 0
        # 舊版 DB 可能沒有 indexed_at 欄位，ALTER TABLE 補上（冪等）
        try:
            c.execute("ALTER TABLE drawing_index ADD COLUMN indexed_at TEXT")
        except Exception:
            pass  # 欄位已存在
        row = c.execute(
            "SELECT MAX(indexed_at), COUNT(*) FROM drawing_index"
        ).fetchone()
        c.close()
        return (row[0] or None), (row[1] or 0)
    except Exception:
        return None, 0


@app.route('/api/drawing/reindex/status', methods=['GET'])
def drawing_reindex_status():
    """回傳索引重建進度狀態（含 DB 持久化的 last_run）"""
    with _reindex_lock:
        state = dict(_reindex_state)
    # 若記憶體中 last_run 為空（剛重啟），依序從 JSON → DB 讀取並快取
    if not state.get('last_run'):
        ts, cnt = _load_last_reindex_state()    # 優先：JSON 狀態檔
        if not ts:
            ts, cnt = _pdm_last_indexed()        # 備援：DB indexed_at
        if ts:
            with _reindex_lock:
                _reindex_state['last_run']   = ts
                _reindex_state['last_count'] = cnt
            state['last_run']   = ts
            state['last_count'] = cnt
    return jsonify(state)


@app.route('/api/zume/status')
def zume_status():
    """回傳技術資料清單的目前狀態（件數、最後匯入檔案）"""
    try:
        con = sqlite3.connect(ZUME_DB_PATH)
        total = con.execute('SELECT COUNT(*) FROM drawings').fetchone()[0]
        last  = con.execute('SELECT filename, imported_at FROM import_log ORDER BY imported_at DESC LIMIT 1').fetchone()
        con.close()
        return jsonify({'total': total, 'last_file': last[0] if last else None, 'last_at': last[1] if last else None})
    except Exception as e:
        return jsonify({'total': 0, 'last_file': None, 'last_at': None, 'error': str(e)})


@app.route('/api/zume/scan', methods=['POST'])
def zume_scan():
    """重新掃描下載/桌面等資料夾並匯入最新 CSV（強制重掃）"""
    files = _find_zume_csvs()
    if not files:
        dirs = '、'.join(_zume_csv_search_dirs()) or os.path.join(os.path.expanduser('~'), 'Downloads')
        return jsonify({'success': False,
                        'error': f'找不到 zume-n_data_list_*.csv\n請先從 ZUMEN 匯出清單 CSV 到「下載」資料夾。\n已搜尋：{dirs}'}), 404
    latest = files[0]
    fname  = os.path.basename(latest)
    recs   = _parse_zume_csv(latest)
    if not recs:
        return jsonify({'success': False, 'error': f'檔案解析失敗或無資料：{fname}'}), 400
    con = sqlite3.connect(ZUME_DB_PATH)
    _insert_zume_recs(con, recs)
    con.execute('INSERT OR REPLACE INTO import_log(filename,imported_at,count) VALUES(?,datetime("now"),?)',
                (fname, len(recs)))
    con.commit()
    total = con.execute('SELECT COUNT(*) FROM drawings').fetchone()[0]
    con.close()
    return jsonify({'success': True, 'imported': len(recs), 'total': total, 'file': fname})


@app.route('/api/zume/import', methods=['POST'])
def zume_import():
    """手動上傳 CSV/XLSX 匯入（保留作為備用方式）"""
    f = request.files.get('file')
    if not f:
        return jsonify({'success': False, 'error': '未上傳檔案'}), 400
    fname_lower = f.filename.lower()
    recs = []
    try:
        if fname_lower.endswith('.csv'):
            content = f.read().decode('utf-8-sig')
            reader  = csv.reader(io.StringIO(content))
            headers = [h.strip() for h in next(reader)]
            idx     = _zume_header_indices(headers)
            if idx['part_no'] is None or idx['url'] is None:
                return jsonify({'success': False, 'error': f'找不到「圖號」或「URL」欄'}), 400
            for row in reader:
                rec = _zume_row_to_rec(row, idx)
                if rec:
                    recs.append(rec)
        elif fname_lower.endswith(('.xlsx', '.xlsm')):
            if not _OPENPYXL_OK:
                return jsonify({'success': False, 'error': '請改用 CSV 格式'}), 400
            wb = openpyxl.load_workbook(io.BytesIO(f.read()), read_only=True)
            ws = wb.active
            headers = [str(c.value or '').strip() for c in next(ws.iter_rows(min_row=1, max_row=1))]
            idx     = _zume_header_indices(headers)
            if idx['part_no'] is None or idx['url'] is None:
                return jsonify({'success': False, 'error': '找不到「圖號」或「URL」欄'}), 400
            for row in ws.iter_rows(min_row=2, values_only=True):
                rec = _zume_row_to_rec(row, idx)
                if rec:
                    recs.append(rec)
            wb.close()
        else:
            return jsonify({'success': False, 'error': '僅支援 .csv 或 .xlsx'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': f'解析失敗：{e}'}), 500
    if not recs:
        return jsonify({'success': False, 'error': '檔案無有效資料'}), 400
    fname = f.filename
    con = sqlite3.connect(ZUME_DB_PATH)
    _insert_zume_recs(con, recs)
    con.execute('INSERT OR REPLACE INTO import_log(filename,imported_at,count) VALUES(?,datetime("now"),?)',
                (fname, len(recs)))
    con.commit()
    total = con.execute('SELECT COUNT(*) FROM drawings').fetchone()[0]
    con.close()
    return jsonify({'success': True, 'imported': len(recs), 'total': total})


@app.route('/api/zume/lookup', methods=['POST'])
def zume_lookup():
    """查詢品號對應的 zume-n.com URL；找不到則回傳搜尋 URL"""
    try:
        data     = request.get_json() or {}
        item_nos = data.get('item_nos', [])
        if not item_nos:
            return jsonify({'success': False, 'error': '請提供品號'}), 400
        con = sqlite3.connect(ZUME_DB_PATH)
        result = []
        for no in item_nos[:20]:
            no = no.strip()
            row = con.execute('SELECT part_name,url FROM drawings WHERE part_no=?', (no,)).fetchone()
            if row:
                result.append({'no': no, 'name': row[0], 'url': row[1], 'found': True})
            else:
                fallback_url = 'https://zume-n.com/freeword_search?query=' + urllib.parse.quote(no, safe='') + '&searchType=drawing'
                result.append({'no': no, 'name': '', 'url': fallback_url, 'found': False})
        con.close()
        return jsonify({'success': True, 'items': result})
    except Exception as e:
        return jsonify({'success': False, 'error': f'查詢失敗：{str(e)}'}), 500


@app.route('/api/zume/open', methods=['POST'])
def zume_open():
    """以系統預設瀏覽器開啟 zume-n.com（優先用清單的直接 URL）"""
    try:
        data     = request.get_json() or {}
        item_nos = data.get('item_nos', [])
        if not item_nos:
            return jsonify({'success': False, 'error': '請提供品號'}), 400
        con = sqlite3.connect(ZUME_DB_PATH)
        opened = []; not_found = []
        for no in item_nos[:10]:
            no = no.strip()
            row = con.execute('SELECT url FROM drawings WHERE part_no=?', (no,)).fetchone()
            url = row[0] if row else ('https://zume-n.com/freeword_search?query=' + urllib.parse.quote(no, safe='') + '&searchType=drawing')
            webbrowser.open(url, new=2)
            if row:
                opened.append(no)
            else:
                not_found.append(no)
        con.close()
        return jsonify({'success': True, 'opened': opened, 'not_found': not_found})
    except Exception as e:
        return jsonify({'success': False, 'error': f'開啟失敗：{str(e)}'}), 500


@app.route('/api/zume/list')
def zume_list():
    """回傳所有匯入的 ZUMEN 圖面（含生產線別/群組/分類/廠商）+ 篩選選項 + 最後匯入資訊"""
    try:
        con = sqlite3.connect(ZUME_DB_PATH)
        con.row_factory = sqlite3.Row
        rows = [{
            'part_no':    r['part_no'],
            'part_name':  r['part_name'] or '',
            'url':        r['url'],
            'line':       r['line'] or '',
            'prod_group': r['prod_group'] or '',
            'category':   r['category'] or '',
            'vendor':     r['vendor'] or '',
        } for r in con.execute(
            'SELECT part_no, part_name, url, line, prod_group, category, vendor '
            'FROM drawings ORDER BY part_no')]
        # 各欄不重複值（給前端下拉用）
        def distinct(col):
            return sorted({(r[0] or '').strip() for r in
                           con.execute(f'SELECT DISTINCT {col} FROM drawings') if (r[0] or '').strip()})
        options = {c: distinct(c) for c in _ZUME_EXTRA_COLS}
        last = con.execute(
            'SELECT filename, imported_at FROM import_log ORDER BY imported_at DESC LIMIT 1').fetchone()
        con.close()
        return jsonify({'success': True, 'count': len(rows), 'drawings': rows,
                        'options': options,
                        'last_file': last[0] if last else None,
                        'last_at':   last[1] if last else None})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取失敗：{str(e)}'}), 500


@app.route('/equipment')
def equipment_page():
    return render_template('equipment.html', app_version=APP_VERSION)


@app.route('/api/equipment/machines')
def equipment_machines():
    try:
        machines = servcloud_get_machines()
        return jsonify({'success': True, 'data': machines})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/equipment/history')
def equipment_history():
    date_str    = request.args.get('date', '')           # YYYY-MM-DD
    start_hour  = int(request.args.get('start_hour', 7))
    end_hour    = int(request.args.get('end_hour', 20))
    machine_ids = request.args.getlist('machines')

    if not date_str or not machine_ids:
        return jsonify({'success': False, 'error': '請提供日期和機台'}), 400

    date_api = date_str.replace('-', '')

    try:
        history = servcloud_get_history(machine_ids, date_api)
        all_machines = servcloud_get_machines()
        name_map = {m['device_id']: m['device_name'] for m in all_machines}

        range_start_s = start_hour * 3600   # 秒
        range_end_s   = end_hour   * 3600
        total_min     = (range_end_s - range_start_s) // 60

        # Group segments by machine (秒精度)
        buckets = {mid: [] for mid in machine_ids}
        for rec in history:
            mid = rec.get('machine_id', '')
            if mid not in buckets:
                continue
            st = rec.get('start_time', '')
            et = rec.get('end_time',   '')
            if len(st) < 14 or len(et) < 14:
                continue
            s_sec = int(st[8:10])*3600 + int(st[10:12])*60 + int(st[12:14])
            e_sec = int(et[8:10])*3600 + int(et[10:12])*60 + int(et[12:14])
            # Handle overnight
            if e_sec < s_sec:
                e_sec += 24 * 3600
            s_clip = max(s_sec, range_start_s)
            e_clip = min(e_sec, range_end_s)
            if s_clip >= e_clip:
                continue
            status = rec.get('status', '0')
            status_name, color = _STATUS_MAP.get(status, ('未知', '#9e9e9e'))
            dur_sec = e_clip - s_clip
            buckets[mid].append({
                'status':      status,
                'status_name': status_name,
                'color':       color,
                'start_min':   (s_clip - range_start_s) / 60,   # float，用於圖表定位
                'end_min':     (e_clip - range_start_s) / 60,
                'start_str':   f"{s_clip//3600%24:02d}:{s_clip%3600//60:02d}:{s_clip%60:02d}",
                'end_str':     f"{e_clip//3600%24:02d}:{e_clip%3600//60:02d}:{e_clip%60:02d}",
                'dur_sec':     dur_sec,
            })

        result = [
            {'id': mid, 'name': name_map.get(mid, mid), 'segments': buckets[mid]}
            for mid in machine_ids
        ]
        return jsonify({
            'success': True, 'machines': result,
            'date': date_str, 'start_hour': start_hour,
            'end_hour': end_hour, 'total_minutes': total_min
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500


# ── BOM 查詢系統 ───────────────────────────────────────────────────
# 資料庫：192.168.1.140 / YC01（Computech ERP）
# 品號主檔：INVMB  材料BOM：BOMMD（MCPI11模組）
_ERP_SQL_SERVER   = getattr(config, 'ERP_SQL_SERVER',   '192.168.1.140')
_ERP_SQL_DATABASE = getattr(config, 'ERP_SQL_DATABASE', 'YC01')
_ERP_SQL_USERNAME = getattr(config, 'ERP_SQL_USERNAME', 'sa')
_ERP_SQL_PASSWORD = getattr(config, 'ERP_SQL_PASSWORD', 'dsc55877948')

# 屬性代碼對照表（INVMB.MB025）
_ATTR_MAP = {'M': 'M:自製件', 'S': 'S:訂外加工', 'P': 'P:採購件',
             'R': 'R:委外', 'F': 'F:虛擬件'}


def get_erp_conn():
    """建立 ERP SQL Server 連線（pyodbc）"""
    try:
        import pyodbc
    except ImportError:
        raise RuntimeError('未安裝 pyodbc，請執行: pip install pyodbc')

    cs = (f'DRIVER={{ODBC Driver 17 for SQL Server}};'
          f'SERVER={_ERP_SQL_SERVER};DATABASE={_ERP_SQL_DATABASE};'
          f'UID={_ERP_SQL_USERNAME};PWD={_ERP_SQL_PASSWORD};'
          f'TrustServerCertificate=yes;Connection Timeout=8;')
    return pyodbc.connect(cs, timeout=8)


@app.route('/routing')
def routing_page():
    return render_template('routing.html', app_version=APP_VERSION)


@app.route('/api/routing/search')
def routing_search():
    """搜尋有途程的品號（INVMB + BOMME，多關鍵字空格分隔，品號/品名同時模糊搜尋）
       line 參數：指定線別篩選（MF006），空字串代表全部"""
    q    = request.args.get('q', '').strip()
    line = request.args.get('line', '').strip()   # '000-1','000-2','000-3' 或 ''=全部
    if not q:
        return jsonify({'success': False, 'error': '請輸入品號或品名'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()
        # 多關鍵字：以空格分割，每個關鍵字都要符合（AND 邏輯）
        keywords = [kw for kw in q.split() if kw]
        where_parts = []
        params = []
        for kw in keywords:
            like = f'%{kw}%'
            where_parts.append("(RTRIM(b.MB001) LIKE ? OR b.MB002 LIKE ?)")
            params.extend([like, like])
        keyword_clause = ' AND '.join(where_parts) if where_parts else '1=1'

        # 線別篩選：只顯示「最新版次」途程中包含該線別至少一步的品號
        if line:
            line_clause = """AND EXISTS (
                SELECT 1 FROM BOMMF f2
                WHERE RTRIM(f2.MF001) = RTRIM(b.MB001)
                  AND f2.MF002 = (
                      SELECT MAX(e2.ME002) FROM BOMME e2
                      WHERE RTRIM(e2.ME001) = RTRIM(b.MB001)
                  )
                  AND RTRIM(f2.MF006) = ?
            )"""
            params.append(line)
        else:
            line_clause = ''

        cur.execute(f"""
            SELECT TOP 200
                RTRIM(b.MB001)            AS item_no,
                RTRIM(ISNULL(b.MB002,'')) AS item_name,
                RTRIM(ISNULL(b.MB003,'')) AS spec,
                ISNULL(b.MB004,'')        AS unit,
                ISNULL(b.MB025,'')        AS attr_code
            FROM INVMB b
            WHERE EXISTS (
                SELECT 1 FROM BOMME e WHERE RTRIM(e.ME001) = RTRIM(b.MB001)
            )
            AND {keyword_clause}
            {line_clause}
            ORDER BY b.MB001
        """, params)
        rows = cur.fetchall()
        conn.close()
        results = [{
            'item_no':   r[0],
            'item_name': r[1],
            'spec':      r[2],
            'unit':      r[3],
            'attr':      _ATTR_MAP.get(str(r[4]).strip(), str(r[4]).strip()),
        } for r in rows]
        return jsonify({'success': True, 'data': results, 'count': len(results)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/bom')
def bom_page():
    # 舊連結相容：直接導向產品途程查詢頁（BOM子分頁）
    return redirect('/routing?tab=bom')


@app.route('/api/bom/search')
def bom_search():
    """搜尋有材料BOM的品號（INVMB + BOMMD，品號/品名模糊搜尋，可篩選品號開頭）"""
    q      = request.args.get('q', '').strip()
    prefix = request.args.get('prefix', '').strip()   # '5', '6', '56' 或空
    if not q:
        return jsonify({'success': False, 'error': '請輸入品號或品名'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()
        like = f'%{q}%'

        # 品號開頭篩選條件
        if prefix == '5':
            prefix_cond = "AND RTRIM(b.MB001) LIKE '5%'"
        elif prefix == '6':
            prefix_cond = "AND RTRIM(b.MB001) LIKE '6%'"
        elif prefix in ('56', '65'):
            prefix_cond = "AND (RTRIM(b.MB001) LIKE '5%' OR RTRIM(b.MB001) LIKE '6%')"
        else:
            prefix_cond = ''

        cur.execute(f"""
            SELECT TOP 200
                RTRIM(b.MB001)          AS item_no,
                RTRIM(ISNULL(b.MB002,'')) AS item_name,
                RTRIM(ISNULL(b.MB003,'')) AS spec,
                ISNULL(b.MB004,'')      AS unit,
                ISNULL(b.MB025,'')      AS attr_code
            FROM INVMB b
            WHERE EXISTS (
                SELECT 1 FROM BOMMD d WHERE RTRIM(d.MD001) = RTRIM(b.MB001)
            )
            {prefix_cond}
            AND (RTRIM(b.MB001) LIKE ? OR b.MB002 LIKE ?)
            ORDER BY b.MB001
        """, (like, like))
        rows = cur.fetchall()
        conn.close()
        results = [{
            'item_no':   r[0],
            'item_name': r[1],
            'spec':      r[2],
            'unit':      r[3],
            'attr':      _ATTR_MAP.get(str(r[4]).strip(), str(r[4]).strip()),
        } for r in rows]
        return jsonify({'success': True, 'data': results, 'count': len(results)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/bom/detail')
def bom_detail():
    """取得指定品號的材料BOM明細（BOMMD JOIN INVMB，對應 MCPI11 BOM用量資料）"""
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()

        # 取品號主檔（母件資訊）
        cur.execute("""
            SELECT RTRIM(MB001), RTRIM(ISNULL(MB002,'')),
                   RTRIM(ISNULL(MB003,'')), ISNULL(MB004,''),
                   ISNULL(MB005,''), ISNULL(MB025,'')
            FROM INVMB WHERE RTRIM(MB001) = ?
        """, (item_no,))
        hrow = cur.fetchone()
        header = {}
        if hrow:
            header = {
                'item_no':   hrow[0],
                'item_name': hrow[1],
                'spec':      hrow[2],
                'unit':      hrow[3],
                'sub_unit':  hrow[4],
                'attr':      _ATTR_MAP.get(str(hrow[5]).strip(), str(hrow[5]).strip()),
            }

        # 取材料BOM明細（BOMMD JOIN INVMB 取子件品名/規格/屬性，並加總 INVMC 現有量為庫存數量）
        cur.execute("""
            SELECT
                d.MD002                           AS seq,
                RTRIM(d.MD003)                    AS child_no,
                RTRIM(ISNULL(b.MB002,''))         AS child_name,
                RTRIM(ISNULL(b.MB003,''))         AS spec,
                ISNULL(d.MD004,'')                AS unit,
                ISNULL(d.MD005,'')                AS sub_unit,
                ISNULL(b.MB025,'')                AS attr_code,
                CAST(ISNULL(d.MD006,0) AS VARCHAR(20)) AS qty,
                CAST(ISNULL((SELECT SUM(c.MC007) FROM INVMC c WHERE RTRIM(c.MC001) = RTRIM(d.MD003)),0) AS VARCHAR(20)) AS stock_qty
            FROM BOMMD d
            LEFT JOIN INVMB b ON RTRIM(b.MB001) = RTRIM(d.MD003)
            WHERE RTRIM(d.MD001) = ?
            ORDER BY d.MD002
        """, (item_no,))
        rows = cur.fetchall()
        conn.close()

        detail = []
        for r in rows:
            attr_code = str(r[6] or '').strip()
            detail.append({
                'seq':        str(r[0] or '').strip(),
                'child_no':   str(r[1] or '').strip(),
                'child_name': str(r[2] or '').strip(),
                'spec':       str(r[3] or '').strip(),
                'unit':       str(r[4] or '').strip(),
                'sub_unit':   str(r[5] or '').strip(),
                'attr':       _ATTR_MAP.get(attr_code, attr_code),
                'qty':        str(r[7] or '').strip(),
                'stock_qty':  str(r[8] or '').strip(),
            })

        return jsonify({'success': True, 'header': header, 'detail': detail})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/bom/routing')
def bom_routing():
    """取得指定品號的產品途程（BOMME + BOMMF，最新版次，對應 BOMMI07）"""
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()

        # 品號主檔
        cur.execute("""
            SELECT RTRIM(MB001), RTRIM(ISNULL(MB002,'')),
                   RTRIM(ISNULL(MB003,'')), ISNULL(MB004,''), ISNULL(MB025,'')
            FROM INVMB WHERE RTRIM(MB001) = ?
        """, (item_no,))
        hrow = cur.fetchone()
        header = {}
        if hrow:
            header = {
                'item_no':   hrow[0],
                'item_name': hrow[1],
                'spec':      hrow[2],
                'unit':      hrow[3],
                'attr':      _ATTR_MAP.get(str(hrow[4]).strip(), str(hrow[4]).strip()),
            }

        # 庫存數量（加總 INVMC.MC007 現有量）
        cur.execute("""
            SELECT CAST(ISNULL(SUM(MC007),0) AS VARCHAR(20))
            FROM INVMC WHERE RTRIM(MC001) = ?
        """, (item_no,))
        srow = cur.fetchone()
        if header:
            header['stock_qty'] = str(srow[0] or '0').strip() if srow else '0'

        # 最新途程版次
        cur.execute("""
            SELECT MAX(ME002) FROM BOMME WHERE RTRIM(ME001) = ?
        """, (item_no,))
        vrow = cur.fetchone()
        latest_ver = (vrow[0] or '').strip() if vrow else ''
        if header:
            header['version'] = latest_ver

        # 途程明細（BOMMF JOIN CMSMW 取製程名稱/機台代號/資源群組代號）
        routing = []
        if latest_ver:
            _NATURE = {'1': '1:廠內', '2': '2:訂外', '3': '3:外包'}
            cur.execute("""
                SELECT
                    f.MF003,
                    RTRIM(ISNULL(f.MF004,'')),
                    RTRIM(ISNULL(w.MW002,''))   AS proc_name,
                    f.MF005,
                    RTRIM(ISNULL(f.MF006,'')),
                    RTRIM(ISNULL(f.MF007,'')),
                    RTRIM(ISNULL(f.MF008,'')),
                    RTRIM(ISNULL(f.MF040,'')),
                    RTRIM(ISNULL(f.MF034,'')),
                    CAST(ISNULL(f.MF010,0) AS DECIMAL(10,3)),
                    RTRIM(ISNULL(x.MX003,''))   AS machine_name
                FROM BOMMF f
                LEFT JOIN CMSMW  w ON RTRIM(w.MW001) = RTRIM(f.MF004)
                LEFT JOIN CMSMX  x ON RTRIM(x.MX001) = RTRIM(f.MF040)
                WHERE RTRIM(f.MF001) = ? AND f.MF002 = ?
                ORDER BY f.MF003
            """, (item_no, latest_ver))
            rows = cur.fetchall()
            for r in rows:
                nature_code = str(r[3] or '').strip()
                vmh_raw = r[9]
                if vmh_raw is None or float(vmh_raw) == 0:
                    vmh = ''
                else:
                    vmh = '{:g}'.format(float(vmh_raw))
                routing.append({
                    'seq':          str(r[0] or '').strip(),
                    'proc_code':    str(r[1] or '').strip(),
                    'proc_name':    str(r[2] or '').strip(),
                    'nature':       _NATURE.get(nature_code, nature_code),
                    'vendor_no':    str(r[4] or '').strip(),
                    'vendor_name':  str(r[5] or '').strip(),
                    'description':  str(r[6] or '').strip(),
                    'machine_code': str(r[7] or '').strip(),
                    'res_group':    str(r[8] or '').strip(),
                    'var_man_hrs':  vmh,
                    'machine_name': str(r[10] or '').strip(),
                })

        conn.close()
        has_routing = bool(latest_ver and routing)
        return jsonify({
            'success': True,
            'has_routing': has_routing,
            'header': header,
            'routing': routing,
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/inventory/history')
def inventory_history():
    """取得指定品號的庫存異動歷史（INVTB + INVTK + MOCTE + MOCTG UNION）。
    回傳：日期、單別/單號/序號、單據名稱、庫別、庫別名稱、入出庫、異動數量、製令參考。
    """
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    months = request.args.get('months', '12').strip()
    try:
        months_int = max(1, min(60, int(months)))
    except ValueError:
        months_int = 12
    # 日期下限（CREATE_DATE 為 YYYYMMDD 字串）
    from datetime import datetime, timedelta
    start_date = (datetime.now() - timedelta(days=months_int * 31)).strftime('%Y%m%d')

    try:
        conn = get_erp_conn()
        cur  = conn.cursor()

        sql = """
            SELECT date_, doc_type, doc_no, doc_seq, warehouse, qty, src_table
            FROM (
                SELECT
                    CREATE_DATE                       AS date_,
                    RTRIM(TB001)                      AS doc_type,
                    RTRIM(TB002)                      AS doc_no,
                    RTRIM(TB003)                      AS doc_seq,
                    RTRIM(ISNULL(TB016,''))           AS warehouse,
                    CAST(ISNULL(TB007,0) AS DECIMAL(18,3)) AS qty,
                    'INVTB'                           AS src_table
                FROM INVTB
                WHERE RTRIM(TB004) = ? AND CREATE_DATE >= ?
                UNION ALL
                SELECT
                    CREATE_DATE,
                    RTRIM(TK001), RTRIM(TK002), RTRIM(TK003),
                    RTRIM(ISNULL(TK017,'')),
                    CAST(ISNULL(TK007,0) AS DECIMAL(18,3)),
                    'INVTK'
                FROM INVTK
                WHERE RTRIM(TK004) = ? AND CREATE_DATE >= ?
                UNION ALL
                SELECT
                    CREATE_DATE,
                    RTRIM(TE001), RTRIM(TE002), RTRIM(TE003),
                    RTRIM(ISNULL(TE008,'')),
                    CAST(ISNULL(TE005,0) AS DECIMAL(18,3)),
                    'MOCTE'
                FROM MOCTE
                WHERE RTRIM(TE004) = ? AND CREATE_DATE >= ?
                UNION ALL
                SELECT
                    CREATE_DATE,
                    RTRIM(TG001), RTRIM(TG002), RTRIM(TG003),
                    RTRIM(ISNULL(TG010,'')),
                    CAST(ISNULL(TG011,0) AS DECIMAL(18,3)),
                    'MOCTG'
                FROM MOCTG
                WHERE RTRIM(TG004) = ? AND CREATE_DATE >= ?
            ) u
            ORDER BY date_ DESC, doc_no DESC, doc_seq DESC
        """
        cur.execute(sql, (item_no, start_date, item_no, start_date, item_no, start_date, item_no, start_date))
        rows = cur.fetchall()

        # 一次撈出單別/庫別對照
        cur.execute("SELECT RTRIM(MQ001), RTRIM(ISNULL(MQ002,'')) FROM CMSMQ")
        doc_map = {r[0]: r[1] for r in cur.fetchall()}
        cur.execute("SELECT RTRIM(MC001), RTRIM(ISNULL(MC002,'')) FROM CMSMC")
        wh_map  = {r[0]: r[1] for r in cur.fetchall()}
        conn.close()

        # 入出庫判斷：依來源表 + 單別名稱
        # MOCTE=出庫(領料); MOCTG=入庫(移轉入庫); INVTK=調整(數量通常為0)
        # INVTB 依名稱關鍵字判斷
        def in_out(src, doc_type, doc_name, qty):
            if src == 'MOCTE':
                return '出庫'
            if src == 'MOCTG':
                return '入庫'
            if src == 'INVTK':
                return '調整'
            n = doc_name or ''
            if '入庫' in n:
                return '入庫'
            if '出庫' in n or '領料' in n or '報廢' in n:
                return '出庫'
            return '調整'

        history = []
        for r in rows:
            date_str = (r[0] or '').strip()
            if date_str and len(date_str) == 8:
                date_str = f'{date_str[:4]}/{date_str[4:6]}/{date_str[6:8]}'
            doc_type = (r[1] or '').strip()
            doc_no   = (r[2] or '').strip()
            doc_seq  = (r[3] or '').strip()
            wh       = (r[4] or '').strip()
            qty      = float(r[5] or 0)
            src      = r[6]
            doc_name = doc_map.get(doc_type, '')
            history.append({
                'date':       date_str,
                'doc_type':   doc_type,
                'doc_no':     doc_no,
                'doc_seq':    doc_seq,
                'doc_name':   doc_name,
                'warehouse':  wh,
                'wh_name':    wh_map.get(wh, ''),
                'in_out':     in_out(src, doc_type, doc_name, qty),
                'qty':        '{:,.3f}'.format(qty) if qty else '0',
                'src':        src,
            })

        return jsonify({'success': True, 'item_no': item_no, 'count': len(history), 'history': history})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/inventory/future')
def inventory_future():
    """取得指定品號的未來異動量（未完成製令）。
    - 預計生產（入庫）：MOCTA WHERE TA006=品號 AND 狀態未完且尚有未完工量
    - 預計領料（出庫）：MOCTB WHERE TB003=品號（子件）AND 狀態未完且尚有未領量
    """
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        from datetime import datetime, timedelta
        # 過濾掉太舊（>30 天前）的孤兒未結紀錄；保留近期含已逾期但未完成的製令
        cutoff = (datetime.now() - timedelta(days=30)).strftime('%Y%m%d')

        conn = get_erp_conn()
        cur  = conn.cursor()

        # 預計生產 (MOCTA) - 入庫；只顯示預計入庫日 >= 今天的紀錄
        cur.execute("""
            SELECT
                TA010                                       AS date_,
                RTRIM(TA001)                                AS doc_type,
                RTRIM(TA002)                                AS doc_no,
                CAST((TA015 - ISNULL(TA017,0)) AS DECIMAL(18,3)) AS qty,
                RTRIM(ISNULL(TA020,''))                     AS warehouse,
                RTRIM(ISNULL(TA019,''))                     AS plant,
                RTRIM(ISNULL(TA011,''))                     AS status,
                RTRIM(ISNULL(TA026,''))                     AS so_no,
                RTRIM(ISNULL(TA027,''))                     AS so_seq,
                RTRIM(ISNULL(TA028,''))                     AS so_line,
                RTRIM(ISNULL(TA029,''))                     AS remark
            FROM MOCTA
            WHERE RTRIM(TA006) = ?
              AND ISNULL(TA011,'') NOT IN ('Y','y')
              AND (TA015 - ISNULL(TA017,0)) > 0
              AND ISNULL(TA010,'') >= ?
        """, (item_no, cutoff))
        produce_rows = cur.fetchall()

        # 預計領料 (MOCTB) - 出庫；只顯示預計領料日 >= 今天的紀錄
        cur.execute("""
            SELECT
                TB015                                       AS date_,
                RTRIM(TB001)                                AS doc_type,
                RTRIM(TB002)                                AS doc_no,
                CAST((TB004 - ISNULL(TB005,0)) AS DECIMAL(18,3)) AS qty,
                RTRIM(ISNULL(TB009,''))                     AS warehouse,
                RTRIM(ISNULL(TB011,''))                     AS status,
                RTRIM(ISNULL(TB014,''))                     AS parent_item,
                RTRIM(ISNULL(TB006,''))                     AS proc_code
            FROM MOCTB
            WHERE RTRIM(TB003) = ?
              AND ISNULL(TB011,'') NOT IN ('Y','y')
              AND (TB004 - ISNULL(TB005,0)) > 0
              AND ISNULL(TB015,'') >= ?
        """, (item_no, cutoff))
        require_rows = cur.fetchall()

        # 廠別/庫別名稱對照
        cur.execute("SELECT RTRIM(MB001), RTRIM(ISNULL(MB002,'')) FROM CMSMB")
        plant_map = {r[0]: r[1] for r in cur.fetchall()}
        cur.execute("SELECT RTRIM(MC001), RTRIM(ISNULL(MC002,'')) FROM CMSMC")
        wh_map = {r[0]: r[1] for r in cur.fetchall()}
        conn.close()

        def fmt_date(d):
            s = (d or '').strip() if isinstance(d, str) else str(d or '').strip()
            return f'{s[:4]}/{s[4:6]}/{s[6:8]}' if len(s) == 8 else s

        future = []
        for r in produce_rows:
            qty = float(r[3] or 0)
            so_no, so_seq, so_line = r[7], r[8], r[9]
            ref_parts = []
            if so_no: ref_parts.append(f'{so_no}-{so_seq}' if so_seq else so_no)
            if so_line: ref_parts.append(so_line)
            ref = ' '.join(ref_parts).strip()
            remark = (r[10] or '').strip()
            future.append({
                'date':       fmt_date(r[0]),
                'date_raw':   (r[0] or '').strip(),
                'type':       '預計生',
                'doc_type':   r[1],
                'doc_no':     r[2],
                'in_qty':     '{:,.3f}'.format(qty).rstrip('0').rstrip('.') if qty else '',
                'out_qty':    '',
                'warehouse':  r[4],
                'wh_name':    wh_map.get(r[4], ''),
                'plant':      r[5],
                'plant_name': plant_map.get(r[5], ''),
                'status':     r[6],
                'remark':     f'{r[1]}-{r[2]}' + (f' {ref}' if ref else '') + (f' {remark}' if remark else ''),
            })
        for r in require_rows:
            qty = float(r[3] or 0)
            future.append({
                'date':       fmt_date(r[0]),
                'date_raw':   (r[0] or '').strip(),
                'type':       '預計領',
                'doc_type':   r[1],
                'doc_no':     r[2],
                'in_qty':     '',
                'out_qty':    '{:,.3f}'.format(qty).rstrip('0').rstrip('.') if qty else '',
                'warehouse':  r[4],
                'wh_name':    wh_map.get(r[4], ''),
                'plant':      '',
                'plant_name': '',
                'status':     r[5],
                'remark':     f'{r[1]}-{r[2]} {r[6]}-{item_no}-{r[7]}'.strip(),
            })

        future.sort(key=lambda x: (x['date_raw'] or '99999999', x['doc_no']))
        return jsonify({'success': True, 'item_no': item_no, 'count': len(future), 'future': future})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/cache/refresh', methods=['POST'])
def refresh_cache():
    """手動清除快取，下次查詢時會重新抓取 SSRS 資料"""
    with _cache_lock:
        _cache.clear()
    return jsonify({'success': True, 'message': '快取已清除'})


# ════════════════════════════════════════════════════════════════
#  申請單功能
# ════════════════════════════════════════════════════════════════

# 各表單類型定義（type_key → 顯示名稱、範本檔名、欄位清單）
_FORM_TYPES = {
    'pp015': {
        'name': 'PP-M-015 內部業務聯絡單',
        'template': 'pp015.docx',
        'fields': [
            {'key': 'date',     'label': '日期',   'type': 'date',           'required': True},
            {'key': 'priority', 'label': '速別',   'type': 'select',         'required': False, 'options': ['普通件', '速件', '最速件']},
            {'key': 'depts', 'label': '受文單位', 'type': 'checkbox_group', 'required': False,
             'options': [
                 '總經理', '副總',
                 '總經理室', '生技', '資管', '美工設計',
                 '資材部',  '採購', '生管', '倉管',
                 '管理部',  '財會', '總務人資',
                 '加工部',  '成品部',
                 '業務部',  '研發部', '品保部',
             ],
             'groups': [                                  # 每個子列表 = 一排
                 ['總經理', '副總'],
                 ['總經理室', '生技', '資管', '美工設計'],
                 ['資材部',  '採購', '生管', '倉管'],
                 ['管理部',  '財會', '總務人資'],
                 ['加工部',  '成品部'],
                 ['業務部',  '研發部', '品保部'],
             ],
             'sub_opts': ['生技', '資管', '美工設計', '採購', '生管', '倉管', '財會', '總務人資', '成品部'],
             'defaults': ['總經理', '加工部']},
            {'key': 'author',   'label': '承辦人', 'type': 'text',           'required': False},
            {'key': 'subject',  'label': '主旨',   'type': 'text',           'required': True},
            {'key': 'content',  'label': '說明',   'type': 'textarea',       'required': True},
        ],
    },
    'hr028': {
        'name': 'P-HR-028-01A 人事懲戒核定書',
        'template': 'hr028.docx',
        'fields': [
            {'key': 'date',     'label': '通知日期',   'type': 'date',     'required': True},
            {'key': 'emp_no',   'label': '員工編號',   'type': 'text',     'required': False},
            {'key': 'name',     'label': '姓名',       'type': 'text',     'required': True},
            {'key': 'title',    'label': '職稱',       'type': 'text',     'required': False},
            {'key': 'punish',     'label': '懲戒分類',   'type': 'select',   'required': True, 'options': ['大過', '小過', '申誡', '警告']},
            {'key': 'pun_count', 'label': '懲戒次數',   'type': 'select',   'required': False, 'options': ['1', '2', '3']},
            {'key': 'reason',   'label': '懲戒事由',   'type': 'textarea', 'required': True},
            {'key': 'evidence', 'label': '檢附之證據', 'type': 'textarea', 'required': False},
            {'key': 'method',   'label': '懲戒方式',   'type': 'text',     'required': False},
        ],
    },
    'hr029': {
        'name': '人事獎勵核定書',
        'template': 'hr029.docx',
        'fields': [
            {'key': 'date',     'label': '通知日期',   'type': 'date',     'required': True},
            {'key': 'emp_no',   'label': '員工編號',   'type': 'text',     'required': False},
            {'key': 'name',     'label': '姓名',       'type': 'text',     'required': True},
            {'key': 'title',    'label': '職稱',       'type': 'text',     'required': False},
            {'key': 'award',     'label': '獎勵分類',   'type': 'select',   'required': True, 'options': ['職位晉昇', '大功', '小功', '嘉獎', '優點']},
            {'key': 'award_count', 'label': '獎勵次數', 'type': 'select',   'required': False, 'options': ['1', '2', '3']},
            {'key': 'reason',   'label': '獎勵事由',   'type': 'textarea', 'required': True},
            {'key': 'evidence', 'label': '檢附之證據', 'type': 'textarea', 'required': False},
            {'key': 'method',   'label': '獎勵方式',   'type': 'text',     'required': False},
        ],
    },
    'pp017': {
        'name': 'PP-M-017 報廢申請單',
        'template': 'pp017.docx',
        'fields': [
            {'key': 'date',      'label': '申請日期',  'type': 'date',     'required': True},
            {'key': 'category',  'label': '類別',      'type': 'select',   'required': True, 'options': ['機器', '模治具', '量規儀器', '其他']},
            {'key': 'applicant', 'label': '申請人',    'type': 'text',     'required': True},
            {'key': 'item',      'label': '品名/規格', 'type': 'text',     'required': True},
            {'key': 'item_no',   'label': '編號/品號', 'type': 'text',     'required': False},
            {'key': 'qty',       'label': '數量',      'type': 'text',     'required': False},
            {'key': 'reason',    'label': '報廢理由',  'type': 'textarea', 'required': True},
        ],
    },
    'gcd': {
        'name': '公出單',
        'template': 'gcd.docx',
        'store_path': r'\\192.168.1.99\加工部-資料夾\【技術資料】\W.表單.活動\工出單',
        'fields': [
            {'key': 'date',     'label': '日期',     'type': 'date',     'required': True},
            {'key': 'name',     'label': '姓名',     'type': 'text',     'required': True},
            {'key': 'location', 'label': '地點',     'type': 'text',     'required': True},
            {'key': 'reason',   'label': '外出原因', 'type': 'textarea', 'required': True},
        ],
    },
}

# 各表單輸出檔名規則（form_type → 產生檔名主體的函式）
def _application_filename_base(form_type, ctx):
    if form_type == 'pp017':
        return f'{ctx.get("item", "")}_報廢單'.strip('_')
    if form_type == 'hr028':
        return f'人事懲處({ctx.get("name", "")})'
    if form_type == 'hr029':
        return f'人事獎勵({ctx.get("name", "")})'
    if form_type == 'gcd':
        return f'{ctx.get("name", "")}_公出單'.strip('_')
    return ctx.get('subject', '') or '內部業務聯絡單'


@app.route('/application')
def application_page():
    """申請單主頁"""
    return render_template('application.html', form_types=_FORM_TYPES)


@app.route('/api/application/list')
def application_list():
    """列出 NAS 申請單資料夾的既有檔案"""
    import datetime as _dt
    year = request.args.get('year', str(_dt.date.today().year))
    keyword = request.args.get('keyword', '').strip()
    base = getattr(config, 'APPLICATION_STORE_PATH', '')
    if not base:
        return jsonify({'error': '未設定 APPLICATION_STORE_PATH'}), 500

    target = os.path.join(base, year)
    try:
        entries = os.scandir(target)
    except Exception as e:
        return jsonify({'files': [], 'error': str(e)})

    files = []
    for e in entries:
        if not e.is_file():
            continue
        name_lower = e.name.lower()
        if not (name_lower.endswith('.doc') or name_lower.endswith('.docx')):
            continue
        if keyword and keyword.lower() not in e.name.lower():
            continue
        stat = e.stat()
        files.append({
            'filename': e.name,
            'path': e.path,
            'size': stat.st_size,
            'mtime': _dt.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
        })
    files.sort(key=lambda x: x['mtime'], reverse=True)
    return jsonify({'files': files, 'year': year, 'total': len(files)})


@app.route('/api/application/create', methods=['POST'])
def application_create():
    """接收表單資料 + 照片，產生 .docx 存到 NAS"""
    import datetime as _dt
    import tempfile
    import shutil

    if not _DOCXTPL_OK:
        return jsonify({'ok': False, 'error': 'docxtpl 套件未安裝，請執行 pip install docxtpl'}), 500

    form_type = request.form.get('form_type', '')
    if form_type not in _FORM_TYPES:
        return jsonify({'ok': False, 'error': f'未知表單類型: {form_type}'}), 400

    form_def = _FORM_TYPES[form_type]
    template_path = os.path.join(
        getattr(config, 'FORM_TEMPLATES_DIR', None) or os.path.join(_APP_DIR, 'static', 'form_templates'),
        form_def['template']
    )

    if not os.path.exists(template_path):
        return jsonify({
            'ok': False,
            'error': f'找不到 Word 範本：{template_path}，請先將 .docx 範本放入 static/form_templates/'
        }), 500

    # 收集欄位值（checkbox_group 用 getlist；其餘用 get）
    context = {}
    for f in form_def['fields']:
        if f['type'] == 'checkbox_group':
            # checkbox 值以陣列送入（name="depts[]"），不放入 context（由下方衍生欄位處理）
            pass
        else:
            context[f['key']] = request.form.get(f['key'], '')

    # 衍生欄位：核取記號（■/□）與民國日期，供範本 placeholder 使用
    def _mark(selected, target):
        return '■' if selected == target else '□'

    def _fmt_date(iso, use_roc=False):
        """ISO 日期轉中文格式；use_roc=True 時用民國年，否則用西元年"""
        try:
            y, m, d = iso.split('-')
            if use_roc:
                return f'中華民國 {int(y) - 1911} 年 {int(m):02d} 月 {int(d):02d} 日'
            return f'{y} 年 {int(m):02d} 月 {int(d):02d} 日'
        except Exception:
            return ''

    # 所有表單統一用西元年
    context['roc_date'] = _fmt_date(context.get('date', ''))

    def _fmt_date_slash(iso):
        """ISO 日期轉 YYYY/M/D 格式"""
        try:
            y, m, d = iso.split('-')
            return f'{int(y)}/{int(m)}/{int(d)}'
        except Exception:
            return ''

    context['gcd_date'] = _fmt_date_slash(context.get('date', ''))

    if form_type == 'pp015':
        pri = context.get('priority', '')
        context['pri1'] = _mark(pri, '最速件')
        context['pri2'] = _mark(pri, '速件')
        context['pri3'] = _mark(pri, '普通件')
        # 受文單位（可複選）→ ■/□ 記號
        depts_sel = request.form.getlist('depts[]')
        # 主部門 → ■/□
        main_dept_map = [
            ('chk_gm',    '總經理'),
            ('chk_vp',    '副總'),
            ('chk_gmo',   '總經理室'),
            ('chk_mat',   '資材部'),
            ('chk_adm',   '管理部'),
            ('chk_proc',  '加工部'),
            ('chk_sales', '業務部'),
            ('chk_rd',    '研發部'),
            ('chk_qa',    '品保部'),
        ]
        for key, dept in main_dept_map:
            context[key] = '■' if dept in depts_sel else '□'
        # 子部門 → ●/○
        sub_dept_map = [
            ('chk_jt',     '生技'),
            ('chk_it',     '資管'),
            ('chk_design', '美工設計'),
            ('chk_pur',    '採購'),
            ('chk_pm',     '生管'),
            ('chk_wh',     '倉管'),
            ('chk_fin',    '財會'),
            ('chk_hr',     '總務人資'),
            ('chk_fp',     '成品部'),
        ]
        for key, dept in sub_dept_map:
            context[key] = '●' if dept in depts_sel else '○'
    elif form_type == 'hr028':
        pun = context.get('punish', '')
        cnt = context.get('pun_count', '')
        # 只保留 4 種懲戒
        context['pun_d'] = _mark(pun, '大過')
        context['pun_e'] = _mark(pun, '小過')
        context['pun_f'] = _mark(pun, '申誡')
        context['pun_g'] = _mark(pun, '警告')
        # 次數欄位（只填入對應懲戒分類的次數）
        context['cnt_d'] = cnt if pun == '大過' else ''
        context['cnt_e'] = cnt if pun == '小過' else ''
        context['cnt_f'] = cnt if pun == '申誡' else ''
        context['cnt_g'] = cnt if pun == '警告' else ''
    elif form_type == 'hr029':
        awd = context.get('award', '')
        cnt = context.get('award_count', '')
        context['award_a'] = _mark(awd, '職位晉昇')
        context['award_b'] = _mark(awd, '大功')
        context['award_c'] = _mark(awd, '小功')
        context['award_d'] = _mark(awd, '嘉獎')
        context['award_e'] = _mark(awd, '優點')
        # 次數欄位（只填入對應獎勵分類的次數）
        context['award_b_cnt'] = cnt if awd == '大功' else ''
        context['award_c_cnt'] = cnt if awd == '小功' else ''
        context['award_d_cnt'] = cnt if awd == '嘉獎' else ''
        context['award_e_cnt'] = cnt if awd == '優點' else ''
    elif form_type == 'pp017':
        cat = context.get('category', '')
        context['cat_a'] = _mark(cat, '機器')
        context['cat_b'] = _mark(cat, '模治具')
        context['cat_c'] = _mark(cat, '量規儀器')
        context['cat_d'] = _mark(cat, '其他')
    elif form_type == 'gcd':
        veh = context.get('vehicle', '')
        context['veh_car']     = _mark(veh, '汽車')
        context['veh_moto']    = _mark(veh, '機車')
        context['veh_priv']    = _mark(veh, '私車公用')
        context['veh_company'] = _mark(veh, '公司車')
        context['veh_other']   = _mark(veh, '其它')
        context.setdefault('emp_no', '')
        context.setdefault('km', '')
        context.setdefault('plate', '')

    # 處理照片（暫存 → 轉成 InlineImage）
    tmp_dir = tempfile.mkdtemp(prefix='pds_upload_')
    photo_objects = []
    try:
        photos = request.files.getlist('photos[]')
        for idx, photo in enumerate(photos):
            if photo and photo.filename:
                ext = os.path.splitext(photo.filename)[1].lower() or '.jpg'
                tmp_path = os.path.join(tmp_dir, f'photo_{idx}{ext}')
                photo.save(tmp_path)
                photo_objects.append(tmp_path)

        # 組裝 docxtpl context（照片用 InlineImage）
        tpl = DocxTemplate(template_path)
        for i, p in enumerate(photo_objects):
            context[f'pic_{i+1}'] = InlineImage(tpl, p, width=Mm(120))
        # 未填照片位置補空字串，避免 KeyError
        for i in range(len(photo_objects) + 1, 6):
            context.setdefault(f'pic_{i}', '')

        tpl.render(context)

        # 處理附件檔案（呈現於文件第2頁）
        attachments = request.files.getlist('attachments[]')
        attach_paths = []
        for idx, att in enumerate(attachments):
            if att and att.filename:
                ext = os.path.splitext(att.filename)[1].lower() or '.bin'
                tmp_path = os.path.join(tmp_dir, f'attach_{idx}{ext}')
                att.save(tmp_path)
                attach_paths.append((tmp_path, att.filename))

        if attach_paths:
            doc = tpl.get_docx()
            doc.add_page_break()
            doc.add_heading('附件', level=1)
            for p, orig_name in attach_paths:
                ext = os.path.splitext(orig_name)[1].lower()
                if ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp'):
                    try:
                        doc.add_picture(p, width=Mm(160))
                    except Exception:
                        doc.add_paragraph(orig_name)
                else:
                    doc.add_paragraph(orig_name)

        # 決定輸出檔名（依表單類型套用對應命名規則）
        date_str = context.get('date', '').replace('-', '') or _dt.date.today().strftime('%Y%m%d')
        subject = _application_filename_base(form_type, context)
        subject = re.sub(r'[\\/:*?"<>|]', '_', subject)[:40]  # 去除非法字元，限長 40
        out_filename = f'{date_str} {subject}.docx'

        # 存到 NAS（若失敗則 fallback 存桌面）
        year = date_str[:4]
        store_base = form_def.get('store_path') or getattr(config, 'APPLICATION_STORE_PATH', '')
        out_dir = os.path.join(store_base, year)
        _nas_warning = None
        try:
            os.makedirs(out_dir, exist_ok=True)
            out_path = os.path.join(out_dir, out_filename)
            tpl.save(out_path)
        except (PermissionError, OSError) as _e:
            # NAS 無法存取（權限、網路、檔案被鎖定等）→ 改存桌面
            _desktop = os.path.join(os.environ.get('USERPROFILE', os.path.expanduser('~')), 'Desktop')
            os.makedirs(_desktop, exist_ok=True)
            # 若桌面同名檔案已存在，加時間戳避免衝突
            _base, _ext = os.path.splitext(out_filename)
            _desk_name = out_filename
            if os.path.exists(os.path.join(_desktop, _desk_name)):
                _desk_name = f"{_base}_{int(time.time())}{_ext}"
            out_path = os.path.join(_desktop, _desk_name)
            tpl.save(out_path)
            _nas_warning = f'NAS 無法存取（{type(_e).__name__}），檔案已改存至桌面：{out_path}'

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    resp = {'ok': True, 'filename': out_filename, 'path': out_path}
    if _nas_warning:
        resp['warning'] = _nas_warning
    return jsonify(resp)


def _docx_to_html(path):
    """將 .docx 轉成可預覽的 HTML（保留段落、表格、內嵌圖片）"""
    import base64
    from markupsafe import escape
    from docx import Document as _Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = _Document(path)

    def _img_tags(elm):
        tags = ''
        for blip in elm.iter(qn('a:blip')):
            rId = blip.get(qn('r:embed'))
            if not rId:
                continue
            try:
                image_part = doc.part.related_parts[rId]
                b64 = base64.b64encode(image_part.blob).decode('ascii')
                tags += (f'<img src="data:{image_part.content_type};base64,{b64}" '
                         f'style="max-width:100%;display:block;margin:.3rem 0;">')
            except Exception:
                pass
        return tags

    def _para_html(p):
        text = p.text.strip()
        imgs = _img_tags(p._p)
        if not text and not imgs:
            return ''
        align = ''
        try:
            if int(p.alignment) == 1:  # WD_ALIGN_PARAGRAPH.CENTER
                align = 'text-align:center;'
        except Exception:
            pass
        body = str(escape(text)).replace('\n', '<br>') if text else ''
        return f'<p style="margin:.25rem 0;{align}">{body}{imgs}</p>'

    def _table_html(tbl):
        rows_html = ''
        for row in tbl.rows:
            cells_html = ''
            for cell in row.cells:
                cell_text = str(escape(cell.text)).replace('\n', '<br>')
                cell_imgs = _img_tags(cell._tc)
                cells_html += (f'<td style="border:1px solid #E0DACC;padding:.3rem .5rem;'
                                f'vertical-align:top;">{cell_text}{cell_imgs}</td>')
            rows_html += f'<tr>{cells_html}</tr>'
        return f'<table style="width:100%;border-collapse:collapse;margin:.4rem 0;font-size:0.92rem;">{rows_html}</table>'

    parts = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            html = _para_html(Paragraph(child, doc))
            if html:
                parts.append(html)
        elif child.tag == qn('w:tbl'):
            parts.append(_table_html(Table(child, doc)))
    return ''.join(parts)


@app.route('/api/application/view')
def application_view():
    """預覽既有申請單檔案內容（轉成 HTML）"""
    path = request.args.get('path', '')
    if not path or not os.path.isfile(path):
        return jsonify({'ok': False, 'error': '檔案不存在'}), 404
    if not path.lower().endswith('.docx'):
        return jsonify({'ok': False, 'error': '僅支援預覽 .docx 格式，請點擊開啟檔案查看'}), 400
    try:
        return jsonify({'ok': True, 'html': _docx_to_html(path)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/application/open', methods=['POST'])
def application_open():
    """用 Shell 開啟指定 NAS 上的申請單檔案"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400
    try:
        os.startfile(path)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/management')
def management_page():
    """管理頁主頁"""
    return render_template('management.html', form_types=_FORM_TYPES)


@app.route('/api/attendance/list')
def attendance_list():
    """讀取 Google 試算表『M9出勤表』資料"""
    cache_key = 'attendance_list'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.ATTENDANCE_SHEET_ID, gid=config.ATTENDANCE_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    try:
        emp_map = fetch_employee_name_map()
        for r in rows:
            for k in list(r.keys()):
                emp_id = r.get(k, '').strip()
                if emp_id and emp_id in emp_map:
                    r[k + '_name'] = emp_map[emp_id]
        result = {'success': True, 'count': len(rows), 'headers': list(rows[0].keys()) if rows else [], 'rows': rows}
        cache_set(cache_key, result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/leave/list')
def leave_list():
    """讀取 Google 試算表『請假單』資料"""
    cache_key = 'leave_list'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'employee_name_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.LEAVE_SHEET_ID, config.LEAVE_SHEET_NAME)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    try:
        name_map = fetch_employee_name_map()
    except Exception:
        name_map = {}

    leaves = []
    for r in rows:
        emp_id = (r.get('請假人員') or '').strip()
        leaves.append({
            'date':   (r.get('請假日期') or '').strip(),
            'employee': emp_id,
            'name':   name_map.get(emp_id, ''),
            'type':   (r.get('假別') or '').strip(),
            'start':  (r.get('起始時間') or '').strip(),
            'end':    (r.get('結束時間') or '').strip(),
            'remark': (r.get('備註說明') or '').strip(),
        })

    def _sort_key(item):
        try:
            y, m, d = item['date'].split('/')
            return (int(y), int(m), int(d))
        except Exception:
            return (0, 0, 0)

    leaves.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(leaves), 'leaves': leaves}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/overtime/list')
def overtime_list():
    """讀取 Google 試算表『加班統計』資料"""
    cache_key = 'overtime_list'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.OVERTIME_SHEET_ID, config.OVERTIME_SHEET_NAME)
    except Exception as e:
        return jsonify({'success': False, 'error': f'無法連線至 Google 試算表：{str(e)}'}), 502

    try:
        name_map = fetch_employee_name_map()
    except Exception:
        name_map = {}

    overtimes = []
    for r in rows:
        emp_id = (r.get('加班人員ID') or r.get('加班人員') or '').strip()
        overtimes.append({
            'date':   (r.get('加班日期') or '').strip(),
            'employee': emp_id,
            'name':   name_map.get(emp_id, ''),
            'hours':  (r.get('加班時數') or '').strip(),
            'remark': (r.get('備註') or '').strip(),
        })

    def _sort_key(item):
        try:
            y, m, d = item['date'].split('/')
            return (int(y), int(m), int(d))
        except Exception:
            return (0, 0, 0)

    overtimes.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(overtimes), 'overtimes': overtimes}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/employee/list')
def employee_list():
    """讀取 Google 試算表『員工登錄系統』(1.1員工登錄系統分頁)，回傳人員名冊"""
    try:
        roster = fetch_employee_roster()
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502
    return jsonify({'success': True, 'count': len(roster), 'employees': roster})


@app.route('/api/purchase/list')
def purchase_list():
    """讀取 Google 試算表『採購登入表』（主表分頁）資料"""
    cache_key = 'purchase_list'
    if request.args.get('refresh'):
        cache_clear(cache_key)
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.PURCHASE_SHEET_ID, sheet_name=config.PURCHASE_SHEET_NAME)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    records = []
    for r in rows:
        raw_qty = (r.get('數量') or '').replace(',', '').strip()
        try:    qty = float(raw_qty)
        except ValueError: qty = 0
        raw_unit = (r.get('單價') or '').replace(',', '').strip()
        try:    unit_price = float(raw_unit)
        except ValueError: unit_price = 0
        raw_sub = (r.get('小計') or '').replace(',', '').strip()
        try:    subtotal = float(raw_sub)
        except ValueError: subtotal = 0
        date = (r.get('日期') or '').strip()
        if not date and not (r.get('品名') or '').strip():
            continue  # 跳過空白列
        records.append({
            'date':       date,
            'status':     (r.get('狀態') or '').strip(),
            'name':       (r.get('品名') or '').strip(),
            'spec':       (r.get('規格') or '').strip(),
            'vendor':     (r.get('供應商') or '').strip(),
            'vendor_no':  (r.get('供應商代號') or '').strip(),
            'qty':        qty,
            'unit_price': unit_price,
            'subtotal':   subtotal,
            'account':    (r.get('會計科目') or '').strip(),
            'doc_no':     (r.get('單據號碼') or '').strip(),
            'remark':     (r.get('備註') or '').strip(),
        })

    def _sort_key(item):
        try:
            parts = item['date'].replace('-', '/').split('/')
            return tuple(int(p) for p in parts)
        except Exception:
            return (0, 0, 0)

    records.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(records), 'purchases': records}
    cache_set(cache_key, result)
    return jsonify(result)


# ── 批成本計算（共用區 Excel 範本：刀具/刀表/批成本計算） ──────────────

def _load_batchcost_wb(read_only=False):
    """開啟批成本計算範本檔案（共用網路路徑）"""
    return openpyxl.load_workbook(config.BATCH_COST_FILE_PATH, data_only=read_only)


def _save_batchcost_wb(wb):
    """儲存，若檔案被佔用（例如使用者自己開著 Excel）則 fallback 存到桌面，回傳警告訊息（無警告回 None）"""
    try:
        wb.save(config.BATCH_COST_FILE_PATH)
        return None
    except (PermissionError, OSError) as e:
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        os.makedirs(desktop, exist_ok=True)
        fallback = os.path.join(desktop, f'批成本計算_{int(time.time())}.xlsx')
        wb.save(fallback)
        return f'原始檔案無法寫入（{type(e).__name__}，可能被開啟中），已改存到：{fallback}'


@app.route('/api/batch_cost/lookup_order')
def batch_cost_lookup_order():
    """依製令號碼查詢品號/品名/製程代號（優先從 P2 表查，再備援 SSRS；允許不打 "-" 直接輸入數字）"""
    order = request.args.get('order', '').strip()
    if not order:
        return jsonify({'success': False, 'error': '請輸入製令號碼'}), 400
    # 製令格式為「4位單位代碼-序號」，若使用者輸入全數字沒打 "-"，自動補上方便比對
    order_norm = order.replace('-', '')
    if order.isdigit() and len(order) > 4:
        order_dashed = f'{order[:4]}-{order[4:]}'
    else:
        order_dashed = order

    records = None
    # 1. 優先從 K1_P2（生產報工統計 P2）查詢
    try:
        p2_rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        records = []
        for p2r in p2_rows:
            wo = (p2r.get('製令') or '').strip()
            if wo == order or wo == order_dashed or wo.replace('-', '') == order_norm:
                sec_str = (p2r.get('出站總工時(秒)') or '').strip().replace(',', '')
                try:
                    seconds = float(sec_str) if sec_str else 0
                except ValueError:
                    seconds = 0
                qty_str = (p2r.get('出站數量') or '').strip().replace(',', '')
                try:
                    qty = float(qty_str) if qty_str else 0
                except ValueError:
                    qty = 0
                records.append({
                    '單別': wo,
                    '品號': (p2r.get('品號') or '').strip(),
                    '品名': (p2r.get('品名') or '').strip(),
                    '製程代號': (p2r.get('製程') or '').strip(),
                    '製程名稱': (p2r.get('製程名稱') or '').strip(),
                    '秒數': seconds,
                    '出站數量': qty,
                    '報工人員': (p2r.get('報工人員') or '').strip(),
                    '出站時間': (p2r.get('欄位格式化') or p2r.get('出站時間') or '').strip(),
                })
    except Exception:
        records = None

    # 2. 備援：如果 P2 查不到，試試 SSRS 未完工製令（用補上 "-" 的版本比對，substring 才能命中）
    if not records:
        try:
            records = search_orders(order_id=order_dashed)
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 502

    if not records:
        return jsonify({'success': False, 'error': '查無此製令'})

    # 如果來自 SSRS（沒有製程名稱和秒數），則再查 P2 表補充
    if not any(r.get('製程名稱') for r in records):
        p2_map = {}
        try:
            p2_rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
            for p2r in p2_rows:
                wo = (p2r.get('製令') or '').strip()
                proc = (p2r.get('製程') or '').strip()
                if wo:
                    proc_name = (p2r.get('製程名稱') or '').strip()
                    sec_str = (p2r.get('出站總工時(秒)') or '').strip().replace(',', '')
                    try:
                        seconds = float(sec_str) if sec_str else 0
                    except ValueError:
                        seconds = 0
                    key = (wo, proc)
                    if key not in p2_map:
                        p2_map[key] = {'proc_name': proc_name, 'seconds': seconds}
        except Exception:
            p2_map = {}
    else:
        p2_map = {}

    # 從『生產日報表P5.3』查機台名稱/機台代號（依製令+製程代號比對；P5.3 的製令欄位帶序號尾碼如 -0010，需先去除才能比對）
    p53_map = {}
    try:
        p53_rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, gid=config.PROD_REPORT_SHEET_GID)
        for p53r in p53_rows:
            wo_full = (p53r.get('製令') or '').strip()
            wo = wo_full.rsplit('-', 1)[0] if wo_full.count('-') >= 2 else wo_full
            proc = (p53r.get('製程代號') or '').strip()
            if wo:
                key = (wo, proc)
                if key not in p53_map:
                    p53_map[key] = {
                        'machine_name': (p53r.get('機台名稱') or '').strip(),
                        'machine_code': (p53r.get('機台代號') or '').strip(),
                    }
    except Exception:
        p53_map = {}

    seen = set()
    rows = []
    for r in records:
        part_no = (r.get('品號') or '').strip()
        proc_code = (r.get('製程代號') or '').strip()
        key = (part_no, proc_code)
        if key in seen:
            continue
        seen.add(key)
        order_full = r.get('單別', '').strip()
        # 如果已經從 P2 查到，直接用；否則從 p2_map 補充
        proc_name = r.get('製程名稱', '')
        seconds = r.get('秒數', 0)
        if not proc_name and order_full in p2_map:
            p2_info = p2_map.get((order_full, proc_code), {})
            proc_name = proc_name or p2_info.get('proc_name', '')
            seconds = seconds or p2_info.get('seconds', 0)
        out_qty = r.get('出站數量', 0)
        p53_info = p53_map.get((order_full, proc_code), {})
        rows.append({
            'order':    order_full,
            'part_no':  part_no,
            'name':     (r.get('品名') or '').rstrip('|').strip(),
            'proc_code': proc_code,
            'proc_name': proc_name,
            'seconds':  seconds,
            'out_qty':  out_qty,
            'machine_name': p53_info.get('machine_name', ''),
            'machine_code': p53_info.get('machine_code', ''),
            'operator': r.get('報工人員', ''),
            'out_time': r.get('出站時間', ''),
        })
    return jsonify({'success': True, 'rows': rows})


@app.route('/api/batch_cost/tool_catalog')
def batch_cost_tool_catalog():
    """讀『刀具』分頁，回傳每把刀的單價/刃數對照表"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        ws = wb[config.BATCH_COST_TOOL_SHEET]
        tools = []
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            name = row[idx.get('品名', 0)] if idx.get('品名') is not None else None
            if not name:
                continue
            tools.append({
                'row':    row_no,
                'name':   str(name).strip(),
                'item':   str(row[idx['項目']] or '').strip() if '項目' in idx else '',
                'vendor': str(row[idx['供應商']] or '').strip() if '供應商' in idx else '',
                'price':  row[idx['單價']] if '單價' in idx else 0,
                'edges':  row[idx['刃數']] if '刃數' in idx else 0,
            })
        wb.close()
        return jsonify({'success': True, 'tools': tools})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_catalog/save', methods=['POST'])
def batch_cost_tool_catalog_save():
    """新增或編輯『刀具資料』分頁裡的一把刀（有 row 就編輯該列，沒有就新增一列）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    name   = (data.get('name') or '').strip()
    item   = (data.get('item') or '').strip()
    vendor = (data.get('vendor') or '').strip()
    price  = data.get('price') or 0
    edges  = data.get('edges') or 0
    if not name:
        return jsonify({'success': False, 'error': '請輸入刀具名稱'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOL_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        values = {'品名': name, '項目': item, '供應商': vendor, '單價': price, '刃數': edges}

        if row_no:
            target_row = int(row_no)
        else:
            target_row = ws.max_row + 1
        for col_name, val in values.items():
            if col_name in idx:
                ws.cell(row=target_row, column=idx[col_name] + 1, value=val)

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_catalog/delete', methods=['POST'])
def batch_cost_tool_catalog_delete():
    """刪除『刀具資料』分頁裡的一列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    if not row_no:
        return jsonify({'success': False, 'error': '缺少列號'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOL_SHEET]
        ws.delete_rows(int(row_no), 1)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map')
def batch_cost_tool_map():
    """讀『刀表』分頁，依品號查該品號目前的刀具配置（T1~T39）"""
    part_no = request.args.get('part_no', '').strip()
    if not part_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        slot_cols = [h for h in header if h and re.match(r'^T\d+$', str(h))]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[idx.get('品號', -1)]:
                continue
            if str(row[idx['品號']]).strip() != part_no:
                continue
            slots = []
            for slot in slot_cols:
                v = row[idx[slot]]
                if v:
                    slots.append({'slot': slot, 'tool': str(v).strip()})
            wb.close()
            return jsonify({'success': True, 'found': True,
                             'machine_type': str(row[idx.get('加工機型', -1)] or '').strip() if '加工機型' in idx else '',
                             'slots': slots})
        wb.close()
        return jsonify({'success': True, 'found': False, 'slots': []})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map/save', methods=['POST'])
def batch_cost_tool_map_save():
    """新增或覆寫『刀表』分頁裡某品號的刀具配置"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    part_no = (data.get('part_no') or '').strip()
    machine_type = (data.get('machine_type') or '').strip()
    slots = data.get('slots') or []   # [{slot:'T1', tool:'...'}, ...]
    if not part_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        today = date.today().strftime('%Y/%m/%d')

        target_row = None
        for row_cells in ws.iter_rows(min_row=2):
            cell = row_cells[idx['品號']]
            if cell.value and str(cell.value).strip() == part_no:
                target_row = row_cells
                break

        if target_row is None:
            new_idx = ws.max_row + 1
            ws.cell(row=new_idx, column=idx['建立日期'] + 1, value=today)
            ws.cell(row=new_idx, column=idx['品號'] + 1, value=part_no)
            ws.cell(row=new_idx, column=idx['加工機型'] + 1, value=machine_type)
            for s in slots:
                if s.get('slot') in idx:
                    ws.cell(row=new_idx, column=idx[s['slot']] + 1, value=s.get('tool', ''))
        else:
            row_no = target_row[0].row
            ws.cell(row=row_no, column=idx['修改日期'] + 1, value=today)
            if machine_type:
                ws.cell(row=row_no, column=idx['加工機型'] + 1, value=machine_type)
            slot_cols = [h for h in header if h and re.match(r'^T\d+$', str(h))]
            filled = {s['slot']: s.get('tool', '') for s in slots if s.get('slot')}
            for slot in slot_cols:
                ws.cell(row=row_no, column=idx[slot] + 1, value=filled.get(slot, ''))

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map/list')
def batch_cost_tool_map_list():
    """讀『刀表』分頁全部品號的刀具配置清單（維護頁用）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        slot_cols = [h for h in header if h and re.match(r'^T\d+$', str(h))]
        items = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            part_no = row[idx.get('品號', -1)] if '品號' in idx else None
            if not part_no:
                continue
            slots = []
            for slot in slot_cols:
                v = row[idx[slot]]
                if v:
                    slots.append({'slot': slot, 'tool': str(v).strip()})
            created = row[idx['建立日期']] if '建立日期' in idx else ''
            modified = row[idx['修改日期']] if '修改日期' in idx else ''
            items.append({
                'part_no': str(part_no).strip(),
                'machine_type': str(row[idx.get('加工機型', -1)] or '').strip() if '加工機型' in idx else '',
                'created': str(created) if created else '',
                'modified': str(modified) if modified else '',
                'slots': slots,
            })
        wb.close()
        return jsonify({'success': True, 'items': items})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map/delete', methods=['POST'])
def batch_cost_tool_map_delete():
    """刪除『刀表』分頁裡某品號的整列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    part_no = (data.get('part_no') or '').strip()
    if not part_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        target_row_no = None
        for row_cells in ws.iter_rows(min_row=2):
            cell = row_cells[idx['品號']]
            if cell.value and str(cell.value).strip() == part_no:
                target_row_no = row_cells[0].row
                break
        if target_row_no is None:
            return jsonify({'success': False, 'error': '查無此品號的刀表'}), 404
        ws.delete_rows(target_row_no, 1)

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/api/batch_cost/save', methods=['POST'])
def batch_cost_save():
    """把一筆批成本計算結果 append 進『批成本計算』分頁（不存在時自動建立；依表頭欄名對應寫入，不依賴固定欄位順序）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    order      = (data.get('order') or '').strip()
    part_no    = (data.get('part_no') or '').strip()
    name       = (data.get('name') or '').strip()
    proc_code  = (data.get('proc_code') or '').strip()
    proc_name  = (data.get('proc_name') or '').strip()
    qty        = data.get('qty') or 0
    hours      = (data.get('hours') or '').strip()
    tool_cost  = data.get('tool_cost') or 0
    labor_cost = data.get('labor_cost') or 0
    seconds    = data.get('seconds') or 0
    rate       = data.get('rate') or 0
    machine_name = (data.get('machine_name') or '').strip()
    machine_code = (data.get('machine_code') or '').strip()
    out_time     = (data.get('out_time') or '').strip()
    tool_usage   = data.get('tool_usage') or []   # [{slot, tool, price, edges, usage, subtotal}, ...]
    if not order or not part_no:
        return jsonify({'success': False, 'error': '缺少製令或品號'}), 400

    qty_num = float(qty) if qty else 0
    tool_cost_per_unit = round((tool_cost / qty_num) if qty_num else 0, 2)
    total_cost_per_unit = round(labor_cost + tool_cost_per_unit, 2)

    default_header = ['建立日期', '製令', '製程代號', '製程名稱', '品號', '品名',
                       '每秒鐘生產費用（元）', '完成數量', '刀具費用', '加工秒數', '加工費用',
                       '刀具成本', '加工費用(含刀具成本)', '機台名稱', '機台代號']
    # 建立日期優先採用生產報工統計P2的出站時間，查不到才退回今天日期
    record_date = out_time or date.today().strftime('%Y/%m/%d')
    value_map = {
        '建立日期': record_date, '製令': order, '製程代號': proc_code, '製程名稱': proc_name,
        '品號': part_no, '品名': name, '每秒鐘生產費用（元）': rate, '完成數量': qty,
        '生產數量': qty,  # 相容舊版表頭命名
        '刀具費用': tool_cost, '加工秒數': seconds,
        '加工費用': labor_cost, '總加工費用': labor_cost,  # 相容舊版表頭命名
        '刀具成本': tool_cost_per_unit,
        '加工費用(含刀具成本)': total_cost_per_unit,
        '機台名稱': machine_name, '機台代號': machine_code,
        '加工時間': hours, '備註': hours,
    }

    try:
        wb = _load_batchcost_wb(read_only=False)
        if config.BATCH_COST_RECORD_SHEET not in wb.sheetnames:
            ws = wb.create_sheet(config.BATCH_COST_RECORD_SHEET)
            ws.append(default_header)
            header = default_header
        else:
            ws = wb[config.BATCH_COST_RECORD_SHEET]
            header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            if not any(header):
                header = default_header
                for col, h in enumerate(header, start=1):
                    ws.cell(row=1, column=col, value=h)

        new_row = [value_map.get(h, '') for h in header]
        ws.append(new_row)

        # 同步把每把刀的使用次數明細寫入『製令與刀具壽命』分頁（不存在時自動建立）
        used_tools = [t for t in tool_usage if (t.get('usage') or 0) > 0]
        if used_tools:
            lifespan_default_header = ['建立日期', '製令', '製程代號', '製程名稱', '品號', '品名',
                                        '刀號', '刀具名稱', '單價', '使用次數', '小計']
            if config.BATCH_COST_LIFESPAN_SHEET not in wb.sheetnames:
                ws2 = wb.create_sheet(config.BATCH_COST_LIFESPAN_SHEET)
                ws2.append(lifespan_default_header)
                header2 = lifespan_default_header
            else:
                ws2 = wb[config.BATCH_COST_LIFESPAN_SHEET]
                header2 = [c.value for c in next(ws2.iter_rows(min_row=1, max_row=1))]
                if not any(header2):
                    header2 = lifespan_default_header
                    for col, h in enumerate(header2, start=1):
                        ws2.cell(row=1, column=col, value=h)
            for t in used_tools:
                tool_value_map = {
                    '建立日期': record_date, '製令': order, '製程代號': proc_code, '製程名稱': proc_name,
                    '品號': part_no, '品名': name,
                    '刀號': t.get('slot', ''), '刀具名稱': t.get('tool', ''),
                    '單價': t.get('price', 0), '使用次數': t.get('usage', 0), '小計': t.get('subtotal', 0),
                }
                ws2.append([tool_value_map.get(h, '') for h in header2])

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/record/list')
def batch_cost_record_list():
    """讀『批成本計算』分頁全部明細（依目前表頭欄名動態回傳，不假設固定欄位）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        if config.BATCH_COST_RECORD_SHEET not in wb.sheetnames:
            return jsonify({'success': True, 'columns': [], 'records': []})
        ws = wb[config.BATCH_COST_RECORD_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        columns = [h for h in header if h]
        records = []
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue
            rec = {'row': row_no}
            for col_idx, h in enumerate(header):
                if not h:
                    continue
                v = row[col_idx] if col_idx < len(row) else ''
                rec[h] = str(v) if v is not None else ''
            records.append(rec)
        wb.close()

        def _date_key(rec):
            date_str = (rec.get('建立日期') or '').split(' ')[0]
            try:
                parts = date_str.replace('-', '/').split('/')
                return tuple(int(p) for p in parts)
            except (ValueError, TypeError):
                return (0, 0, 0)

        records.sort(key=_date_key, reverse=True)  # 依建立日期由新到舊排序
        return jsonify({'success': True, 'columns': columns, 'records': records})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/lifespan/list')
def batch_cost_lifespan_list():
    """讀『製令與刀具壽命』分頁全部明細（依目前表頭欄名動態回傳，供批成本明細展開列使用）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        if config.BATCH_COST_LIFESPAN_SHEET not in wb.sheetnames:
            return jsonify({'success': True, 'columns': [], 'records': []})

        # 依刀具名稱查供應商/項目（來自刀具資料分頁）
        tool_extra = {}
        if config.BATCH_COST_TOOL_SHEET in wb.sheetnames:
            tws = wb[config.BATCH_COST_TOOL_SHEET]
            theader = [c.value for c in next(tws.iter_rows(min_row=1, max_row=1))]
            tidx = {h: i for i, h in enumerate(theader) if h}
            for trow in tws.iter_rows(min_row=2, values_only=True):
                tname = trow[tidx.get('品名', -1)] if '品名' in tidx else None
                if not tname:
                    continue
                tname = str(tname).strip()
                if tname not in tool_extra:
                    tool_extra[tname] = {
                        'item':   str(trow[tidx['項目']] or '').strip() if '項目' in tidx else '',
                        'vendor': str(trow[tidx['供應商']] or '').strip() if '供應商' in tidx else '',
                    }

        ws = wb[config.BATCH_COST_LIFESPAN_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        columns = [h for h in header if h]
        records = []
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue
            rec = {'row': row_no}
            for col_idx, h in enumerate(header):
                if not h:
                    continue
                v = row[col_idx] if col_idx < len(row) else ''
                rec[h] = str(v) if v is not None else ''
            extra = tool_extra.get((rec.get('刀具名稱') or '').strip(), {})
            rec['項目'] = extra.get('item', '')
            rec['供應商'] = extra.get('vendor', '')
            records.append(rec)
        wb.close()
        if '項目' not in columns:
            columns.append('項目')
        if '供應商' not in columns:
            columns.append('供應商')
        return jsonify({'success': True, 'columns': columns, 'records': records})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/record/delete', methods=['POST'])
def batch_cost_record_delete():
    """刪除『批成本計算』分頁裡的一列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    if not row_no:
        return jsonify({'success': False, 'error': '缺少列號'}), 400
    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_RECORD_SHEET]
        ws.delete_rows(int(row_no), 1)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/batch_cost')
def batch_cost_page():
    """批成本計算頁面"""
    return render_template('batch_cost.html', app_version=APP_VERSION)


@app.route('/api/scrap/list')
def scrap_list():
    """讀取 Google 試算表『報廢統計』資料"""
    cache_key = 'scrap_list'
    if request.args.get('refresh'):
        cache_clear(cache_key)
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.SCRAP_SHEET_ID, gid=config.SCRAP_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    CATEGORY_MAP = {'A': 'CNC', 'B': '鑽床', 'C': '沖床', 'L': '車床'}
    scraps = []
    for r in rows:
        raw_amount = (r.get('報廢金額') or '').replace(',', '').strip()
        try:
            amount = float(raw_amount)
        except ValueError:
            amount = 0
        raw_qty = (r.get('出庫異動數量') or '').replace(',', '').strip()
        try:
            qty = float(raw_qty)
        except ValueError:
            qty = 0
        cat = (r.get('分類') or '').strip().upper()
        scraps.append({
            'date':     (r.get('異動日期') or '').strip(),
            'doc_no':   (r.get('單別-單號') or '').strip(),
            'part_no':  (r.get('品號') or '').strip(),
            'name':     (r.get('品名') or '').strip(),
            'spec':     (r.get('規格') or '').strip(),
            'dept':     (r.get('部門名稱') or '').strip(),
            'remark':   (r.get('備註') or '').strip(),
            'qty':      qty,
            'amount':   amount,
            'category': cat,
            'cat_name': CATEGORY_MAP.get(cat, cat),
        })

    def _sort_key(item):
        try:
            parts = item['date'].replace('-', '/').split('/')
            return tuple(int(p) for p in parts)
        except Exception:
            return (0, 0, 0)

    scraps.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(scraps), 'scraps': scraps}
    cache_set(cache_key, result)
    return jsonify(result)


def _build_category_map():
    """讀取 K1_P2.ref 試算表，建立 {品號+製程: A/B/C} 對照表（含 Flask 快取）"""
    cache_key = 'category_map'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached.get('map', {})
    try:
        cat_rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        cat_map = {}
        for r in cat_rows:
            part_no      = (r.get('品號')   or '').strip()
            process_code = (r.get('製程')   or '').strip()
            category     = (r.get('A/B/C分類') or r.get('A/B/C') or '').strip()
            if part_no and process_code and category:
                cat_map[part_no + process_code] = category
        result = {'success': True, 'map': cat_map}
        cache_set(cache_key, result)
        return cat_map
    except Exception:
        return {}


@app.route('/api/prod_report/list')
def prod_report_list():
    """讀取 Google 試算表『生產日報表』資料（含 ABC 分類）"""
    cache_key = 'prod_report_list'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'prod_report_monthly_stats_v1', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, gid=config.PROD_REPORT_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    cat_map = _build_category_map()

    records = []
    for r in rows:
        raw_qty = (r.get('生產數') or '').replace(',', '').strip()
        try:    qty = float(raw_qty)
        except ValueError: qty = 0
        raw_sec = (r.get('秒數') or '').replace(',', '').strip()
        try:    seconds = float(raw_sec)
        except ValueError: seconds = 0
        part_no      = (r.get('品號') or '').strip()
        process_code = (r.get('製程代號') or '').strip()
        machine_name = (r.get('機台名稱') or '').strip()
        machine_code = (r.get('機台代號') or '').strip()

        # 分類判斷：① 試算表直接欄位 → ② 品號+製程對照表 → ③ 機台名稱/代號關鍵字
        cat = (r.get('A/B/C') or r.get('類別') or r.get('分類') or '').strip()
        if cat not in ('A', 'B', 'C', 'L'):
            cat = cat_map.get(part_no + process_code, '')
        if cat not in ('A', 'B', 'C', 'L'):
            kw = (machine_name + machine_code + process_code).upper()
            if 'CNC' in kw or '加工中心' in kw or '綜合加工' in kw:
                cat = 'A'
            elif '鑽' in kw:
                cat = 'B'
            elif '沖' in kw or '衝' in kw:
                cat = 'C'
            elif '車' in kw:
                cat = 'L'

        records.append({
            'std_time':     (r.get('標工') or '').strip(),
            'date':         (r.get('生產日期') or '').strip(),
            'operator':     (r.get('生產人員') or '').strip(),
            'work_order':   (r.get('製令') or '').strip(),
            'part_no':      part_no,
            'name':         (r.get('品名') or '').strip(),
            'process_code': process_code,
            'qty':          qty,
            'seconds':      seconds,
            'machine_code': machine_code,
            'machine_name': machine_name,
            'remark':       (r.get('備註') or '').strip(),
            'category':     cat,
        })

    def _sort_key(item):
        try:
            parts = item['date'].replace('-', '/').split('/')
            return tuple(int(p) for p in parts)
        except Exception:
            return (0, 0, 0)

    records.sort(key=_sort_key, reverse=True)
    result = {'success': True, 'count': len(records), 'records': records}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/prod_report/monthly_stats')
def prod_report_monthly_stats():
    """生產日報表圖表資料：讀『P5.3生產日報表data_ref』分頁，
    依 A/B/C/L 分類 × 年月 彙整生產數（資料來源與『生產報工統計P2』的出站數量不同）"""
    cache_key = 'prod_report_monthly_stats_v1'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'prod_report_list', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, gid=config.PROD_REPORT_CHART_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    cat_map = _build_category_map()
    from collections import defaultdict
    month_data = defaultdict(lambda: {
        'A': 0, 'B': 0, 'C': 0, 'L': 0, 'other': 0,
        'records': 0, 'uncat': 0,
    })

    for r in rows:
        date_str = (r.get('生產日期') or '').strip()
        if not date_str:
            continue
        clean = date_str.split(' ')[0].replace('-', '/')
        parts = clean.split('/')
        if len(parts) < 2:
            continue
        try:
            ym = f'{int(parts[0])}-{int(parts[1]):02d}'
        except Exception:
            continue

        raw_qty = (r.get('生產數') or '').replace(',', '').strip()
        try:
            qty = int(float(raw_qty)) if raw_qty else 0
        except ValueError:
            qty = 0

        part_no      = (r.get('品號') or '').strip()
        process_code = (r.get('製程代號') or '').strip()
        machine_code = (r.get('機台代號') or '').strip()
        machine_name = (r.get('機台名稱') or '').strip()

        # 分類：① 品號+製程對照表 → ② 機台名稱/代號/製程關鍵字
        cat = cat_map.get(part_no + process_code, '')
        if cat not in ('A', 'B', 'C', 'L'):
            kw = (machine_name + machine_code + process_code).upper()
            if 'CNC' in kw or '加工中心' in kw or '綜合加工' in kw:
                cat = 'A'
            elif '鑽' in kw:
                cat = 'B'
            elif '沖' in kw or '衝' in kw:
                cat = 'C'
            elif '車' in kw:
                cat = 'L'

        d = month_data[ym]
        if cat in ('A', 'B', 'C', 'L'):
            if qty > 0:
                d[cat]       += qty
                d['records'] += 1
        else:
            d['uncat'] += 1
            if qty > 0:
                d['other'] += qty

    months = []
    for ym in sorted(month_data.keys()):
        d = month_data[ym]
        classified = d['A'] + d['B'] + d['C'] + d['L']
        months.append({
            'month': ym,
            'A': d['A'], 'B': d['B'], 'C': d['C'], 'L': d['L'],
            'other': d['other'],
            'total': classified + d['other'],
            'records': d['records'],
            'uncat': d['uncat'],
        })

    result = {'success': True, 'months': months}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/category/map')
def category_map():
    """讀取 Google 試算表，回傳 {品號+製程代號: A/B/C分類} 對照表"""
    cache_key = 'category_map'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表'}), 502

    cat_map = {}
    for r in rows:
        # 欄位依序：製令(A)、預計數量(B)、品號(C)、品名(D)...製程代號(F)...A/B/C(Q)
        # fetch_google_sheet_csv 已用第一列 header 做 key，直接用 index 取值更安全
        vals = list(r.values())
        if len(vals) < 17:
            continue
        part_no      = vals[2].strip()   # C欄：品號
        process_code = vals[5].strip()   # F欄：製程代號
        category     = vals[16].strip()  # Q欄：A/B/C 分類
        if part_no and process_code and category:
            cat_map[part_no + process_code] = category

    result = {'success': True, 'map': cat_map}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/k1p2/list')
def k1p2_list():
    """讀取 K1_P2.ref 全部明細記錄"""
    cache_key = 'k1p2_list'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'k1p2_monthly_stats_v3', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)
    try:
        rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        records = []
        for r in rows:
            out_str = (r.get('欄位格式化') or r.get('出站時間') or '').strip()
            ym = ''
            if out_str:
                parts = out_str.replace('-', '/').split('/')
                if len(parts) >= 2:
                    try: ym = f'{int(parts[0])}-{int(parts[1]):02d}'
                    except: pass
            records.append({
                'wo':           (r.get('製令')           or '').strip(),
                'plan_qty':     (r.get('預計產量')        or '').strip(),
                'part_no':      (r.get('品號')            or '').strip(),
                'name':         (r.get('品名')            or '').strip(),
                'seq':          (r.get('加工順序')        or '').strip(),
                'process':      (r.get('製程')            or '').strip(),
                'process_name': (r.get('製程名稱')        or '').strip(),
                'qty':          (r.get('出站數量')        or '').strip(),
                'seconds':      (r.get('出站總工時(秒)')  or '').strip(),
                'in_time':      (r.get('進站時間')        or '').strip(),
                'out_time':     out_str,
                'operator':     (r.get('報工人員')        or '').strip(),
                'category':     (r.get('A/B/C分類') or r.get('A/B/C') or '').strip(),
                'ym':           ym,
            })
        records.sort(key=lambda x: x.get('out_time', ''), reverse=True)
        result = {'success': True, 'count': len(records), 'records': records}
        cache_set(cache_key, result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/k1p2/monthly_stats')
def k1p2_monthly_stats():
    """K1_P2.ref 月份分類統計：出站數量 × A/B/C/L × 年月"""
    cache_key = 'k1p2_monthly_stats_v3'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'k1p2_list', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)
    try:
        rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        from collections import defaultdict
        month_data = defaultdict(lambda: {
            'A': 0, 'B': 0, 'C': 0, 'L': 0, 'other': 0,
            'sec_A': 0, 'sec_B': 0, 'sec_C': 0, 'sec_L': 0,
            'records': 0, 'uncat': 0, 'seconds': 0,
        })
        for r in rows:
            # 優先用欄位格式化（YYYY/MM/DD），備援出站時間（含中文上午/下午）
            date_str = (r.get('欄位格式化') or r.get('出站時間') or '').strip()
            qty_str  = (r.get('出站數量') or '').strip()
            # 相容半型 () 和全型（）括號
            sec_str  = (r.get('出站總工時(秒)') or r.get('出站總工時（秒）') or '').strip()
            cat      = (r.get('A/B/C') or r.get('A/B/C分類') or '').strip()
            try:
                qty = int(float(qty_str.replace(',', ''))) if qty_str else 0
            except Exception:
                qty = 0
            try:
                sec = int(float(sec_str.replace(',', ''))) if sec_str else 0
            except Exception:
                sec = 0
            if not date_str:
                continue
            # 日期解析 → 年月（先取日期部份再分割，避免時間字串干擾）
            ym = ''
            clean_date = date_str.split(' ')[0]
            if '/' in clean_date:
                parts = clean_date.split('/')
                if len(parts) >= 2:
                    try:
                        ym = f'{int(parts[0])}-{int(parts[1]):02d}'
                    except Exception:
                        pass
            elif '-' in clean_date:
                parts = clean_date.split('-')
                if len(parts) >= 2:
                    try:
                        ym = f'{parts[0]}-{int(parts[1]):02d}'
                    except Exception:
                        pass
            if not ym:
                continue
            month_data[ym]['seconds'] += sec
            if cat in ('A', 'B', 'C', 'L'):
                month_data[ym][f'sec_{cat}'] += sec   # 秒數：不論 qty 是否為 0 都累加
                if qty > 0:
                    month_data[ym][cat]       += qty
                    month_data[ym]['records'] += 1
            else:
                month_data[ym]['uncat'] += 1
                if qty > 0:
                    month_data[ym]['other'] += qty

        months = []
        for ym in sorted(month_data.keys()):
            d = month_data[ym]
            classified = d['A'] + d['B'] + d['C'] + d['L']
            months.append({
                'month': ym,
                'A': d['A'], 'B': d['B'], 'C': d['C'], 'L': d['L'],
                'sec_A': d['sec_A'], 'sec_B': d['sec_B'],
                'sec_C': d['sec_C'], 'sec_L': d['sec_L'],
                'other': d['other'],
                'total': classified + d['other'],
                'records': d['records'],
                'uncat': d['uncat'],
                'seconds': d['seconds'],
            })

        result = {'success': True, 'months': months}
        cache_set(cache_key, result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/jig/list')
def jig_list():
    """讀取治檢具清單索引（PDM 資料夾資料卡，build_jig_index.py 建立）"""
    conn = get_pdm_db()
    if not conn:
        return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

    try:
        cur = conn.execute(
            'SELECT folder_name, folder_path, product_model, item_name, '
            '       handler, submitter, status, apply_date, unit, remarks '
            'FROM jig_index ORDER BY folder_name DESC'
        )
        jigs = [dict(r) for r in cur.fetchall()]
        row = conn.execute('SELECT MAX(indexed_at) FROM jig_index').fetchone()
        last_updated = row[0] if row and row[0] else None
    except sqlite3.OperationalError:
        jigs = []
        last_updated = None
    finally:
        conn.close()

    return jsonify({'success': True, 'count': len(jigs), 'jigs': jigs,
                    'last_updated': last_updated})


@app.route('/api/jig/rebuild', methods=['POST'])
def jig_rebuild():
    """重建治檢具索引（在背景執行 build_jig_index.py --deploy）"""
    import subprocess, re as _re
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_jig_index.py')
    if not os.path.exists(script):
        return jsonify({'success': False, 'error': 'build_jig_index.py 不存在'}), 500
    try:
        script_dir = os.path.dirname(script)
        result = subprocess.run(
            [sys.executable, script, '--deploy'],
            capture_output=True, text=True, timeout=300,
            cwd=script_dir
        )
        if result.returncode != 0:
            err = (result.stderr or result.stdout)[-500:]
            return jsonify({'success': False, 'error': err}), 500
        m = _re.search(r'寫入：(\d+)', result.stdout)
        count = int(m.group(1)) if m else 0
        return jsonify({'success': True, 'count': count})
    except subprocess.TimeoutExpired:
        return jsonify({'success': False, 'error': '更新逾時（超過 5 分鐘）'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/jig/open-folder', methods=['POST'])
def jig_open_folder():
    """用 Shell 開啟治檢具資料夾"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400

    jig_root = os.path.normpath(config.JIG_VAULT_PATH)
    norm_path = os.path.normpath(path)
    if not norm_path.lower().startswith(jig_root.lower()):
        return jsonify({'ok': False, 'error': '無效的路徑'}), 400
    if not os.path.isdir(norm_path):
        return jsonify({'ok': False, 'error': '資料夾不存在'}), 404

    try:
        os.startfile(norm_path)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/dcn/list')
def dcn_list():
    """讀取設計變更通知單索引（PDM 資料夾，build_dcn_index.py 建立）"""
    conn = get_pdm_db()
    if not conn:
        return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

    try:
        cur = conn.execute(
            'SELECT folder_name, folder_path, issue_date, product_model, '
            '       submitter, handler, reviewer, approver, reason, description, status '
            'FROM dcn_index ORDER BY folder_name DESC'
        )
        rows = [dict(r) for r in cur.fetchall()]
        row = conn.execute('SELECT MAX(indexed_at) FROM dcn_index').fetchone()
        last_updated = row[0] if row and row[0] else None
    except sqlite3.OperationalError:
        rows = []
        last_updated = None
    finally:
        conn.close()

    return jsonify({'success': True, 'count': len(rows), 'dcns': rows,
                    'last_updated': last_updated})


@app.route('/api/dcn/rebuild', methods=['POST'])
def dcn_rebuild():
    """重建設計變更通知單索引（背景執行 build_dcn_index.py --deploy）"""
    import subprocess, re as _re
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_dcn_index.py')
    if not os.path.exists(script):
        return jsonify({'success': False, 'error': 'build_dcn_index.py 不存在'}), 500
    try:
        script_dir = os.path.dirname(script)
        result = subprocess.run(
            [sys.executable, script, '--deploy'],
            capture_output=True, text=True, timeout=600,
            cwd=script_dir
        )
        if result.returncode != 0:
            err = (result.stderr or result.stdout)[-500:]
            return jsonify({'success': False, 'error': err}), 500
        m = _re.search(r'寫入：(\d+)', result.stdout)
        count = int(m.group(1)) if m else 0
        return jsonify({'success': True, 'count': count})
    except subprocess.TimeoutExpired:
        return jsonify({'success': False, 'error': '更新逾時（超過 10 分鐘）'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/dcn/open-folder', methods=['POST'])
def dcn_open_folder():
    """用 Shell 開啟設計變更通知單資料夾"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400

    dcn_root = os.path.normpath(config.DCN_VAULT_PATH)
    norm_path = os.path.normpath(path)
    if not norm_path.lower().startswith(dcn_root.lower()):
        return jsonify({'ok': False, 'error': '無效的路徑'}), 400
    if not os.path.isdir(norm_path):
        return jsonify({'ok': False, 'error': '資料夾不存在'}), 404

    try:
        os.startfile(norm_path)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _preload_cache():
    """背景預載 SSRS 資料到快取（加速首次搜尋）"""
    try:
        with ThreadPoolExecutor(max_workers=2) as pool:
            pool.submit(fetch_all_unfinished)
            pool.submit(fetch_efficiency_data)
    except Exception:
        pass


if __name__ == '__main__':
    # 獨立執行模式（除錯用）— GUI 版由 main.py 啟動
    threading.Thread(target=_preload_cache, daemon=True).start()
    app.run(host=config.FLASK_HOST, port=config.FLASK_PORT, debug=config.FLASK_DEBUG)
