import csv
import hashlib
import io
import json
import os
import sys
import re
import sqlite3
import time
import threading
import uuid
import webbrowser
import urllib.parse
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime, timedelta
from pathlib import Path
from flask import Flask, render_template, request, jsonify, redirect
import requests
from requests_ntlm import HttpNtlmAuth
from urllib.parse import quote
try:
    import openpyxl
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm, Pt
    _DOCXTPL_OK = True
except ImportError:
    _DOCXTPL_OK = False

# ── PyInstaller 路徑處理 ────────────────────────────────────────
# 執行為 EXE 時：_MEIPASS 為解壓目錄（含 templates）；
# EXE 同目錄為使用者可編輯的 config.py / pdm_search.db
if getattr(sys, 'frozen', False):
    _BASE = sys._MEIPASS          # 唯讀資源（templates 等）
    _APP_DIR = os.path.dirname(sys.executable)  # EXE 所在目錄
    sys.path.insert(0, _APP_DIR)  # 讓 import config 找到 EXE 旁的 config.py
else:
    _BASE = os.path.dirname(os.path.abspath(__file__))
    _APP_DIR = _BASE
    # python311._pth 存在時，嵌入式 Python 不自動加入腳本目錄
    # 需手動插入，讓 import config / templates 等都能正常運作
    if _APP_DIR not in sys.path:
        sys.path.insert(0, _APP_DIR)
    _PARENT = os.path.dirname(_APP_DIR)
    if os.path.exists(os.path.join(_PARENT, 'config.py')) and _PARENT not in sys.path:
        sys.path.insert(0, _PARENT)

import config

# 自主巡檢表路徑
INSPECTION_ROOT = r"\\192.168.1.99\共用區\品保加工共用平台\加工課自主巡檢表(空白)"

# PDM 圖面索引資料庫（可在 config.py 覆寫）
PDM_DB_PATH = getattr(config, 'PDM_DB_PATH',
    r"C:\Users\user\AppData\Local\PDMSearch\pdm_search.db")

# Fallback：若設定路徑不存在，嘗試同目錄下的 pdm_search.db（打包時隨包）
if not os.path.exists(PDM_DB_PATH):
    _bundled_db = os.path.join(_APP_DIR, 'pdm_search.db')
    if os.path.exists(_bundled_db):
        PDM_DB_PATH = _bundled_db

# ── 技術資料清單資料庫（zume-n.com 圖號↔URL 對照表）──────────────────────
ZUME_DB_PATH = os.path.join(_APP_DIR, 'zume_drawings.db')

# ── CNC 程式索引資料庫（build_cnc_program_index.py 建立）──────────────────
CNC_DB_PATH = os.path.join(_APP_DIR, 'cnc_program_index.db')

# ── 設備主檔資料庫（build_equipment_index.py 建立，見 docs/equipment-master.md）──
# 2026-08-07：改讀 config.EQUIPMENT_DB_PATH（網芳共用路徑，多人共用同一份設備資料）；
# 沒設定時退回本機路徑，向下相容舊的 config.py。
EQ_DB_PATH = getattr(config, 'EQUIPMENT_DB_PATH', None) or os.path.join(_APP_DIR, 'equipment.db')

# ── 油品主檔資料庫（build_oil_index.py 建立，見 docs/oil-management.md）──────────
# 跟 equipment.db 一樣放網芳共用，讓多人在不同電腦同時使用；沒設定時退回本機路徑。
OIL_DB_PATH = getattr(config, 'OIL_DB_PATH', None) or os.path.join(_APP_DIR, 'oil.db')

# ── 工作日行事曆（純使用者輸入，沒有任何 build 工具會產生它，見 docs/calendar.md）──
# 執行期資料庫：桌面應用執行中會直接寫入，絕對不可以加進 sync_to_dist.py 的覆蓋清單
CAL_DB_PATH = os.path.join(_APP_DIR, 'calendar.db')

# ZUMEN 圖面欄位 → CSV 標題關鍵字對照（依關鍵字動態偵測欄位，缺的留空）
_ZUME_EXTRA_COLS = ['line', 'prod_group', 'category', 'vendor']
_ZUME_COL_KEYWORDS = {
    'line':       ['生產線別', '線別'],
    'prod_group': ['生產群組', '群組'],
    'category':   ['分類'],
    'vendor':     ['廠商', '客戶'],
}

def _ensure_zume_columns(con):
    """相容遷移：drawings 表若缺新欄位則補上（line/prod_group/category/vendor）"""
    existing = {r[1] for r in con.execute('PRAGMA table_info(drawings)')}
    for col in _ZUME_EXTRA_COLS:
        if col not in existing:
            con.execute(f'ALTER TABLE drawings ADD COLUMN {col} TEXT')

def _init_zume_db():
    con = sqlite3.connect(ZUME_DB_PATH)
    con.execute('''
        CREATE TABLE IF NOT EXISTS drawings (
            part_no   TEXT PRIMARY KEY,
            part_name TEXT,
            url       TEXT NOT NULL
        )
    ''')
    con.execute('''
        CREATE TABLE IF NOT EXISTS import_log (
            filename  TEXT PRIMARY KEY,
            imported_at TEXT,
            count     INTEGER
        )
    ''')
    _ensure_zume_columns(con)
    con.commit(); con.close()

def _zume_header_indices(headers):
    """依標題關鍵字回傳各欄位的索引 dict（找不到的為 None）"""
    idx = {
        'part_no':   next((i for i, h in enumerate(headers) if '圖號' in h), None),
        'url':       next((i for i, h in enumerate(headers) if h.strip().upper() == 'URL'), None),
        'part_name': next((i for i, h in enumerate(headers) if '品名' in h), 1),
    }
    for col, kws in _ZUME_COL_KEYWORDS.items():
        idx[col] = next((i for i, h in enumerate(headers) if any(k in h for k in kws)), None)
    return idx

def _zume_row_to_rec(row, idx):
    """依欄位索引把一列轉成 dict；圖號/URL 缺或無效則回傳 None"""
    def cell(col):
        i = idx.get(col)
        return str(row[i]).strip() if (i is not None and i < len(row) and row[i] is not None) else ''
    no  = cell('part_no')
    url = cell('url')
    if not no or not url.startswith('http'):
        return None
    return {
        'part_no': no, 'part_name': cell('part_name'), 'url': url,
        'line': cell('line'), 'prod_group': cell('prod_group'),
        'category': cell('category'), 'vendor': cell('vendor'),
    }

def _insert_zume_recs(con, recs):
    """把解析後的 dict 清單寫入 drawings 表"""
    con.executemany(
        'INSERT OR REPLACE INTO drawings'
        '(part_no, part_name, url, line, prod_group, category, vendor) '
        'VALUES (:part_no, :part_name, :url, :line, :prod_group, :category, :vendor)',
        recs)

def _parse_zume_csv(filepath):
    """解析 zume-n_data_list_*.csv，回傳 [dict, ...]"""
    recs = []
    try:
        with open(filepath, encoding='utf-8-sig', newline='') as f:
            reader  = csv.reader(f)
            headers = [h.strip() for h in next(reader)]
            idx     = _zume_header_indices(headers)
            if idx['part_no'] is None or idx['url'] is None:
                return []
            for row in reader:
                rec = _zume_row_to_rec(row, idx)
                if rec:
                    recs.append(rec)
    except Exception:
        pass
    return recs

def _zume_csv_search_dirs():
    """回傳要掃描 zume-n_data_list_*.csv 的候選資料夾（去重、僅保留存在者）"""
    home = os.path.expanduser('~')
    cands = [
        os.path.join(home, 'Downloads'),
        os.path.join(home, '下載'),
        os.path.join(home, 'Desktop'),
        os.path.join(home, 'OneDrive', 'Downloads'),
        os.path.join(home, 'OneDrive', '下載'),
        os.path.join(home, 'OneDrive', 'Desktop'),
    ]
    seen, dirs = set(), []
    for d in cands:
        key = os.path.normcase(os.path.abspath(d))
        if key not in seen and os.path.isdir(d):
            seen.add(key); dirs.append(d)
    return dirs


def _find_zume_csvs():
    """跨多個常見資料夾搜尋 zume-n_data_list_*.csv，依修改時間新→舊排序"""
    import glob as _glob
    files = []
    for d in _zume_csv_search_dirs():
        files.extend(_glob.glob(os.path.join(d, 'zume-n_data_list_*.csv')))
    return sorted(files, key=os.path.getmtime, reverse=True)


def _auto_import_zume_csv():
    """啟動時自動掃描下載/桌面等資料夾，匯入最新的 zume-n_data_list_*.csv"""
    files = _find_zume_csvs()
    if not files:
        return
    latest = files[0]
    fname  = os.path.basename(latest)
    con = sqlite3.connect(ZUME_DB_PATH)
    already = con.execute('SELECT count FROM import_log WHERE filename=?', (fname,)).fetchone()
    con.close()
    if already:
        return  # 已匯入過，略過
    recs = _parse_zume_csv(latest)
    if not recs:
        return
    con = sqlite3.connect(ZUME_DB_PATH)
    _insert_zume_recs(con, recs)
    con.execute('INSERT OR REPLACE INTO import_log(filename,imported_at,count) VALUES(?,datetime("now"),?)',
                (fname, len(recs)))
    con.commit(); con.close()

_init_zume_db()
_auto_import_zume_csv()   # 啟動時自動掃描匯入

app = Flask(__name__,
            template_folder=os.path.join(_BASE, 'templates'),
            static_folder=os.path.join(_BASE, 'static'))
app.config['TEMPLATES_AUTO_RELOAD'] = True

@app.after_request
def set_no_cache(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    # CORS：允許本機的 ZUMEN 圖面管理工具（Node, port 3456）呼叫本系統 API
    # （PDM 圖面查詢/開啟、途程查詢整合面板用；僅限 localhost 來源，不對外開放）
    origin = request.headers.get('Origin', '')
    if origin in ('http://localhost:3456', 'http://127.0.0.1:3456'):
        response.headers['Access-Control-Allow-Origin'] = origin
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    # 動態注入版本號到所有 HTML 頁面（繞過模板快取）
    if response.content_type and 'text/html' in response.content_type:
        try:
            html = response.get_data(as_text=True)
            if '</head>' in html and APP_VERSION not in html:
                inject = f'<script>document.title="Maxclaw PDS系統 {APP_VERSION}";</script>\n</head>'
                html = html.replace('</head>', inject)
                response.set_data(html)
        except Exception:
            pass
    return response

# ── 全域錯誤處理（確保所有錯誤都回傳 JSON，而非 HTML）────────
@app.errorhandler(400)
def bad_request(e):
    return jsonify({'success': False, 'error': f'請求錯誤：{str(e)}'}), 400

@app.errorhandler(404)
def not_found(e):
    return jsonify({'success': False, 'error': f'路由不存在：{str(e)}'}), 404

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'success': False, 'error': f'伺服器錯誤：{str(e)}'}), 500

@app.errorhandler(Exception)
def unhandled_exception(e):
    return jsonify({'success': False, 'error': f'未預期的錯誤：{str(e)}'}), 500

# ── TTL 快取（避免每次搜尋都重新呼叫 SSRS）──────────────────
CACHE_TTL = 120  # 快取有效期 120 秒

_cache = {}
_cache_lock = threading.Lock()


def cache_get(key):
    """取得快取資料，過期回傳 None"""
    with _cache_lock:
        entry = _cache.get(key)
        if entry and (time.time() - entry['ts']) < CACHE_TTL:
            return entry['data']
    return None


def cache_set(key, data):
    """設定快取資料"""
    with _cache_lock:
        _cache[key] = {'data': data, 'ts': time.time()}


def cache_clear(*keys):
    """清除指定的快取 key（用於「重新整理」強制重抓最新資料）"""
    with _cache_lock:
        for k in keys:
            _cache.pop(k, None)


def match_and_not(query, text):
    """AND/NOT 模糊搜尋語法
    空格分隔 = AND（全部都要符合）
    -前綴 = NOT（排除）
    例: 'FTB 夾管座'   → 包含 FTB 且包含 夾管座
        'FTB -夾管座'  → 包含 FTB 但不包含 夾管座
    """
    if not query:
        return True
    text_upper = text.upper()
    for token in query.strip().split():
        if token.startswith('-') and len(token) > 1:
            if token[1:].upper() in text_upper:
                return False
        else:
            if token.upper() not in text_upper:
                return False
    return True


def fetch_all_unfinished(release_status='已發放'):
    """取得所有未完工製令（帶 TTL 快取）
    release_status: '已發放'（預設，產線上已在跑的）／'未發放'（ERP製令發放作業裡還沒發放的）／'all'（兩者都查再合併）
    """
    if release_status == 'all':
        return fetch_all_unfinished('已發放') + fetch_all_unfinished('未發放')

    cache_key = f'unfinished_{release_status}'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached
    params = {
        '發放情況': release_status,
        'SFT完工碼': '尚未',
        '加工單位': '*',
    }
    csv_text = fetch_report_csv(config.REPORT_PATHS['unfinished'], params)
    data = parse_csv(csv_text)
    cache_set(cache_key, data)
    return data


_ssrs_session = None
_ssrs_session_lock = threading.Lock()


def get_ssrs_session():
    """建立/重用 SSRS NTLM 認證連線（連線池，避免每次建立新連線）"""
    global _ssrs_session
    with _ssrs_session_lock:
        if _ssrs_session is None:
            _ssrs_session = requests.Session()
            _ssrs_session.auth = HttpNtlmAuth(config.SSRS_USERNAME, config.SSRS_PASSWORD)
        return _ssrs_session


def fetch_report_csv(report_path, params=None):
    """從 SSRS 取得報表 CSV 資料
    NTLM 連線在多執行緒並行（ThreadPoolExecutor）冷啟動時偶發 401（共用 session 的已知問題），
    遇到 401 時丟掉舊 session 重建一次再試，避免偶發失敗"""
    global _ssrs_session
    url = f'{config.SSRS_BASE_URL}?{quote(report_path)}&rs:Command=Render&rs:Format=CSV'
    if params:
        param_str = '&'.join(f'{quote(k)}={quote(str(v))}' for k, v in params.items())
        url += f'&{param_str}'
    session = get_ssrs_session()
    resp = session.get(url, timeout=30)
    if resp.status_code == 401:
        with _ssrs_session_lock:
            _ssrs_session = None
        session = get_ssrs_session()
        resp = session.get(url, timeout=30)
    resp.raise_for_status()
    return resp.content.decode('utf-8-sig')


def fetch_google_sheet_csv(sheet_id, sheet_name=None, gid=None):
    """從 Google 試算表（共用設定：知道連結的人均可檢視）取得指定分頁的 CSV 資料
    可用分頁名稱（sheet_name）或分頁 gid 指定要讀取的分頁"""
    if gid is not None:
        url = f'https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&gid={gid}'
    else:
        url = f'https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={quote(sheet_name)}'
    resp = requests.get(url, timeout=10)
    resp.raise_for_status()
    text = resp.content.decode('utf-8-sig')
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return []
    header = rows[0]
    return [dict(zip(header, row)) for row in rows[1:] if any(cell.strip() for cell in row)]


def fetch_employee_name_map():
    """讀取 Google 試算表『員工登錄系統』，回傳 {工號ID: 姓名} 對照表（含快取）"""
    cache_key = 'employee_name_map'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached

    rows = fetch_google_sheet_csv(config.EMPLOYEE_SHEET_ID, gid=config.EMPLOYEE_SHEET_GID)
    name_map = {}
    for r in rows:
        emp_id = (r.get('工號ID') or '').strip()
        emp_name = (r.get('姓名') or '').strip()
        if emp_id and emp_name:
            name_map[emp_id] = emp_name

    cache_set(cache_key, name_map)
    return name_map


def fetch_employee_roster():
    """讀取 Google 試算表『員工登錄系統』(1.1員工登錄系統分頁)，回傳完整人員名冊（含快取）"""
    cache_key = 'employee_roster'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached

    rows = fetch_google_sheet_csv(config.EMPLOYEE_SHEET_ID, gid=config.EMPLOYEE_SHEET_GID)
    roster = []
    for r in rows:
        emp_id = (r.get('工號ID') or '').strip()
        name   = (r.get('姓名') or '').strip()
        if not emp_id or not name:
            continue
        roster.append({
            'emp_id':      emp_id,
            'name':        name,
            'group':       (r.get('組別') or '').strip(),
            'title':       (r.get('職務') or '').strip(),
            'status':      (r.get('狀態') or '').strip(),
            'nationality': (r.get('國籍') or '').strip(),
            'hire_date':   (r.get('到職日') or '').strip(),
            'leave_date':  (r.get('離職日') or '').strip(),
            'remark':      (r.get('備註') or '').strip(),
        })

    cache_set(cache_key, roster)
    return roster


def fetch_efficiency_all_groups():
    """透過 SSRS SOAP Execution Service 取得加工部總效率(全部)報表 CSV"""
    import base64 as _b64
    import re as _re
    session = get_ssrs_session()
    soap_url = f'http://192.168.1.212/ReportServer/ReportExecution2005.asmx'

    def _soap(body, action):
        return session.post(soap_url, data=body.encode('utf-8'),
            headers={'Content-Type': 'text/xml; charset=utf-8',
                     'SOAPAction': f'"http://schemas.microsoft.com/sqlserver/2005/06/30/reporting/reportingservices/{action}"'},
            timeout=30)

    NS = 'xmlns:rs="http://schemas.microsoft.com/sqlserver/2005/06/30/reporting/reportingservices"'
    ENV_OPEN = f'<?xml version="1.0" encoding="utf-8"?><soap:Envelope xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/" {NS}>'

    # Step 1: LoadReport
    r1 = _soap(f'{ENV_OPEN}<soap:Body><rs:LoadReport><rs:Report>/加工課/加工部總效率</rs:Report></rs:LoadReport></soap:Body></soap:Envelope>', 'LoadReport')
    r1.raise_for_status()
    eid = _re.search(r'<ExecutionID>([^<]+)</ExecutionID>', r1.text)
    if not eid:
        raise RuntimeError('LoadReport: 無法取得 ExecutionID')
    exec_id = eid.group(1)

    # Step 2: SetExecutionParameters ORG=全部
    r2 = _soap(f'{ENV_OPEN}<soap:Header><rs:ExecutionHeader><rs:ExecutionID>{exec_id}</rs:ExecutionID></rs:ExecutionHeader></soap:Header><soap:Body><rs:SetExecutionParameters><rs:Parameters><rs:ParameterValue><rs:Name>ORG</rs:Name><rs:Value>全部</rs:Value></rs:ParameterValue></rs:Parameters><rs:ParameterLanguage>zh-TW</rs:ParameterLanguage></rs:SetExecutionParameters></soap:Body></soap:Envelope>', 'SetExecutionParameters')
    r2.raise_for_status()

    # Step 3: Render as CSV
    r3 = _soap(f'{ENV_OPEN}<soap:Header><rs:ExecutionHeader><rs:ExecutionID>{exec_id}</rs:ExecutionID></rs:ExecutionHeader></soap:Header><soap:Body><rs:Render><rs:Format>CSV</rs:Format><rs:DeviceInfo></rs:DeviceInfo></rs:Render></soap:Body></soap:Envelope>', 'Render')
    r3.raise_for_status()
    b64 = _re.search(r'<Result>([^<]+)</Result>', r3.text)
    if not b64:
        raise RuntimeError('Render: 無法取得 CSV 結果')
    csv_bytes = _b64.b64decode(b64.group(1))
    return csv_bytes.decode('utf-8-sig', errors='replace')


def parse_csv(csv_text):
    """解析 CSV 為 dict list"""
    reader = csv.reader(io.StringIO(csv_text))
    rows = list(reader)
    if not rows:
        return []
    header = rows[0]
    results = []
    for row in rows[1:]:
        if len(row) < len(header):
            continue
        results.append(dict(zip(header, row)))
    return results


UNIT_OPTIONS = {
    '000-1': '加工課',
    '000-3': '製三課',
    '000-2': '成品部',
}
KNOWN_UNITS = set(UNIT_OPTIONS.keys())


def search_orders(order_id='', product_name='', unit='000-1', release_status='已發放'):
    """從未完工製令報表搜尋，支援製令號碼、品名模糊搜尋和生產線篩選
    unit: '000-1'/'000-2'/'000-3'/'other'/'*'
    release_status: '已發放'（預設）／'未發放'／'all'
    """
    all_records = fetch_all_unfinished(release_status)
    filtered = []
    for r in all_records:
        if order_id and order_id.upper() not in r.get('單別', '').upper():
            continue
        if product_name and not match_and_not(product_name, r.get('品名', '')):
            continue
        # 生產線篩選
        if unit and unit != '*':
            rec_unit = r.get('加工單位', '').strip()
            if unit == 'other':
                if rec_unit in KNOWN_UNITS:
                    continue
            else:
                if rec_unit != unit:
                    continue
        filtered.append(r)
    return filtered


def fetch_std_time(order_type, order_num):
    """從加工課-生產日報表取得標準工時和製程名稱（帶 TTL 快取）"""
    cache_key = f'std_{order_type}-{order_num}'
    cached = cache_get(cache_key)
    if cached is not None:
        return cached
    try:
        params = {'p1': order_type, 'p2': order_num}
        csv_text = fetch_report_csv(config.REPORT_PATHS['daily'], params)
        reader = csv.reader(io.StringIO(csv_text))
        rows = list(reader)
        result = {}
        for row in rows[1:]:
            if len(row) < 6:
                continue
            std_text = row[1]       # "標準工時：44 PCS/H"
            proc_code_text = row[3] # "製程代號：N19"
            proc_name_text = row[5] # "製程名稱：鑽+攻+倒"
            m_std = re.search(r'標準工時：(.+)', std_text)
            m_proc = re.search(r'製程代號：(.+)', proc_code_text)
            m_name = re.search(r'製程名稱：(.+)', proc_name_text)
            if m_proc:
                code = m_proc.group(1).strip()
                result[code] = {
                    'std_time': m_std.group(1).strip() if m_std else '',
                    'process_name': m_name.group(1).strip() if m_name else '',
                }
        cache_set(cache_key, result)
        return result
    except Exception:
        return {}


def fetch_efficiency_data():
    """從加工部總效率報表取得 SFT 預計開工日/完工日/標工/即時生產量（帶 TTL 快取）
    回傳 dict: { '製令號碼': { '製程代號': {...} } }
    """
    cached = cache_get('efficiency')
    if cached is not None:
        return cached
    try:
        csv_text = fetch_report_csv(config.REPORT_PATHS['efficiency'], {})
        records = parse_csv(csv_text)
        result = {}
        for r in records:
            oid = r.get('製令', '').strip()
            proc = r.get('製程', '').strip()
            if not oid:
                continue
            if oid not in result:
                result[oid] = {}
            result[oid][proc] = {
                'SFT預計開工日': r.get('SFT預計開工日', '').strip(),
                'SFT預計完工日': r.get('SFT預計完工日', '').strip(),
                '標工_秒': r.get('ERP變動人時3', '').strip(),
                '生產機台': r.get('生產機台', '').strip(),
                '即時生產量': r.get('現在產量', '').strip(),
            }
        cache_set('efficiency', result)
        return result
    except Exception:
        return {}


# ── ServCloud 設備稼動 ─────────────────────────────────────────────
_servcloud_session = None
_servcloud_session_lock = threading.Lock()
SERVCLOUD_BASE = getattr(config, 'SERVCLOUD_BASE', 'http://192.168.1.69:58080/ServCloud')
SERVCLOUD_USER = getattr(config, 'SERVCLOUD_USER', 'adminstd')
SERVCLOUD_PASS = getattr(config, 'SERVCLOUD_PASS', 'adminstd')

_STATUS_MAP = {
    '0':  ('離線', '#bdbdbd'),
    '11': ('運轉', '#43a047'),
    '12': ('待機', '#fb8c00'),
    '13': ('警報', '#e53935'),
    'B':  ('離線', '#bdbdbd'),
}

def _sc_login(s):
    s.post(f'{SERVCLOUD_BASE}/api/user/login',
           data={'id': SERVCLOUD_USER, 'password': SERVCLOUD_PASS}, timeout=10)

def get_servcloud_session():
    global _servcloud_session
    with _servcloud_session_lock:
        if _servcloud_session is None:
            s = requests.Session()
            _sc_login(s)
            _servcloud_session = s
        return _servcloud_session

def _sc_post(path, payload, retry=True):
    """POST to ServCloud; auto re-login on session expiry (type==2)."""
    session = get_servcloud_session()
    resp = session.post(f'{SERVCLOUD_BASE}{path}', json=payload, timeout=30)
    data = resp.json()
    if data.get('type') == 2 and retry:
        global _servcloud_session
        with _servcloud_session_lock:
            _sc_login(session)
            _servcloud_session = session
        return _sc_post(path, payload, retry=False)
    return data

def _sc_parse(raw):
    """ServCloud data 欄位可能是 JSON 字串或已解析的 list/dict，統一回傳 list。"""
    if isinstance(raw, list):
        return raw
    if isinstance(raw, str):
        return json.loads(raw)
    return []

def servcloud_get_machines():
    data = _sc_post('/api/getdata/db',
                    {'table': 'm_device', 'columns': ['device_id', 'device_name']})
    machines = _sc_parse(data.get('data', []))
    return [m for m in machines if m.get('device_name') and m['device_name'] != m['device_id']]

def servcloud_get_history(machine_ids, date_str):
    """date_str: YYYYMMDD。最多 10 台/次，自動分批。"""
    all_recs = []
    for i in range(0, len(machine_ids), 10):
        batch = machine_ids[i:i+10]
        data = _sc_post('/api/hippo/simple', {
            'space': 'machine_status_history',
            'index': {'machine_id': batch},
            'indexRange': {'key': 'date', 'start': date_str, 'end': date_str},
            'columns': ['machine_id', 'status', 'start_time', 'end_time', 'duration']
        })
        if data.get('type') == 0 and data.get('data'):
            all_recs.extend(_sc_parse(data['data']))
    return all_recs


APP_VERSION = 'V20260820'

@app.route('/ver')
def ver_check():
    """版本確認（純文字，繞過模板和快取）"""
    return f'<h1>{APP_VERSION}</h1><p>template_folder={app.template_folder}</p>'

@app.route('/')
def index():
    return render_template('index.html', app_version=APP_VERSION)


@app.route('/drawing')
def drawing_page():
    return render_template('drawing.html')


@app.route('/production')
def production_page():
    return render_template('production.html', app_version=APP_VERSION)


@app.route('/api/production', methods=['GET'])
def production_data():
    """生產中製令 API：用 SOAP 取得加工部總效率(全部)"""
    try:
        csv_text = fetch_efficiency_all_groups()
        records = parse_csv(csv_text)
        # 回傳原始欄位名稱供前端自動對應
        raw_columns = list(records[0].keys()) if records else []
        return jsonify({'success': True, 'data': records, 'count': len(records), 'columns': raw_columns})
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500


@app.route('/api/production/debug')
def production_debug():
    """除錯：顯示 SSRS CSV 原始欄位名稱"""
    try:
        csv_text = fetch_report_csv(config.REPORT_PATHS['efficiency'], {})
        reader = csv.reader(io.StringIO(csv_text))
        rows = list(reader)
        if not rows:
            return '<h2>CSV 空白</h2>'
        header = rows[0]
        html = '<h2>SSRS CSV 欄位名稱</h2><ol>'
        for i, col in enumerate(header):
            html += f'<li><b>{col}</b> (repr: {repr(col)})</li>'
        html += '</ol>'
        if len(rows) > 1:
            html += '<h3>第一筆資料</h3><table border="1" cellpadding="4"><tr><th>欄位</th><th>值</th></tr>'
            for col, val in zip(header, rows[1]):
                html += f'<tr><td>{col}</td><td>{val}</td></tr>'
            html += '</table>'
        return html
    except Exception as e:
        import traceback
        return f'<pre>Error: {traceback.format_exc()}</pre>', 500


@app.route('/api/query', methods=['GET'])
def query():
    """製令查詢 API：支援製令號碼和品名模糊搜尋"""
    order_id = request.args.get('order_id', '').strip()
    product_name_q = request.args.get('product_name', '').strip()
    unit = request.args.get('unit', '000-1').strip()
    release_status = request.args.get('release', '已發放').strip() or '已發放'

    try:
        # ── 第 1 步：取未完工製令 + 效率報表（並行） ──
        with ThreadPoolExecutor(max_workers=2) as pool:
            fut_orders = pool.submit(search_orders, order_id, product_name_q, unit, release_status)
            fut_eff = pool.submit(fetch_efficiency_data)
            records = fut_orders.result()
            eff_data = fut_eff.result()

        if not records:
            return jsonify({'success': False, 'error': '查無符合條件的資料'})

        # ── 第 2 步：日報表標工查詢（並行，限最多 20 筆） ──
        std_time_cache = {}
        unique_oids = list({r.get('單別', '') for r in records})
        oids_to_fetch = []
        for oid in unique_oids[:20]:
            oparts = oid.split('-', 1)
            if len(oparts) == 2:
                # 先檢查快取，只查沒快取的
                cache_key = f'std_{oparts[0]}-{oparts[1]}'
                cached = cache_get(cache_key)
                if cached is not None:
                    std_time_cache[oid] = cached
                else:
                    oids_to_fetch.append((oid, oparts[0], oparts[1]))

        if oids_to_fetch:
            with ThreadPoolExecutor(max_workers=min(8, len(oids_to_fetch))) as pool:
                futures = {
                    pool.submit(fetch_std_time, otype, onum): oid
                    for oid, otype, onum in oids_to_fetch
                }
                for fut in as_completed(futures):
                    oid = futures[fut]
                    try:
                        std_time_cache[oid] = fut.result()
                    except Exception:
                        std_time_cache[oid] = {}

        # ── 第 3 步：組合結果 ──
        data = []
        for r in records:
            rid = r.get('單別', '')
            proc_code = r.get('製程代號', '').strip()

            daily_info = std_time_cache.get(rid, {}).get(proc_code, {})
            eff_info = eff_data.get(rid, {}).get(proc_code, {})

            product_name = r.get('品名', '').rstrip('|').strip()

            std_time = daily_info.get('std_time', '')
            if not std_time and eff_info.get('標工_秒'):
                std_time = eff_info['標工_秒'] + ' 秒'
            # 標工：原始格式為 "105 PCS/H"（每小時產出件數），去掉 "PCS" 避免看起來像數量欄，保留 "/H" 標示這是速率
            if std_time:
                std_time = re.sub(r'\s*PCS\s*/\s*H', '/H', std_time, flags=re.IGNORECASE).strip()
                if not re.search(r'\d', std_time):
                    std_time = ''

            realtime_qty = eff_info.get('即時生產量', '') or r.get('即時生產量', '')

            data.append({
                '製令': rid,
                '品號': r.get('品號', ''),
                '品名': product_name,
                '加工順序': r.get('加工順序', ''),
                '製程代號': proc_code,
                '製程名稱': daily_info.get('process_name', ''),
                '發放狀態': r.get('情況', ''),
                '標工': std_time,
                '預計生產量': r.get('預計生產數', ''),
                '預計開工日': eff_info.get('SFT預計開工日', '') or r.get('製程預計開工日', ''),
                # SFT效率報表只有正在生產的製程；未開工製程用製程預計開工日作為估算完工日（has_sft_finish=False 表示估算值）
                '預計完工日': eff_info.get('SFT預計完工日', '') or r.get('製程預計開工日', ''),
                'has_sft_finish': bool(eff_info.get('SFT預計完工日', '')),
                '即時生產量': realtime_qty,
                '生產機台': eff_info.get('生產機台', ''),
            })

        return jsonify({'success': True, 'data': data})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/detail', methods=['GET'])
def detail():
    """製令詳細資訊 API：用料、製程、移轉單"""
    order_id = request.args.get('order_id', '').strip()
    if not order_id:
        return jsonify({'success': False, 'error': '請輸入製令號碼'}), 400

    try:
        params = {'製令': order_id}
        csv_text = fetch_report_csv(config.REPORT_PATHS['all'], params)

        # 報表回傳三段式 CSV（空行分隔）：材料、製程、移轉單
        csv_text = csv_text.replace('\r\n', '\n').replace('\r', '\n')
        sections = [s.strip() for s in csv_text.split('\n\n') if s.strip()]
        result = {'materials': [], 'processes': [], 'transfers': [], 'product': None, 'order': None}

        for section in sections:
            reader = csv.reader(io.StringIO(section))
            rows = list(reader)
            if len(rows) < 2:
                continue
            header = rows[0]
            data = []
            for row in rows[1:]:
                if len(row) >= len(header) and any(cell.strip() for cell in row):
                    data.append(dict(zip(header, row)))

            if '材料品號' in header:
                result['materials'] = data
            elif '完工否' in header:
                result['processes'] = data
            elif '移轉單' in header:
                result['transfers'] = data

        # 製程明細改直連 ERP SFCTA 覆蓋 SSRS 版本：SSRS 報表沒有「加工順序」跟
        # 預計/實際開工日，SFCR06 頁面（月度負荷等功能）已證實這張表撈得到這些欄位。
        # 查詢失敗就沿用上面 SSRS 版本（欄位較少但不會整頁掛掉）。
        oid_parts = order_id.split('-', 1)
        if len(oid_parts) == 2:
            try:
                erp_conn = get_erp_conn()
                erp_cur = erp_conn.cursor()
                erp_cur.execute("""
                    SELECT a.TA003, a.TA004, RTRIM(ISNULL(w.MW002,'')),
                           RTRIM(ISNULL(a.TA007,'')),
                           a.TA008, a.TA009, a.TA010, a.TA011,
                           a.TA030, a.TA031, RTRIM(ISNULL(a.TA024,''))
                    FROM SFCTA a LEFT JOIN CMSMW w ON RTRIM(w.MW001) = RTRIM(a.TA004)
                    WHERE a.TA001 = ? AND a.TA002 = ?
                    ORDER BY a.TA003
                """, oid_parts)
                # 同一個 cursor 下一個 execute 會把這批結果沖掉，所以先撈完再查下一段
                proc_rows = erp_cur.fetchall()

                def _fmt_date(s):
                    s = (s or '').strip()
                    return f'{s[:4]}/{s[4:6]}/{s[6:8]}' if len(s) == 8 else ''

                # 成品本身的庫存數（跟產品途程明細那頁的「庫存數」同一個算法：
                # INVMC.MC007 現有量加總）。品號從製令主檔 MOCTA.TA006 撈，不靠前端傳，
                # 這樣不管從哪個畫面打開詳細資訊都看得到。
                # TA026/TA027 = 這張製令的來源訂單（訂單單別/訂單單號，COPTC/COPTD 的
                # TC001/TC002、TD001/TD002）——不是每張製令都有（庫存單這類沒有客戶訂單來源
                # 的製令，這兩欄是空字串），前端靠這個決定要不要顯示「展開訂單」按鈕。
                erp_cur.execute("""
                    SELECT RTRIM(a.TA006), RTRIM(ISNULL(b.MB002,'')), RTRIM(ISNULL(b.MB004,'')),
                           CAST(ISNULL((SELECT SUM(c.MC007) FROM INVMC c
                                         WHERE RTRIM(c.MC001) = RTRIM(a.TA006)),0) AS VARCHAR(20)),
                           RTRIM(ISNULL(a.TA026,'')), RTRIM(ISNULL(a.TA027,''))
                      FROM MOCTA a LEFT JOIN INVMB b ON RTRIM(b.MB001) = RTRIM(a.TA006)
                     WHERE a.TA001 = ? AND a.TA002 = ?
                """, oid_parts)
                prow = erp_cur.fetchone()
                if prow and (prow[0] or '').strip():
                    result['product'] = {
                        'item_no':   prow[0].strip(),
                        'item_name': (prow[1] or '').strip(),
                        'unit':      (prow[2] or '').strip(),
                        'stock_qty': str(prow[3] or '0').strip(),
                    }
                    so_type, so_no = (prow[4] or '').strip(), (prow[5] or '').strip()
                    if so_type and so_no:
                        result['order'] = {'type': so_type, 'no': so_no}

                result['processes'] = [{
                    '加工順序': r[0],
                    '製程代號': r[1],
                    '製程名稱': r[2],
                    '加工廠商': r[3],
                    '預計開工日': _fmt_date(r[4]),
                    '預計完工日': _fmt_date(r[5]),
                    '投入數量': r[6],
                    '完成數量': r[7],
                    '實際開工日': _fmt_date(r[8]),
                    '實際完工日': _fmt_date(r[9]),
                    '製程敘述': r[10],
                } for r in proc_rows]
                erp_conn.close()
            except Exception:
                pass

        return jsonify({'success': True, 'data': result})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── 列印暫存（pywebview 無法用 window.open，改用 URL 方式）──
_print_cache = {}
_print_cache_lock = threading.Lock()


def _render_print_html(records):
    """將列印資料渲染成 HTML"""
    from datetime import datetime
    today = datetime.now().strftime('%Y/%m/%d %H:%M')

    std_time_cache = {}
    for r in records:
        order_id = r.get('order_id', '')
        if order_id and order_id not in std_time_cache:
            parts = order_id.split('-')
            if len(parts) >= 2:
                std_time_cache[order_id] = fetch_std_time(parts[0], parts[1])

    items = []
    for r in records:
        order_id = r.get('order_id', '')
        process_seq = r.get('process_seq', '')
        process_code = r.get('process_code', '')
        process_name = r.get('process_name', '')
        unit = r.get('unit', '').strip()
        product_name = r.get('product_name', '')
        qty = r.get('qty', '')

        proc_info = std_time_cache.get(order_id, {}).get(process_code, {})
        std_time = proc_info.get('std_time', '') if isinstance(proc_info, dict) else ''
        actual_process_name = proc_info.get('process_name', '') if isinstance(proc_info, dict) else ''
        if not actual_process_name:
            actual_process_name = process_name

        barcode = f"{order_id};{process_seq};{process_code};{unit}"

        items.append({
            'order_id': f"{order_id}-{process_seq}",
            'product_name': product_name,
            'process_code': process_code,
            'process_name': actual_process_name,
            'qty': qty,
            'std_time': std_time,
            'barcode': barcode,
            'machine': r.get('machine', ''),
        })

    return render_template('print_report.html',
                           items=items,
                           items_json=json.dumps(items, ensure_ascii=False),
                           today=today)


@app.route('/print', methods=['POST'])
def print_report():
    """生產日報表列印頁面 — POST 暫存 HTML，回傳 key 供 GET 取用"""
    try:
        records = request.get_json() or []
    except Exception:
        return jsonify({'success': False, 'error': '資料格式錯誤'}), 400

    html = _render_print_html(records)
    key = str(uuid.uuid4())
    with _print_cache_lock:
        _print_cache[key] = html
    return jsonify({'success': True, 'key': key})


@app.route('/print/view/<key>', methods=['GET'])
def print_view(key):
    """取得暫存的列印 HTML"""
    with _print_cache_lock:
        html = _print_cache.get(key)
        # 列出目前快取的 keys 用於除錯
        cached_keys = list(_print_cache.keys())
    if html:
        return html
    return f'列印資料不存在。要求 key={key}，目前快取 keys={cached_keys}', 404


@app.route('/api/inspection', methods=['GET'])
def find_inspection():
    """根據品號查找自主巡檢表"""
    product_no = request.args.get('product_no', '').strip()
    if not product_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        root = Path(INSPECTION_ROOT)
        if not root.exists():
            return jsonify({'success': False, 'error': '無法存取巡檢表網路資料夾'}), 500

        found = []
        for f in root.rglob('*.xls*'):
            fname = f.stem
            if product_no.upper() in fname.upper():
                found.append({
                    'file_path': str(f),
                    'file_name': f.name,
                    'dir': str(f.parent),
                    'need_inspect': '需送檢' in fname,
                })

        if found:
            return jsonify({'success': True, 'files': found, 'count': len(found)})
        else:
            return jsonify({'success': False, 'error': '無對應自主檢查表', 'files': []})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _insp_grid(path):
    """把檢查表第一個工作表讀成 2D 值陣列（.xls 用 xlrd、.xlsx/.xlsm 用 openpyxl）"""
    if path.lower().endswith('.xls'):
        import xlrd
        ws = xlrd.open_workbook(path).sheets()[0]
        return [[ws.cell_value(r, c) for c in range(ws.ncols)] for r in range(ws.nrows)]
    else:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        ws = wb.worksheets[0]
        grid = [[('' if c.value is None else c.value) for c in row] for row in ws.iter_rows()]
        wb.close()
        return grid


def _insp_layout(path):
    """解析自主檢查表的版面（用標籤錨點動態定位，新舊版型通吃）。
    回傳 (layout dict, grid)：header 各欄位 (row,col)（0-based）、檢驗項目區塊的資料列範圍。
    新版範本(240327/C01) 與舊表(M01) 的列位、數量欄(J/K) 都不同，不能寫死。"""
    grid = _insp_grid(path)
    nrows = len(grid)
    ncols = max((len(r) for r in grid), default=0)

    def val(r, c):
        try:
            return str(grid[r][c]).strip()
        except Exception:
            return ''

    lay = {'sheet_index': 0, 'nrows': nrows}
    # 圖號（A 欄找「圖號」標籤，值在 B）
    for r in range(min(nrows, 8)):
        if '圖號' in val(r, 0):
            lay['fig'] = (r, 1)
            break
    # 表頭列（A 欄=機台別），值在下一列；數量/投料單號/架模人員欄位用標題文字找
    for r in range(min(nrows, 10)):
        if val(r, 0).replace(' ', '') == '機台別':
            vr = r + 1
            lay['machine'] = (vr, 0)
            lay['model']   = (vr, 1)
            lay['proc']    = (vr, 2)
            for c in range(ncols):
                h = val(r, c).replace(' ', '')
                if '架模' in h:
                    lay['operator'] = (vr, c)
                elif '投料單號' in h.replace('\n', ''):
                    lay['order'] = (vr, c)
                elif h.startswith('數量') or '數量(Pcs)' in h:
                    lay['qty'] = (vr, c)
            break
    # 檢驗項目區塊：標題列「檢驗項目」→ 跳過次標題列（首件/1/2…）→ 資料列
    # 直到遇到 A 欄含「作業員」「判定」「備註」為止（中間空列也算可填欄位）
    blocks = []
    r = 0
    while r < nrows:
        if val(r, 0).replace(' ', '') == '檢驗項目':
            start = r + 2
            end = start
            while end < nrows:
                a = val(end, 0)
                if any(k in a for k in ('作業員', '判定', '備註', '現場主管')):
                    break
                end += 1
            blocks.append({'start': start, 'end': end})   # [start, end) 0-based
            r = end
        r += 1
    lay['blocks'] = blocks[:2]
    return lay, grid


@app.route('/api/inspection/read', methods=['GET'])
def inspection_read():
    """讀取自主檢查表內容（GUI 填寫用）：表頭欄位＋兩個檢驗項目區塊"""
    file_path = request.args.get('file_path', '').strip()
    if not file_path:
        return jsonify({'success': False, 'error': '請提供檔案路徑'}), 400
    # 僅允許巡檢表資料夾內的檔案
    if not os.path.normpath(file_path).lower().startswith(os.path.normpath(INSPECTION_ROOT).lower()):
        return jsonify({'success': False, 'error': '無效的檔案路徑'}), 400
    try:
        lay, grid = _insp_layout(file_path)

        def gval(r, c):
            try:
                v = grid[r][c]
            except Exception:
                return ''
            if isinstance(v, float) and v == int(v):
                v = int(v)
            return str(v).strip()

        def cell(key):
            if key not in lay:
                return ''
            r, c = lay[key]
            return gval(r, c)

        blocks = []
        for b in lay.get('blocks', []):
            rows = []
            for r in range(b['start'], b['end']):
                rows.append({'item': gval(r, 0), 'spec': gval(r, 1)})
            blocks.append(rows)

        return jsonify({'success': True,
                        'header': {
                            'fig':      cell('fig'),
                            'machine':  cell('machine'),
                            'model':    cell('model'),
                            'proc':     cell('proc'),
                            'operator': cell('operator'),
                            'order':    cell('order'),
                            'qty':      cell('qty'),
                        },
                        'blocks': blocks})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取失敗：{e}'}), 500


@app.route('/api/inspection/save', methods=['POST'])
def inspection_save():
    """把 GUI 填寫的內容寫回檢查表（Excel COM，保留 .xls 格式與公式）。
    用 DispatchEx 開獨立 Excel 程序，不干擾使用者已開啟的 Excel。"""
    d = request.get_json(silent=True) or {}
    file_path  = (d.get('file_path') or '').strip()
    header     = d.get('header') or {}
    blocks     = d.get('blocks') or []
    open_after = bool(d.get('open_after'))

    if not file_path:
        return jsonify({'success': False, 'error': '請提供檔案路徑'}), 400
    if not os.path.normpath(file_path).lower().startswith(os.path.normpath(INSPECTION_ROOT).lower()):
        return jsonify({'success': False, 'error': '無效的檔案路徑'}), 400

    try:
        lay, _ = _insp_layout(file_path)
    except Exception as e:
        return jsonify({'success': False, 'error': f'解析版面失敗：{e}'}), 500

    try:
        import pythoncom
        import win32com.client
        pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
        excel = None
        try:
            excel = win32com.client.DispatchEx('Excel.Application')
            excel.Visible = False
            excel.DisplayAlerts = False
            wb = excel.Workbooks.Open(file_path)
            if wb.ReadOnly:
                wb.Close(False)
                excel.Quit()
                return jsonify({'success': False, 'error': '檔案被其他人開啟中（唯讀），請稍後再存'}), 409
            ws = wb.Worksheets(1)

            def put(key, value):
                if key in lay and value is not None:
                    r, c = lay[key]
                    ws.Cells(r + 1, c + 1).Value = value  # COM 是 1-based

            put('fig',      header.get('fig'))
            put('machine',  header.get('machine'))
            put('model',    header.get('model'))
            put('proc',     header.get('proc'))
            put('operator', header.get('operator'))
            put('order',    header.get('order'))
            put('qty',      header.get('qty'))

            for bi, b in enumerate(lay.get('blocks', [])):
                rows = blocks[bi] if bi < len(blocks) else []
                for i, r in enumerate(range(b['start'], b['end'])):
                    if i >= len(rows):
                        break
                    ws.Cells(r + 1, 1).Value = rows[i].get('item') or ''
                    ws.Cells(r + 1, 2).Value = rows[i].get('spec') or ''

            wb.Save()
            if open_after:
                excel.Visible = True
                excel.DisplayAlerts = True
            else:
                wb.Close(False)
                excel.Quit()
            return jsonify({'success': True,
                            'message': f'已寫入 {os.path.basename(file_path)}' + ('（已開啟）' if open_after else '')})
        except Exception:
            # 失敗時收掉隱藏的 Excel 程序，避免殘留
            try:
                if excel is not None and not excel.Visible:
                    excel.DisplayAlerts = False
                    excel.Quit()
            except Exception:
                pass
            raise
        finally:
            pythoncom.CoUninitialize()
    except Exception as e:
        return jsonify({'success': False, 'error': f'寫入失敗：{e}（若檔案開啟中請先關閉）'}), 500


@app.route('/api/inspection/templates', methods=['GET'])
def inspection_templates():
    """列出巡檢表空白範本與機台分類資料夾（新增檢查表用）。
    - 機台分類 = 根目錄下【*】開頭的資料夾，各附其子資料夾（系列）清單
    - 範本 = 「空白表單(已套用公式)」內的 .xls*
    """
    try:
        root = Path(INSPECTION_ROOT)
        if not root.exists():
            return jsonify({'success': False, 'error': '無法存取巡檢表網路資料夾'}), 500

        machines = []
        for d in sorted(root.iterdir()):
            if d.is_dir() and d.name.startswith('【'):
                subs = sorted(s.name for s in d.iterdir() if s.is_dir())
                machines.append({'name': d.name, 'series': subs})

        templates = []
        tpl_dir = root / '空白表單(已套用公式)'
        if tpl_dir.is_dir():
            for f in sorted(tpl_dir.glob('*.xls*')):
                if not f.name.startswith('~$'):
                    templates.append({'name': f.name, 'path': str(f)})

        return jsonify({'success': True, 'machines': machines, 'templates': templates})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/inspection/create', methods=['POST'])
def inspection_create():
    """從空白範本建立新的自主檢查表：
    複製範本到 機台分類[/系列資料夾]/檔名.xls，清除唯讀屬性後回傳路徑。
    不覆蓋既有檔案（409）。"""
    d = request.get_json(silent=True) or {}
    template  = (d.get('template') or '').strip()      # 範本完整路徑
    machine   = (d.get('machine') or '').strip()       # 機台分類資料夾名
    series    = (d.get('series') or '').strip()        # 系列子資料夾（可空，可新建）
    file_name = (d.get('file_name') or '').strip()     # 檔名（不含副檔名）

    if not template or not machine or not file_name:
        return jsonify({'success': False, 'error': '範本、機台分類、檔名為必填'}), 400
    # 檔名不允許路徑符號
    if any(c in file_name for c in r'\/:*?"<>|'):
        return jsonify({'success': False, 'error': '檔名含有不允許的字元'}), 400

    try:
        root = Path(INSPECTION_ROOT)
        tpl = Path(template)
        # 範本必須位於巡檢表資料夾內（防止任意路徑複製）
        if not str(tpl).lower().startswith(str(root).lower()) or not tpl.is_file():
            return jsonify({'success': False, 'error': '無效的範本路徑'}), 400
        target_dir = root / machine
        if not target_dir.is_dir():
            return jsonify({'success': False, 'error': f'機台分類資料夾不存在：{machine}'}), 400
        if series:
            if any(c in series for c in r'\/:*?"<>|'):
                return jsonify({'success': False, 'error': '系列資料夾含有不允許的字元'}), 400
            target_dir = target_dir / series
            target_dir.mkdir(exist_ok=True)

        target = target_dir / (file_name + tpl.suffix)
        if target.exists():
            return jsonify({'success': False, 'error': f'檔案已存在：{target.name}（請改檔名或直接開啟既有檔案）'}), 409

        import shutil
        shutil.copyfile(tpl, target)          # copyfile 不帶屬性，避免唯讀跟過來
        try:
            os.chmod(target, 0o666)
        except Exception:
            pass
        return jsonify({'success': True, 'file_path': str(target), 'file_name': target.name})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/inspection/open', methods=['POST'])
def open_inspection():
    """開啟自主巡檢表，並將製令、預計產量寫入 Excel 指定欄位
    - 單張製令：G6 寫製令號碼（垂直置中），J6+K6 寫預計產量
    - 多張製令：G6 寫所有製令號碼（換行，垂直置中），J6/K6 不寫
    """
    import time as _time
    data = request.get_json() or {}
    file_path = data.get('file_path', '').strip()
    order_ids = data.get('order_ids', [])       # 製令號碼列表
    qty = data.get('qty', '').strip()            # 預計產量（僅單張時寫入）

    if not file_path:
        return jsonify({'success': False, 'error': '請提供檔案路徑'}), 400

    try:
        import pythoncom
        import win32com.client

        # 使用 STA threading model，避免 RPC_E_CALL_REJECTED 錯誤
        pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
        try:
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = True
            wb = excel.Workbooks.Open(file_path)

            # ── COM retry 機制：Excel 開啟後可能忙碌，需重試 ──
            max_retries = 3
            ws = None
            for attempt in range(max_retries):
                try:
                    _time.sleep(2)  # 等待 Excel 載入完成
                    ws = wb.ActiveSheet

                    # xlVAlignCenter = -4108
                    if order_ids:
                        cell_g6 = ws.Range("G6")
                        cell_g6.Value = '\n'.join(order_ids)
                        cell_g6.WrapText = True
                        cell_g6.VerticalAlignment = -4108   # 垂直置中
                        cell_g6.ShrinkToFit = False
                        # 自動調整文字大小以符合儲存格（避免跨到下一排）
                        if len(order_ids) > 1:
                            cell_g6.ShrinkToFit = True

                        # 同時填入 J6 和 K6 預計產量
                        if qty:
                            for cell_name in ["J6", "K6"]:
                                cell = ws.Range(cell_name)
                                cell.Value = qty
                                cell.VerticalAlignment = -4108

                    break  # 成功，跳出重試迴圈
                except pythoncom.com_error as ce:
                    if attempt < max_retries - 1:
                        _time.sleep(2)  # 等待後重試
                    else:
                        raise  # 最後一次仍失敗，拋出錯誤

            return jsonify({'success': True, 'message': f'已開啟並寫入 {os.path.basename(file_path)}'})
        finally:
            pythoncom.CoUninitialize()
    except ImportError:
        # pywin32 未安裝（如 Embeddable Python 環境）→ 直接用 os.startfile 開啟
        try:
            os.startfile(file_path)
            return jsonify({'success': True,
                            'message': f'已開啟 {os.path.basename(file_path)}（此電腦無法自動填入製令號碼，請手動輸入）'})
        except Exception as e2:
            return jsonify({'success': False, 'error': f'無法開啟檔案：{e2}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def get_pdm_db():
    """取得 PDMSearch SQLite 連線（自動嘗試多個路徑）"""
    candidates = [
        PDM_DB_PATH,                                                          # config.py 設定值
        os.path.join(_APP_DIR, 'pdm_search.db'),                              # 隨包附帶 (_app/)
        os.path.join(os.path.dirname(_APP_DIR), 'pdm_search.db'),             # 上一層 (製令查詢/)
        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'PDMSearch', 'pdm_search.db'),  # 使用者 AppData
    ]
    for path in candidates:
        if path and os.path.isfile(path):
            conn = sqlite3.connect(path)
            conn.row_factory = sqlite3.Row
            return conn
    return None


@app.route('/api/debug/paths')
def debug_paths():
    """診斷用：顯示 PDM 路徑解析結果"""
    candidates = [
        ('config.py 設定值', PDM_DB_PATH),
        ('_APP_DIR/pdm_search.db', os.path.join(_APP_DIR, 'pdm_search.db')),
        ('上一層/pdm_search.db', os.path.join(os.path.dirname(_APP_DIR), 'pdm_search.db')),
        ('LOCALAPPDATA', os.path.join(os.environ.get('LOCALAPPDATA', ''), 'PDMSearch', 'pdm_search.db')),
        ('sys.executable 同層', os.path.join(os.path.dirname(sys.executable), 'pdm_search.db')),
        ('__file__ 目錄', os.path.dirname(os.path.abspath(__file__))),
    ]
    result = {
        '_APP_DIR': _APP_DIR,
        '_BASE': _BASE,
        'PDM_DB_PATH': PDM_DB_PATH,
        'sys.executable': sys.executable,
        '__file__': os.path.abspath(__file__),
        'cwd': os.getcwd(),
        'candidates': []
    }
    for label, path in candidates:
        result['candidates'].append({
            'label': label,
            'path': path,
            'exists': os.path.isfile(path) if not label.endswith('目錄') else os.path.isdir(path),
        })
    # 列出 _APP_DIR 裡有什麼檔案
    try:
        result['_APP_DIR_files'] = os.listdir(_APP_DIR)
    except Exception as e:
        result['_APP_DIR_files'] = str(e)
    return jsonify(result)


@app.route('/api/drawing', methods=['GET'])
def find_drawing():
    """根據品號從 PDMSearch 索引資料庫查找圖面"""
    product_no = request.args.get('product_no', '').strip()
    if not product_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        conn = get_pdm_db()
        if not conn:
            return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

        cur = conn.cursor()
        cur.execute(
            """SELECT file_path, pin_hao, tu_mian_pin_ming, ji_xing, xing_hao, modified_at
               FROM drawing_index WHERE UPPER(pin_hao) = UPPER(?)""",
            (product_no,)
        )
        rows = cur.fetchall()
        conn.close()

        if rows:
            files = _rows_to_drawing_list(rows)
            return jsonify({'success': True, 'files': files, 'count': len(files)})
        else:
            return jsonify({'success': False, 'error': f'PDM 索引中找不到品號 {product_no} 的圖面', 'files': []})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/drawing/search', methods=['GET'])
def search_drawing():
    """圖面模糊搜尋 API：支援品號或品名模糊搜尋（AND/NOT 語法）"""
    q = request.args.get('q', '').strip()
    if not q:
        return jsonify({'success': False, 'error': '請輸入品號或品名'}), 400

    try:
        conn = get_pdm_db()
        if not conn:
            return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

        cur = conn.cursor()
        # 取出所有的關鍵字（排除 NOT 語法用的）做 SQL LIKE 初篩
        tokens = q.strip().split()
        positive = [t for t in tokens if not t.startswith('-')]
        conditions = []
        params = []
        for t in positive[:3]:  # 最多取前3個正向關鍵字做 SQL 篩選
            conditions.append(
                "(UPPER(pin_hao) LIKE ? OR UPPER(tu_mian_pin_ming) LIKE ? "
                "OR UPPER(file_path) LIKE ? OR UPPER(ji_xing) LIKE ?)"
            )
            like = f'%{t.upper()}%'
            params.extend([like, like, like, like])

        where = ' AND '.join(conditions) if conditions else '1=1'
        cur.execute(
            f"""SELECT file_path, pin_hao, tu_mian_pin_ming, ji_xing, xing_hao, modified_at
                FROM drawing_index WHERE {where} LIMIT 200""",
            params
        )
        rows = cur.fetchall()
        conn.close()

        # 再用 match_and_not 做精確的 AND/NOT 過濾（機型 ji_xing 也納入比對範圍）
        files = []
        for row in rows:
            combined = (f"{row['pin_hao'] or ''} {row['tu_mian_pin_ming'] or ''} "
                        f"{row['file_path'] or ''} {row['ji_xing'] or ''}")
            if match_and_not(q, combined):
                files.append(_row_to_drawing_dict(row))

        if files:
            return jsonify({'success': True, 'files': files, 'count': len(files)})
        else:
            return jsonify({'success': False, 'error': '查無符合條件的圖面', 'files': []})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _row_to_drawing_dict(row):
    """將 SQLite Row 轉為圖面 dict"""
    fp = row['file_path']
    return {
        'file_path': fp,
        'file_name': os.path.basename(fp),
        'dir': os.path.dirname(fp),
        'pin_hao': row['pin_hao'] or '',
        'product_name': row['tu_mian_pin_ming'] or '',
        'ji_xing': row['ji_xing'] or '',
        'xing_hao': row['xing_hao'] or '',
        'modified_at': row['modified_at'] or '',
    }


def _rows_to_drawing_list(rows):
    """將多筆 SQLite Row 轉為圖面 dict list"""
    return [_row_to_drawing_dict(row) for row in rows]


def pdm_get_file(file_path):
    """用 PowerShell 呼叫 PDM COM API 取出檔案（Python COM 的 GetFileCopy 有型別問題，改用 PowerShell）"""
    import subprocess

    # PowerShell 腳本：透過 PDM COM 取出檔案到本機快取
    ps_script = f'''
$ErrorActionPreference = "Stop"
try {{
    $vault = New-Object -ComObject "ConisioLib.EdmVault"
    $vault.LoginAuto("MAXCLAW", 0)
    $path = "{file_path.replace(chr(92), chr(92)+chr(92))}"
    $dir = [System.IO.Path]::GetDirectoryName($path)
    $fname = [System.IO.Path]::GetFileName($path)
    $folder = $vault.GetFolderFromPath($dir)
    if ($null -eq $folder) {{
        Write-Error "PDM 找不到資料夾: $dir"
        exit 1
    }}
    $fileObj = $folder.GetFile($fname)
    if ($null -eq $fileObj) {{
        Write-Error "PDM 找不到檔案: $fname"
        exit 1
    }}
    $fileObj.GetFileCopy(0)
    Write-Output "OK"
}} catch {{
    Write-Error $_.Exception.Message
    exit 1
}}
'''
    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-Command', ps_script],
            capture_output=True, text=True, timeout=30
        )

        if result.returncode == 0 and 'OK' in result.stdout:
            return True, 'OK'
        else:
            err = result.stderr.strip()
            if not err:
                err = result.stdout.strip() or '未知錯誤'
            return False, err
    except subprocess.TimeoutExpired:
        return False, 'PDM 取出逾時'
    except Exception as e:
        return False, str(e)


@app.route('/api/drawing/open', methods=['POST'])
def open_drawing():
    """開啟 PDM 圖面檔案"""
    data = request.get_json() or {}
    file_path = data.get('file_path', '').strip()

    if not file_path:
        return jsonify({'success': False, 'error': '請提供圖面路徑'}), 400

    try:
        if os.path.exists(file_path):
            # 檔案已在本機（或可直接存取的網路路徑）
            # 嘗試 PDM 取最新版，但失敗不阻擋開啟
            pdm_get_file(file_path)
            os.startfile(file_path)
            return jsonify({'success': True, 'message': f'已開啟 {os.path.basename(file_path)}'})
        else:
            # 本機沒有此檔案，嘗試透過 PDM COM 取出
            ok, msg = pdm_get_file(file_path)
            if ok:
                os.startfile(file_path)
                return jsonify({'success': True, 'message': f'已開啟 {os.path.basename(file_path)}'})
            else:
                # PDM 取出失敗（可能未安裝 PDM）→ 仍嘗試直接開啟
                try:
                    os.startfile(file_path)
                    return jsonify({'success': True, 'message': f'已開啟 {os.path.basename(file_path)}'})
                except Exception:
                    pdm_hint = '（此電腦未安裝 SolidWorks PDM，無法自動取出圖面）' if 'ConisioLib' in msg or 'EdmVault' in msg or '無法建立' in msg else ''
                    return jsonify({'success': False,
                                    'error': f'無法開啟圖面：{msg}{pdm_hint}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── PDM 索引重建 ───────────────────────────────────────────────────────────────

import threading as _threading
import subprocess as _subprocess
import re as _re

_reindex_state = {
    'running':    False,
    'phase':      'idle',   # idle | scanning | indexing | done | error
    'scanned':    0,
    'total':      0,
    'indexed':    0,
    'message':    '',
    'error':      '',
    'last_run':   None,
    'last_count': 0,
}
_reindex_lock = _threading.Lock()

# 持久化狀態檔：重啟後仍可讀取最後一次 reindex 結果
_REINDEX_STATE_FILE = os.path.join(_APP_DIR, 'last_reindex.json')

# P2 管理指標上下限：改存這個 JSON 檔（原本只存瀏覽器 localStorage，
# 清快取／換瀏覽器／換機器就會消失，2026-07-11 改伺服器端持久化）
_P2_TARGETS_FILE = os.path.join(_APP_DIR, 'p2_targets.json')
_P2_TARGETS_DEFAULT = {
    'A': {'upper': 75, 'lower': 70},
    'B': {'upper': 15, 'lower': 13},
    'C': {'upper': 12, 'lower': 10},
    'L': {'upper': 70, 'lower': 65},
}


@app.route('/api/p2/targets', methods=['GET'])
def p2_targets_get():
    try:
        with open(_P2_TARGETS_FILE, encoding='utf-8') as f:
            targets = json.load(f)
    except Exception:
        targets = _P2_TARGETS_DEFAULT
    return jsonify({'success': True, 'targets': targets})


@app.route('/api/p2/targets', methods=['POST'])
def p2_targets_save():
    body = request.get_json(force=True, silent=True) or {}
    raw = body.get('targets')
    if not isinstance(raw, dict):
        return jsonify({'success': False, 'error': '格式錯誤'}), 400
    targets = {}
    for cat in ('A', 'B', 'C', 'L'):
        item = raw.get(cat) or {}
        def _num(v):
            try:
                return float(v) if v is not None and v != '' else None
            except (TypeError, ValueError):
                return None
        targets[cat] = {'upper': _num(item.get('upper')), 'lower': _num(item.get('lower'))}
    try:
        with open(_P2_TARGETS_FILE, 'w', encoding='utf-8') as f:
            json.dump(targets, f, ensure_ascii=False)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    return jsonify({'success': True, 'targets': targets})


def _save_last_reindex_state(ts: str, cnt: int):
    """reindex 完成後將時間與筆數寫入 JSON（重啟後仍保留）"""
    try:
        with open(_REINDEX_STATE_FILE, 'w', encoding='utf-8') as f:
            json.dump({'last_run': ts, 'last_count': cnt}, f)
    except Exception:
        pass


def _load_last_reindex_state():
    """從 JSON 讀取最後一次 reindex 結果（比 DB indexed_at 更可靠）"""
    try:
        with open(_REINDEX_STATE_FILE, encoding='utf-8') as f:
            d = json.load(f)
            return d.get('last_run'), int(d.get('last_count', 0))
    except Exception:
        return None, 0


def _run_reindex(update_only: bool):
    """背景執行索引重建，解析 stdout 更新進度狀態"""
    global _reindex_state
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_pdm_index.py')
    cmd = [sys.executable, script] + (['--update'] if update_only else [])

    with _reindex_lock:
        _reindex_state.update(running=True, phase='scanning', scanned=0,
                              total=0, indexed=0, message='啟動中...', error='')

    try:
        # 子行程預設用主控台編碼（cp950）輸出中文，跟這裡的 utf-8 解碼對不上會變亂碼，
        # 導致下面的中文正規表達式比對失敗；強制子行程輸出改用 UTF-8
        child_env = os.environ.copy()
        child_env['PYTHONIOENCODING'] = 'utf-8'
        # 子行程 stdout 接到管線（非終端機）時 Python 預設整批緩衝，進度輸出會卡到緩衝區滿
        # 或行程結束才一次送出，畫面看起來像卡住不動；設這個讓它逐行即時送出
        child_env['PYTHONUNBUFFERED'] = '1'
        proc = _subprocess.Popen(
            cmd, stdout=_subprocess.PIPE, stderr=_subprocess.STDOUT,
            text=True, encoding='utf-8', errors='replace', env=child_env
        )
        for line in proc.stdout:
            line = line.rstrip()
            if not line:
                continue
            with _reindex_lock:
                _reindex_state['message'] = line

            # 解析掃描進度：「掃描中... 1,200 個檔案」
            m = _re.search(r'掃描中.*?(\d[\d,]*)\s*個', line)
            if m:
                with _reindex_lock:
                    _reindex_state['phase']   = 'scanning'
                    _reindex_state['scanned'] = int(m.group(1).replace(',', ''))
                continue

            # 解析找到總數：「找到 X 個 .SLDDRW」
            m = _re.search(r'找到\s*(\d[\d,]*)\s*個', line)
            if m:
                with _reindex_lock:
                    _reindex_state['total'] = int(m.group(1).replace(',', ''))
                continue

            # 解析寫入進度：「 33%  1000/3000  新增:X」
            m = _re.search(r'(\d+)%.*?(\d[\d,]*)/(\d[\d,]*)', line)
            if m:
                with _reindex_lock:
                    _reindex_state['phase']   = 'indexing'
                    _reindex_state['indexed'] = int(m.group(2).replace(',', ''))
                    _reindex_state['total']   = int(m.group(3).replace(',', ''))
                continue

            # 完成行：「完成！新增:X  更新:X  ...」
            m = _re.search(r'完成.*新增:(\d[\d,]*)', line)
            if m:
                with _reindex_lock:
                    _reindex_state['indexed'] = int(m.group(1).replace(',', ''))

        proc.wait()

        if proc.returncode == 0:
            # 取得資料庫最新筆數
            try:
                conn_tmp = get_pdm_db()
                if conn_tmp:
                    cnt = conn_tmp.execute('SELECT COUNT(*) FROM drawing_index').fetchone()[0]
                    conn_tmp.close()
                else:
                    cnt = _reindex_state['indexed']
            except Exception:
                cnt = _reindex_state['indexed']

            ts_done = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            with _reindex_lock:
                _reindex_state.update(
                    running=False, phase='done', last_count=cnt,
                    last_run=ts_done,
                    message=f'更新完成，共 {cnt:,} 筆圖面資料'
                )
            # 持久化到 JSON，重啟後仍可讀取
            _save_last_reindex_state(ts_done, cnt)
        else:
            with _reindex_lock:
                _reindex_state.update(running=False, phase='error',
                                      error='重建失敗，請查看伺服器日誌')
    except Exception as exc:
        with _reindex_lock:
            _reindex_state.update(running=False, phase='error', error=str(exc))


@app.route('/api/drawing/reindex', methods=['POST'])
def drawing_reindex():
    """啟動 PDM 圖面索引重建（背景執行）"""
    with _reindex_lock:
        if _reindex_state['running']:
            return jsonify({'success': False, 'error': '索引重建已在執行中，請稍候'}), 409

    data        = request.get_json() or {}
    update_only = data.get('update_only', True)   # 預設增量更新

    t = _threading.Thread(target=_run_reindex, args=(update_only,), daemon=True)
    t.start()
    return jsonify({'success': True, 'message': '索引重建已啟動'})


def _pdm_last_indexed():
    """從 drawing_index 讀取最後一次 indexed_at（持久，不受重啟影響）"""
    try:
        c = get_pdm_db()
        if not c:
            return None, 0
        # 舊版 DB 可能沒有 indexed_at 欄位，ALTER TABLE 補上（冪等）
        try:
            c.execute("ALTER TABLE drawing_index ADD COLUMN indexed_at TEXT")
        except Exception:
            pass  # 欄位已存在
        row = c.execute(
            "SELECT MAX(indexed_at), COUNT(*) FROM drawing_index"
        ).fetchone()
        c.close()
        return (row[0] or None), (row[1] or 0)
    except Exception:
        return None, 0


@app.route('/api/drawing/reindex/status', methods=['GET'])
def drawing_reindex_status():
    """回傳索引重建進度狀態（含 DB 持久化的 last_run）"""
    with _reindex_lock:
        state = dict(_reindex_state)
    # 若記憶體中 last_run 為空（剛重啟），依序從 JSON → DB 讀取並快取
    if not state.get('last_run'):
        ts, cnt = _load_last_reindex_state()    # 優先：JSON 狀態檔
        if not ts:
            ts, cnt = _pdm_last_indexed()        # 備援：DB indexed_at
        if ts:
            with _reindex_lock:
                _reindex_state['last_run']   = ts
                _reindex_state['last_count'] = cnt
            state['last_run']   = ts
            state['last_count'] = cnt
    return jsonify(state)


# ══════════════════════════════════════════════════════════
#  ZUMEN 線上圖面管理 GUI（Node.js 工具整合，選配）
# ══════════════════════════════════════════════════════════

def _zumen_gui_alive():
    """檢查本機 ZUMEN GUI（Node 伺服器）是否在跑"""
    try:
        r = requests.get(f'http://127.0.0.1:{config.ZUMEN_GUI_PORT}/', timeout=2)
        return r.status_code < 500
    except Exception:
        return False


@app.route('/api/zumen_gui/status')
def zumen_gui_status():
    """ZUMEN GUI 狀態：installed=工具有沒有裝在這台電腦、running=伺服器活著沒、env=帳密檔存在沒"""
    gui_dir = getattr(config, 'ZUMEN_GUI_DIR', '')
    installed = bool(gui_dir) and os.path.isfile(os.path.join(gui_dir, 'server.js')) \
        and os.path.isdir(os.path.join(gui_dir, 'node_modules'))
    has_env = installed and os.path.isfile(os.path.join(gui_dir, '.env'))
    return jsonify({
        'installed': installed,
        'has_env': has_env,
        'running': _zumen_gui_alive(),
        'url': f'http://localhost:{config.ZUMEN_GUI_PORT}',
    })


def _zumen_node_exe(gui_dir):
    """優先用 ZUMEN 工具內建的免安裝 Node.js（_node/node.exe，內網無外網電腦也能跑），
    沒有內建版才退回系統 PATH 的 node（開發機可能已全域安裝）"""
    bundled = os.path.join(gui_dir, '_node', 'node.exe')
    return bundled if os.path.isfile(bundled) else 'node'


@app.route('/api/zumen_gui/start', methods=['POST'])
def zumen_gui_start():
    """啟動 ZUMEN GUI 的 Node 伺服器（背景、隱藏視窗）。啟動後前端自行輪詢 status 等就緒。"""
    gui_dir = getattr(config, 'ZUMEN_GUI_DIR', '')
    if not gui_dir or not os.path.isfile(os.path.join(gui_dir, 'server.js')):
        return jsonify({'success': False, 'error': f'這台電腦未安裝 ZUMEN 圖面管理工具（{gui_dir}）'}), 404
    if not os.path.isdir(os.path.join(gui_dir, 'node_modules')):
        return jsonify({'success': False, 'error': '工具尚未安裝相依套件，請在該目錄執行 npm install'}), 409
    if not os.path.isfile(os.path.join(gui_dir, '.env')):
        return jsonify({'success': False,
                        'error': '缺少 .env 帳密檔：請複製 .env.example 為 .env，填入 ZUMEN_EMAIL / ZUMEN_PASSWORD'}), 409
    if _zumen_gui_alive():
        return jsonify({'success': True, 'already_running': True})
    try:
        import subprocess
        node_exe = _zumen_node_exe(gui_dir)
        # CREATE_NO_WINDOW=0x08000000：背景啟動不彈黑窗；DETACHED 讓它不隨 Flask 結束
        subprocess.Popen(
            [node_exe, 'server.js'], cwd=gui_dir,
            creationflags=0x08000000 | 0x00000008,  # CREATE_NO_WINDOW | DETACHED_PROCESS
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        return jsonify({'success': True, 'already_running': False})
    except FileNotFoundError:
        return jsonify({'success': False, 'error': '找不到 node，請確認這台電腦已安裝 Node.js，或工具資料夾內附的 _node 免安裝版是否完整'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/zumen_gui/open', methods=['POST'])
def zumen_gui_open():
    """用 Edge/Chrome --app 模式開啟 ZUMEN GUI（獨立應用視窗，無網址列，跟 PDS 主視窗同模式）"""
    if not _zumen_gui_alive():
        return jsonify({'success': False, 'error': 'ZUMEN 伺服器未啟動'}), 409
    url = f'http://localhost:{config.ZUMEN_GUI_PORT}'
    for base in ('ProgramFiles(x86)', 'ProgramFiles', 'LOCALAPPDATA'):
        p = os.path.join(os.environ.get(base, ''), 'Microsoft', 'Edge', 'Application', 'msedge.exe')
        if p and os.path.isfile(p):
            import subprocess
            subprocess.Popen([p, f'--app={url}'])
            return jsonify({'success': True})
    for base in ('ProgramFiles', 'ProgramFiles(x86)', 'LOCALAPPDATA'):
        p = os.path.join(os.environ.get(base, ''), 'Google', 'Chrome', 'Application', 'chrome.exe')
        if p and os.path.isfile(p):
            import subprocess
            subprocess.Popen([p, f'--app={url}'])
            return jsonify({'success': True})
    # 找不到 Edge/Chrome 就退回預設瀏覽器
    webbrowser.open(url)
    return jsonify({'success': True, 'fallback': True})


@app.route('/api/zume/status')
def zume_status():
    """回傳技術資料清單的目前狀態（件數、最後匯入檔案）"""
    try:
        con = sqlite3.connect(ZUME_DB_PATH)
        total = con.execute('SELECT COUNT(*) FROM drawings').fetchone()[0]
        last  = con.execute('SELECT filename, imported_at FROM import_log ORDER BY imported_at DESC LIMIT 1').fetchone()
        con.close()
        return jsonify({'total': total, 'last_file': last[0] if last else None, 'last_at': last[1] if last else None})
    except Exception as e:
        return jsonify({'total': 0, 'last_file': None, 'last_at': None, 'error': str(e)})


@app.route('/api/zume/scan', methods=['POST'])
def zume_scan():
    """重新掃描下載/桌面等資料夾並匯入最新 CSV（強制重掃）"""
    files = _find_zume_csvs()
    if not files:
        dirs = '、'.join(_zume_csv_search_dirs()) or os.path.join(os.path.expanduser('~'), 'Downloads')
        return jsonify({'success': False,
                        'error': f'找不到 zume-n_data_list_*.csv\n請先從 ZUMEN 匯出清單 CSV 到「下載」資料夾。\n已搜尋：{dirs}'}), 404
    latest = files[0]
    fname  = os.path.basename(latest)
    recs   = _parse_zume_csv(latest)
    if not recs:
        return jsonify({'success': False, 'error': f'檔案解析失敗或無資料：{fname}'}), 400
    con = sqlite3.connect(ZUME_DB_PATH)
    _insert_zume_recs(con, recs)
    con.execute('INSERT OR REPLACE INTO import_log(filename,imported_at,count) VALUES(?,datetime("now"),?)',
                (fname, len(recs)))
    con.commit()
    total = con.execute('SELECT COUNT(*) FROM drawings').fetchone()[0]
    con.close()
    return jsonify({'success': True, 'imported': len(recs), 'total': total, 'file': fname})


@app.route('/api/zume/import', methods=['POST'])
def zume_import():
    """手動上傳 CSV/XLSX 匯入（保留作為備用方式）"""
    f = request.files.get('file')
    if not f:
        return jsonify({'success': False, 'error': '未上傳檔案'}), 400
    fname_lower = f.filename.lower()
    recs = []
    try:
        if fname_lower.endswith('.csv'):
            content = f.read().decode('utf-8-sig')
            reader  = csv.reader(io.StringIO(content))
            headers = [h.strip() for h in next(reader)]
            idx     = _zume_header_indices(headers)
            if idx['part_no'] is None or idx['url'] is None:
                return jsonify({'success': False, 'error': f'找不到「圖號」或「URL」欄'}), 400
            for row in reader:
                rec = _zume_row_to_rec(row, idx)
                if rec:
                    recs.append(rec)
        elif fname_lower.endswith(('.xlsx', '.xlsm')):
            if not _OPENPYXL_OK:
                return jsonify({'success': False, 'error': '請改用 CSV 格式'}), 400
            wb = openpyxl.load_workbook(io.BytesIO(f.read()), read_only=True)
            ws = wb.active
            headers = [str(c.value or '').strip() for c in next(ws.iter_rows(min_row=1, max_row=1))]
            idx     = _zume_header_indices(headers)
            if idx['part_no'] is None or idx['url'] is None:
                return jsonify({'success': False, 'error': '找不到「圖號」或「URL」欄'}), 400
            for row in ws.iter_rows(min_row=2, values_only=True):
                rec = _zume_row_to_rec(row, idx)
                if rec:
                    recs.append(rec)
            wb.close()
        else:
            return jsonify({'success': False, 'error': '僅支援 .csv 或 .xlsx'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': f'解析失敗：{e}'}), 500
    if not recs:
        return jsonify({'success': False, 'error': '檔案無有效資料'}), 400
    fname = f.filename
    con = sqlite3.connect(ZUME_DB_PATH)
    _insert_zume_recs(con, recs)
    con.execute('INSERT OR REPLACE INTO import_log(filename,imported_at,count) VALUES(?,datetime("now"),?)',
                (fname, len(recs)))
    con.commit()
    total = con.execute('SELECT COUNT(*) FROM drawings').fetchone()[0]
    con.close()
    return jsonify({'success': True, 'imported': len(recs), 'total': total})


@app.route('/api/zume/lookup', methods=['POST'])
def zume_lookup():
    """查詢品號對應的 zume-n.com URL；找不到則回傳搜尋 URL"""
    try:
        data     = request.get_json() or {}
        item_nos = data.get('item_nos', [])
        if not item_nos:
            return jsonify({'success': False, 'error': '請提供品號'}), 400
        con = sqlite3.connect(ZUME_DB_PATH)
        result = []
        for no in item_nos[:20]:
            no = no.strip()
            row = con.execute('SELECT part_name,url FROM drawings WHERE part_no=?', (no,)).fetchone()
            if row:
                result.append({'no': no, 'name': row[0], 'url': row[1], 'found': True})
            else:
                fallback_url = 'https://zume-n.com/freeword_search?query=' + urllib.parse.quote(no, safe='') + '&searchType=drawing'
                result.append({'no': no, 'name': '', 'url': fallback_url, 'found': False})
        con.close()
        return jsonify({'success': True, 'items': result})
    except Exception as e:
        return jsonify({'success': False, 'error': f'查詢失敗：{str(e)}'}), 500


@app.route('/api/zume/open', methods=['POST'])
def zume_open():
    """以系統預設瀏覽器開啟 zume-n.com（優先用清單的直接 URL）"""
    try:
        data     = request.get_json() or {}
        item_nos = data.get('item_nos', [])
        if not item_nos:
            return jsonify({'success': False, 'error': '請提供品號'}), 400
        con = sqlite3.connect(ZUME_DB_PATH)
        opened = []; not_found = []
        for no in item_nos[:10]:
            no = no.strip()
            row = con.execute('SELECT url FROM drawings WHERE part_no=?', (no,)).fetchone()
            url = row[0] if row else ('https://zume-n.com/freeword_search?query=' + urllib.parse.quote(no, safe='') + '&searchType=drawing')
            webbrowser.open(url, new=2)
            if row:
                opened.append(no)
            else:
                not_found.append(no)
        con.close()
        return jsonify({'success': True, 'opened': opened, 'not_found': not_found})
    except Exception as e:
        return jsonify({'success': False, 'error': f'開啟失敗：{str(e)}'}), 500


@app.route('/api/zume/list')
def zume_list():
    """回傳所有匯入的 ZUMEN 圖面（含生產線別/群組/分類/廠商）+ 篩選選項 + 最後匯入資訊"""
    try:
        con = sqlite3.connect(ZUME_DB_PATH)
        con.row_factory = sqlite3.Row
        rows = [{
            'part_no':    r['part_no'],
            'part_name':  r['part_name'] or '',
            'url':        r['url'],
            'line':       r['line'] or '',
            'prod_group': r['prod_group'] or '',
            'category':   r['category'] or '',
            'vendor':     r['vendor'] or '',
        } for r in con.execute(
            'SELECT part_no, part_name, url, line, prod_group, category, vendor '
            'FROM drawings ORDER BY part_no')]
        # 各欄不重複值（給前端下拉用）
        def distinct(col):
            return sorted({(r[0] or '').strip() for r in
                           con.execute(f'SELECT DISTINCT {col} FROM drawings') if (r[0] or '').strip()})
        options = {c: distinct(c) for c in _ZUME_EXTRA_COLS}
        last = con.execute(
            'SELECT filename, imported_at FROM import_log ORDER BY imported_at DESC LIMIT 1').fetchone()
        con.close()
        return jsonify({'success': True, 'count': len(rows), 'drawings': rows,
                        'options': options,
                        'last_file': last[0] if last else None,
                        'last_at':   last[1] if last else None})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取失敗：{str(e)}'}), 500


@app.route('/equipment')
def equipment_page():
    return render_template('equipment.html', app_version=APP_VERSION)


@app.route('/api/equipment/machines')
def equipment_machines():
    try:
        machines = servcloud_get_machines()
        return jsonify({'success': True, 'data': machines})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/equipment/history')
def equipment_history():
    date_str    = request.args.get('date', '')           # YYYY-MM-DD
    start_hour  = int(request.args.get('start_hour', 7))
    end_hour    = int(request.args.get('end_hour', 20))
    machine_ids = request.args.getlist('machines')

    if not date_str or not machine_ids:
        return jsonify({'success': False, 'error': '請提供日期和機台'}), 400

    date_api = date_str.replace('-', '')

    try:
        history = servcloud_get_history(machine_ids, date_api)
        all_machines = servcloud_get_machines()
        name_map = {m['device_id']: m['device_name'] for m in all_machines}

        range_start_s = start_hour * 3600   # 秒
        range_end_s   = end_hour   * 3600
        total_min     = (range_end_s - range_start_s) // 60

        # Group segments by machine (秒精度)
        buckets = {mid: [] for mid in machine_ids}
        for rec in history:
            mid = rec.get('machine_id', '')
            if mid not in buckets:
                continue
            st = rec.get('start_time', '')
            et = rec.get('end_time',   '')
            if len(st) < 14 or len(et) < 14:
                continue
            s_sec = int(st[8:10])*3600 + int(st[10:12])*60 + int(st[12:14])
            e_sec = int(et[8:10])*3600 + int(et[10:12])*60 + int(et[12:14])
            # Handle overnight
            if e_sec < s_sec:
                e_sec += 24 * 3600
            s_clip = max(s_sec, range_start_s)
            e_clip = min(e_sec, range_end_s)
            if s_clip >= e_clip:
                continue
            status = rec.get('status', '0')
            status_name, color = _STATUS_MAP.get(status, ('未知', '#9e9e9e'))
            dur_sec = e_clip - s_clip
            buckets[mid].append({
                'status':      status,
                'status_name': status_name,
                'color':       color,
                'start_min':   (s_clip - range_start_s) / 60,   # float，用於圖表定位
                'end_min':     (e_clip - range_start_s) / 60,
                'start_str':   f"{s_clip//3600%24:02d}:{s_clip%3600//60:02d}:{s_clip%60:02d}",
                'end_str':     f"{e_clip//3600%24:02d}:{e_clip%3600//60:02d}:{e_clip%60:02d}",
                'dur_sec':     dur_sec,
            })

        result = [
            {'id': mid, 'name': name_map.get(mid, mid), 'segments': buckets[mid]}
            for mid in machine_ids
        ]
        return jsonify({
            'success': True, 'machines': result,
            'date': date_str, 'start_hour': start_hour,
            'end_hour': end_hour, 'total_minutes': total_min
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500


# ── BOM 查詢系統 ───────────────────────────────────────────────────
# 資料庫：192.168.1.140 / YC01（Computech ERP）
# 品號主檔：INVMB  材料BOM：BOMMD（MCPI11模組）
_ERP_SQL_SERVER   = getattr(config, 'ERP_SQL_SERVER',   '192.168.1.140')
_ERP_SQL_DATABASE = getattr(config, 'ERP_SQL_DATABASE', 'YC01')
_ERP_SQL_USERNAME = getattr(config, 'ERP_SQL_USERNAME', 'sa')
_ERP_SQL_PASSWORD = getattr(config, 'ERP_SQL_PASSWORD', 'dsc55877948')

# 屬性代碼對照表（INVMB.MB025）
_ATTR_MAP = {'M': 'M:自製件', 'S': 'S:訂外加工', 'P': 'P:採購件',
             'R': 'R:委外', 'F': 'F:虛擬件'}


def get_erp_conn():
    """建立 ERP SQL Server 連線（pyodbc）"""
    try:
        import pyodbc
    except ImportError:
        raise RuntimeError('未安裝 pyodbc，請執行: pip install pyodbc')

    cs = (f'DRIVER={{ODBC Driver 17 for SQL Server}};'
          f'SERVER={_ERP_SQL_SERVER};DATABASE={_ERP_SQL_DATABASE};'
          f'UID={_ERP_SQL_USERNAME};PWD={_ERP_SQL_PASSWORD};'
          f'TrustServerCertificate=yes;Connection Timeout=8;')
    return pyodbc.connect(cs, timeout=8)


@app.route('/routing')
def routing_page():
    return render_template('routing.html', app_version=APP_VERSION)


@app.route('/api/routing/search')
def routing_search():
    """搜尋有途程的品號（INVMB + BOMME，多關鍵字空格分隔，品號/品名同時模糊搜尋）
       line 參數：指定線別篩選（MF006），空字串代表全部"""
    q    = request.args.get('q', '').strip()
    line = request.args.get('line', '').strip()   # '000-1','000-2','000-3' 或 ''=全部
    if not q:
        return jsonify({'success': False, 'error': '請輸入品號或品名'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()
        # 多關鍵字：以空格分割，每個關鍵字都要符合（AND 邏輯）
        keywords = [kw for kw in q.split() if kw]
        where_parts = []
        params = []
        for kw in keywords:
            like = f'%{kw}%'
            where_parts.append("(RTRIM(b.MB001) LIKE ? OR b.MB002 LIKE ?)")
            params.extend([like, like])
        keyword_clause = ' AND '.join(where_parts) if where_parts else '1=1'

        # 線別篩選：只顯示「最新版次」途程中包含該線別至少一步的品號
        if line:
            line_clause = """AND EXISTS (
                SELECT 1 FROM BOMMF f2
                WHERE RTRIM(f2.MF001) = RTRIM(b.MB001)
                  AND f2.MF002 = (
                      SELECT MAX(e2.ME002) FROM BOMME e2
                      WHERE RTRIM(e2.ME001) = RTRIM(b.MB001)
                  )
                  AND RTRIM(f2.MF006) = ?
            )"""
            params.append(line)
        else:
            line_clause = ''

        cur.execute(f"""
            SELECT TOP 200
                RTRIM(b.MB001)            AS item_no,
                RTRIM(ISNULL(b.MB002,'')) AS item_name,
                RTRIM(ISNULL(b.MB003,'')) AS spec,
                ISNULL(b.MB004,'')        AS unit,
                ISNULL(b.MB025,'')        AS attr_code
            FROM INVMB b
            WHERE EXISTS (
                SELECT 1 FROM BOMME e WHERE RTRIM(e.ME001) = RTRIM(b.MB001)
            )
            AND {keyword_clause}
            {line_clause}
            ORDER BY b.MB001
        """, params)
        rows = cur.fetchall()
        conn.close()
        results = [{
            'item_no':   r[0],
            'item_name': r[1],
            'spec':      r[2],
            'unit':      r[3],
            'attr':      _ATTR_MAP.get(str(r[4]).strip(), str(r[4]).strip()),
        } for r in rows]
        return jsonify({'success': True, 'data': results, 'count': len(results)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/bom')
def bom_page():
    # 舊連結相容：直接導向產品途程查詢頁（BOM子分頁）
    return redirect('/routing?tab=bom')


@app.route('/api/bom/search')
def bom_search():
    """搜尋有材料BOM的品號（INVMB + BOMMD，品號/品名模糊搜尋，可篩選品號開頭）"""
    q      = request.args.get('q', '').strip()
    prefix = request.args.get('prefix', '').strip()   # '5', '6', '56' 或空
    if not q:
        return jsonify({'success': False, 'error': '請輸入品號或品名'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()
        like = f'%{q}%'

        # 品號開頭篩選條件
        if prefix == '5':
            prefix_cond = "AND RTRIM(b.MB001) LIKE '5%'"
        elif prefix == '6':
            prefix_cond = "AND RTRIM(b.MB001) LIKE '6%'"
        elif prefix in ('56', '65'):
            prefix_cond = "AND (RTRIM(b.MB001) LIKE '5%' OR RTRIM(b.MB001) LIKE '6%')"
        else:
            prefix_cond = ''

        cur.execute(f"""
            SELECT TOP 200
                RTRIM(b.MB001)          AS item_no,
                RTRIM(ISNULL(b.MB002,'')) AS item_name,
                RTRIM(ISNULL(b.MB003,'')) AS spec,
                ISNULL(b.MB004,'')      AS unit,
                ISNULL(b.MB025,'')      AS attr_code
            FROM INVMB b
            WHERE EXISTS (
                SELECT 1 FROM BOMMD d WHERE RTRIM(d.MD001) = RTRIM(b.MB001)
            )
            {prefix_cond}
            AND (RTRIM(b.MB001) LIKE ? OR b.MB002 LIKE ?)
            ORDER BY b.MB001
        """, (like, like))
        rows = cur.fetchall()
        conn.close()
        results = [{
            'item_no':   r[0],
            'item_name': r[1],
            'spec':      r[2],
            'unit':      r[3],
            'attr':      _ATTR_MAP.get(str(r[4]).strip(), str(r[4]).strip()),
        } for r in rows]
        return jsonify({'success': True, 'data': results, 'count': len(results)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/bom/detail')
def bom_detail():
    """取得指定品號的材料BOM明細（BOMMD JOIN INVMB，對應 MCPI11 BOM用量資料）"""
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()

        # 取品號主檔（母件資訊）
        cur.execute("""
            SELECT RTRIM(MB001), RTRIM(ISNULL(MB002,'')),
                   RTRIM(ISNULL(MB003,'')), ISNULL(MB004,''),
                   ISNULL(MB005,''), ISNULL(MB025,'')
            FROM INVMB WHERE RTRIM(MB001) = ?
        """, (item_no,))
        hrow = cur.fetchone()
        header = {}
        if hrow:
            header = {
                'item_no':   hrow[0],
                'item_name': hrow[1],
                'spec':      hrow[2],
                'unit':      hrow[3],
                'sub_unit':  hrow[4],
                'attr':      _ATTR_MAP.get(str(hrow[5]).strip(), str(hrow[5]).strip()),
            }

        # 取材料BOM明細（BOMMD JOIN INVMB 取子件品名/規格/屬性，並加總 INVMC 現有量為庫存數量）
        cur.execute("""
            SELECT
                d.MD002                           AS seq,
                RTRIM(d.MD003)                    AS child_no,
                RTRIM(ISNULL(b.MB002,''))         AS child_name,
                RTRIM(ISNULL(b.MB003,''))         AS spec,
                ISNULL(d.MD004,'')                AS unit,
                ISNULL(d.MD005,'')                AS sub_unit,
                ISNULL(b.MB025,'')                AS attr_code,
                CAST(ISNULL(d.MD006,0) AS VARCHAR(20)) AS qty,
                CAST(ISNULL((SELECT SUM(c.MC007) FROM INVMC c WHERE RTRIM(c.MC001) = RTRIM(d.MD003)),0) AS VARCHAR(20)) AS stock_qty
            FROM BOMMD d
            LEFT JOIN INVMB b ON RTRIM(b.MB001) = RTRIM(d.MD003)
            WHERE RTRIM(d.MD001) = ?
            ORDER BY d.MD002
        """, (item_no,))
        rows = cur.fetchall()
        conn.close()

        detail = []
        for r in rows:
            attr_code = str(r[6] or '').strip()
            detail.append({
                'seq':        str(r[0] or '').strip(),
                'child_no':   str(r[1] or '').strip(),
                'child_name': str(r[2] or '').strip(),
                'spec':       str(r[3] or '').strip(),
                'unit':       str(r[4] or '').strip(),
                'sub_unit':   str(r[5] or '').strip(),
                'attr':       _ATTR_MAP.get(attr_code, attr_code),
                'qty':        str(r[7] or '').strip(),
                'stock_qty':  str(r[8] or '').strip(),
            })

        return jsonify({'success': True, 'header': header, 'detail': detail})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/bom/routing')
def bom_routing():
    """取得指定品號的產品途程（BOMME + BOMMF，最新版次，對應 BOMMI07）"""
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    try:
        conn = get_erp_conn()
        cur  = conn.cursor()

        # 品號主檔
        cur.execute("""
            SELECT RTRIM(MB001), RTRIM(ISNULL(MB002,'')),
                   RTRIM(ISNULL(MB003,'')), ISNULL(MB004,''), ISNULL(MB025,'')
            FROM INVMB WHERE RTRIM(MB001) = ?
        """, (item_no,))
        hrow = cur.fetchone()
        header = {}
        if hrow:
            header = {
                'item_no':   hrow[0],
                'item_name': hrow[1],
                'spec':      hrow[2],
                'unit':      hrow[3],
                'attr':      _ATTR_MAP.get(str(hrow[4]).strip(), str(hrow[4]).strip()),
            }

        # 庫存數量（加總 INVMC.MC007 現有量）
        cur.execute("""
            SELECT CAST(ISNULL(SUM(MC007),0) AS VARCHAR(20))
            FROM INVMC WHERE RTRIM(MC001) = ?
        """, (item_no,))
        srow = cur.fetchone()
        if header:
            header['stock_qty'] = str(srow[0] or '0').strip() if srow else '0'

        # 最新途程版次
        cur.execute("""
            SELECT MAX(ME002) FROM BOMME WHERE RTRIM(ME001) = ?
        """, (item_no,))
        vrow = cur.fetchone()
        latest_ver = (vrow[0] or '').strip() if vrow else ''
        if header:
            header['version'] = latest_ver

        # 途程明細（BOMMF JOIN CMSMW 取製程名稱/機台代號/資源群組代號）
        routing = []
        if latest_ver:
            _NATURE = {'1': '1:廠內', '2': '2:訂外', '3': '3:外包'}
            cur.execute("""
                SELECT
                    f.MF003,
                    RTRIM(ISNULL(f.MF004,'')),
                    RTRIM(ISNULL(w.MW002,''))   AS proc_name,
                    f.MF005,
                    RTRIM(ISNULL(f.MF006,'')),
                    RTRIM(ISNULL(f.MF007,'')),
                    RTRIM(ISNULL(f.MF008,'')),
                    RTRIM(ISNULL(f.MF040,'')),
                    RTRIM(ISNULL(f.MF034,'')),
                    CAST(ISNULL(f.MF010,0) AS DECIMAL(10,3)),
                    RTRIM(ISNULL(x.MX003,''))   AS machine_name
                FROM BOMMF f
                LEFT JOIN CMSMW  w ON RTRIM(w.MW001) = RTRIM(f.MF004)
                LEFT JOIN CMSMX  x ON RTRIM(x.MX001) = RTRIM(f.MF040)
                WHERE RTRIM(f.MF001) = ? AND f.MF002 = ?
                ORDER BY f.MF003
            """, (item_no, latest_ver))
            rows = cur.fetchall()
            for r in rows:
                nature_code = str(r[3] or '').strip()
                vmh_raw = r[9]
                if vmh_raw is None or float(vmh_raw) == 0:
                    vmh = ''
                else:
                    vmh = '{:g}'.format(float(vmh_raw))
                routing.append({
                    'seq':          str(r[0] or '').strip(),
                    'proc_code':    str(r[1] or '').strip(),
                    'proc_name':    str(r[2] or '').strip(),
                    'nature':       _NATURE.get(nature_code, nature_code),
                    'vendor_no':    str(r[4] or '').strip(),
                    'vendor_name':  str(r[5] or '').strip(),
                    'description':  str(r[6] or '').strip(),
                    'machine_code': str(r[7] or '').strip(),
                    'res_group':    str(r[8] or '').strip(),
                    'var_man_hrs':  vmh,
                    'machine_name': str(r[10] or '').strip(),
                })

        conn.close()
        has_routing = bool(latest_ver and routing)
        return jsonify({
            'success': True,
            'has_routing': has_routing,
            'header': header,
            'routing': routing,
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


_proc_codes_cache = {'ts': 0, 'data': None}


@app.route('/api/proc_codes')
def proc_codes():
    """全部製程代號主檔（CMSMW，MW001=代號/MW002=名稱），供治檢具申請單「製程代號」
    欄位的自動完成清單用。這張表跟 /api/bom/routing、SFCR06 用的是同一份代碼表。

    治檢具申請單的製程代號欄位（management.html ja2-proc-list）原本是手 key 的固定
    30 筆清單，缺很多實際存在的代號（2026-08-26 實測：品號 0310302000E17007 的真實
    製程是 M46，固定清單裡沒有這一筆，使用者只能看到清單裡湊巧存在的 M42），
    改成直接查主檔，快取 1 小時（純代碼表，不會頻繁變動，不用像製令資料那樣短 TTL）。
    """
    now = time.time()
    if not _proc_codes_cache['data'] or now - _proc_codes_cache['ts'] > 3600:
        try:
            conn = get_erp_conn()
            cur = conn.cursor()
            cur.execute("""
                SELECT RTRIM(MW001), RTRIM(ISNULL(MW002,''))
                FROM CMSMW WHERE RTRIM(ISNULL(MW001,'')) <> '' ORDER BY MW001
            """)
            rows = cur.fetchall()
            conn.close()
            _proc_codes_cache.update(ts=now, data=[{'code': r[0], 'name': r[1]} for r in rows])
        except Exception as e:
            if not _proc_codes_cache['data']:
                return jsonify({'success': False, 'error': str(e)}), 502
            # 已有舊快取：這次查詢失敗就沿用舊資料，不要讓使用者連清單都看不到
    return jsonify({'success': True, 'data': _proc_codes_cache['data']})


@app.route('/api/inventory/history')
def inventory_history():
    """取得指定品號的庫存異動歷史（INVTB + INVTK + MOCTE + MOCTG UNION）。
    回傳：日期、單別/單號/序號、單據名稱、庫別、庫別名稱、入出庫、異動數量、製令參考。
    """
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    months = request.args.get('months', '12').strip()
    try:
        months_int = max(1, min(60, int(months)))
    except ValueError:
        months_int = 12
    # 日期下限（CREATE_DATE 為 YYYYMMDD 字串）
    from datetime import datetime, timedelta
    start_date = (datetime.now() - timedelta(days=months_int * 31)).strftime('%Y%m%d')

    try:
        conn = get_erp_conn()
        cur  = conn.cursor()

        sql = """
            SELECT date_, doc_type, doc_no, doc_seq, warehouse, qty, src_table
            FROM (
                SELECT
                    CREATE_DATE                       AS date_,
                    RTRIM(TB001)                      AS doc_type,
                    RTRIM(TB002)                      AS doc_no,
                    RTRIM(TB003)                      AS doc_seq,
                    RTRIM(ISNULL(TB016,''))           AS warehouse,
                    CAST(ISNULL(TB007,0) AS DECIMAL(18,3)) AS qty,
                    'INVTB'                           AS src_table
                FROM INVTB
                WHERE RTRIM(TB004) = ? AND CREATE_DATE >= ?
                UNION ALL
                SELECT
                    CREATE_DATE,
                    RTRIM(TK001), RTRIM(TK002), RTRIM(TK003),
                    RTRIM(ISNULL(TK017,'')),
                    CAST(ISNULL(TK007,0) AS DECIMAL(18,3)),
                    'INVTK'
                FROM INVTK
                WHERE RTRIM(TK004) = ? AND CREATE_DATE >= ?
                UNION ALL
                SELECT
                    CREATE_DATE,
                    RTRIM(TE001), RTRIM(TE002), RTRIM(TE003),
                    RTRIM(ISNULL(TE008,'')),
                    CAST(ISNULL(TE005,0) AS DECIMAL(18,3)),
                    'MOCTE'
                FROM MOCTE
                WHERE RTRIM(TE004) = ? AND CREATE_DATE >= ?
                UNION ALL
                SELECT
                    CREATE_DATE,
                    RTRIM(TG001), RTRIM(TG002), RTRIM(TG003),
                    RTRIM(ISNULL(TG010,'')),
                    CAST(ISNULL(TG011,0) AS DECIMAL(18,3)),
                    'MOCTG'
                FROM MOCTG
                WHERE RTRIM(TG004) = ? AND CREATE_DATE >= ?
            ) u
            ORDER BY date_ DESC, doc_no DESC, doc_seq DESC
        """
        cur.execute(sql, (item_no, start_date, item_no, start_date, item_no, start_date, item_no, start_date))
        rows = cur.fetchall()

        # 一次撈出單別/庫別對照
        cur.execute("SELECT RTRIM(MQ001), RTRIM(ISNULL(MQ002,'')) FROM CMSMQ")
        doc_map = {r[0]: r[1] for r in cur.fetchall()}
        cur.execute("SELECT RTRIM(MC001), RTRIM(ISNULL(MC002,'')) FROM CMSMC")
        wh_map  = {r[0]: r[1] for r in cur.fetchall()}
        conn.close()

        # 入出庫判斷：依來源表 + 單別名稱
        # MOCTE=出庫(領料); MOCTG=入庫(移轉入庫); INVTK=調整(數量通常為0)
        # INVTB 依名稱關鍵字判斷
        def in_out(src, doc_type, doc_name, qty):
            if src == 'MOCTE':
                return '出庫'
            if src == 'MOCTG':
                return '入庫'
            if src == 'INVTK':
                return '調整'
            n = doc_name or ''
            if '入庫' in n:
                return '入庫'
            if '出庫' in n or '領料' in n or '報廢' in n:
                return '出庫'
            return '調整'

        history = []
        for r in rows:
            date_str = (r[0] or '').strip()
            if date_str and len(date_str) == 8:
                date_str = f'{date_str[:4]}/{date_str[4:6]}/{date_str[6:8]}'
            doc_type = (r[1] or '').strip()
            doc_no   = (r[2] or '').strip()
            doc_seq  = (r[3] or '').strip()
            wh       = (r[4] or '').strip()
            qty      = float(r[5] or 0)
            src      = r[6]
            doc_name = doc_map.get(doc_type, '')
            history.append({
                'date':       date_str,
                'doc_type':   doc_type,
                'doc_no':     doc_no,
                'doc_seq':    doc_seq,
                'doc_name':   doc_name,
                'warehouse':  wh,
                'wh_name':    wh_map.get(wh, ''),
                'in_out':     in_out(src, doc_type, doc_name, qty),
                'qty':        '{:,.3f}'.format(qty) if qty else '0',
                'src':        src,
            })

        return jsonify({'success': True, 'item_no': item_no, 'count': len(history), 'history': history})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/inventory/future')
def inventory_future():
    """取得指定品號的未來異動量（未完成製令）。
    - 預計生產（入庫）：MOCTA WHERE TA006=品號 AND 狀態未完且尚有未完工量
    - 預計領料（出庫）：MOCTB WHERE TB003=品號（子件）AND 狀態未完且尚有未領量
    """
    item_no = request.args.get('item_no', '').strip()
    if not item_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        from datetime import datetime, timedelta
        # 過濾掉太舊（>30 天前）的孤兒未結紀錄；保留近期含已逾期但未完成的製令
        cutoff = (datetime.now() - timedelta(days=30)).strftime('%Y%m%d')

        conn = get_erp_conn()
        cur  = conn.cursor()

        # 預計生產 (MOCTA) - 入庫；只顯示預計入庫日 >= 今天的紀錄
        cur.execute("""
            SELECT
                TA010                                       AS date_,
                RTRIM(TA001)                                AS doc_type,
                RTRIM(TA002)                                AS doc_no,
                CAST((TA015 - ISNULL(TA017,0)) AS DECIMAL(18,3)) AS qty,
                RTRIM(ISNULL(TA020,''))                     AS warehouse,
                RTRIM(ISNULL(TA019,''))                     AS plant,
                RTRIM(ISNULL(TA011,''))                     AS status,
                RTRIM(ISNULL(TA026,''))                     AS so_no,
                RTRIM(ISNULL(TA027,''))                     AS so_seq,
                RTRIM(ISNULL(TA028,''))                     AS so_line,
                RTRIM(ISNULL(TA029,''))                     AS remark
            FROM MOCTA
            WHERE RTRIM(TA006) = ?
              AND ISNULL(TA011,'') NOT IN ('Y','y')
              AND (TA015 - ISNULL(TA017,0)) > 0
              AND ISNULL(TA010,'') >= ?
        """, (item_no, cutoff))
        produce_rows = cur.fetchall()

        # 預計領料 (MOCTB) - 出庫；只顯示預計領料日 >= 今天的紀錄
        cur.execute("""
            SELECT
                TB015                                       AS date_,
                RTRIM(TB001)                                AS doc_type,
                RTRIM(TB002)                                AS doc_no,
                CAST((TB004 - ISNULL(TB005,0)) AS DECIMAL(18,3)) AS qty,
                RTRIM(ISNULL(TB009,''))                     AS warehouse,
                RTRIM(ISNULL(TB011,''))                     AS status,
                RTRIM(ISNULL(TB014,''))                     AS parent_item,
                RTRIM(ISNULL(TB006,''))                     AS proc_code
            FROM MOCTB
            WHERE RTRIM(TB003) = ?
              AND ISNULL(TB011,'') NOT IN ('Y','y')
              AND (TB004 - ISNULL(TB005,0)) > 0
              AND ISNULL(TB015,'') >= ?
        """, (item_no, cutoff))
        require_rows = cur.fetchall()

        # 廠別/庫別名稱對照
        cur.execute("SELECT RTRIM(MB001), RTRIM(ISNULL(MB002,'')) FROM CMSMB")
        plant_map = {r[0]: r[1] for r in cur.fetchall()}
        cur.execute("SELECT RTRIM(MC001), RTRIM(ISNULL(MC002,'')) FROM CMSMC")
        wh_map = {r[0]: r[1] for r in cur.fetchall()}
        conn.close()

        def fmt_date(d):
            s = (d or '').strip() if isinstance(d, str) else str(d or '').strip()
            return f'{s[:4]}/{s[4:6]}/{s[6:8]}' if len(s) == 8 else s

        future = []
        for r in produce_rows:
            qty = float(r[3] or 0)
            so_no, so_seq, so_line = r[7], r[8], r[9]
            ref_parts = []
            if so_no: ref_parts.append(f'{so_no}-{so_seq}' if so_seq else so_no)
            if so_line: ref_parts.append(so_line)
            ref = ' '.join(ref_parts).strip()
            remark = (r[10] or '').strip()
            future.append({
                'date':       fmt_date(r[0]),
                'date_raw':   (r[0] or '').strip(),
                'type':       '預計生',
                'doc_type':   r[1],
                'doc_no':     r[2],
                'in_qty':     '{:,.3f}'.format(qty).rstrip('0').rstrip('.') if qty else '',
                'out_qty':    '',
                'warehouse':  r[4],
                'wh_name':    wh_map.get(r[4], ''),
                'plant':      r[5],
                'plant_name': plant_map.get(r[5], ''),
                'status':     r[6],
                'remark':     f'{r[1]}-{r[2]}' + (f' {ref}' if ref else '') + (f' {remark}' if remark else ''),
            })
        for r in require_rows:
            qty = float(r[3] or 0)
            future.append({
                'date':       fmt_date(r[0]),
                'date_raw':   (r[0] or '').strip(),
                'type':       '預計領',
                'doc_type':   r[1],
                'doc_no':     r[2],
                'in_qty':     '',
                'out_qty':    '{:,.3f}'.format(qty).rstrip('0').rstrip('.') if qty else '',
                'warehouse':  r[4],
                'wh_name':    wh_map.get(r[4], ''),
                'plant':      '',
                'plant_name': '',
                'status':     r[5],
                'remark':     f'{r[1]}-{r[2]} {r[6]}-{item_no}-{r[7]}'.strip(),
            })

        future.sort(key=lambda x: (x['date_raw'] or '99999999', x['doc_no']))
        return jsonify({'success': True, 'item_no': item_no, 'count': len(future), 'future': future})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/order')
def order_detail():
    """客戶訂單詳情（訂單主檔 COPTC + 明細 COPTD），給製令詳細資訊彈窗的「展開訂單」用。

    製令主檔 MOCTA 的 TA026/TA027（訂單單別/訂單單號）唯一對得上的訂單資料表，
    2026-08-20 唯讀探測實機驗證：COPTC 是訂單頭（1 張訂單 1 列），COPTD 是訂單品項明細
    （TD003 序號，跟 ERP 原生畫面「客戶訂單建立作業(COPI06)」逐欄核對過，見 docs/erp-order.md）。
    """
    order_type = request.args.get('type', '').strip()
    order_no = request.args.get('no', '').strip()
    if not order_type or not order_no:
        return jsonify({'success': False, 'error': '請提供訂單單別與訂單單號'}), 400

    try:
        conn = get_erp_conn()
        cur = conn.cursor()

        def fmt_date(s):
            s = (s or '').strip()
            return f'{s[:4]}/{s[4:6]}/{s[6:8]}' if len(s) == 8 else ''

        cur.execute("""
            SELECT RTRIM(ISNULL(TC004,'')), RTRIM(ISNULL(TC053,'')), TC003, RTRIM(ISNULL(TC040,''))
            FROM COPTC WHERE TC001 = ? AND TC002 = ?
        """, (order_type, order_no))
        hrow = cur.fetchone()
        if not hrow:
            return jsonify({'success': False, 'error': f'查無訂單 {order_type}-{order_no}'}), 404

        confirmer_id = (hrow[3] or '').strip()
        confirmer_name = confirmer_id
        if confirmer_id:
            try:
                confirmer_name = fetch_employee_name_map().get(confirmer_id, confirmer_id)
            except Exception:
                pass

        header = {
            'type': order_type, 'no': order_no,
            'customer_no': hrow[0], 'customer_name': hrow[1],
            'order_date': fmt_date(hrow[2]), 'confirmer': confirmer_name,
        }

        cur.execute("SELECT RTRIM(MC001), RTRIM(ISNULL(MC002,'')) FROM CMSMC")
        wh_map = {r[0]: r[1] for r in cur.fetchall()}

        cur.execute("""
            SELECT RTRIM(ISNULL(TD003,'')), RTRIM(ISNULL(TD004,'')), RTRIM(ISNULL(TD005,'')),
                   RTRIM(ISNULL(TD006,'')), RTRIM(ISNULL(TD007,'')), TD008, TD009, TD013
            FROM COPTD WHERE TD001 = ? AND TD002 = ? ORDER BY TD003
        """, (order_type, order_no))

        def fmt_qty(v):
            try:
                f = float(v or 0)
            except (TypeError, ValueError):
                return str(v or '')
            return '{:,.3f}'.format(f).rstrip('0').rstrip('.')

        lines = [{
            'seq': r[0], 'item_no': r[1], 'item_name': r[2], 'spec': r[3],
            'warehouse': r[4], 'wh_name': wh_map.get(r[4], ''),
            'order_qty': fmt_qty(r[5]), 'delivered_qty': fmt_qty(r[6]),
            'due_date': fmt_date(r[7]),
        } for r in cur.fetchall()]
        conn.close()

        return jsonify({'success': True, 'header': header, 'lines': lines})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/cache/refresh', methods=['POST'])
def refresh_cache():
    """手動清除快取，下次查詢時會重新抓取 SSRS 資料"""
    with _cache_lock:
        _cache.clear()
    return jsonify({'success': True, 'message': '快取已清除'})


# ════════════════════════════════════════════════════════════════
#  申請單功能
# ════════════════════════════════════════════════════════════════

# 各表單類型定義（type_key → 顯示名稱、範本檔名、欄位清單）
_FORM_TYPES = {
    'pp015': {
        'name': 'PP-M-015 內部業務聯絡單',
        'template': 'pp015.docx',
        'fields': [
            {'key': 'date',     'label': '日期',   'type': 'date',           'required': True},
            {'key': 'priority', 'label': '速別',   'type': 'select',         'required': False, 'options': ['普通件', '速件', '最速件']},
            {'key': 'depts', 'label': '受文單位', 'type': 'checkbox_group', 'required': False,
             'options': [
                 '總經理', '副總',
                 '總經理室', '生技', '資管', '美工設計',
                 '資材部',  '採購', '生管', '倉管',
                 '管理部',  '財會', '總務人資',
                 '加工部',  '成品部',
                 '業務部',  '研發部', '品保部',
             ],
             'groups': [                                  # 每個子列表 = 一排
                 ['總經理', '副總'],
                 ['總經理室', '生技', '資管', '美工設計'],
                 ['資材部',  '採購', '生管', '倉管'],
                 ['管理部',  '財會', '總務人資'],
                 ['加工部',  '成品部'],
                 ['業務部',  '研發部', '品保部'],
             ],
             'sub_opts': ['生技', '資管', '美工設計', '採購', '生管', '倉管', '財會', '總務人資', '成品部'],
             'defaults': ['總經理', '加工部']},
            {'key': 'author',   'label': '承辦人', 'type': 'text',           'required': False},
            {'key': 'subject',  'label': '主旨',   'type': 'text',           'required': True},
            {'key': 'content',  'label': '說明',   'type': 'textarea',       'required': True},
        ],
    },
    'hr028': {
        'name': 'P-HR-028-01A 人事懲戒核定書',
        'template': 'hr028.docx',
        'fields': [
            {'key': 'date',     'label': '通知日期',   'type': 'date',     'required': True},
            {'key': 'emp_no',   'label': '員工編號',   'type': 'text',     'required': False},
            {'key': 'name',     'label': '姓名',       'type': 'text',     'required': True},
            {'key': 'title',    'label': '職稱',       'type': 'text',     'required': False},
            {'key': 'punish',     'label': '懲戒分類',   'type': 'select',   'required': True, 'options': ['大過', '小過', '申誡', '警告']},
            {'key': 'pun_count', 'label': '懲戒次數',   'type': 'select',   'required': False, 'options': ['1', '2', '3']},
            {'key': 'reason',   'label': '懲戒事由',   'type': 'textarea', 'required': True},
            {'key': 'evidence', 'label': '檢附之證據', 'type': 'textarea', 'required': False},
            {'key': 'method',   'label': '懲戒方式',   'type': 'text',     'required': False},
        ],
    },
    'hr029': {
        'name': '人事獎勵核定書',
        'template': 'hr029.docx',
        'fields': [
            {'key': 'date',     'label': '通知日期',   'type': 'date',     'required': True},
            {'key': 'emp_no',   'label': '員工編號',   'type': 'text',     'required': False},
            {'key': 'name',     'label': '姓名',       'type': 'text',     'required': True},
            {'key': 'title',    'label': '職稱',       'type': 'text',     'required': False},
            {'key': 'award',     'label': '獎勵分類',   'type': 'select',   'required': True, 'options': ['職位晉昇', '大功', '小功', '嘉獎', '優點']},
            {'key': 'award_count', 'label': '獎勵次數', 'type': 'select',   'required': False, 'options': ['1', '2', '3']},
            {'key': 'reason',   'label': '獎勵事由',   'type': 'textarea', 'required': True},
            {'key': 'evidence', 'label': '檢附之證據', 'type': 'textarea', 'required': False},
            {'key': 'method',   'label': '獎勵方式',   'type': 'text',     'required': False},
        ],
    },
    'pp017': {
        'name': 'PP-M-017 報廢申請單',
        'template': 'pp017.docx',
        'fields': [
            {'key': 'date',      'label': '申請日期',  'type': 'date',     'required': True},
            {'key': 'category',  'label': '類別',      'type': 'select',   'required': True, 'options': ['機器', '模治具', '量規儀器', '其他']},
            {'key': 'applicant', 'label': '申請人',    'type': 'text',     'required': True},
            {'key': 'item',      'label': '品名/規格', 'type': 'text',     'required': True},
            {'key': 'item_no',   'label': '編號/品號', 'type': 'text',     'required': False},
            {'key': 'qty',       'label': '數量',      'type': 'text',     'required': False},
            {'key': 'reason',    'label': '報廢理由',  'type': 'textarea', 'required': True},
        ],
    },
    'ppq006': {
        'name': 'PP-Q-006 異常處理單',
        'template': 'ppq006.docx',
        'store_path': r'\\192.168.1.99\加工部-資料夾\【技術資料】\W.表單.活動\4.異常處理單',
        'photos_page2': True,          # 照片改附在文件第2頁（範本第1頁沒有照片欄位）
        'fields': [
            {'key': 'date',        'label': '受理日期',     'type': 'date',     'required': True},
            {'key': 'mo_no',       'label': '生產製令',     'type': 'text',     'required': False},
            {'key': 'item_no',     'label': '產品編號',     'type': 'text',     'required': False},
            {'key': 'item_name',   'label': '品名',         'type': 'text',     'required': True},
            {'key': 'qty',         'label': '批量',         'type': 'text',     'required': False},
            {'key': 'defect_rate', 'label': '不合格率',     'type': 'text',     'required': False},
            {'key': 'abn_type',    'label': '異常情形',     'type': 'text',     'required': True},
            {'key': 'abn_desc',    'label': '異常內容說明', 'type': 'textarea', 'required': False},
            {'key': 'temp_action', 'label': '臨時矯正對策', 'type': 'textarea', 'required': False},
            {'key': 'abn_hours',   'label': '異常工時',     'type': 'text',     'required': False},
            {'key': 'handler',     'label': '經辦人',       'type': 'text',     'required': True},
        ],
    },
    'gcd': {
        'name': '公出單',
        'template': 'gcd.docx',
        'store_path': r'\\192.168.1.99\加工部-資料夾\【技術資料】\W.表單.活動\工出單',
        'fields': [
            {'key': 'date',     'label': '日期',     'type': 'date',     'required': True},
            {'key': 'name',     'label': '姓名',     'type': 'text',     'required': True},
            {'key': 'location', 'label': '地點',     'type': 'text',     'required': True},
            {'key': 'reason',   'label': '外出原因', 'type': 'textarea', 'required': True},
        ],
    },
}

# 各表單輸出檔名規則（form_type → 產生檔名主體的函式）
def _application_filename_base(form_type, ctx):
    if form_type == 'pp017':
        return f'{ctx.get("item", "")}_報廢單'.strip('_')
    if form_type == 'hr028':
        return f'人事懲處({ctx.get("name", "")})'
    if form_type == 'hr029':
        return f'人事獎勵({ctx.get("name", "")})'
    if form_type == 'gcd':
        return f'{ctx.get("name", "")}_公出單'.strip('_')
    if form_type == 'ppq006':
        return f'{ctx.get("item_name", "")}_異常處理單'.strip('_')
    return ctx.get('subject', '') or '內部業務聯絡單'


# 各表單檔名裡的固定字尾（來自 _application_filename_base），用來在搜尋時
# 精準辨識檔案類型——多個表單共用同一個預設資料夾（pp015/hr028/hr029/pp017），
# 光靠資料夾分不出誰是誰，要靠檔名字尾。pp015 沒有固定字尾（主旨自由填寫），
# 用排除法：不含其他任何字尾的才算 pp015。
_FORM_FILE_MARKERS = {
    'pp017':  '_報廢單',
    'hr028':  '人事懲處(',
    'hr029':  '人事獎勵(',
    'gcd':    '_公出單',
    'ppq006': '_異常處理單',
}


def _file_matches_type(form_type, filename):
    marker = _FORM_FILE_MARKERS.get(form_type)
    if marker:
        return marker in filename
    return not any(m in filename for m in _FORM_FILE_MARKERS.values())


@app.route('/application')
def application_page():
    """申請單主頁"""
    return render_template('application.html', form_types=_FORM_TYPES)


@app.route('/api/application/order_lookup')
def application_order_lookup():
    """異常處理單「品名」欄位打字搜尋目前的製令，選一筆代入生產製令/產品編號/批量。

    直接呼叫 search_orders()（未完工製令報表，跟製令查詢首頁同一份資料源、支援
    CLAUDE.md 講的空格=AND／-前綴=排除語法），不像 /api/query 額外查 SFT／效率報表
    ——那兩支是 SOAP/SSRS 呼叫，這裡只是打字自動完成用，跑那些純粹拖慢速度。
    """
    q = request.args.get('q', '').strip()
    if not q:
        return jsonify({'success': False, 'error': '請輸入品名'}), 400
    try:
        records = search_orders(order_id='', product_name=q, unit='*', release_status='已發放')
        # 同一張製令會因為多道製程出現好幾列，只取每張製令第一筆代表
        seen = set()
        results = []
        for r in records:
            mo_no = r.get('單別', '').strip()
            if not mo_no or mo_no in seen:
                continue
            seen.add(mo_no)
            results.append({
                'mo_no':      mo_no,
                'item_no':    r.get('品號', '').strip(),
                'item_name':  r.get('品名', '').rstrip('|').strip(),
                'qty':        r.get('預計生產數', '').strip(),
            })
            if len(results) >= 30:
                break
        return jsonify({'success': True, 'data': results, 'count': len(results)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/application/list')
def application_list():
    """列出 NAS 申請單資料夾的既有檔案"""
    import datetime as _dt
    year = request.args.get('year', str(_dt.date.today().year))
    keyword = request.args.get('keyword', '').strip()
    type_filter = request.args.get('type', '').strip()
    if type_filter and type_filter not in _FORM_TYPES:
        return jsonify({'files': [], 'error': f'未知表單類型: {type_filter}'}), 400
    base = getattr(config, 'APPLICATION_STORE_PATH', '')
    if not base:
        return jsonify({'error': '未設定 APPLICATION_STORE_PATH'}), 500

    if type_filter:
        # 指定類型：只掃該類型自己的資料夾（多個表單共用預設資料夾時，
        # 靠 _file_matches_type 的檔名字尾再篩掉同資料夾的其他類型）
        bases = [_FORM_TYPES[type_filter].get('store_path') or base]
    else:
        # 未指定：預設資料夾 + 各表單自己的 store_path 全部掃
        bases, seen = [base], {base}
        for fd in _FORM_TYPES.values():
            sp = fd.get('store_path')
            if sp and sp not in seen:
                seen.add(sp)
                bases.append(sp)

    files = []
    errors = []
    for b in bases:
        target = os.path.join(b, year)
        try:
            entries = list(os.scandir(target))
        except Exception as e:
            errors.append(str(e))
            continue
        for e in entries:
            if not e.is_file():
                continue
            name_lower = e.name.lower()
            if not (name_lower.endswith('.doc') or name_lower.endswith('.docx')):
                continue
            if keyword and keyword.lower() not in e.name.lower():
                continue
            if type_filter and not _file_matches_type(type_filter, e.name):
                continue
            stat = e.stat()
            files.append({
                'filename': e.name,
                'path': e.path,
                'size': stat.st_size,
                'mtime': _dt.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
            })

    # 全部資料夾都掃不到才算錯誤（某個年度資料夾還沒建立是正常的）
    if not files and errors:
        return jsonify({'files': [], 'error': errors[0]})

    files.sort(key=lambda x: x['mtime'], reverse=True)
    return jsonify({'files': files, 'year': year, 'total': len(files)})


def _expand_newline_runs(doc):
    """把 run 內的換行字元展開成 Word 的 <w:br/>

    docxtpl 直接塞多行字串時，換行字元在 Word 裡不會斷行（會黏成一行）。
    python-docx 的 run.text setter 會自動把 \\n 轉成 <w:br/>，所以重設一次即可。
    """
    def _walk(parent):
        for p in parent.paragraphs:
            for r in p.runs:
                if '\n' in r.text:
                    r.text = r.text
        for t in parent.tables:
            for row in t.rows:
                for cell in row.cells:
                    _walk(cell)
    _walk(doc)


def _strip_trailing_empty_paragraphs(doc):
    """刪掉文件尾端的空白段落

    範本尾端常留一堆空段落，會把內容擠到第2頁，附件照片就變成第3頁。
    先清乾淨再 add_page_break()，照片才會正好落在第2頁。
    """
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    for child in reversed(list(body.iterchildren())):
        if child.tag == _qn('w:sectPr'):
            continue
        is_empty_p = (child.tag == _qn('w:p')
                      and not ''.join(child.itertext()).strip()
                      and next(child.iter(_qn('a:blip')), None) is None)
        if not is_empty_p:
            break
        body.remove(child)


def _add_attach_title(doc, text):
    """加附件標題（不用 add_heading，避免範本沒有 Heading 樣式時出錯）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return p


def _add_fitted_picture(doc, path):
    """插入照片並限制尺寸：直式限高、橫式限寬，避免撐破版面"""
    try:
        from PIL import Image
        with Image.open(path) as im:
            w, h = im.size
        if h > w:
            doc.add_picture(path, height=Mm(150))
            return
    except Exception:
        pass
    doc.add_picture(path, width=Mm(160))


@app.route('/api/application/create', methods=['POST'])
def application_create():
    """接收表單資料 + 照片，產生 .docx 存到 NAS"""
    import datetime as _dt
    import tempfile
    import shutil

    if not _DOCXTPL_OK:
        return jsonify({'ok': False, 'error': 'docxtpl 套件未安裝，請執行 pip install docxtpl'}), 500

    form_type = request.form.get('form_type', '')
    if form_type not in _FORM_TYPES:
        return jsonify({'ok': False, 'error': f'未知表單類型: {form_type}'}), 400

    form_def = _FORM_TYPES[form_type]
    template_path = os.path.join(
        getattr(config, 'FORM_TEMPLATES_DIR', None) or os.path.join(_APP_DIR, 'static', 'form_templates'),
        form_def['template']
    )

    if not os.path.exists(template_path):
        return jsonify({
            'ok': False,
            'error': f'找不到 Word 範本：{template_path}，請先將 .docx 範本放入 static/form_templates/'
        }), 500

    # 收集欄位值（checkbox_group 用 getlist；其餘用 get）
    context = {}
    for f in form_def['fields']:
        if f['type'] == 'checkbox_group':
            # checkbox 值以陣列送入（name="depts[]"），不放入 context（由下方衍生欄位處理）
            pass
        else:
            context[f['key']] = request.form.get(f['key'], '')

    # 衍生欄位：核取記號（■/□）與民國日期，供範本 placeholder 使用
    def _mark(selected, target):
        return '■' if selected == target else '□'

    def _fmt_date(iso, use_roc=False):
        """ISO 日期轉中文格式；use_roc=True 時用民國年，否則用西元年"""
        try:
            y, m, d = iso.split('-')
            if use_roc:
                return f'中華民國 {int(y) - 1911} 年 {int(m):02d} 月 {int(d):02d} 日'
            return f'{y} 年 {int(m):02d} 月 {int(d):02d} 日'
        except Exception:
            return ''

    # 所有表單統一用西元年
    context['roc_date'] = _fmt_date(context.get('date', ''))

    def _fmt_date_slash(iso):
        """ISO 日期轉 YYYY/M/D 格式"""
        try:
            y, m, d = iso.split('-')
            return f'{int(y)}/{int(m)}/{int(d)}'
        except Exception:
            return ''

    context['gcd_date'] = _fmt_date_slash(context.get('date', ''))
    context['slash_date'] = context['gcd_date']

    if form_type == 'pp015':
        pri = context.get('priority', '')
        context['pri1'] = _mark(pri, '最速件')
        context['pri2'] = _mark(pri, '速件')
        context['pri3'] = _mark(pri, '普通件')
        # 受文單位（可複選）→ ■/□ 記號
        depts_sel = request.form.getlist('depts[]')
        # 主部門 → ■/□
        main_dept_map = [
            ('chk_gm',    '總經理'),
            ('chk_vp',    '副總'),
            ('chk_gmo',   '總經理室'),
            ('chk_mat',   '資材部'),
            ('chk_adm',   '管理部'),
            ('chk_proc',  '加工部'),
            ('chk_sales', '業務部'),
            ('chk_rd',    '研發部'),
            ('chk_qa',    '品保部'),
        ]
        for key, dept in main_dept_map:
            context[key] = '■' if dept in depts_sel else '□'
        # 子部門 → ●/○
        sub_dept_map = [
            ('chk_jt',     '生技'),
            ('chk_it',     '資管'),
            ('chk_design', '美工設計'),
            ('chk_pur',    '採購'),
            ('chk_pm',     '生管'),
            ('chk_wh',     '倉管'),
            ('chk_fin',    '財會'),
            ('chk_hr',     '總務人資'),
            ('chk_fp',     '成品部'),
        ]
        for key, dept in sub_dept_map:
            context[key] = '●' if dept in depts_sel else '○'
    elif form_type == 'hr028':
        pun = context.get('punish', '')
        cnt = context.get('pun_count', '')
        # 只保留 4 種懲戒
        context['pun_d'] = _mark(pun, '大過')
        context['pun_e'] = _mark(pun, '小過')
        context['pun_f'] = _mark(pun, '申誡')
        context['pun_g'] = _mark(pun, '警告')
        # 次數欄位（只填入對應懲戒分類的次數）
        context['cnt_d'] = cnt if pun == '大過' else ''
        context['cnt_e'] = cnt if pun == '小過' else ''
        context['cnt_f'] = cnt if pun == '申誡' else ''
        context['cnt_g'] = cnt if pun == '警告' else ''
    elif form_type == 'hr029':
        awd = context.get('award', '')
        cnt = context.get('award_count', '')
        context['award_a'] = _mark(awd, '職位晉昇')
        context['award_b'] = _mark(awd, '大功')
        context['award_c'] = _mark(awd, '小功')
        context['award_d'] = _mark(awd, '嘉獎')
        context['award_e'] = _mark(awd, '優點')
        # 次數欄位（只填入對應獎勵分類的次數）
        context['award_b_cnt'] = cnt if awd == '大功' else ''
        context['award_c_cnt'] = cnt if awd == '小功' else ''
        context['award_d_cnt'] = cnt if awd == '嘉獎' else ''
        context['award_e_cnt'] = cnt if awd == '優點' else ''
    elif form_type == 'pp017':
        cat = context.get('category', '')
        context['cat_a'] = _mark(cat, '機器')
        context['cat_b'] = _mark(cat, '模治具')
        context['cat_c'] = _mark(cat, '量規儀器')
        context['cat_d'] = _mark(cat, '其他')
    elif form_type == 'gcd':
        veh = context.get('vehicle', '')
        context['veh_car']     = _mark(veh, '汽車')
        context['veh_moto']    = _mark(veh, '機車')
        context['veh_priv']    = _mark(veh, '私車公用')
        context['veh_company'] = _mark(veh, '公司車')
        context['veh_other']   = _mark(veh, '其它')
        context.setdefault('emp_no', '')
        context.setdefault('km', '')
        context.setdefault('plate', '')

    # 照片放第2頁的表單（範本第1頁沒有 {{pic_N}} 照片欄位）
    photos_page2 = form_def.get('photos_page2', False)

    # 處理照片（暫存 → 轉成 InlineImage）
    tmp_dir = tempfile.mkdtemp(prefix='pds_upload_')
    photo_objects = []
    try:
        photos = request.files.getlist('photos[]')
        for idx, photo in enumerate(photos):
            if photo and photo.filename:
                ext = os.path.splitext(photo.filename)[1].lower() or '.jpg'
                tmp_path = os.path.join(tmp_dir, f'photo_{idx}{ext}')
                photo.save(tmp_path)
                photo_objects.append(tmp_path)

        # 組裝 docxtpl context（照片用 InlineImage）
        tpl = DocxTemplate(template_path)
        if not photos_page2:
            for i, p in enumerate(photo_objects):
                context[f'pic_{i+1}'] = InlineImage(tpl, p, width=Mm(120))
            # 未填照片位置補空字串，避免 KeyError
            for i in range(len(photo_objects) + 1, 6):
                context.setdefault(f'pic_{i}', '')

        tpl.render(context)
        # ★ render 之後一律用 tpl.docx，不能用 tpl.get_docx()——後者會偵測到
        #   is_rendered 而重新載入乾淨範本，把套版結果整個丟掉（實測踩過）。
        rendered_doc = tpl.docx
        # textarea 的換行在 docx 裡預設不會斷行，要展開成 <w:br/>
        _expand_newline_runs(rendered_doc)

        # 處理附件檔案（呈現於文件第2頁）
        attachments = request.files.getlist('attachments[]')
        attach_paths = []
        for idx, att in enumerate(attachments):
            if att and att.filename:
                ext = os.path.splitext(att.filename)[1].lower() or '.bin'
                tmp_path = os.path.join(tmp_dir, f'attach_{idx}{ext}')
                att.save(tmp_path)
                attach_paths.append((tmp_path, att.filename))

        # 照片也走第2頁附件流程（異常處理單）
        if photos_page2 and photo_objects:
            attach_paths = [(p, os.path.basename(p)) for p in photo_objects] + attach_paths

        if attach_paths:
            doc = rendered_doc
            _strip_trailing_empty_paragraphs(doc)
            doc.add_page_break()
            _add_attach_title(doc, '附件照片' if photos_page2 else '附件')
            for p, orig_name in attach_paths:
                ext = os.path.splitext(orig_name)[1].lower()
                if ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp'):
                    try:
                        _add_fitted_picture(doc, p)
                    except Exception:
                        doc.add_paragraph(orig_name)
                else:
                    doc.add_paragraph(orig_name)

        # 決定輸出檔名（依表單類型套用對應命名規則）
        date_str = context.get('date', '').replace('-', '') or _dt.date.today().strftime('%Y%m%d')
        subject = _application_filename_base(form_type, context)
        subject = re.sub(r'[\\/:*?"<>|]', '_', subject)[:40]  # 去除非法字元，限長 40
        out_filename = f'{date_str} {subject}.docx'

        # 存到 NAS（若失敗則 fallback 存桌面）
        year = date_str[:4]
        store_base = form_def.get('store_path') or getattr(config, 'APPLICATION_STORE_PATH', '')
        out_dir = os.path.join(store_base, year)
        _nas_warning = None
        try:
            os.makedirs(out_dir, exist_ok=True)
            out_path = os.path.join(out_dir, out_filename)
            tpl.save(out_path)
        except (PermissionError, OSError) as _e:
            # NAS 無法存取（權限、網路、檔案被鎖定等）→ 改存桌面
            _desktop = os.path.join(os.environ.get('USERPROFILE', os.path.expanduser('~')), 'Desktop')
            os.makedirs(_desktop, exist_ok=True)
            # 若桌面同名檔案已存在，加時間戳避免衝突
            _base, _ext = os.path.splitext(out_filename)
            _desk_name = out_filename
            if os.path.exists(os.path.join(_desktop, _desk_name)):
                _desk_name = f"{_base}_{int(time.time())}{_ext}"
            out_path = os.path.join(_desktop, _desk_name)
            tpl.save(out_path)
            _nas_warning = f'NAS 無法存取（{type(_e).__name__}），檔案已改存至桌面：{out_path}'

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    resp = {'ok': True, 'filename': out_filename, 'path': out_path}
    if _nas_warning:
        resp['warning'] = _nas_warning
    return jsonify(resp)


def _docx_to_html(path):
    """將 .docx 轉成可預覽的 HTML（保留段落、表格、內嵌圖片）"""
    import base64
    from markupsafe import escape
    from docx import Document as _Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = _Document(path)

    def _img_tags(elm):
        tags = ''
        for blip in elm.iter(qn('a:blip')):
            rId = blip.get(qn('r:embed'))
            if not rId:
                continue
            try:
                image_part = doc.part.related_parts[rId]
                b64 = base64.b64encode(image_part.blob).decode('ascii')
                tags += (f'<img src="data:{image_part.content_type};base64,{b64}" '
                         f'style="max-width:100%;display:block;margin:.3rem 0;">')
            except Exception:
                pass
        return tags

    def _para_html(p):
        text = p.text.strip()
        imgs = _img_tags(p._p)
        if not text and not imgs:
            return ''
        align = ''
        try:
            if int(p.alignment) == 1:  # WD_ALIGN_PARAGRAPH.CENTER
                align = 'text-align:center;'
        except Exception:
            pass
        body = str(escape(text)).replace('\n', '<br>') if text else ''
        return f'<p style="margin:.25rem 0;{align}">{body}{imgs}</p>'

    def _table_html(tbl):
        rows_html = ''
        for row in tbl.rows:
            cells_html = ''
            for cell in row.cells:
                cell_text = str(escape(cell.text)).replace('\n', '<br>')
                cell_imgs = _img_tags(cell._tc)
                cells_html += (f'<td style="border:1px solid #E0DACC;padding:.3rem .5rem;'
                                f'vertical-align:top;">{cell_text}{cell_imgs}</td>')
            rows_html += f'<tr>{cells_html}</tr>'
        return f'<table style="width:100%;border-collapse:collapse;margin:.4rem 0;font-size:0.92rem;">{rows_html}</table>'

    parts = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            html = _para_html(Paragraph(child, doc))
            if html:
                parts.append(html)
        elif child.tag == qn('w:tbl'):
            parts.append(_table_html(Table(child, doc)))
    return ''.join(parts)


@app.route('/api/application/view')
def application_view():
    """預覽既有申請單檔案內容（轉成 HTML）"""
    path = request.args.get('path', '')
    if not path or not os.path.isfile(path):
        return jsonify({'ok': False, 'error': '檔案不存在'}), 404
    if not path.lower().endswith('.docx'):
        return jsonify({'ok': False, 'error': '僅支援預覽 .docx 格式，請點擊開啟檔案查看'}), 400
    try:
        return jsonify({'ok': True, 'html': _docx_to_html(path)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/application/open', methods=['POST'])
def application_open():
    """用 Shell 開啟指定 NAS 上的申請單檔案"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400
    try:
        os.startfile(path)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/application/mail', methods=['POST'])
def application_mail():
    """把申請單檔案帶進 Outlook 新郵件草稿（附加檔案後開啟視窗，不自動寄出）

    只負責「開草稿」：收件者、內文都可以再改，按下寄出的動作留給使用者。
    """
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path or not os.path.isfile(path):
        return jsonify({'ok': False, 'error': '檔案不存在'}), 404

    subject = data.get('subject') or os.path.splitext(os.path.basename(path))[0]
    body = data.get('body') or ''
    to = data.get('to') or ''

    import pythoncom
    pythoncom.CoInitialize()
    try:
        import win32com.client
        # 用 Dispatch（不是 DispatchEx）才會接到使用者已登入的 Outlook 設定檔
        outlook = win32com.client.Dispatch('Outlook.Application')
        mail = outlook.CreateItem(0)          # 0 = olMailItem
        mail.Subject = subject
        mail.Body = body
        if to:
            mail.To = to
        mail.Attachments.Add(os.path.abspath(path))
        mail.Display(False)                   # 非強制回應視窗，讓使用者自己按寄出
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': f'開啟 Outlook 失敗：{e}'}), 500
    finally:
        pythoncom.CoUninitialize()


@app.route('/management')
def management_page():
    """管理頁主頁"""
    return render_template('management.html', form_types=_FORM_TYPES)


@app.route('/api/attendance/list')
def attendance_list():
    """讀取 Google 試算表『M9出勤表』資料"""
    cache_key = 'attendance_list'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.ATTENDANCE_SHEET_ID, gid=config.ATTENDANCE_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    try:
        emp_map = fetch_employee_name_map()
        for r in rows:
            for k in list(r.keys()):
                emp_id = r.get(k, '').strip()
                if emp_id and emp_id in emp_map:
                    r[k + '_name'] = emp_map[emp_id]
        result = {'success': True, 'count': len(rows), 'headers': list(rows[0].keys()) if rows else [], 'rows': rows}
        cache_set(cache_key, result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/leave/list')
def leave_list():
    """讀取 Google 試算表『請假單』資料"""
    cache_key = 'leave_list'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'employee_name_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.LEAVE_SHEET_ID, config.LEAVE_SHEET_NAME)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    try:
        name_map = fetch_employee_name_map()
    except Exception:
        name_map = {}

    leaves = []
    for r in rows:
        emp_id = (r.get('請假人員') or '').strip()
        leaves.append({
            'date':   (r.get('請假日期') or '').strip(),
            'employee': emp_id,
            'name':   name_map.get(emp_id, ''),
            'type':   (r.get('假別') or '').strip(),
            'start':  (r.get('起始時間') or '').strip(),
            'end':    (r.get('結束時間') or '').strip(),
            'remark': (r.get('備註說明') or '').strip(),
        })

    def _sort_key(item):
        try:
            y, m, d = item['date'].split('/')
            return (int(y), int(m), int(d))
        except Exception:
            return (0, 0, 0)

    leaves.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(leaves), 'leaves': leaves}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/overtime/list')
def overtime_list():
    """讀取 Google 試算表『加班統計』資料"""
    cache_key = 'overtime_list'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.OVERTIME_SHEET_ID, config.OVERTIME_SHEET_NAME)
    except Exception as e:
        return jsonify({'success': False, 'error': f'無法連線至 Google 試算表：{str(e)}'}), 502

    try:
        name_map = fetch_employee_name_map()
    except Exception:
        name_map = {}

    overtimes = []
    for r in rows:
        emp_id = (r.get('加班人員ID') or r.get('加班人員') or '').strip()
        overtimes.append({
            'date':   (r.get('加班日期') or '').strip(),
            'employee': emp_id,
            'name':   name_map.get(emp_id, ''),
            'hours':  (r.get('加班時數') or '').strip(),
            'remark': (r.get('備註') or '').strip(),
        })

    def _sort_key(item):
        try:
            y, m, d = item['date'].split('/')
            return (int(y), int(m), int(d))
        except Exception:
            return (0, 0, 0)

    overtimes.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(overtimes), 'overtimes': overtimes}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/employee/list')
def employee_list():
    """讀取 Google 試算表『員工登錄系統』(1.1員工登錄系統分頁)，回傳人員名冊"""
    try:
        roster = fetch_employee_roster()
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502
    return jsonify({'success': True, 'count': len(roster), 'employees': roster})


def _purchase_date_key(s):
    """把日期字串（'2026/7/5' 或 '2026-7-5' 等寫法皆可）正規化成 'YYYY-MM-DD'，
    跟 <input type=date> 的 value 格式一致，可直接字串比較。"""
    s = (s or '').strip().replace('-', '/')
    parts = s.split('/')
    if len(parts) < 3:
        return ''
    try:
        y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
    except ValueError:
        return ''
    return f'{y:04d}-{m:02d}-{d:02d}'


@app.route('/api/purchase/list')
def purchase_list():
    """讀取 Google 試算表『採購登入表』（主表分頁）資料。
    整份累積到 1.5萬+ 筆，全部丟給前端會撐爆 localStorage 快取配額（實測炸過），
    所以支援可選的 date_from/date_to（YYYY-MM-DD）縮小回傳範圍；快取本身仍然
    存「全部」資料（避免不同日期區間各打一次 Google Sheet），篩選在取出快取之後做。"""
    cache_key = 'purchase_list'
    if request.args.get('refresh'):
        cache_clear(cache_key)
    cached = cache_get(cache_key)
    if cached is not None:
        records = cached['purchases']
        total_count = cached['count']
    else:
        try:
            rows = fetch_google_sheet_csv(config.PURCHASE_SHEET_ID, sheet_name=config.PURCHASE_SHEET_NAME)
        except Exception:
            return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

        records = []
        for r in rows:
            raw_qty = (r.get('數量') or '').replace(',', '').strip()
            try:    qty = float(raw_qty)
            except ValueError: qty = 0
            raw_unit = (r.get('單價') or '').replace(',', '').strip()
            try:    unit_price = float(raw_unit)
            except ValueError: unit_price = 0
            raw_sub = (r.get('小計') or '').replace(',', '').strip()
            try:    subtotal = float(raw_sub)
            except ValueError: subtotal = 0
            date = (r.get('日期') or '').strip()
            if not date and not (r.get('品名') or '').strip():
                continue  # 跳過空白列
            records.append({
                'date':       date,
                'status':     (r.get('狀態') or '').strip(),
                'name':       (r.get('品名') or '').strip(),
                'spec':       (r.get('規格') or '').strip(),
                'vendor':     (r.get('供應商') or '').strip(),
                'vendor_no':  (r.get('供應商代號') or '').strip(),
                'qty':        qty,
                'unit_price': unit_price,
                'subtotal':   subtotal,
                'account':    (r.get('會計科目') or '').strip(),
                'doc_no':     (r.get('單據號碼') or '').strip(),
                'remark':     (r.get('備註') or '').strip(),
            })

        def _sort_key(item):
            try:
                parts = item['date'].replace('-', '/').split('/')
                return tuple(int(p) for p in parts)
            except Exception:
                return (0, 0, 0)

        records.sort(key=_sort_key, reverse=True)
        total_count = len(records)
        cache_set(cache_key, {'success': True, 'count': total_count, 'purchases': records})

    date_from = request.args.get('date_from', '').strip()
    date_to   = request.args.get('date_to', '').strip()
    if date_from or date_to:
        def _in_range(r):
            k = _purchase_date_key(r['date'])
            if not k:
                return True  # 日期格式怪異看不懂的列，寧可保留給使用者看到也不要悄悄濾掉
            if date_from and k < date_from:
                return False
            if date_to and k > date_to:
                return False
            return True
        records = [r for r in records if _in_range(r)]

    return jsonify({'success': True, 'count': len(records), 'total_count': total_count, 'purchases': records})


# ── 批成本計算（共用區 Excel 範本：刀具/刀表/批成本計算） ──────────────

def _load_batchcost_wb(read_only=False):
    """開啟批成本計算範本檔案（共用網路路徑）"""
    return openpyxl.load_workbook(config.BATCH_COST_FILE_PATH, data_only=read_only)


def _save_batchcost_wb(wb):
    """儲存，若檔案被佔用（例如使用者自己開著 Excel）則 fallback 存到桌面，回傳警告訊息（無警告回 None）"""
    try:
        wb.save(config.BATCH_COST_FILE_PATH)
        return None
    except (PermissionError, OSError) as e:
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        os.makedirs(desktop, exist_ok=True)
        fallback = os.path.join(desktop, f'批成本計算_{int(time.time())}.xlsx')
        wb.save(fallback)
        return f'原始檔案無法寫入（{type(e).__name__}，可能被開啟中），已改存到：{fallback}'


@app.route('/api/batch_cost/lookup_order')
def batch_cost_lookup_order():
    """依製令號碼查詢品號/品名/製程代號（優先從 P2 表查，再備援 SSRS；允許不打 "-" 直接輸入數字）"""
    order = request.args.get('order', '').strip()
    if not order:
        return jsonify({'success': False, 'error': '請輸入製令號碼'}), 400
    # 製令格式為「4位單位代碼-序號」，若使用者輸入全數字沒打 "-"，自動補上方便比對
    order_norm = order.replace('-', '')
    if order.isdigit() and len(order) > 4:
        order_dashed = f'{order[:4]}-{order[4:]}'
    else:
        order_dashed = order

    records = None
    # 1. 優先從 K1_P2（生產報工統計 P2）查詢
    try:
        p2_rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        records = []
        for p2r in p2_rows:
            wo = (p2r.get('製令') or '').strip()
            if wo == order or wo == order_dashed or wo.replace('-', '') == order_norm:
                sec_str = (p2r.get('出站總工時(秒)') or '').strip().replace(',', '')
                try:
                    seconds = float(sec_str) if sec_str else 0
                except ValueError:
                    seconds = 0
                qty_str = (p2r.get('出站數量') or '').strip().replace(',', '')
                try:
                    qty = float(qty_str) if qty_str else 0
                except ValueError:
                    qty = 0
                records.append({
                    '單別': wo,
                    '品號': (p2r.get('品號') or '').strip(),
                    '品名': (p2r.get('品名') or '').strip(),
                    '製程代號': (p2r.get('製程') or '').strip(),
                    '製程名稱': (p2r.get('製程名稱') or '').strip(),
                    '秒數': seconds,
                    '出站數量': qty,
                    '報工人員': (p2r.get('報工人員') or '').strip(),
                    '出站時間': (p2r.get('欄位格式化') or p2r.get('出站時間') or '').strip(),
                })
    except Exception:
        records = None

    # 2. 備援：如果 P2 查不到，試試 SSRS 未完工製令（用補上 "-" 的版本比對，substring 才能命中）
    if not records:
        try:
            records = search_orders(order_id=order_dashed)
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 502

    if not records:
        return jsonify({'success': False, 'error': '查無此製令'})

    # 如果來自 SSRS（沒有製程名稱和秒數），則再查 P2 表補充
    if not any(r.get('製程名稱') for r in records):
        p2_map = {}
        try:
            p2_rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
            for p2r in p2_rows:
                wo = (p2r.get('製令') or '').strip()
                proc = (p2r.get('製程') or '').strip()
                if wo:
                    proc_name = (p2r.get('製程名稱') or '').strip()
                    sec_str = (p2r.get('出站總工時(秒)') or '').strip().replace(',', '')
                    try:
                        seconds = float(sec_str) if sec_str else 0
                    except ValueError:
                        seconds = 0
                    key = (wo, proc)
                    if key not in p2_map:
                        p2_map[key] = {'proc_name': proc_name, 'seconds': seconds}
        except Exception:
            p2_map = {}
    else:
        p2_map = {}

    # 從『生產日報表P5.3』查機台名稱/機台代號（依製令+製程代號比對；P5.3 的製令欄位帶序號尾碼如 -0010，需先去除才能比對）
    p53_map = {}
    try:
        p53_rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, sheet_name=config.PROD_REPORT_SHEET_NAME)
        for p53r in p53_rows:
            wo_full = (p53r.get('製令') or '').strip()
            wo = wo_full.rsplit('-', 1)[0] if wo_full.count('-') >= 2 else wo_full
            proc = (p53r.get('製程代號') or '').strip()
            if wo:
                key = (wo, proc)
                if key not in p53_map:
                    p53_map[key] = {
                        'machine_name': (p53r.get('機台名稱') or '').strip(),
                        'machine_code': (p53r.get('機台代號') or '').strip(),
                    }
    except Exception:
        p53_map = {}

    seen = set()
    rows = []
    for r in records:
        part_no = (r.get('品號') or '').strip()
        proc_code = (r.get('製程代號') or '').strip()
        key = (part_no, proc_code)
        if key in seen:
            continue
        seen.add(key)
        order_full = r.get('單別', '').strip()
        # 如果已經從 P2 查到，直接用；否則從 p2_map 補充
        proc_name = r.get('製程名稱', '')
        seconds = r.get('秒數', 0)
        if not proc_name and order_full in p2_map:
            p2_info = p2_map.get((order_full, proc_code), {})
            proc_name = proc_name or p2_info.get('proc_name', '')
            seconds = seconds or p2_info.get('seconds', 0)
        out_qty = r.get('出站數量', 0)
        p53_info = p53_map.get((order_full, proc_code), {})
        rows.append({
            'order':    order_full,
            'part_no':  part_no,
            'name':     (r.get('品名') or '').rstrip('|').strip(),
            'proc_code': proc_code,
            'proc_name': proc_name,
            'seconds':  seconds,
            'out_qty':  out_qty,
            'machine_name': p53_info.get('machine_name', ''),
            'machine_code': p53_info.get('machine_code', ''),
            'operator': r.get('報工人員', ''),
            'out_time': r.get('出站時間', ''),
        })
    return jsonify({'success': True, 'rows': rows})


@app.route('/api/batch_cost/tool_catalog')
def batch_cost_tool_catalog():
    """讀『刀具』分頁，回傳每把刀的單價/刃數對照表"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        ws = wb[config.BATCH_COST_TOOL_SHEET]
        tools = []
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            name = row[idx.get('品名', 0)] if idx.get('品名') is not None else None
            if not name:
                continue
            tools.append({
                'row':    row_no,
                'name':   str(name).strip(),
                'item':   str(row[idx['項目']] or '').strip() if '項目' in idx else '',
                'vendor': str(row[idx['供應商']] or '').strip() if '供應商' in idx else '',
                'price':  row[idx['單價']] if '單價' in idx else 0,
                'edges':  row[idx['刃數']] if '刃數' in idx else 0,
            })
        wb.close()
        return jsonify({'success': True, 'tools': tools})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_catalog/save', methods=['POST'])
def batch_cost_tool_catalog_save():
    """新增或編輯『刀具資料』分頁裡的一把刀（有 row 就編輯該列，沒有就新增一列）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    name   = (data.get('name') or '').strip()
    item   = (data.get('item') or '').strip()
    vendor = (data.get('vendor') or '').strip()
    price  = data.get('price') or 0
    edges  = data.get('edges') or 0
    if not name:
        return jsonify({'success': False, 'error': '請輸入刀具名稱'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOL_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        values = {'品名': name, '項目': item, '供應商': vendor, '單價': price, '刃數': edges}

        if row_no:
            target_row = int(row_no)
        else:
            target_row = ws.max_row + 1
        for col_name, val in values.items():
            if col_name in idx:
                ws.cell(row=target_row, column=idx[col_name] + 1, value=val)

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_catalog/delete', methods=['POST'])
def batch_cost_tool_catalog_delete():
    """刪除『刀具資料』分頁裡的一列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    if not row_no:
        return jsonify({'success': False, 'error': '缺少列號'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOL_SHEET]
        ws.delete_rows(int(row_no), 1)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map')
def batch_cost_tool_map():
    """讀『刀表』分頁，依品號查該品號目前的刀具配置（T1~T39）"""
    part_no = request.args.get('part_no', '').strip()
    if not part_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        slot_cols = [h for h in header if h and re.match(r'^T\d+$', str(h))]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[idx.get('品號', -1)]:
                continue
            if str(row[idx['品號']]).strip() != part_no:
                continue
            slots = []
            for slot in slot_cols:
                v = row[idx[slot]]
                if v:
                    slots.append({'slot': slot, 'tool': str(v).strip()})
            wb.close()
            return jsonify({'success': True, 'found': True,
                             'machine_type': str(row[idx.get('加工機型', -1)] or '').strip() if '加工機型' in idx else '',
                             'slots': slots})
        wb.close()
        return jsonify({'success': True, 'found': False, 'slots': []})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map/save', methods=['POST'])
def batch_cost_tool_map_save():
    """新增或覆寫『刀表』分頁裡某品號的刀具配置"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    part_no = (data.get('part_no') or '').strip()
    machine_type = (data.get('machine_type') or '').strip()
    slots = data.get('slots') or []   # [{slot:'T1', tool:'...'}, ...]
    if not part_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        today = date.today().strftime('%Y/%m/%d')

        target_row = None
        for row_cells in ws.iter_rows(min_row=2):
            cell = row_cells[idx['品號']]
            if cell.value and str(cell.value).strip() == part_no:
                target_row = row_cells
                break

        if target_row is None:
            new_idx = ws.max_row + 1
            ws.cell(row=new_idx, column=idx['建立日期'] + 1, value=today)
            ws.cell(row=new_idx, column=idx['品號'] + 1, value=part_no)
            ws.cell(row=new_idx, column=idx['加工機型'] + 1, value=machine_type)
            for s in slots:
                if s.get('slot') in idx:
                    ws.cell(row=new_idx, column=idx[s['slot']] + 1, value=s.get('tool', ''))
        else:
            row_no = target_row[0].row
            ws.cell(row=row_no, column=idx['修改日期'] + 1, value=today)
            if machine_type:
                ws.cell(row=row_no, column=idx['加工機型'] + 1, value=machine_type)
            slot_cols = [h for h in header if h and re.match(r'^T\d+$', str(h))]
            filled = {s['slot']: s.get('tool', '') for s in slots if s.get('slot')}
            for slot in slot_cols:
                ws.cell(row=row_no, column=idx[slot] + 1, value=filled.get(slot, ''))

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map/list')
def batch_cost_tool_map_list():
    """讀『刀表』分頁全部品號的刀具配置清單（維護頁用）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        slot_cols = [h for h in header if h and re.match(r'^T\d+$', str(h))]
        items = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            part_no = row[idx.get('品號', -1)] if '品號' in idx else None
            if not part_no:
                continue
            slots = []
            for slot in slot_cols:
                v = row[idx[slot]]
                if v:
                    slots.append({'slot': slot, 'tool': str(v).strip()})
            created = row[idx['建立日期']] if '建立日期' in idx else ''
            modified = row[idx['修改日期']] if '修改日期' in idx else ''
            items.append({
                'part_no': str(part_no).strip(),
                'machine_type': str(row[idx.get('加工機型', -1)] or '').strip() if '加工機型' in idx else '',
                'created': str(created) if created else '',
                'modified': str(modified) if modified else '',
                'slots': slots,
            })
        wb.close()
        return jsonify({'success': True, 'items': items})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/tool_map/delete', methods=['POST'])
def batch_cost_tool_map_delete():
    """刪除『刀表』分頁裡某品號的整列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    part_no = (data.get('part_no') or '').strip()
    if not part_no:
        return jsonify({'success': False, 'error': '請提供品號'}), 400

    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_TOOLMAP_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        idx = {h: i for i, h in enumerate(header) if h}
        target_row_no = None
        for row_cells in ws.iter_rows(min_row=2):
            cell = row_cells[idx['品號']]
            if cell.value and str(cell.value).strip() == part_no:
                target_row_no = row_cells[0].row
                break
        if target_row_no is None:
            return jsonify({'success': False, 'error': '查無此品號的刀表'}), 404
        ws.delete_rows(target_row_no, 1)

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/api/batch_cost/save', methods=['POST'])
def batch_cost_save():
    """把一筆批成本計算結果 append 進『批成本計算』分頁（不存在時自動建立；依表頭欄名對應寫入，不依賴固定欄位順序）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    order      = (data.get('order') or '').strip()
    part_no    = (data.get('part_no') or '').strip()
    name       = (data.get('name') or '').strip()
    proc_code  = (data.get('proc_code') or '').strip()
    proc_name  = (data.get('proc_name') or '').strip()
    qty        = data.get('qty') or 0
    hours      = (data.get('hours') or '').strip()
    tool_cost  = data.get('tool_cost') or 0
    labor_cost = data.get('labor_cost') or 0
    seconds    = data.get('seconds') or 0
    rate       = data.get('rate') or 0
    electricity_cost = data.get('electricity_cost') or 0   # 電費（單位，已是每件成本，非批總額）
    oil_cost          = data.get('oil_cost') or 0            # 油費（單位，同上）
    machine_name = (data.get('machine_name') or '').strip()
    machine_code = (data.get('machine_code') or '').strip()
    out_time     = (data.get('out_time') or '').strip()
    tool_usage   = data.get('tool_usage') or []   # [{slot, tool, price, edges, usage, subtotal}, ...]
    if not order or not part_no:
        return jsonify({'success': False, 'error': '缺少製令或品號'}), 400

    qty_num = float(qty) if qty else 0
    tool_cost_per_unit = round((tool_cost / qty_num) if qty_num else 0, 2)
    total_cost_per_unit = round(labor_cost + tool_cost_per_unit, 2)
    grand_total_per_unit = round(total_cost_per_unit + electricity_cost + oil_cost, 2)

    default_header = ['建立日期', '製令', '製程代號', '製程名稱', '品號', '品名',
                       '每秒鐘生產費用（元）', '完成數量', '刀具費用', '加工秒數', '加工費用',
                       '刀具成本', '加工費用(含刀具成本)', '電費', '油費', '總成本(含刀具+油電)',
                       '機台名稱', '機台代號']
    # 建立日期優先採用生產報工統計P2的出站時間，查不到才退回今天日期
    record_date = out_time or date.today().strftime('%Y/%m/%d')
    value_map = {
        '建立日期': record_date, '製令': order, '製程代號': proc_code, '製程名稱': proc_name,
        '品號': part_no, '品名': name, '每秒鐘生產費用（元）': rate, '完成數量': qty,
        '生產數量': qty,  # 相容舊版表頭命名
        '刀具費用': tool_cost, '加工秒數': seconds,
        '加工費用': labor_cost, '總加工費用': labor_cost,  # 相容舊版表頭命名
        '刀具成本': tool_cost_per_unit,
        '加工費用(含刀具成本)': total_cost_per_unit,
        '電費': electricity_cost, '油費': oil_cost,
        '總成本(含刀具+油電)': grand_total_per_unit,
        '機台名稱': machine_name, '機台代號': machine_code,
        '加工時間': hours, '備註': hours,
    }

    try:
        wb = _load_batchcost_wb(read_only=False)
        if config.BATCH_COST_RECORD_SHEET not in wb.sheetnames:
            ws = wb.create_sheet(config.BATCH_COST_RECORD_SHEET)
            ws.append(default_header)
            header = default_header
        else:
            ws = wb[config.BATCH_COST_RECORD_SHEET]
            header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            if not any(header):
                header = default_header
                for col, h in enumerate(header, start=1):
                    ws.cell(row=1, column=col, value=h)

        new_row = [value_map.get(h, '') for h in header]
        ws.append(new_row)

        # 同步把每把刀的使用次數明細寫入『製令與刀具壽命』分頁（不存在時自動建立）
        used_tools = [t for t in tool_usage if (t.get('usage') or 0) > 0]
        if used_tools:
            lifespan_default_header = ['建立日期', '製令', '製程代號', '製程名稱', '品號', '品名',
                                        '刀號', '刀具名稱', '單價', '使用次數', '小計']
            if config.BATCH_COST_LIFESPAN_SHEET not in wb.sheetnames:
                ws2 = wb.create_sheet(config.BATCH_COST_LIFESPAN_SHEET)
                ws2.append(lifespan_default_header)
                header2 = lifespan_default_header
            else:
                ws2 = wb[config.BATCH_COST_LIFESPAN_SHEET]
                header2 = [c.value for c in next(ws2.iter_rows(min_row=1, max_row=1))]
                if not any(header2):
                    header2 = lifespan_default_header
                    for col, h in enumerate(header2, start=1):
                        ws2.cell(row=1, column=col, value=h)
            for t in used_tools:
                tool_value_map = {
                    '建立日期': record_date, '製令': order, '製程代號': proc_code, '製程名稱': proc_name,
                    '品號': part_no, '品名': name,
                    '刀號': t.get('slot', ''), '刀具名稱': t.get('tool', ''),
                    '單價': t.get('price', 0), '使用次數': t.get('usage', 0), '小計': t.get('subtotal', 0),
                }
                ws2.append([tool_value_map.get(h, '') for h in header2])

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/record/list')
def batch_cost_record_list():
    """讀『批成本計算』分頁全部明細（依目前表頭欄名動態回傳，不假設固定欄位）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        if config.BATCH_COST_RECORD_SHEET not in wb.sheetnames:
            return jsonify({'success': True, 'columns': [], 'records': []})
        ws = wb[config.BATCH_COST_RECORD_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        columns = [h for h in header if h]
        records = []
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue
            rec = {'row': row_no}
            for col_idx, h in enumerate(header):
                if not h:
                    continue
                v = row[col_idx] if col_idx < len(row) else ''
                rec[h] = str(v) if v is not None else ''
            records.append(rec)
        wb.close()

        def _date_key(rec):
            date_str = (rec.get('建立日期') or '').split(' ')[0]
            try:
                parts = date_str.replace('-', '/').split('/')
                return tuple(int(p) for p in parts)
            except (ValueError, TypeError):
                return (0, 0, 0)

        records.sort(key=_date_key, reverse=True)  # 依建立日期由新到舊排序
        return jsonify({'success': True, 'columns': columns, 'records': records})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/lifespan/list')
def batch_cost_lifespan_list():
    """讀『製令與刀具壽命』分頁全部明細（依目前表頭欄名動態回傳，供批成本明細展開列使用）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        if config.BATCH_COST_LIFESPAN_SHEET not in wb.sheetnames:
            return jsonify({'success': True, 'columns': [], 'records': []})

        # 依刀具名稱查供應商/項目（來自刀具資料分頁）
        tool_extra = {}
        if config.BATCH_COST_TOOL_SHEET in wb.sheetnames:
            tws = wb[config.BATCH_COST_TOOL_SHEET]
            theader = [c.value for c in next(tws.iter_rows(min_row=1, max_row=1))]
            tidx = {h: i for i, h in enumerate(theader) if h}
            for trow in tws.iter_rows(min_row=2, values_only=True):
                tname = trow[tidx.get('品名', -1)] if '品名' in tidx else None
                if not tname:
                    continue
                tname = str(tname).strip()
                if tname not in tool_extra:
                    tool_extra[tname] = {
                        'item':   str(trow[tidx['項目']] or '').strip() if '項目' in tidx else '',
                        'vendor': str(trow[tidx['供應商']] or '').strip() if '供應商' in tidx else '',
                    }

        ws = wb[config.BATCH_COST_LIFESPAN_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        columns = [h for h in header if h]
        records = []
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue
            rec = {'row': row_no}
            for col_idx, h in enumerate(header):
                if not h:
                    continue
                v = row[col_idx] if col_idx < len(row) else ''
                rec[h] = str(v) if v is not None else ''
            extra = tool_extra.get((rec.get('刀具名稱') or '').strip(), {})
            rec['項目'] = extra.get('item', '')
            rec['供應商'] = extra.get('vendor', '')
            records.append(rec)
        wb.close()
        if '項目' not in columns:
            columns.append('項目')
        if '供應商' not in columns:
            columns.append('供應商')
        return jsonify({'success': True, 'columns': columns, 'records': records})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/record/delete', methods=['POST'])
def batch_cost_record_delete():
    """刪除『批成本計算』分頁裡的一列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    if not row_no:
        return jsonify({'success': False, 'error': '缺少列號'}), 400
    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_RECORD_SHEET]
        ws.delete_rows(int(row_no), 1)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/api/batch_cost/record/update_machine', methods=['POST'])
def batch_cost_record_update_machine():
    """修改『批成本計算』分頁裡一筆記錄的機台代號/機台名稱，只改這兩欄、不動其他欄位。
    用途：記錄當初依「製令+製程代號」比對生產日報表P5.3帶錯機台時，不用整筆刪除、
    重新輸入完成數量+逐刀使用次數跑一次批成本計算，只要在明細頁直接改機台即可。
    依表頭欄名動態對應（批成本相關分頁的硬規則，見 docs/batch-cost.md）。"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    code = (data.get('machine_code') or '').strip()
    name = (data.get('machine_name') or '').strip()
    if not row_no:
        return jsonify({'success': False, 'error': '缺少列號'}), 400
    if not code:
        return jsonify({'success': False, 'error': '請輸入機台代號'}), 400
    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_RECORD_SHEET]
        header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        col_idx = {h: i + 1 for i, h in enumerate(header) if h}
        if '機台代號' not in col_idx:
            return jsonify({'success': False, 'error': '批成本計算分頁找不到「機台代號」欄位'}), 500
        ws.cell(row=int(row_no), column=col_idx['機台代號'], value=code)
        if '機台名稱' in col_idx:
            ws.cell(row=int(row_no), column=col_idx['機台名稱'], value=name)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'修改失敗：{e}'}), 502


@app.route('/api/batch_cost/machine_list')
def batch_cost_machine_list():
    """從『生產日報表P5.3』抓出現過的機台代號/機台名稱清單（去重，供設備耗用設定的機台代號下拉選單用）"""
    try:
        rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, sheet_name=config.PROD_REPORT_SHEET_NAME)
        seen = {}
        for r in rows:
            code = (r.get('機台代號') or '').strip()
            name = (r.get('機台名稱') or '').strip()
            if not code or code in seen:
                continue
            seen[code] = name
        machines = [{'code': c, 'name': n} for c, n in seen.items()]
        machines.sort(key=lambda m: m['code'])
        return jsonify({'success': True, 'machines': machines})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取機台清單失敗：{e}'}), 502


@app.route('/api/batch_cost/equipment')
def batch_cost_equipment():
    """讀『設備耗用設定』（每台設備電流/電壓/每日耗油費用/每秒鐘生產費用）＋『批成本全域設定』（每度電價）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    try:
        wb = _load_batchcost_wb(read_only=True)
        equipment = []
        if config.BATCH_COST_EQUIPMENT_SHEET in wb.sheetnames:
            ws = wb[config.BATCH_COST_EQUIPMENT_SHEET]
            header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            idx = {h: i for i, h in enumerate(header) if h}
            for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                code = row[idx.get('機台代號', 0)] if idx.get('機台代號') is not None else None
                if not code:
                    continue
                equipment.append({
                    'row':      row_no,
                    'code':     str(code).strip(),
                    'name':     str(row[idx['機台名稱']] or '').strip() if '機台名稱' in idx else '',
                    'model':    str(row[idx['機台型號']] or '').strip() if '機台型號' in idx else '',
                    'current':  row[idx['平均電流(A)']] if '平均電流(A)' in idx else 0,
                    'voltage':  row[idx['電壓(V)']] if '電壓(V)' in idx else 0,
                    'oil_cost': row[idx['每日耗油費用(元)']] if '每日耗油費用(元)' in idx else 0,
                    'hours':    row[idx['每日工作時數']] if '每日工作時數' in idx else 8,
                    'rate':     row[idx['每秒鐘生產費用(元)']] if '每秒鐘生產費用(元)' in idx else 0.2,
                })

        elec_price = 0
        if config.BATCH_COST_GLOBAL_SHEET in wb.sheetnames:
            gws = wb[config.BATCH_COST_GLOBAL_SHEET]
            gheader = [c.value for c in next(gws.iter_rows(min_row=1, max_row=1))]
            gidx = {h: i for i, h in enumerate(gheader) if h}
            if '每度電價' in gidx:
                grow = next(gws.iter_rows(min_row=2, max_row=2, values_only=True), None)
                if grow:
                    elec_price = grow[gidx['每度電價']] or 0
        wb.close()
        return jsonify({'success': True, 'equipment': equipment, 'elec_price': elec_price})
    except Exception as e:
        return jsonify({'success': False, 'error': f'讀取範本檔案失敗：{e}'}), 502


@app.route('/api/batch_cost/equipment/save', methods=['POST'])
def batch_cost_equipment_save():
    """新增或編輯『設備耗用設定』分頁裡的一台設備（有 row 就編輯該列，沒有就新增一列）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no   = data.get('row')
    code     = (data.get('code') or '').strip()
    name     = (data.get('name') or '').strip()
    model    = (data.get('model') or '').strip()
    current  = data.get('current') or 0
    voltage  = data.get('voltage') or 0
    oil_cost = data.get('oil_cost') or 0
    hours    = data.get('hours') or 8
    rate     = data.get('rate') if data.get('rate') not in (None, '') else 0.2
    if not code:
        return jsonify({'success': False, 'error': '請輸入機台代號'}), 400

    default_header = ['機台代號', '機台名稱', '機台型號', '平均電流(A)', '電壓(V)', '每日耗油費用(元)', '每日工作時數', '每秒鐘生產費用(元)']
    values = {
        '機台代號': code, '機台名稱': name, '機台型號': model, '平均電流(A)': current,
        '電壓(V)': voltage, '每日耗油費用(元)': oil_cost, '每日工作時數': hours,
        '每秒鐘生產費用(元)': rate,
    }
    try:
        wb = _load_batchcost_wb(read_only=False)
        if config.BATCH_COST_EQUIPMENT_SHEET not in wb.sheetnames:
            ws = wb.create_sheet(config.BATCH_COST_EQUIPMENT_SHEET)
            ws.append(default_header)
            header = default_header
        else:
            ws = wb[config.BATCH_COST_EQUIPMENT_SHEET]
            header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            if not any(header):
                header = default_header
                for col, h in enumerate(header, start=1):
                    ws.cell(row=1, column=col, value=h)

        # 自我修復：程式後來新增的欄位（例如「機台型號」）如果既有檔案的表頭裡還沒有，
        # 自動補到表頭最後一欄，不動既有欄位順序，也不影響使用者手動調整過的排列
        for col_name in values:
            if col_name not in header:
                header.append(col_name)
                ws.cell(row=1, column=len(header), value=col_name)

        target_row = int(row_no) if row_no else ws.max_row + 1
        for col_name, val in values.items():
            if col_name in header:
                ws.cell(row=target_row, column=header.index(col_name) + 1, value=val)

        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/api/batch_cost/equipment/delete', methods=['POST'])
def batch_cost_equipment_delete():
    """刪除『設備耗用設定』分頁裡的一列"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    row_no = data.get('row')
    if not row_no:
        return jsonify({'success': False, 'error': '缺少列號'}), 400
    try:
        wb = _load_batchcost_wb(read_only=False)
        ws = wb[config.BATCH_COST_EQUIPMENT_SHEET]
        ws.delete_rows(int(row_no), 1)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'刪除失敗：{e}'}), 502


@app.route('/api/batch_cost/global/save', methods=['POST'])
def batch_cost_global_save():
    """儲存『批成本全域設定』分頁的每度電價（單列，所有設備共用，不用逐台設定）"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
    data = request.get_json(force=True) or {}
    elec_price = data.get('elec_price')
    if elec_price is None:
        return jsonify({'success': False, 'error': '缺少電價'}), 400
    try:
        wb = _load_batchcost_wb(read_only=False)
        if config.BATCH_COST_GLOBAL_SHEET not in wb.sheetnames:
            ws = wb.create_sheet(config.BATCH_COST_GLOBAL_SHEET)
            ws.append(['每度電價'])
            ws.append([elec_price])
        else:
            ws = wb[config.BATCH_COST_GLOBAL_SHEET]
            header = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            if not any(header):
                header = ['每度電價']
                ws.cell(row=1, column=1, value='每度電價')
            idx = {h: i for i, h in enumerate(header) if h}
            col = idx.get('每度電價', 0) + 1
            ws.cell(row=2, column=col, value=elec_price)
        warning = _save_batchcost_wb(wb)
        return jsonify({'success': True, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 502


@app.route('/batch_cost')
def batch_cost_page():
    """批成本計算頁面"""
    return render_template('batch_cost.html', app_version=APP_VERSION)


# ══════════════════════════════════════════════════════════
#  設備管理（設備主檔／編碼／照片）  詳見 docs/equipment-master.md
# ══════════════════════════════════════════════════════════

def _eq_conn():
    """開啟設備主檔資料庫；索引未建立時回傳 None。

    EQ_DB_PATH 現在多半指向網芳（多人共用），timeout=15 讓遇到別人正在寫入時等待
    重試而不是立刻丟「database is locked」；刻意不開 WAL 模式——WAL 需要 shared memory
    mmap，Windows SMB 網芳不可靠支援，網芳上還是用預設的 rollback journal 比較安全。"""
    if not os.path.exists(EQ_DB_PATH):
        return None
    conn = sqlite3.connect(EQ_DB_PATH, timeout=15)
    conn.row_factory = sqlite3.Row
    _eq_ensure_model_column(conn)
    return conn


def _eq_ensure_model_column(conn):
    """加「型號」欄位（2026-08-02）。equipment.db 是執行期資料庫，早就有資料在跑，
    不能靠 CREATE TABLE IF NOT EXISTS 生出新欄位，要用 ALTER TABLE 補一次，
    PRAGMA table_info 很便宜，每次開連線檢查一次成本可忽略。"""
    cols = {r['name'] for r in conn.execute('PRAGMA table_info(equipment)')}
    if 'model' not in cols:
        conn.execute('ALTER TABLE equipment ADD COLUMN model TEXT')
        conn.commit()


def _eq_ensure_downtime(conn):
    """建立停機時數表（第一次用到時才建，不用另外做 schema 遷移）"""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS eq_downtime (
            rep_key    TEXT PRIMARY KEY,   -- 對應採購列的內容雜湊，見 _eq_repair_key
            code       TEXT,               -- 設備編碼（冗餘，供查詢與人工稽核）
            date       TEXT,               -- 維修日期（冗餘，同上）
            hours      REAL,               -- 停機時數
            pur_seq    TEXT,               -- 採購表序號（冗餘，失聯時的救援線索）
            updated_at TEXT
        )""")
    conn.execute('CREATE INDEX IF NOT EXISTS idx_eq_downtime_code ON eq_downtime(code)')
    conn.commit()


def _eq_repair_key(rec):
    """把一筆採購維修資料轉成穩定的識別碼。

    用「日期＋品名＋規格＋廠商＋小計」的內容雜湊，**不是**採購表的序號——
    序號雖然全表唯一，但只要有人在試算表中間插一列，底下所有序號都會往下推，
    停機時數就會整批對到別筆維修單（無聲的錯誤資料）。內容雜湊碰到插列不受影響；
    最壞情況是有人事後改了品名／金額導致該筆失聯（時數顯示不出來，看得出來、可補填），
    失敗模式比錯誤對應安全得多。2026-08-01 實測 795 筆維修候選資料無重複。
    """
    raw = '|'.join([(rec.get(k) or '') if isinstance(rec.get(k), str) else str(rec.get(k) or '')
                    for k in ('date', 'name', 'spec', 'vendor', 'subtotal')])
    return hashlib.md5(raw.encode('utf-8')).hexdigest()[:16]


def _eq_safe_path(root, relpath):
    """將相對路徑轉成絕對路徑，並確保仍在指定 root 底下（防止路徑跳脫）"""
    root = os.path.normpath(root)
    full = os.path.normpath(os.path.join(root, relpath or ''))
    if full.lower() != root.lower() and not full.lower().startswith(root.lower() + os.sep):
        return None
    return full


@app.route('/equipment_master')
def equipment_master_page():
    """設備管理頁面（設備主檔，與 /equipment 設備稼動是不同功能）"""
    return render_template('equipment_master.html', app_version=APP_VERSION)


@app.route('/api/equipment_master/search')
def equipment_master_search():
    """設備清單查詢（空格=AND、-前綴=NOT，比對編碼/舊編號/類型/屬性/廠商/備註/規格）"""
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立，請先執行匯入'}), 500
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT e.*,
                   g.name  AS group_name,
                   -- 編碼壞掉的設備（needs_fix）類型/屬性碼是硬湊的，查字典會查到別台的名稱，
                   -- 這時要顯示 Excel 原文，見 build_equipment_index.py 的 type_name_raw 註解
                   CASE WHEN e.needs_fix = 1 AND IFNULL(e.type_name_raw,'') <> ''
                        THEN e.type_name_raw ELSE IFNULL(t.name,'') END AS type_name,
                   CASE WHEN e.needs_fix = 1 AND IFNULL(e.attr_name_raw,'') <> ''
                        THEN e.attr_name_raw ELSE IFNULL(a.name,'') END AS attr_name,
                   (SELECT COUNT(*) FROM eq_photo p     WHERE p.code = e.code) AS photo_count,
                   (SELECT COUNT(*) FROM eq_tech_file f WHERE f.code = e.code) AS tech_count,
                   (SELECT p.relpath FROM eq_photo p
                     WHERE p.code = e.code ORDER BY p.is_cover DESC, p.relpath LIMIT 1) AS cover
              FROM equipment e
              LEFT JOIN eq_group g ON g.code = e.group_code
              LEFT JOIN eq_type  t ON t.group_code = e.group_code AND t.code = e.type_code
              LEFT JOIN eq_attr  a ON a.group_code = e.group_code AND a.type_code = e.type_code
                                  AND a.code = e.attr_code
             ORDER BY e.code""").fetchall()]
        specs = {}
        for r in conn.execute('SELECT code, spec_name, spec_value FROM eq_spec ORDER BY code, sort'):
            specs.setdefault(r['code'], []).append(f"{r['spec_name']}{r['spec_value']}")
    finally:
        conn.close()

    q = request.args.get('q', '').strip()
    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def haystack(r):
            return ' '.join([r.get(k) or '' for k in
                             ('code', 'old_code', 'group_name', 'type_name', 'attr_name',
                              'vendor', 'model', 'remark', 'note', 'location', 'status')]
                            + specs.get(r['code'], [])).lower()

        rows = [r for r in rows if all(m in haystack(r) for m in must)
                and not any(m in haystack(r) for m in must_not)]

    return jsonify({'success': True, 'count': len(rows), 'data': rows})


@app.route('/api/equipment_master/detail')
def equipment_master_detail():
    """單台設備詳情：規格、照片清單、技術資料檔案、異動歷程"""
    code = request.args.get('code', '').strip()
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute("""
            SELECT e.*, g.name AS group_name,
                   CASE WHEN e.needs_fix = 1 AND IFNULL(e.type_name_raw,'') <> ''
                        THEN e.type_name_raw ELSE IFNULL(t.name,'') END AS type_name,
                   CASE WHEN e.needs_fix = 1 AND IFNULL(e.attr_name_raw,'') <> ''
                        THEN e.attr_name_raw ELSE IFNULL(a.name,'') END AS attr_name
              FROM equipment e
              LEFT JOIN eq_group g ON g.code = e.group_code
              LEFT JOIN eq_type  t ON t.group_code = e.group_code AND t.code = e.type_code
              LEFT JOIN eq_attr  a ON a.group_code = e.group_code AND a.type_code = e.type_code
                                  AND a.code = e.attr_code
             WHERE e.code = ?""", (code,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': f'查無設備 {code}'}), 404
        detail = dict(row)
        detail['specs'] = [dict(r) for r in conn.execute(
            'SELECT spec_name, spec_value FROM eq_spec WHERE code=? ORDER BY sort', (code,))]
        detail['photos'] = [dict(r) for r in conn.execute(
            'SELECT relpath, filename, is_cover FROM eq_photo WHERE code=? '
            'ORDER BY is_cover DESC, relpath', (code,))]
        detail['tech_files'] = [dict(r) for r in conn.execute(
            'SELECT relpath, filename, ext, size, mtime FROM eq_tech_file WHERE code=? '
            'ORDER BY relpath', (code,))]
        # 歷程新的排前面（最近發生的事最常要看）；同一天的用 rowid 讓後寫入的排前面
        # eq_history 沒有另外定義 PK，直接用 sqlite 內建 rowid 當識別碼給編輯/刪除用即可
        detail['history'] = [dict(r) for r in conn.execute(
            'SELECT rowid AS id, date, action, detail, user FROM eq_history WHERE code=? '
            'ORDER BY date DESC, rowid DESC', (code,))]
    finally:
        conn.close()
    return jsonify({'success': True, 'data': detail})


@app.route('/api/equipment_master/stats')
def equipment_master_stats():
    """群組／保管位置／狀態的台數統計，供徽章篩選列顯示數字"""
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': True, 'groups': [], 'locations': [], 'statuses': []})
    try:
        groups = [dict(r) for r in conn.execute("""
            SELECT e.group_code AS code, COALESCE(g.name,'') AS name, COUNT(*) AS cnt
              FROM equipment e LEFT JOIN eq_group g ON g.code = e.group_code
             GROUP BY e.group_code ORDER BY e.group_code""")]
        locations = [dict(r) for r in conn.execute(
            'SELECT location AS name, COUNT(*) AS cnt FROM equipment '
            'GROUP BY location ORDER BY cnt DESC')]
        statuses = [dict(r) for r in conn.execute(
            'SELECT status AS name, COUNT(*) AS cnt FROM equipment '
            'GROUP BY status ORDER BY cnt DESC')]
        fix = conn.execute('SELECT COUNT(*) FROM equipment WHERE needs_fix=1').fetchone()[0]
        orphan = conn.execute('SELECT COUNT(*) FROM eq_photo WHERE code IS NULL').fetchone()[0]
    finally:
        conn.close()
    return jsonify({'success': True, 'groups': groups, 'locations': locations,
                    'statuses': statuses, 'needs_fix': fix, 'orphan_photos': orphan})


@app.route('/api/equipment_master/history/list')
def equipment_master_history_list():
    """跨設備瀏覽全部異動歷程（左側「異動歷程」子頁用），支援與其他清單一致的搜尋語法
    （空格=AND、-前綴=NOT，比對設備編碼/動作/內容/群組名稱）。動作/來源/群組的篩選
    下拉選單由前端拿到全部資料後自己算選項，不用另開一支 stats API。"""
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': True, 'count': 0, 'data': []})
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT h.rowid AS id, h.code, h.date, h.action, h.detail, h.user,
                   e.group_code, COALESCE(g.name,'') AS group_name
              FROM eq_history h
              LEFT JOIN equipment e ON e.code = h.code
              LEFT JOIN eq_group g  ON g.code = e.group_code
             ORDER BY h.date DESC, h.rowid DESC""")]
    finally:
        conn.close()

    q = request.args.get('q', '').strip()
    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def haystack(r):
            return ' '.join([r.get(k) or '' for k in ('code', 'action', 'detail', 'group_name')]).lower()

        rows = [r for r in rows if all(m in haystack(r) for m in must)
                and not any(m in haystack(r) for m in must_not)]

    return jsonify({'success': True, 'count': len(rows), 'data': rows})


_EQ_REPAIR_CODE_RE = re.compile(r'[A-Z]\d{2}-\d{3}')

# 採購登入表的欄位位置（用位置不用表頭名：這份表第 0 列是篩選控制列、表頭在第 1 列，
# 直接餵給 csv.DictReader 會把空白的第 0 列當表頭）
_PUR_COL = {'seq': 0, 'status': 1, 'date': 2, 'name': 3, 'spec': 4, 'vendor': 5,
            'vendor_no': 6, 'qty': 7, 'unit_price': 8, 'subtotal': 9,
            'account': 10, 'doc_no': 11, 'remark': 12}


def _fetch_purchase_rows_full():
    """抓採購登入表的**完整**資料。

    一定要用 export?format=csv，不能用 gviz/tq?tqx=out:csv——2026-07-29 實測發現
    gviz 匯出會**跟著試算表當下的篩選狀態走**：使用者在 Google Sheets 上開了篩選時
    只回傳篩選後的 615 列（維修 613 筆），沒開篩選時回傳 12,213 列（維修只認得 57 筆），
    同一支程式在不同時間抓到的資料量天差地遠。export 端點永遠回傳完整原始資料
    （16,958 列、維修 793 筆、95 家廠商），不受任何人的篩選操作影響。
    """
    url = (f'https://docs.google.com/spreadsheets/d/{config.PURCHASE_SHEET_ID}'
           f'/export?format=csv&gid=0')
    resp = requests.get(url, timeout=120)
    resp.raise_for_status()
    resp.encoding = 'utf-8'
    rows = list(csv.reader(io.StringIO(resp.text)))

    # 表頭列位置不寫死：找第一列同時含「日期」與「品名」的當表頭，資料從下一列開始
    start = 2
    for i, row in enumerate(rows[:10]):
        if any('日期' in c for c in row) and any('品名' in c for c in row):
            start = i + 1
            break

    def _num(row, key):
        try:
            return float((row[_PUR_COL[key]] or '').replace(',', '').strip())
        except (ValueError, IndexError):
            return 0

    def _cell(row, key):
        i = _PUR_COL[key]
        return (row[i] or '').strip() if i < len(row) else ''

    out = []
    for row in rows[start:]:
        if len(row) <= _PUR_COL['account']:
            continue
        if not _cell(row, 'date') and not _cell(row, 'name'):
            continue                      # 跳過空白列
        out.append({
            'date': _cell(row, 'date'), 'status': _cell(row, 'status'),
            'name': _cell(row, 'name'), 'spec': _cell(row, 'spec'),
            'vendor': _cell(row, 'vendor'), 'account': _cell(row, 'account'),
            'remark': _cell(row, 'remark'), 'pur_seq': _cell(row, 'seq'),
            'qty': _num(row, 'qty'), 'unit_price': _num(row, 'unit_price'),
            'subtotal': _num(row, 'subtotal'),
        })
    return out


def _eq_repair_is_maintenance(rec):
    """判定一筆採購資料算不算維修：使用者自己標的「維修」分類，或會計科目歸在修繕保養。
    這兩個欄位都是人工填的，比用關鍵字猜品名可靠得多。"""
    return (rec.get('status') or '').strip() == '維修' \
        or '修繕' in (rec.get('account') or '') or '保養' in (rec.get('account') or '')


def _eq_repair_match(rec, codes, old_map):
    """把一筆採購資料對應到設備編碼，對應不到回 None。

    只認「設備主檔真的有的編碼」——採購表裡刀具料號長得很像設備編碼
    （SVF-A**B03-100**單斜固定鉗口、BT40-100-001 → T40-100），直接用正則抓會誤判，
    所以比對出來還要確認存在於 equipment 表才算數。
    """
    text = ' '.join([rec.get('spec') or '', rec.get('name') or '', rec.get('remark') or ''])
    for m in _EQ_REPAIR_CODE_RE.finditer(text):
        if m.group() in codes:
            return m.group()
    # 沒寫編碼時退而比對舊編號（例如「詠設.綜銑15」）。長的先比，避免「車1」誤命中「車15」；
    # 同一個舊編號被多台共用時（例如砂輪機都叫 B10）無法判斷是哪一台，寧可不猜。
    for old in sorted(old_map, key=len, reverse=True):
        if old and old in text and len(old_map[old]) == 1:
            return old_map[old][0]
    return None


@app.route('/api/equipment_master/repair/list')
def equipment_master_repair_list():
    """設備維修記錄：從採購登入表挑出維修類支出，比對到設備編碼後回傳。

    資料來源是既有的 /api/purchase/list 那份 Google 試算表（PURCHASE_SHEET_ID），
    這裡沿用同一套 fetch + 快取，不另外連線。只回傳「對應得到設備」的記錄
    （共用件、飲水機濾心之類歸不到某一台的不列入，見 docs/equipment-master.md）。
    """
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': True, 'count': 0, 'data': []})
    try:
        # 已售出／報廢的設備不列出來：機器已經不是我們的，那些維修是死歷史，
        # 而且清單上每一列都可以填停機時數，填了卻永遠不會被妥善率採計（妥善率只算使用中），
        # 留著只會讓人白填。對應不到設備的列本來就會被濾掉，所以從 codes 拿掉就等於排除
        codes = {r['code'] for r in conn.execute(
            "SELECT code FROM equipment WHERE status NOT IN ('已售出','報廢')")}
        old_map = {}
        for r in conn.execute("SELECT code, old_code, status FROM equipment WHERE old_code<>''"):
            if r['status'] in ('已售出', '報廢'):
                continue
            old_map.setdefault(r['old_code'].strip(), []).append(r['code'])
        names = {r['code']: r for r in conn.execute("""
            SELECT e.code, IFNULL(t.name,'') AS type_name, e.group_code
              FROM equipment e
              LEFT JOIN eq_type t ON t.group_code = e.group_code AND t.code = e.type_code""")}
    finally:
        conn.close()

    cache_key = 'equipment_repair_raw'
    if request.args.get('refresh'):
        cache_clear(cache_key)
    cached = cache_get(cache_key)
    if cached is not None:
        purchases = cached['purchases']
    else:
        try:
            purchases = _fetch_purchase_rows_full()
        except Exception:
            return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502
        cache_set(cache_key, {'purchases': purchases})

    # 已登記的停機時數：key 是採購列的內容雜湊，跟著每一列一起回傳給前端
    conn = _eq_conn()
    downtime = {}
    if conn is not None:
        try:
            _eq_ensure_downtime(conn)
            downtime = {r['rep_key']: r['hours'] for r in conn.execute(
                'SELECT rep_key, hours FROM eq_downtime')}
        finally:
            conn.close()

    out = []
    for rec in purchases:
        if not _eq_repair_is_maintenance(rec):
            continue
        code = _eq_repair_match(rec, codes, old_map)
        if not code:
            continue
        info = names.get(code) or {}
        rep_key = _eq_repair_key(rec)
        out.append({**rec, 'code': code, 'rep_key': rep_key,
                    'downtime': downtime.get(rep_key),
                    'group_code': info['group_code'] if info else '',
                    'type_name': info['type_name'] if info else ''})

    q = request.args.get('q', '').strip()
    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def haystack(r):
            return ' '.join([r.get(k) or '' for k in
                             ('code', 'name', 'spec', 'vendor', 'account', 'remark', 'type_name')]).lower()

        out = [r for r in out if all(m in haystack(r) for m in must)
               and not any(m in haystack(r) for m in must_not)]

    out.sort(key=lambda r: _purchase_date_key(r['date']) or '', reverse=True)
    return jsonify({'success': True, 'count': len(out),
                    'total_cost': sum(r['subtotal'] for r in out), 'data': out})


@app.route('/api/equipment_master/downtime/save', methods=['POST'])
def equipment_master_downtime_save():
    """登記某一筆維修的停機時數（時數留空或 0 就把該筆刪掉）"""
    d = request.get_json(force=True) or {}
    rep_key = (d.get('rep_key') or '').strip()
    if not rep_key:
        return jsonify({'success': False, 'error': '缺少維修單識別碼'}), 400

    raw = str(d.get('hours', '')).strip()
    if raw == '':
        hours = 0.0
    else:
        try:
            hours = float(raw)
        except ValueError:
            return jsonify({'success': False, 'error': '停機時數必須是數字'}), 400
        if hours < 0:
            return jsonify({'success': False, 'error': '停機時數不能是負數'}), 400

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        _eq_ensure_downtime(conn)
        if hours <= 0:
            conn.execute('DELETE FROM eq_downtime WHERE rep_key=?', (rep_key,))
        else:
            conn.execute(
                'INSERT INTO eq_downtime (rep_key, code, date, hours, pur_seq, updated_at) '
                'VALUES (?,?,?,?,?,?) ON CONFLICT(rep_key) DO UPDATE SET '
                'hours=excluded.hours, code=excluded.code, date=excluded.date, '
                'pur_seq=excluded.pur_seq, updated_at=excluded.updated_at',
                (rep_key, (d.get('code') or '').strip(), (d.get('date') or '').strip(),
                 hours, (d.get('pur_seq') or '').strip(),
                 datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'hours': hours})


# ══════════════════════════════════════════════════════════════════════
# 工作日行事曆（管理頁「行事曆」子頁，見 docs/calendar.md）
#
# cal_day 只存「跟預設規則不一樣的日子」，不是每天都存一筆：
#   預設規則 = 週一~週五上班、週六日放假
#   例外     = 平日放假（國定假日/尾牙/颱風假）或 假日上班（補班）
# 這樣不用先建一整年的資料就能直接用，行事曆空的時候行為跟以前完全一樣。
# ══════════════════════════════════════════════════════════════════════

def _cal_conn():
    """開啟行事曆資料庫（不存在就建一個空的——這是使用者輸入的資料，不由任何索引工具產生）"""
    conn = sqlite3.connect(CAL_DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("""
        CREATE TABLE IF NOT EXISTS cal_day (
            date       TEXT PRIMARY KEY,   -- YYYY-MM-DD
            kind       TEXT NOT NULL,      -- 'work'=上班, 'off'=放假
            note       TEXT,               -- 假日名稱，例如 中秋節／颱風假／補班
            updated_at TEXT
        )""")
    conn.commit()
    return conn


def _cal_default_kind(d):
    """預設規則：週一~週五上班，週六日放假"""
    return 'work' if d.weekday() < 5 else 'off'


def _cal_overrides(start, end):
    """取出區間內所有例外日，回傳 {'YYYY-MM-DD': sqlite3.Row}"""
    conn = _cal_conn()
    try:
        return {r['date']: r for r in conn.execute(
            'SELECT date, kind, note FROM cal_day WHERE date BETWEEN ? AND ?',
            (start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')))}
    finally:
        conn.close()


def _cal_workdays(start, end):
    """區間內（含頭尾）的工作天數：先看行事曆例外，沒設定的用預設規則"""
    if start > end:
        return 0
    ov = _cal_overrides(start, end)
    n = 0
    for i in range((end - start).days + 1):
        d = start + timedelta(days=i)
        row = ov.get(d.strftime('%Y-%m-%d'))
        if (row['kind'] if row else _cal_default_kind(d)) == 'work':
            n += 1
    return n


@app.route('/api/calendar/month')
def calendar_month():
    """某個月的每一天：日期、星期、上班/放假、備註、是不是例外日"""
    ym = (request.args.get('ym') or '').strip()
    try:
        year, month = int(ym[:4]), int(ym[5:7])
        first = date(year, month, 1)
    except (ValueError, IndexError):
        today = date.today()
        first = today.replace(day=1)
    last = (first.replace(day=28) + timedelta(days=4)).replace(day=1) - timedelta(days=1)

    ov = _cal_overrides(first, last)
    days = []
    for i in range((last - first).days + 1):
        d = first + timedelta(days=i)
        key = d.strftime('%Y-%m-%d')
        row = ov.get(key)
        default = _cal_default_kind(d)
        days.append({'date': key, 'day': d.day, 'dow': (d.weekday() + 1) % 7,   # 0=週日
                     'kind': row['kind'] if row else default,
                     'note': (row['note'] if row else '') or '',
                     'is_override': row is not None, 'default_kind': default})
    return jsonify({'success': True, 'ym': first.strftime('%Y-%m'),
                    'first_dow': (first.weekday() + 1) % 7, 'days': days,
                    'workdays': sum(1 for x in days if x['kind'] == 'work'),
                    'offdays': sum(1 for x in days if x['kind'] == 'off')})


@app.route('/api/calendar/set', methods=['POST'])
def calendar_set():
    """設定某一天上班/放假。跟預設規則相同且沒有備註時直接刪掉該筆，表裡只留真正的例外"""
    d = request.get_json(force=True) or {}
    day = (d.get('date') or '').strip()
    kind = (d.get('kind') or '').strip()
    note = (d.get('note') or '').strip()
    try:
        dt = datetime.strptime(day, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'success': False, 'error': '日期格式須為 YYYY-MM-DD'}), 400
    if kind not in ('work', 'off'):
        return jsonify({'success': False, 'error': "kind 只能是 'work' 或 'off'"}), 400

    conn = _cal_conn()
    try:
        if kind == _cal_default_kind(dt) and not note:
            conn.execute('DELETE FROM cal_day WHERE date=?', (day,))
        else:
            conn.execute(
                'INSERT INTO cal_day (date, kind, note, updated_at) VALUES (?,?,?,?) '
                'ON CONFLICT(date) DO UPDATE SET kind=excluded.kind, note=excluded.note, '
                'updated_at=excluded.updated_at',
                (day, kind, note, datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'date': day, 'kind': kind, 'note': note,
                    'is_override': not (kind == _cal_default_kind(dt) and not note)})


@app.route('/api/calendar/reset_month', methods=['POST'])
def calendar_reset_month():
    """把某個月的所有例外日清掉，整個月回到預設的週一~週五"""
    ym = ((request.get_json(force=True) or {}).get('ym') or '').strip()
    if not re.match(r'^\d{4}-\d{2}$', ym):
        return jsonify({'success': False, 'error': '月份格式須為 YYYY-MM'}), 400
    conn = _cal_conn()
    try:
        n = conn.execute('DELETE FROM cal_day WHERE date LIKE ?', (ym + '-%',)).rowcount
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'removed': n})


@app.route('/api/calendar/workdays')
def calendar_workdays():
    """任意區間的工作天數，給其他功能（妥善率…）呼叫"""
    def _parse(s, fallback):
        try:
            return datetime.strptime(s, '%Y-%m-%d').date()
        except (ValueError, TypeError):
            return fallback
    today = date.today()
    start = _parse(request.args.get('start'), today.replace(day=1))
    end = _parse(request.args.get('end'), today)
    return jsonify({'success': True, 'start': start.strftime('%Y-%m-%d'),
                    'end': end.strftime('%Y-%m-%d'), 'workdays': _cal_workdays(start, end)})


def _eq_workdays(start, end):
    """計算 start~end（含頭尾）之間的工作天數。

    以管理頁「行事曆」的設定為準（國定假日、颱風假、補班都可在那裡調），
    沒設定的日子退回預設規則：週一~週五算工作日、週六日不算。見 _cal_workdays。
    """
    return _cal_workdays(start, end)


@app.route('/api/equipment_master/availability')
def equipment_master_availability():
    """設備妥善率：(應有稼動時數 - 停機時數) / 應有稼動時數

    應有稼動時數 = 使用中台數 × 每日工時 × 平日天數（週一~週五）
    停機時數     = 期間內該群組設備已登記的停機時數合計（eq_downtime）

    `location` 保管位置是分子分母**一起**套用的過濾條件（預設只看加工課自己的設備）：
    詠設的機器不算我們的停機，它們的產能自然也不該留在分母裡灌高妥善率。

    ※ 不要改用採購表「規格/說明」欄有沒有寫「詠設」來判斷——那欄是人工打的，
      2026-08-01 實測 A 群組 355 筆維修裡，位置屬詠設的 90 筆只有 39 筆有寫「詠設」
      （漏 51 筆），反而有 3 筆加工課的機器被誤標成詠設，還有 6 台機器同一台兩種標法
      都出現過。設備主檔的 location 欄才是準的。
    """
    group = (request.args.get('group') or '').strip()
    location = (request.args.get('location') or '').strip()
    try:
        hours_per_day = float(request.args.get('hours_per_day') or 8)
    except ValueError:
        hours_per_day = 8.0

    today = date.today()
    def _parse(s, fallback):
        try:
            return datetime.strptime(s, '%Y-%m-%d').date()
        except (ValueError, TypeError):
            return fallback
    start = _parse(request.args.get('start'), today.replace(day=1))
    end = _parse(request.args.get('end'), today)

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        _eq_ensure_downtime(conn)
        # 分母只算「使用中」設備——已售出／報廢的機器本來就不該列入應有稼動
        conds, args = ["e.status='使用中'"], []
        if group:
            conds.append('e.group_code=?')
            args.append(group)
        if location:
            conds.append('e.location=?')
            args.append(location)
        # 一併撈舊編號與機台類型，停機明細要顯示（使用者認機器是靠「綜銑07」這種舊稱呼）
        machines = [dict(r) for r in conn.execute(
            f"""SELECT e.code, e.group_code, e.old_code,
                       CASE WHEN e.needs_fix = 1 AND IFNULL(e.type_name_raw,'') <> ''
                            THEN e.type_name_raw ELSE IFNULL(t.name,'') END AS type_name
                  FROM equipment e
                  LEFT JOIN eq_type t ON t.group_code = e.group_code AND t.code = e.type_code
                 WHERE {' AND '.join(conds)}
                 ORDER BY e.code""", args)]
        rows = [dict(r) for r in conn.execute(
            'SELECT d.rep_key, d.code, d.date, d.hours, e.group_code '
            '  FROM eq_downtime d JOIN equipment e ON e.code = d.code '
            ' WHERE d.hours > 0')]
    finally:
        conn.close()

    codes = {m['code'] for m in machines}
    per_machine = {}
    total_down = 0.0
    for r in rows:
        if r['code'] not in codes:
            continue                       # 不在本次統計範圍（別的群組或非使用中）
        # _purchase_date_key 回傳 'YYYY-MM-DD'（有連字號），比字串時兩邊格式必須一致，
        # 寫成 '%Y%m%d' 會因為 '0' > '-' 而永遠比不中，停機時數全部被濾掉
        key = _purchase_date_key(r['date'])
        if not key or not (start.strftime('%Y-%m-%d') <= key <= end.strftime('%Y-%m-%d')):
            continue
        total_down += r['hours'] or 0
        per_machine[r['code']] = per_machine.get(r['code'], 0) + (r['hours'] or 0)

    workdays = _eq_workdays(start, end)
    planned = len(machines) * hours_per_day * workdays
    rate = ((planned - total_down) / planned * 100) if planned > 0 else None

    info = {m['code']: m for m in machines}
    detail = sorted(({'code': c, 'hours': h,
                      'old_code': (info.get(c) or {}).get('old_code') or '',
                      'type_name': (info.get(c) or {}).get('type_name') or '',
                      'rate': (1 - h / (hours_per_day * workdays)) * 100 if workdays else None}
                     for c, h in per_machine.items()),
                    key=lambda x: x['hours'], reverse=True)

    return jsonify({'success': True, 'group': group, 'location': location,
                    'start': start.strftime('%Y-%m-%d'), 'end': end.strftime('%Y-%m-%d'),
                    'machines': len(machines), 'workdays': workdays,
                    'hours_per_day': hours_per_day, 'planned_hours': planned,
                    'downtime_hours': total_down, 'rate': rate,
                    'affected': len(per_machine), 'detail': detail})


@app.route('/api/equipment_master/dict')
def equipment_master_dict():
    """編碼字典：群組／機台類型／加工屬性三層，供新增設備的連動下拉選單使用"""
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': True, 'groups': [], 'types': [], 'attrs': []})
    try:
        groups = [dict(r) for r in conn.execute(
            'SELECT code, name FROM eq_group ORDER BY code')]
        types = [dict(r) for r in conn.execute(
            'SELECT group_code, code, name FROM eq_type ORDER BY group_code, code')]
        attrs = [dict(r) for r in conn.execute(
            'SELECT group_code, type_code, code, name FROM eq_attr '
            'ORDER BY group_code, type_code, code')]
    finally:
        conn.close()
    return jsonify({'success': True, 'groups': groups, 'types': types, 'attrs': attrs})


def _next_num(conn, sql, params, width):
    """回傳指定範圍內「還沒被用掉的最小號碼」，補零到 width 位。
    報廢設備的號碼不放回重用（sql 查的是全部設備，不篩狀態），避免新舊設備同號。"""
    used = {r[0] for r in conn.execute(sql, params) if r[0] and str(r[0]).isdigit()}
    used = {int(x) for x in used}
    n = 1
    while n in used:
        n += 1
    return str(n).zfill(width)


@app.route('/api/equipment_master/next_code')
def equipment_master_next_code():
    """自動配號：回傳該群組下一個可用的機台類型碼、該類型下一個屬性碼、該屬性下一個流水號"""
    g = request.args.get('group_code', '').strip().upper()
    t = request.args.get('type_code', '').strip()
    a = request.args.get('attr_code', '').strip()
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        out = {'success': True}
        if g:
            # 類型碼要同時避開字典裡與設備上已用的號碼（字典可能有還沒設備的類型）
            out['next_type_code'] = _next_num(
                conn, 'SELECT code FROM eq_type WHERE group_code=? '
                      'UNION SELECT type_code FROM equipment WHERE group_code=?', (g, g), 2)
        if g and t:
            out['next_attr_code'] = _next_num(
                conn, 'SELECT code FROM eq_attr WHERE group_code=? AND type_code=? '
                      'UNION SELECT attr_code FROM equipment WHERE group_code=? AND type_code=?',
                (g, t, g, t), 1)
        if g and t and a:
            out['next_seq'] = _next_num(
                conn, 'SELECT seq FROM equipment WHERE group_code=? AND type_code=? AND attr_code=?',
                (g, t, a), 2)
            out['code'] = f'{g}{t}-{a}{out["next_seq"]}'
            out['siblings'] = conn.execute(
                'SELECT COUNT(*) FROM equipment WHERE group_code=? AND type_code=? AND attr_code=?',
                (g, t, a)).fetchone()[0]
    finally:
        conn.close()
    return jsonify(out)


_EQ_CODE_RE = re.compile(r'^[A-Z]\d{2}-\d{3}$')

# 編輯設備時要自動留下歷程的欄位：欄位名 → (顯示名稱, 歷程動作)
# 只挑「事後會想知道什麼時候改的」欄位；純打字修正（例如備註）不記，避免歷程被洗版
_EQ_TRACKED_FIELDS = {
    'location': ('保管位置', '移轉'),
    'status':   ('狀態',     '狀態變更'),
    'vendor':   ('廠商',     '資料修改'),
    'model':    ('型號',     '資料修改'),
    'old_code': ('舊設備編號', '資料修改'),
    'buy_date': ('採購時間', '資料修改'),
}


def _eq_log_changes(cur, code, before, after):
    """比對編輯前後，把有變動的追蹤欄位寫成異動歷程。回傳寫入筆數。"""
    n = 0
    for field, (label, action) in _EQ_TRACKED_FIELDS.items():
        old = (before[field] if before[field] is not None else '')
        new = after.get(field) or ''
        if str(old).strip() == str(new).strip():
            continue
        detail = f'{label}：{old or "（空白）"} → {new or "（空白）"}'
        cur.execute("INSERT INTO eq_history (code, date, action, detail, user) "
                    "VALUES (?, date('now','localtime'), ?, ?, 'system')", (code, action, detail))
        n += 1
    return n


@app.route('/api/equipment_master/save', methods=['POST'])
def equipment_master_save():
    """新增或編輯設備。

    重要：任何在系統內改過的設備都會把 source 標成 'manual'，
    之後重新匯入 Excel 時不會被蓋回去（見 docs/equipment-master.md 第二節）。
    """
    d = request.get_json(silent=True) or {}
    is_new = bool(d.get('is_new'))
    orig_code = (d.get('orig_code') or '').strip().upper()
    g = (d.get('group_code') or '').strip().upper()
    t = (d.get('type_code') or '').strip()
    a = (d.get('attr_code') or '').strip()
    seq = (d.get('seq') or '').strip()

    if not (g and t and a and seq):
        return jsonify({'success': False, 'error': '群組／機台類型／加工屬性／流水號都必須填寫'}), 400
    if not (len(g) == 1 and g.isalpha()):
        return jsonify({'success': False, 'error': '群組分類必須是 1 個英文字母'}), 400
    if not (len(t) == 2 and t.isdigit() and len(a) == 1 and a.isdigit() and len(seq) == 2 and seq.isdigit()):
        return jsonify({'success': False, 'error': '機台類型須 2 位數字、加工屬性 1 位、流水號 2 位'}), 400

    code = f'{g}{t}-{a}{seq}'
    if not _EQ_CODE_RE.match(code):
        return jsonify({'success': False, 'error': f'產生的編碼「{code}」格式不正確'}), 400

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        cur = conn.cursor()
        exists = cur.execute('SELECT code FROM equipment WHERE code=?', (code,)).fetchone()
        if is_new and exists:
            return jsonify({'success': False, 'error': f'編碼 {code} 已存在，請重新配號'}), 400
        if not is_new:
            if not orig_code or not cur.execute(
                    'SELECT code FROM equipment WHERE code=?', (orig_code,)).fetchone():
                return jsonify({'success': False, 'error': f'找不到要編輯的設備 {orig_code}'}), 404
            if code != orig_code and exists:
                return jsonify({'success': False, 'error': f'編碼 {code} 已被其他設備使用'}), 400

        # 字典：使用者新填的類型／屬性名稱一併登記，之後其他設備可以直接選
        if d.get('type_name'):
            cur.execute('INSERT INTO eq_type (group_code, code, name, sort) VALUES (?,?,?,?) '
                        'ON CONFLICT(group_code, code) DO UPDATE SET name=excluded.name',
                        (g, t, d['type_name'].strip(), int(t)))
        if d.get('attr_name'):
            cur.execute('INSERT INTO eq_attr (group_code, type_code, code, name) VALUES (?,?,?,?) '
                        'ON CONFLICT(group_code, type_code, code) DO UPDATE SET name=excluded.name',
                        (g, t, a, d['attr_name'].strip()))
        if not cur.execute('SELECT code FROM eq_group WHERE code=?', (g,)).fetchone():
            cur.execute('INSERT INTO eq_group (code, name, sort) VALUES (?,?,?)',
                        (g, (d.get('group_name') or '').strip(), ord(g)))

        fields = (g, t, a, seq,
                  (d.get('old_code') or '').strip(), (d.get('vendor') or '').strip(),
                  (d.get('model') or '').strip(),
                  (d.get('buy_date') or '').strip(), (d.get('remark') or '').strip(),
                  (d.get('note') or '').strip(), (d.get('location') or '').strip(),
                  (d.get('status') or '使用中').strip())

        if is_new:
            cur.execute(
                'INSERT INTO equipment (code, group_code, type_code, attr_code, seq, old_code, '
                'vendor, model, buy_date, remark, note, location, status, source, origin, '
                'needs_fix, fix_reason) '
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,'manual','manual',0,'')", (code,) + fields)
            cur.execute("INSERT INTO eq_history (code, date, action, detail, user) "
                        "VALUES (?, date('now','localtime'), '新增', ?, 'system')",
                        (code, f'於系統內新增設備 {code}'))
        else:
            # 改動前先把原值撈起來，存檔後才比對得出哪些欄位變了（要在 UPDATE 之前）
            before = cur.execute(
                'SELECT location, status, vendor, model, old_code, buy_date FROM equipment WHERE code=?',
                (orig_code,)).fetchone()
            if code != orig_code:
                # 重新編碼：主鍵換掉，照片/技資/規格/歷程的對應要一起搬，否則會變孤兒資料
                for tbl in ('eq_spec', 'eq_photo', 'eq_tech_file', 'eq_history'):
                    cur.execute(f'UPDATE {tbl} SET code=? WHERE code=?', (code, orig_code))
                cur.execute('UPDATE equipment SET code=? WHERE code=?', (code, orig_code))
                cur.execute("INSERT INTO eq_history (code, date, action, detail, user) "
                            "VALUES (?, date('now','localtime'), '重新編碼', ?, 'system')",
                            (code, f'{orig_code} → {code}'))
            _eq_log_changes(cur, code, before, {
                'location': fields[10], 'status': fields[11], 'vendor': fields[5],
                'model': fields[6], 'old_code': fields[4], 'buy_date': fields[7]})
            cur.execute(
                'UPDATE equipment SET group_code=?, type_code=?, attr_code=?, seq=?, old_code=?, '
                'vendor=?, model=?, buy_date=?, remark=?, note=?, location=?, status=?, '
                "source='manual', needs_fix=0, fix_reason='', "
                "updated_at=datetime('now','localtime') WHERE code=?", fields + (code,))

        # 規格整組覆寫（前端每次都送完整清單）
        cur.execute('DELETE FROM eq_spec WHERE code=?', (code,))
        for i, s in enumerate(d.get('specs') or []):
            name = (s.get('spec_name') or '').strip()
            value = (s.get('spec_value') or '').strip()
            if name or value:
                cur.execute('INSERT INTO eq_spec (code, spec_name, spec_value, sort) '
                            'VALUES (?,?,?,?)', (code, name, value, i))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': f'儲存失敗：{e}'}), 500
    finally:
        conn.close()

    return jsonify({'success': True, 'code': code,
                    'message': f'已{"新增" if is_new else "更新"}設備 {code}'})


# 可以在清單表格裡直接改的欄位（白名單，不接受任意欄位名；仿油品清單「快速編輯」的做法）
_EQ_INLINE_FIELDS = {'vendor': '廠商', 'model': '型號', 'buy_date': '採購時間'}


@app.route('/api/equipment_master/inline_save', methods=['POST'])
def equipment_master_inline_save():
    """清單表格內直接編輯單一欄位（廠商／型號／採購日）。密碼在前端擋
    （跟編碼鎖定、油品清單快速編輯同一組 maxclaw），這裡只做欄位白名單與寫入。

    這三個欄位都在 _EQ_TRACKED_FIELDS 裡，所以沿用跟整份編輯表單同一套 _eq_log_changes
    寫進異動歷程——但那個函式是拿完整的 after dict 跟 before 逐欄比對，只丟
    {field: value} 進去的話，其餘追蹤欄位會被當成「改成空白」誤記一筆，所以要把
    before 的其他欄位原封不動地一起帶進去，只換掉正在編輯的那一個。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip().upper()
    field = (d.get('field') or '').strip()
    value = (d.get('value') or '').strip()
    if field not in _EQ_INLINE_FIELDS:
        return jsonify({'success': False, 'error': f'不支援編輯欄位 {field}'}), 400
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        cur = conn.cursor()
        before = cur.execute(
            'SELECT location, status, vendor, model, old_code, buy_date FROM equipment WHERE code=?',
            (code,)).fetchone()
        if before is None:
            return jsonify({'success': False, 'error': f'查無設備 {code}'}), 404
        after = {k: before[k] for k in before.keys()}
        after[field] = value
        _eq_log_changes(cur, code, before, after)
        cur.execute(f"UPDATE equipment SET {field}=?, source='manual', "
                    "updated_at=datetime('now','localtime') WHERE code=?", (value, code))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'code': code, 'field': field, 'value': value})


_EQ_LOCATIONS = {'加工', '詠設', '成品', '生技'}
_EQ_STATUSES = {'使用中', '閒置', '轉賣', '已售出', '報廢'}
# 這幾種動作代表設備當下狀態真的變了，補記錄時要連動更新設備本身的欄位，
# 不能只是寫一筆純文字歷程了事，否則詳情頁上方的「保管位置／狀態」會跟歷程對不上
_EQ_HIST_LOCATION_ACTIONS = {'移轉'}
_EQ_HIST_STATUS_ACTIONS = {'轉賣', '報廢', '狀態變更'}


@app.route('/api/equipment_master/history/add', methods=['POST'])
def equipment_master_history_add():
    """人工補一筆異動歷程（保養、維修、移轉、轉賣、報廢…等）。

    「移轉」會連動更新設備的保管位置，「轉賣／報廢／狀態變更」會連動更新設備狀態——
    這樣詳情頁上方顯示的目前狀態，才會跟歷程記錄一致，不用使用者自己再去編輯表單改一次。
    """
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip().upper()
    detail = (d.get('detail') or '').strip()
    action = (d.get('action') or '其他').strip()
    date = (d.get('date') or '').strip()
    new_location = (d.get('new_location') or '').strip()
    new_status = (d.get('new_status') or '').strip()

    if date and not re.match(r'^\d{4}-\d{2}-\d{2}$', date):
        return jsonify({'success': False, 'error': '日期格式須為 YYYY-MM-DD'}), 400
    if action in _EQ_HIST_LOCATION_ACTIONS:
        if new_location not in _EQ_LOCATIONS:
            return jsonify({'success': False, 'error': '請選擇移轉後的保管位置'}), 400
    elif action in _EQ_HIST_STATUS_ACTIONS:
        if new_status not in _EQ_STATUSES:
            return jsonify({'success': False, 'error': '請選擇異動後的狀態'}), 400
    elif not detail:
        return jsonify({'success': False, 'error': '請填寫內容'}), 400

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute('SELECT location, status FROM equipment WHERE code=?', (code,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': f'查無設備 {code}'}), 404

        auto_parts = []
        if action in _EQ_HIST_LOCATION_ACTIONS and new_location != row['location']:
            conn.execute("UPDATE equipment SET location=?, source='manual', "
                        "updated_at=datetime('now','localtime') WHERE code=?", (new_location, code))
            auto_parts.append(f"保管位置：{row['location'] or '（空白）'} → {new_location}")
        if action in _EQ_HIST_STATUS_ACTIONS and new_status != row['status']:
            conn.execute("UPDATE equipment SET status=?, source='manual', "
                        "updated_at=datetime('now','localtime') WHERE code=?", (new_status, code))
            auto_parts.append(f"狀態：{row['status'] or '（空白）'} → {new_status}")
        full_detail = '；'.join(auto_parts + ([detail] if detail else [])) or action

        if date:
            conn.execute('INSERT INTO eq_history (code, date, action, detail, user) '
                         "VALUES (?,?,?,?,'user')", (code, date, action, full_detail))
        else:
            conn.execute("INSERT INTO eq_history (code, date, action, detail, user) "
                         "VALUES (?, date('now','localtime'), ?, ?, 'user')", (code, action, full_detail))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已新增歷程記錄'})


@app.route('/api/equipment_master/history/edit', methods=['POST'])
def equipment_master_history_edit():
    """編輯一筆既有的歷程記錄（只改文字內容，不重新連動設備目前的保管位置／狀態——
    補改舊記錄不該把設備現在的狀態改回過去的值，見前端 startEditHistory 的註解）"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip().upper()
    hist_id = d.get('id')
    date = (d.get('date') or '').strip()
    action = (d.get('action') or '其他').strip()
    detail = (d.get('detail') or '').strip()

    if not hist_id:
        return jsonify({'success': False, 'error': '缺少記錄編號'}), 400
    if not date or not re.match(r'^\d{4}-\d{2}-\d{2}$', date):
        return jsonify({'success': False, 'error': '請選擇日期'}), 400
    if not detail:
        return jsonify({'success': False, 'error': '請填寫內容'}), 400

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute('SELECT rowid FROM eq_history WHERE rowid=? AND code=?',
                           (hist_id, code)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': '查無這筆歷程記錄'}), 404
        conn.execute('UPDATE eq_history SET date=?, action=?, detail=? WHERE rowid=?',
                     (date, action, detail, hist_id))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已更新歷程記錄'})


@app.route('/api/equipment_master/history/delete', methods=['POST'])
def equipment_master_history_delete():
    """刪除一筆歷程記錄。只能刪人工補登的（user='user'）——系統自動記錄與 Excel
    帶入的異動是稽核軌跡，不開放刪除，避免「狀態變更／移轉」的自動記錄被清空。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip().upper()
    hist_id = d.get('id')
    if not hist_id:
        return jsonify({'success': False, 'error': '缺少記錄編號'}), 400

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute('SELECT user FROM eq_history WHERE rowid=? AND code=?',
                           (hist_id, code)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': '查無這筆歷程記錄'}), 404
        if row['user'] != 'user':
            return jsonify({'success': False, 'error': '系統自動產生的記錄不能刪除'}), 400
        conn.execute('DELETE FROM eq_history WHERE rowid=? AND code=?', (hist_id, code))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已刪除歷程記錄'})


@app.route('/api/equipment_master/delete', methods=['POST'])
def equipment_master_delete():
    """設備下架。預設是軟刪除（狀態改報廢，編碼保留不放回重用）；
    只有系統內新增（source=manual）且明確要求時才真的刪除記錄。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip().upper()
    hard = bool(d.get('hard'))
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute('SELECT source, origin FROM equipment WHERE code=?', (code,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': f'查無設備 {code}'}), 404
        # 這裡一定要看 origin 不能看 source：source 會因為「在系統內編輯過」而變 manual，
        # 拿它當硬刪門檻的話，只要改過一次的 Excel 設備就會變成可以真的刪掉
        if hard and row['origin'] != 'manual':
            return jsonify({'success': False,
                            'error': '這台設備來自 Excel 匯入，只能改為報廢，不能直接刪除'}), 400
        if hard:
            for tbl in ('eq_spec', 'eq_photo', 'eq_tech_file', 'eq_history'):
                conn.execute(f'DELETE FROM {tbl} WHERE code=?', (code,))
            conn.execute('DELETE FROM equipment WHERE code=?', (code,))
            msg = f'已刪除設備 {code}'
        else:
            conn.execute("UPDATE equipment SET status='報廢', source='manual', "
                         "updated_at=datetime('now','localtime') WHERE code=?", (code,))
            conn.execute("INSERT INTO eq_history (code, date, action, detail, user) "
                         "VALUES (?, date('now','localtime'), '報廢', ?, 'system')",
                         (code, (d.get('reason') or '於系統內標記報廢')))
            msg = f'已將 {code} 標記為報廢'
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': msg})


_EQ_HIST_SHEET = '設備歷程'
_EQ_HIST_USER_LABEL = {'system': '系統自動', 'excel': 'Excel匯入', 'user': '人工登錄'}


@app.route('/api/equipment_master/export_history', methods=['POST'])
def equipment_master_export_history():
    """把全部設備的異動歷程匯出成獨立檔案 config.EQUIPMENT_HISTORY_EXPORT_XLSX。

    刻意不寫進 EQUIPMENT_CODING_XLSX（PDM新格式編碼分頁）：那份的「PDM編碼」欄是公式
    （=$A$4&$B$4&"-"&D4&F4），openpyxl 開檔不指定 data_only 再存檔會把公式的快取值
    整個洗掉（openpyxl 不會重算公式），2026-07 用本地副本測試時實際踩到——整欄編碼
    存檔後變空白，要等有人用真的 Excel 開一次存檔公式才會重新算出來。每次都整份重建
    （不是增量），所以不用擔心舊資料殘留；若檔案被佔用（例如有人開著在看）就 fallback
    存到桌面，沿用 batch_cost 那套慣例（見 docs/batch-cost.md `_save_batchcost_wb`）。
    """
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        rows = conn.execute(
            'SELECT code, date, action, detail, user FROM eq_history ORDER BY code, date, rowid').fetchall()
    finally:
        conn.close()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = _EQ_HIST_SHEET
    ws.append(['設備編碼', '日期', '動作', '內容', '來源'])
    for r in rows:
        ws.append([r['code'], r['date'], r['action'], r['detail'],
                   _EQ_HIST_USER_LABEL.get(r['user'], r['user'])])
    ws.freeze_panes = 'A2'
    for col, width in zip('ABCDE', (12, 12, 10, 50, 10)):
        ws.column_dimensions[col].width = width

    path = config.EQUIPMENT_HISTORY_EXPORT_XLSX
    try:
        wb.save(path)
        return jsonify({'success': True, 'message': f'已匯出 {len(rows)} 筆歷程記錄：{path}'})
    except (PermissionError, OSError) as e:
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        os.makedirs(desktop, exist_ok=True)
        fallback = os.path.join(desktop, f'設備歷程_{int(time.time())}.xlsx')
        wb.save(fallback)
        return jsonify({'success': True,
                        'message': f'原始檔案無法寫入（{type(e).__name__}，可能被開啟中），'
                                   f'已改存 {len(rows)} 筆到桌面：{fallback}'})


@app.route('/api/equipment_master/photo')
def equipment_master_photo():
    """輸出設備照片原圖（縮圖快取留待 P3）"""
    from flask import send_file
    full = _eq_safe_path(config.EQUIPMENT_PHOTO_ROOT, request.args.get('relpath', ''))
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '照片不存在'}), 404
    return send_file(full)


_EQ_PHOTO_EXT = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp'}


@app.route('/api/equipment_master/photo/upload', methods=['POST'])
def equipment_master_photo_upload():
    """上傳設備照片（可一次多張）到網芳 EQUIPMENT_PHOTO_ROOT\\<資料夾>\\，立即補進索引。

    資料夾名稱一定要以編碼開頭，build_equipment_index.py 的 scan_photos() 才會自動歸位
    （FOLDER_CODE_RE）：這台設備如果已經有照片資料夾（不論叫 A05-303 還是 A05-303同清），
    沿用同一個，避免同一台設備的照片被拆進兩個資料夾；完全沒有照片時才新建一個以編碼
    命名的資料夾。寫入 eq_photo 時標記 claimed=1，下次重新掃描不會被覆蓋/重新分配封面。
    """
    code = (request.form.get('code') or '').strip().upper()
    files = request.files.getlist('files')
    if not code:
        return jsonify({'success': False, 'error': '缺少設備編碼'}), 400
    if not files or not any(f and f.filename for f in files):
        return jsonify({'success': False, 'error': '請選擇照片'}), 400

    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        if not conn.execute('SELECT code FROM equipment WHERE code=?', (code,)).fetchone():
            return jsonify({'success': False, 'error': f'查無設備 {code}'}), 404

        existing = conn.execute('SELECT folder FROM eq_photo WHERE code=? LIMIT 1', (code,)).fetchone()
        folder_name = existing['folder'] if existing else code
        dest_dir = os.path.join(config.EQUIPMENT_PHOTO_ROOT, folder_name)
        try:
            os.makedirs(dest_dir, exist_ok=True)
        except OSError as e:
            return jsonify({'success': False, 'error': f'無法建立資料夾：{e}'}), 500

        saved, rejected = [], []
        for f in files:
            if not f or not f.filename:
                continue
            ext = os.path.splitext(f.filename)[1].lower()
            if ext not in _EQ_PHOTO_EXT:
                rejected.append(f.filename)
                continue
            filename = _cnc_safe_filename(f.filename)
            dest = os.path.join(dest_dir, filename)
            if os.path.exists(dest):
                stem, ex = os.path.splitext(filename)
                filename = f'{stem}_{int(time.time())}{ex}'
                dest = os.path.join(dest_dir, filename)
            try:
                f.save(dest)
            except OSError as e:
                rejected.append(f'{f.filename}（存檔失敗：{e}）')
                continue
            relpath = os.path.relpath(dest, config.EQUIPMENT_PHOTO_ROOT).replace(os.sep, '/')
            has_photo = conn.execute('SELECT 1 FROM eq_photo WHERE code=? LIMIT 1', (code,)).fetchone()
            conn.execute(
                'INSERT INTO eq_photo (relpath, code, folder, filename, size, mtime, is_cover, claimed) '
                'VALUES (?,?,?,?,?,datetime("now","localtime"),?,1)',
                (relpath, code, folder_name, filename, os.path.getsize(dest), 0 if has_photo else 1))
            saved.append(filename)
        conn.commit()
    finally:
        conn.close()

    if not saved:
        error = '沒有成功上傳的照片'
        if rejected:
            error += '：' + '、'.join(rejected) + '（僅支援 ' + '/'.join(sorted(_EQ_PHOTO_EXT)) + '）'
        return jsonify({'success': False, 'error': error}), 400

    msg = f'已上傳 {len(saved)} 張照片'
    if rejected:
        msg += f'，{len(rejected)} 個檔案格式不支援已略過'
    return jsonify({'success': True, 'message': msg, 'saved': saved})


@app.route('/api/equipment_master/tech_download')
def equipment_master_tech_download():
    """下載技術資料檔案"""
    from flask import send_file
    full = _eq_safe_path(config.EQUIPMENT_TECH_ROOT, request.args.get('relpath', ''))
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404
    return send_file(full, as_attachment=True, download_name=os.path.basename(full))


# ══════════════════════════════════════════════════════════
#  設備保養基準書（P1）  詳見 docs/equipment-maintenance.md
# ══════════════════════════════════════════════════════════
#  基準書掛在「群組+機台類型」層當範本（scope='type'），單台設備用差異表
#  mt_equip_item 覆寫（停用某項／改週期），單台專屬的加項則放 scope='equip'
#  的基準書。刻意不把範本複製給每一台——範本改一個字要同步 189 台是災難，
#  而且事後分不清「這台真的不同」還是「忘了同步」。

# 週期代碼 → (顯示名稱, 排序, 量詞)。日／週點檢不預先展開（見文件 3.5），但基準書本身要能寫
_MT_CYCLES = {
    'day':     ('每日',   1, '日'),
    'week':    ('每週',   2, '週'),
    'month':   ('每月',   3, '個月'),
    'quarter': ('每季',   4, '季'),
    'half':    ('每半年', 5, '半年'),
    'year':    ('每年',   6, '年'),
}
_MT_EXPORT_SHEET = '保養基準書'


def _mt_ensure_tables(conn):
    """建立保養基準書相關資料表（第一次用到時才建，不用另外做 schema 遷移）。

    只在保養相關端點呼叫，不放進 _eq_conn()——EQ_DB_PATH 在網芳上，每開一次連線多跑
    三個 CREATE 是不必要的網路來回。"""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS mt_standard (
            id         INTEGER PRIMARY KEY,
            scope      TEXT,      -- 'type'（類型範本）| 'equip'（單台專屬加項）
            group_code TEXT,      -- scope='type' 用
            type_code  TEXT,
            code       TEXT,      -- scope='equip' 用（設備編碼）
            title      TEXT,
            rev        TEXT,      -- 版次 A/B/C
            rev_date   TEXT,
            author     TEXT,
            status     TEXT,      -- 草稿 / 生效 / 停用
            note       TEXT,
            created_at TEXT, updated_at TEXT
        )""")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS mt_item (
            id         INTEGER PRIMARY KEY,
            std_id     INTEGER,
            seq        INTEGER,   -- 項次
            part       TEXT,      -- 部位：主軸／導軌／潤滑／氣壓／電氣
            name       TEXT,      -- 保養內容
            method     TEXT,      -- 目視／清潔／加油／更換／量測／校正
            criteria   TEXT,      -- 判定基準（基準書跟待辦清單的分水嶺）
            cycle_kind TEXT, cycle_n INTEGER,
            anchor     TEXT,      -- 起算基準日（空=用設備採購日）
            duration_min INTEGER, -- 標準工時（分）
            need_stop  INTEGER,   -- 是否需停機
            owner      TEXT,      -- 操作員／保養員／委外廠商
            tools      TEXT,      -- 工具與耗材
            oil_code   TEXT,      -- 使用油品：油品主檔（oil.db）的代號，優先
            oil        TEXT,      -- 使用油品：主檔沒有這支時的自行輸入文字
            safety     TEXT,      -- 安全注意事項
            attach     TEXT,      -- 附件檔名
            active     INTEGER,
            created_at TEXT, updated_at TEXT
        )""")
    # P2 展開後的保養工作。UNIQUE(code,item_id,period_key) 是整個展開機制的關鍵：
    # 展開作業可以重複執行任意次，只會補上缺的，不會產生重複待辦。不要拿 due_date 當
    # 唯一鍵——調整週期時 due_date 會變，period_key 不會。
    conn.execute("""
        CREATE TABLE IF NOT EXISTS mt_plan (
            id         INTEGER PRIMARY KEY,
            code       TEXT,      -- 設備編碼
            item_id    INTEGER,   -- 來源保養項目
            due_date   TEXT,      -- 應執行日 YYYY-MM-DD
            period_key TEXT,      -- 2026 / 2026-H1 / 2026-Q3 / 2026-08 / 2026-W32 / 2026-08-10
            status     TEXT,      -- 待辦 / 已完成 / 跳過
            done_date  TEXT, done_by TEXT,
            result     TEXT,      -- OK / NG
            minutes    REAL,      -- 實際工時
            remark     TEXT,
            photo      TEXT,
            created_at TEXT, updated_at TEXT
        )""")
    conn.execute('CREATE UNIQUE INDEX IF NOT EXISTS idx_mt_plan_uniq '
                 'ON mt_plan(code, item_id, period_key)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_mt_plan_due ON mt_plan(due_date, status)')
    conn.execute("""
        CREATE TABLE IF NOT EXISTS mt_equip_item (
            code    TEXT,         -- 設備編碼
            item_id INTEGER,      -- 範本項目
            disabled INTEGER,     -- 1 = 這台不做這項
            cycle_kind TEXT, cycle_n INTEGER,   -- NULL = 沿用範本
            anchor  TEXT,         -- NULL = 沿用範本
            note    TEXT,
            updated_at TEXT,
            PRIMARY KEY(code, item_id)
        )""")
    _mt_ensure_columns(conn)
    conn.execute('CREATE INDEX IF NOT EXISTS idx_mt_item_std ON mt_item(std_id)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_mt_std_type ON mt_standard(group_code, type_code)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_mt_std_code ON mt_standard(code)')
    conn.commit()


# 上線後才補的欄位：表名 → {欄位: 型別}
# CREATE TABLE IF NOT EXISTS 對「已存在」的表不會補欄位，正式機那份 mt_item 是舊的，
# 只改上面的 CREATE 沒有用（同 equipment.db 加 model 欄位踩過的坑，見 docs/equipment-master.md）
_MT_ADDED_COLUMNS = {
    'mt_item': {'oil': 'TEXT',          # 2026-08-11 使用油品（自行輸入）
                'oil_code': 'TEXT'},    # 2026-08-11 使用油品（對應 oil.db 的代號）
}


def _mt_ensure_columns(conn):
    """補上舊資料庫沒有的欄位。PRAGMA table_info 很便宜，每次開連線檢查一次成本可忽略。"""
    for table, cols in _MT_ADDED_COLUMNS.items():
        have = {r['name'] for r in conn.execute(f'PRAGMA table_info({table})')}
        for name, decl in cols.items():
            if name not in have:
                conn.execute(f'ALTER TABLE {table} ADD COLUMN {name} {decl}')


def _mt_conn():
    """開設備資料庫並確保保養資料表存在；索引未建立時回傳 None"""
    conn = _eq_conn()
    if conn is not None:
        _mt_ensure_tables(conn)
    return conn


def _mt_now():
    return datetime.now().strftime('%Y-%m-%d %H:%M:%S')


def _mt_cycle_label(kind, n):
    """(month, 3) → 每 3 個月；(quarter, 1) → 每季"""
    c = _MT_CYCLES.get(kind)
    if not c:
        return ''
    try:
        n = int(n or 1)
    except (TypeError, ValueError):
        n = 1
    return c[0] if n <= 1 else f'每 {n} {c[2]}'


def _mt_std_row(r):
    d = dict(r)
    d['id'] = int(d['id'])
    return d


def _mt_item_rows(conn, std_id):
    return [dict(r) for r in conn.execute(
        'SELECT * FROM mt_item WHERE std_id=? ORDER BY seq, id', (std_id,))]


def _mt_fill_oil(items):
    """把項目的 oil_code 對回油品主檔（`oil.db`，見 docs/oil-management.md）。

    油品主檔跟設備主檔是**兩個不同的資料庫檔案**，join 不了，所以在 Python 這邊補；
    一次把用到的代號查完，不要每筆各開一次連線（都在網芳上，來回很貴）。

    **顯示名稱一律以主檔當下的值為準**，不是存基準書時的快照——油品在主檔改了品名或
    改成停用，基準書上要立刻看得到，不然現場照著一份寫著已停用油品的基準書去領料。
    """
    codes = {(it.get('oil_code') or '').strip() for it in items}
    codes.discard('')
    info = {}
    if codes:
        conn = _oil_conn()
        if conn is not None:
            try:
                ph = ','.join('?' * len(codes))
                info = {r['code']: dict(r) for r in conn.execute(
                    f'SELECT code, name, category, supplier, status FROM oil WHERE code IN ({ph})',
                    list(codes))}
            except sqlite3.Error:
                info = {}
            finally:
                conn.close()
    for it in items:
        code = (it.get('oil_code') or '').strip()
        o = info.get(code)
        it['oil_name'] = (o or {}).get('name') or ''
        it['oil_status'] = (o or {}).get('status') or ''
        it['oil_category'] = (o or {}).get('category') or ''
        # 主檔查不到（油品被硬刪或改了代號）要看得出來，不能默默顯示成空白
        it['oil_missing'] = bool(code) and o is None
        # 顯示只用代號：代號本身已經帶中文說明（「半合成切削液」AW-30），
        # 再接上完整品名（STORK金屬加工液 SW5031-030）會把表格欄位撐爆。
        # 品名放 oil_name 給前端當 tooltip，需要時滑過去看。
        it['oil_label'] = code if code else (it.get('oil') or '').strip()
    return items


@app.route('/api/equipment_master/maint/oil_options')
def maint_oil_options():
    """「使用油品」的候選清單，直接讀油品主檔 oil.db。

    停用的油品也一起回傳（前端會標註）：基準書可能引用到已經停用的油，把它藏起來
    只會讓人看不懂為什麼下拉裡選不到現在這支。排序沿用油品清單那套分類優先的規則。"""
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': True, 'oils': [], 'note': '尚未建立油品主檔'})
    try:
        rows = [dict(r) for r in conn.execute(
            'SELECT code, name, category, supplier, status FROM oil')]
    except sqlite3.Error as e:
        return jsonify({'success': True, 'oils': [], 'note': f'油品主檔讀取失敗：{e}'})
    finally:
        conn.close()
    rows.sort(key=_oil_sort_key)
    return jsonify({'success': True, 'oils': rows})


@app.route('/api/equipment_master/maint/tree')
def maint_tree():
    """基準書總覽：每個群組／機台類型有幾台設備、有沒有建基準書、幾個項目。

    未建基準書的類型也要回傳（前端要顯示「尚未建立」讓人去建），所以是以設備與
    編碼字典為主體 LEFT JOIN 基準書，不是列舉 mt_standard。"""
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        groups = {r['code']: dict(r) for r in conn.execute('SELECT code, name FROM eq_group ORDER BY code')}
        # 類型清單以字典為主，並補上字典裡沒有但設備上有的類型（編碼待修正的設備會出現這種）
        types = {}
        for r in conn.execute('SELECT group_code, code, name FROM eq_type ORDER BY group_code, code'):
            types[(r['group_code'], r['code'])] = {'group_code': r['group_code'], 'type_code': r['code'],
                                                   'type_name': r['name'], 'equip_cnt': 0, 'active_cnt': 0}
        for r in conn.execute(
                "SELECT group_code, type_code, COUNT(*) AS c,"
                " SUM(CASE WHEN status IN ('使用中','閒置') THEN 1 ELSE 0 END) AS a"
                ' FROM equipment GROUP BY group_code, type_code'):
            key = (r['group_code'], r['type_code'])
            t = types.setdefault(key, {'group_code': r['group_code'], 'type_code': r['type_code'],
                                       'type_name': '', 'equip_cnt': 0, 'active_cnt': 0})
            t['equip_cnt'] = r['c']
            t['active_cnt'] = r['a'] or 0
        stds = {}
        for r in conn.execute(
                "SELECT s.id, s.scope, s.group_code, s.type_code, s.code, s.title, s.rev, s.status,"
                ' (SELECT COUNT(*) FROM mt_item i WHERE i.std_id=s.id) AS item_cnt'
                ' FROM mt_standard s'):
            d = _mt_std_row(r)
            if d['scope'] == 'type':
                stds[(d['group_code'], d['type_code'])] = d
        for key, t in types.items():
            t.update({'std_id': None, 'item_cnt': 0, 'status': '', 'rev': '', 'title': ''})
            s = stds.get(key)
            if s:
                t.update({'std_id': s['id'], 'item_cnt': s['item_cnt'],
                          'status': s['status'] or '', 'rev': s['rev'] or '', 'title': s['title'] or ''})
        # 單台專屬基準書另外列（掛在設備上，不屬於任何類型範本）
        equip_stds = [_mt_std_row(r) for r in conn.execute(
            "SELECT s.*, (SELECT COUNT(*) FROM mt_item i WHERE i.std_id=s.id) AS item_cnt"
            " FROM mt_standard s WHERE s.scope='equip' ORDER BY s.code")]
        # 已經用過的油品／部位，給表單當下拉候選：41 種類型各自打字很容易把同一種油
        # 打成三種寫法，之後要統計用油量或叫料就對不起來
        oils = [r[0] for r in conn.execute(
            "SELECT DISTINCT oil FROM mt_item WHERE IFNULL(oil,'')<>'' ORDER BY oil")]
        parts = [r[0] for r in conn.execute(
            "SELECT DISTINCT part FROM mt_item WHERE IFNULL(part,'')<>'' ORDER BY part")]
    finally:
        conn.close()
    out = sorted(types.values(), key=lambda x: (x['group_code'] or '', x['type_code'] or ''))
    return jsonify({'success': True, 'groups': list(groups.values()),
                    'types': out, 'equip_stds': equip_stds, 'oils': oils, 'parts': parts,
                    'cycles': [{'kind': k, 'label': v[0]} for k, v in
                               sorted(_MT_CYCLES.items(), key=lambda x: x[1][1])]})


@app.route('/api/equipment_master/maint/standard')
def maint_standard_get():
    """讀一份基準書＋所有項目。可用 id、(group_code,type_code)、或 code（單台專屬）指定。"""
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    sid = request.args.get('id', '').strip()
    g = request.args.get('group_code', '').strip().upper()
    t = request.args.get('type_code', '').strip()
    code = request.args.get('code', '').strip().upper()
    try:
        if sid:
            row = conn.execute('SELECT * FROM mt_standard WHERE id=?', (sid,)).fetchone()
        elif code:
            row = conn.execute("SELECT * FROM mt_standard WHERE scope='equip' AND code=?", (code,)).fetchone()
        elif g and t:
            row = conn.execute("SELECT * FROM mt_standard WHERE scope='type' AND group_code=? AND type_code=?",
                               (g, t)).fetchone()
        else:
            return jsonify({'success': False, 'error': '請指定基準書'}), 400
        if row is None:
            return jsonify({'success': True, 'standard': None, 'items': []})
        std = _mt_std_row(row)
        folder = _mt_std_folder(row)
        items = _mt_fill_oil(_mt_item_rows(conn, std['id']))
        for it in items:
            it['cycle_label'] = _mt_cycle_label(it['cycle_kind'], it['cycle_n'])
            it['attach_files'] = _mt_attach_list(it['attach'], folder)
        # 這份範本適用哪些設備（單台專屬的就是那一台）
        if std['scope'] == 'type':
            equips = [dict(r) for r in conn.execute(
                'SELECT code, old_code, status, location FROM equipment'
                ' WHERE group_code=? AND type_code=? ORDER BY code', (std['group_code'], std['type_code']))]
        else:
            equips = [dict(r) for r in conn.execute(
                'SELECT code, old_code, status, location FROM equipment WHERE code=?', (std['code'],))]
    finally:
        conn.close()
    return jsonify({'success': True, 'standard': std, 'items': items, 'equips': equips, 'folder': folder})


@app.route('/api/equipment_master/maint/standard/save', methods=['POST'])
def maint_standard_save():
    """新增或修改基準書表頭。scope='type' 時同一個類型只能有一份（重複建立會回傳既有那份）。"""
    d = request.get_json(silent=True) or {}
    scope = (d.get('scope') or 'type').strip()
    g = (d.get('group_code') or '').strip().upper()
    t = (d.get('type_code') or '').strip()
    code = (d.get('code') or '').strip().upper()
    if scope == 'type' and not (g and t):
        return jsonify({'success': False, 'error': '請指定群組與機台類型'}), 400
    if scope == 'equip' and not code:
        return jsonify({'success': False, 'error': '請指定設備編碼'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        sid = d.get('id')
        now = _mt_now()
        fields = {
            'title':  (d.get('title') or '').strip(),
            'rev':    (d.get('rev') or '').strip(),
            'rev_date': (d.get('rev_date') or '').strip(),
            'author': (d.get('author') or '').strip(),
            'status': (d.get('status') or '生效').strip(),
            'note':   (d.get('note') or '').strip(),
        }
        if not sid:
            # 同一個類型／同一台設備只允許一份基準書，已存在就直接沿用（避免重複按新增建出兩份）
            exist = conn.execute(
                "SELECT id FROM mt_standard WHERE scope=? AND IFNULL(group_code,'')=? "
                "AND IFNULL(type_code,'')=? AND IFNULL(code,'')=?",
                (scope, g, t, code)).fetchone()
            if exist:
                sid = exist['id']
        if not fields['title']:
            if scope == 'type':
                tn = conn.execute('SELECT name FROM eq_type WHERE group_code=? AND code=?', (g, t)).fetchone()
                fields['title'] = f'{g}{t} {tn["name"] if tn else ""} 保養基準書'.replace('  ', ' ').strip()
            else:
                fields['title'] = f'{code} 專屬保養基準書'
        if sid:
            conn.execute('UPDATE mt_standard SET title=?, rev=?, rev_date=?, author=?, status=?, note=?,'
                         ' updated_at=? WHERE id=?',
                         (fields['title'], fields['rev'], fields['rev_date'], fields['author'],
                          fields['status'], fields['note'], now, sid))
        else:
            cur = conn.execute(
                'INSERT INTO mt_standard (scope, group_code, type_code, code, title, rev, rev_date,'
                ' author, status, note, created_at, updated_at)'
                ' VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
                (scope, g or None, t or None, code or None, fields['title'], fields['rev'],
                 fields['rev_date'], fields['author'], fields['status'], fields['note'], now, now))
            sid = cur.lastrowid
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'id': sid})


@app.route('/api/equipment_master/maint/standard/delete', methods=['POST'])
def maint_standard_delete():
    """刪除整份基準書（連同項目與各台的覆寫設定）"""
    sid = (request.get_json(silent=True) or {}).get('id')
    if not sid:
        return jsonify({'success': False, 'error': '缺少 id'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        ids = [r['id'] for r in conn.execute('SELECT id FROM mt_item WHERE std_id=?', (sid,))]
        if ids:
            conn.execute('DELETE FROM mt_equip_item WHERE item_id IN (%s)' % ','.join('?' * len(ids)), ids)
        conn.execute('DELETE FROM mt_item WHERE std_id=?', (sid,))
        conn.execute('DELETE FROM mt_standard WHERE id=?', (sid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted_items': len(ids)})


@app.route('/api/equipment_master/maint/standard/copy', methods=['POST'])
def maint_standard_copy():
    """把某份基準書的項目複製到另一個機台類型。

    41 種類型 × 10~20 項要人工建，沒有這個功能內容永遠建不完（見文件第九節）。
    目標已有基準書時採「附加」而非覆蓋，項次接在後面，避免一次誤操作洗掉既有內容。"""
    d = request.get_json(silent=True) or {}
    from_id = d.get('from_id')
    g = (d.get('to_group') or '').strip().upper()
    t = (d.get('to_type') or '').strip()
    if not from_id or not (g and t):
        return jsonify({'success': False, 'error': '請指定來源基準書與目標機台類型'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        src = conn.execute('SELECT * FROM mt_standard WHERE id=?', (from_id,)).fetchone()
        if src is None:
            return jsonify({'success': False, 'error': '來源基準書不存在'}), 404
        if src['scope'] == 'type' and src['group_code'] == g and src['type_code'] == t:
            return jsonify({'success': False, 'error': '來源與目標是同一個機台類型'}), 400
        now = _mt_now()
        dst = conn.execute("SELECT id FROM mt_standard WHERE scope='type' AND group_code=? AND type_code=?",
                           (g, t)).fetchone()
        if dst:
            dst_id = dst['id']
        else:
            tn = conn.execute('SELECT name FROM eq_type WHERE group_code=? AND code=?', (g, t)).fetchone()
            title = f'{g}{t} {tn["name"] if tn else ""} 保養基準書'.replace('  ', ' ').strip()
            dst_id = conn.execute(
                "INSERT INTO mt_standard (scope, group_code, type_code, title, rev, status, note,"
                " created_at, updated_at) VALUES ('type',?,?,?,?,?,?,?,?)",
                (g, t, title, src['rev'], '草稿', f'複製自 {src["title"]}', now, now)).lastrowid
        base = conn.execute('SELECT IFNULL(MAX(seq),0) FROM mt_item WHERE std_id=?', (dst_id,)).fetchone()[0]
        rows = conn.execute('SELECT * FROM mt_item WHERE std_id=? ORDER BY seq, id', (from_id,)).fetchall()
        for i, r in enumerate(rows, 1):
            conn.execute(
                'INSERT INTO mt_item (std_id, seq, part, name, method, criteria, cycle_kind, cycle_n,'
                ' anchor, duration_min, need_stop, owner, tools, oil_code, oil, safety, attach, active,'
                ' created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                (dst_id, base + i, r['part'], r['name'], r['method'], r['criteria'], r['cycle_kind'],
                 r['cycle_n'], r['anchor'], r['duration_min'], r['need_stop'], r['owner'], r['tools'],
                 r['oil_code'], r['oil'], r['safety'], r['attach'], r['active'], now, now))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'id': dst_id, 'copied': len(rows)})


def _mt_item_payload(d):
    """把前端送來的項目欄位整理成寫入用的 tuple（不含 std_id/seq）"""
    kind = (d.get('cycle_kind') or '').strip()
    if kind not in _MT_CYCLES:
        kind = 'year'
    try:
        n = max(1, int(d.get('cycle_n') or 1))
    except (TypeError, ValueError):
        n = 1
    try:
        mins = int(d.get('duration_min') or 0) or None
    except (TypeError, ValueError):
        mins = None
    return ((d.get('part') or '').strip(), (d.get('name') or '').strip(),
            (d.get('method') or '').strip(), (d.get('criteria') or '').strip(),
            kind, n, (d.get('anchor') or '').strip() or None, mins,
            1 if d.get('need_stop') else 0, (d.get('owner') or '').strip(),
            (d.get('tools') or '').strip(), (d.get('oil_code') or '').strip(),
            (d.get('oil') or '').strip(), (d.get('safety') or '').strip(),
            (d.get('attach') or '').strip(), 0 if d.get('active') in (0, '0', False) else 1)


@app.route('/api/equipment_master/maint/item/save', methods=['POST'])
def maint_item_save():
    """新增或修改一個保養項目"""
    d = request.get_json(silent=True) or {}
    if not (d.get('name') or '').strip():
        return jsonify({'success': False, 'error': '請填寫保養內容'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        now = _mt_now()
        vals = _mt_item_payload(d)
        iid = d.get('id')
        if iid:
            conn.execute(
                'UPDATE mt_item SET part=?, name=?, method=?, criteria=?, cycle_kind=?, cycle_n=?,'
                ' anchor=?, duration_min=?, need_stop=?, owner=?, tools=?, oil_code=?, oil=?,'
                ' safety=?, attach=?, active=?, updated_at=? WHERE id=?', vals + (now, iid))
        else:
            std_id = d.get('std_id')
            if not std_id:
                return jsonify({'success': False, 'error': '缺少 std_id'}), 400
            seq = (conn.execute('SELECT IFNULL(MAX(seq),0) FROM mt_item WHERE std_id=?',
                                (std_id,)).fetchone()[0] or 0) + 1
            iid = conn.execute(
                'INSERT INTO mt_item (std_id, seq, part, name, method, criteria, cycle_kind, cycle_n,'
                ' anchor, duration_min, need_stop, owner, tools, oil_code, oil, safety, attach, active,'
                ' created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                (std_id, seq) + vals + (now, now)).lastrowid
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'id': iid})


@app.route('/api/equipment_master/maint/item/delete', methods=['POST'])
def maint_item_delete():
    """刪除保養項目（連同各台設備對它的覆寫設定）"""
    iid = (request.get_json(silent=True) or {}).get('id')
    if not iid:
        return jsonify({'success': False, 'error': '缺少 id'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        conn.execute('DELETE FROM mt_equip_item WHERE item_id=?', (iid,))
        conn.execute('DELETE FROM mt_item WHERE id=?', (iid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/equipment_master/maint/item/move', methods=['POST'])
def maint_item_move():
    """項目上移／下移（跟相鄰的那筆交換 seq）"""
    d = request.get_json(silent=True) or {}
    iid, direction = d.get('id'), d.get('dir')
    if not iid or direction not in ('up', 'down'):
        return jsonify({'success': False, 'error': '參數錯誤'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        rows = [dict(r) for r in conn.execute(
            'SELECT id, seq FROM mt_item WHERE std_id=(SELECT std_id FROM mt_item WHERE id=?)'
            ' ORDER BY seq, id', (iid,))]
        idx = next((i for i, r in enumerate(rows) if r['id'] == int(iid)), -1)
        j = idx - 1 if direction == 'up' else idx + 1
        if idx < 0 or j < 0 or j >= len(rows):
            return jsonify({'success': True, 'moved': False})
        rows[idx], rows[j] = rows[j], rows[idx]
        # seq 重新從 1 編號（舊資料可能有重複或空的 seq，交換兩筆的值不一定有效）
        for i, r in enumerate(rows, 1):
            conn.execute('UPDATE mt_item SET seq=? WHERE id=?', (i, r['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'moved': True})


@app.route('/api/equipment_master/maint/for_equip')
def maint_for_equip():
    """某一台設備實際適用的保養項目 = 類型範本 + 這台的覆寫 + 這台專屬的加項。

    範本項目一律回傳（含被這台停用的，前端要顯示成刪除線讓人可以再啟用），
    週期欄位回傳的是套用覆寫後的有效值，另外附上範本原值供對照。"""
    code = request.args.get('code', '').strip().upper()
    if not code:
        return jsonify({'success': False, 'error': '缺少設備編碼'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        eq = conn.execute('SELECT code, group_code, type_code FROM equipment WHERE code=?', (code,)).fetchone()
        if eq is None:
            return jsonify({'success': False, 'error': '設備不存在'}), 404
        std = conn.execute("SELECT * FROM mt_standard WHERE scope='type' AND group_code=? AND type_code=?",
                           (eq['group_code'], eq['type_code'])).fetchone()
        own = conn.execute("SELECT * FROM mt_standard WHERE scope='equip' AND code=?", (code,)).fetchone()
        ov = {r['item_id']: dict(r) for r in conn.execute(
            'SELECT * FROM mt_equip_item WHERE code=?', (code,))}
        items = []
        for src, kind in ((std, 'type'), (own, 'equip')):
            if src is None:
                continue
            folder = _mt_std_folder(src)
            for it in _mt_fill_oil(_mt_item_rows(conn, src['id'])):
                it['attach_files'] = _mt_attach_list(it['attach'], folder)
                it['folder'] = folder
                o = ov.get(it['id']) or {}
                eff_kind = o.get('cycle_kind') or it['cycle_kind']
                eff_n = o.get('cycle_n') or it['cycle_n']
                items.append(dict(it, from_scope=kind,
                                  base_cycle_label=_mt_cycle_label(it['cycle_kind'], it['cycle_n']),
                                  cycle_kind=eff_kind, cycle_n=eff_n,
                                  cycle_label=_mt_cycle_label(eff_kind, eff_n),
                                  anchor=o.get('anchor') or it['anchor'],
                                  disabled=1 if o.get('disabled') else 0,
                                  override_note=o.get('note') or '',
                                  overridden=bool(o.get('cycle_kind') or o.get('cycle_n') or o.get('anchor'))))
        std_out = _mt_std_row(std) if std is not None else None
        own_out = _mt_std_row(own) if own is not None else None
    finally:
        conn.close()
    return jsonify({'success': True, 'standard': std_out, 'own_standard': own_out, 'items': items})


@app.route('/api/equipment_master/maint/equip_item/save', methods=['POST'])
def maint_equip_item_save():
    """單台設備對某個範本項目的覆寫（停用／改週期／備註）。

    全部都是預設值時直接刪掉該筆，表裡不累積無意義資料——同行事曆 cal_day 的做法
    （見 docs/calendar.md），這樣「有沒有被動過手腳」一眼就看得出來。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip().upper()
    iid = d.get('item_id')
    if not code or not iid:
        return jsonify({'success': False, 'error': '參數錯誤'}), 400
    disabled = 1 if d.get('disabled') else 0
    kind = (d.get('cycle_kind') or '').strip() or None
    if kind not in _MT_CYCLES:
        kind = None
    try:
        n = int(d.get('cycle_n') or 0) or None
    except (TypeError, ValueError):
        n = None
    anchor = (d.get('anchor') or '').strip() or None
    note = (d.get('note') or '').strip()
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        if not disabled and not kind and not n and not anchor and not note:
            conn.execute('DELETE FROM mt_equip_item WHERE code=? AND item_id=?', (code, iid))
        else:
            conn.execute(
                'INSERT OR REPLACE INTO mt_equip_item (code, item_id, disabled, cycle_kind, cycle_n,'
                ' anchor, note, updated_at) VALUES (?,?,?,?,?,?,?,?)',
                (code, iid, disabled, kind, n, anchor, note, _mt_now()))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


# 基準書附件允許的副檔名：以原廠保養手冊／注意事項的 PDF 為主，另收圖片與 Office 檔
_MT_ATTACH_EXT = {'.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp',
                  '.doc', '.docx', '.xls', '.xlsx', '.txt'}


def _mt_std_folder(row):
    """基準書附件的資料夾名稱：類型範本用 B01、單台專屬用設備編碼。

    用看得懂的名字而不是 std_id，是因為這是網芳資料夾——沒裝系統的人直接開資料夾
    也要知道那疊 PDF 是哪個機型的（同 EQUIPMENT_PHOTO_ROOT 以編碼命名的理由）。"""
    name = (f'{row["group_code"] or ""}{row["type_code"] or ""}'
            if row['scope'] == 'type' else (row['code'] or ''))
    return _cnc_safe_filename(name) if name else f'std{row["id"]}'


def _mt_attach_list(value, folder):
    """把 attach 欄位（'a.pdf|b.pdf'）拆成清單，並標記檔案是否真的在網芳上。

    舊資料（例如舊維護一覽表匯入的「240329 機械手臂電池更換注意事項」）只是一段文字、
    沒有對應檔案，exists=False，前端就顯示純文字不做成連結。"""
    out = []
    for name in [x.strip() for x in (value or '').split('|') if x.strip()]:
        full = _eq_safe_path(config.EQUIPMENT_MAINT_ROOT, os.path.join(folder, name))
        out.append({'name': name, 'exists': bool(full and os.path.isfile(full))})
    return out


@app.route('/api/equipment_master/maint/attach/upload', methods=['POST'])
def maint_attach_upload():
    """上傳基準書附件（可一次多個）到網芳 EQUIPMENT_MAINT_ROOT\\<基準書資料夾>\\。

    附件掛在「基準書」層而不是「項目」層，所以項目還沒存檔就能先上傳——前端拖進來
    就立刻上傳，回傳的檔名再填進項目的 attach 欄位。撞名時加時間戳記，不覆蓋既有檔案。"""
    std_id = (request.form.get('std_id') or '').strip()
    files = request.files.getlist('files')
    if not std_id:
        return jsonify({'success': False, 'error': '缺少 std_id'}), 400
    if not files or not any(f and f.filename for f in files):
        return jsonify({'success': False, 'error': '請選擇檔案'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute('SELECT * FROM mt_standard WHERE id=?', (std_id,)).fetchone()
    finally:
        conn.close()
    if row is None:
        return jsonify({'success': False, 'error': '基準書不存在'}), 404

    folder = _mt_std_folder(row)
    dest_dir = _eq_safe_path(config.EQUIPMENT_MAINT_ROOT, folder)
    if not dest_dir:
        return jsonify({'success': False, 'error': '資料夾名稱不合法'}), 400
    try:
        os.makedirs(dest_dir, exist_ok=True)
    except OSError as e:
        return jsonify({'success': False, 'error': f'無法建立資料夾（{dest_dir}）：{e}'}), 500

    saved, rejected = [], []
    for f in files:
        if not f or not f.filename:
            continue
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in _MT_ATTACH_EXT:
            rejected.append(f.filename)
            continue
        filename = _cnc_safe_filename(f.filename)
        dest = os.path.join(dest_dir, filename)
        if os.path.exists(dest):
            stem, ex = os.path.splitext(filename)
            filename = f'{stem}_{int(time.time())}{ex}'
            dest = os.path.join(dest_dir, filename)
        try:
            f.save(dest)
        except OSError as e:
            rejected.append(f'{f.filename}（存檔失敗：{e}）')
            continue
        saved.append(filename)

    if not saved:
        err = '沒有成功上傳的檔案'
        if rejected:
            err += '：' + '、'.join(rejected) + '（僅支援 ' + '/'.join(sorted(_MT_ATTACH_EXT)) + '）'
        return jsonify({'success': False, 'error': err}), 400
    msg = f'已上傳 {len(saved)} 個檔案'
    if rejected:
        msg += f'，{len(rejected)} 個格式不支援已略過'
    return jsonify({'success': True, 'message': msg, 'saved': saved, 'folder': folder})


@app.route('/api/equipment_master/maint/attach')
def maint_attach_get():
    """開啟／下載基準書附件。PDF 預設用瀏覽器內建檢視器開（as_attachment=False），
    看完就關比先下載到磁碟再開順手；路徑一律走 _eq_safe_path 擋跳脫。"""
    from flask import send_file
    folder = request.args.get('folder', '')
    name = request.args.get('f', '')
    full = _eq_safe_path(config.EQUIPMENT_MAINT_ROOT, os.path.join(folder, name))
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404
    return send_file(full, as_attachment=False, download_name=os.path.basename(full))


# 舊「設備維護一覽表.xlsx」的週期文字 → 週期代碼
_MT_LEGACY_CYCLE = {
    '每日': ('day', 1), '日': ('day', 1),
    '每週': ('week', 1), '週': ('week', 1), '每周': ('week', 1),
    '每月': ('month', 1), '月': ('month', 1),
    '每季': ('quarter', 1), '季': ('quarter', 1), '三個月': ('month', 3),
    '半年': ('half', 1), '每半年': ('half', 1), '六個月': ('month', 6),
    '一年': ('year', 1), '每年': ('year', 1), '年': ('year', 1),
    '兩年': ('year', 2), '二年': ('year', 2), '三年': ('year', 3),
}


@app.route('/api/equipment_master/maint/import_legacy', methods=['POST'])
def maint_import_legacy():
    """把舊的「設備維護一覽表.xlsx」匯進來（一次性；重複執行不會產生重複資料）。

    總表 → 該台設備的專屬基準書（scope='equip'，因為原表本來就是逐台寫的）；
    維護記錄 → eq_history（`user='user'`）。**刻意不寫成 `user='excel'`**：
    build_equipment_index.py 匯入時會 `DELETE WHERE user='excel'` 再重寫，
    寫成 excel 的話下次重匯就整批消失。"""
    path = getattr(config, 'EQUIPMENT_MAINT_LEGACY_XLSX', '')
    if not path or not os.path.isfile(path):
        return jsonify({'success': False, 'error': f'找不到舊維護一覽表：{path}'}), 404
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    added_items = added_hist = skipped = 0
    notes = []
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        now = _mt_now()
        known = {r['code'] for r in conn.execute('SELECT code FROM equipment')}

        if '總表' in wb.sheetnames:
            for row in wb['總表'].iter_rows(min_row=2, values_only=True):
                cells = list(row) + [None] * 6
                eq_name, code, cycle, content, safety, attach = [
                    (str(c).strip() if c is not None else '') for c in cells[:6]]
                code = code.upper()
                if not code:
                    continue
                if code not in known:
                    skipped += 1
                    notes.append(f'{code}：設備主檔查無此編碼，略過')
                    continue
                kind, n = _MT_LEGACY_CYCLE.get(cycle, ('year', 1))
                if cycle and cycle not in _MT_LEGACY_CYCLE:
                    notes.append(f'{code}：週期「{cycle}」無法辨識，先當成每年')
                std = conn.execute("SELECT id FROM mt_standard WHERE scope='equip' AND code=?",
                                   (code,)).fetchone()
                if std:
                    sid = std['id']
                else:
                    sid = conn.execute(
                        "INSERT INTO mt_standard (scope, code, title, rev, status, note, created_at, updated_at)"
                        " VALUES ('equip',?,?,?,?,?,?,?)",
                        (code, f'{code} 專屬保養基準書', 'A', '生效',
                         f'由舊「設備維護一覽表」匯入（原設備名稱：{eq_name}）', now, now)).lastrowid
                # 沒填維護內容的列照樣建起來（週期是有意義的資訊），但標成停用讓人補
                name = content or '（原表未填維護內容，待補）'
                active = 1 if content else 0
                dup = conn.execute('SELECT id FROM mt_item WHERE std_id=? AND name=?', (sid, name)).fetchone()
                if dup:
                    continue
                seq = (conn.execute('SELECT IFNULL(MAX(seq),0) FROM mt_item WHERE std_id=?',
                                    (sid,)).fetchone()[0] or 0) + 1
                conn.execute(
                    'INSERT INTO mt_item (std_id, seq, name, criteria, cycle_kind, cycle_n, safety,'
                    ' attach, active, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)',
                    (sid, seq, name, '', kind, n, safety, attach, active, now, now))
                added_items += 1

        if '維護記錄' in wb.sheetnames:
            for row in wb['維護記錄'].iter_rows(min_row=2, values_only=True):
                cells = list(row) + [None] * 5
                dt, _eq_name, code, content, vendor = cells[:5]
                code = (str(code).strip().upper() if code else '')
                if not code or code not in known:
                    continue
                if isinstance(dt, datetime):
                    dstr = dt.strftime('%Y-%m-%d')
                else:
                    dstr = (str(dt).strip()[:10] if dt else '')
                detail = ' '.join(x for x in [(str(content).strip() if content else ''),
                                              (f'（廠商：{str(vendor).strip()}）' if vendor else '')] if x)
                detail = (detail or '保養') + '｜舊維護一覽表匯入'
                dup = conn.execute('SELECT rowid FROM eq_history WHERE code=? AND date=? AND detail=?',
                                   (code, dstr, detail)).fetchone()
                if dup:
                    continue
                conn.execute('INSERT INTO eq_history (code, date, action, detail, user)'
                             " VALUES (?,?,'保養',?,'user')", (code, dstr, detail))
                added_hist += 1
        wb.close()
        conn.commit()
    except Exception as e:
        return jsonify({'success': False, 'error': f'匯入失敗：{type(e).__name__} {e}'}), 500
    finally:
        conn.close()
    msg = f'匯入完成：新增 {added_items} 個保養項目、{added_hist} 筆保養歷程'
    if skipped:
        msg += f'，略過 {skipped} 列'
    return jsonify({'success': True, 'message': msg, 'notes': notes})


@app.route('/api/equipment_master/maint/export', methods=['POST'])
def maint_export():
    """把全部基準書匯出成獨立檔案 config.EQUIPMENT_MAINT_EXPORT_XLSX。

    每次整份重建一個新活頁簿（不載入既有檔案）——理由同 export_history：
    本專案的共用 Excel 有公式欄位，開檔存檔會洗掉快取值（見 docs/equipment-master.md）。"""
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT s.scope, s.group_code, s.type_code, s.code AS std_code, s.title, s.rev, s.status,
                   t.name AS type_name, i.*
              FROM mt_standard s
              JOIN mt_item i ON i.std_id = s.id
              LEFT JOIN eq_type t ON t.group_code = s.group_code AND t.code = s.type_code
             ORDER BY s.scope, s.group_code, s.type_code, s.code, i.seq, i.id""")]
    finally:
        conn.close()
    _mt_fill_oil(rows)   # 匯出的油品欄要寫主檔當下的品名，不是存檔當時的快照

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = _MT_EXPORT_SHEET
    ws.append(['適用範圍', '基準書', '版次', '狀態', '項次', '部位', '保養內容', '方法', '判定基準',
               '週期', '標準工時(分)', '需停機', '負責', '工具耗材', '使用油品', '安全注意事項',
               '附件', '啟用'])
    for r in rows:
        scope = (f'{r["group_code"]}{r["type_code"]} {r["type_name"] or ""}'.strip()
                 if r['scope'] == 'type' else f'{r["std_code"]}（單台）')
        ws.append([scope, r['title'], r['rev'], r['status'], r['seq'], r['part'], r['name'], r['method'],
                   r['criteria'], _mt_cycle_label(r['cycle_kind'], r['cycle_n']), r['duration_min'],
                   '是' if r['need_stop'] else '', r['owner'], r['tools'], r['oil_label'], r['safety'],
                   r['attach'], '' if r['active'] else '停用'])
    ws.freeze_panes = 'A2'
    for col, width in zip(['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I',
                           'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R'],
                          (22, 30, 6, 8, 6, 12, 34, 10, 30, 12, 12, 8, 10, 22, 16, 30, 20, 8)):
        ws.column_dimensions[col].width = width

    path = getattr(config, 'EQUIPMENT_MAINT_EXPORT_XLSX', '') or os.path.join(
        os.path.expanduser('~'), 'Desktop', '設備保養基準書.xlsx')
    try:
        wb.save(path)
        return jsonify({'success': True, 'message': f'已匯出 {len(rows)} 個保養項目：{path}'})
    except (PermissionError, OSError) as e:
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        os.makedirs(desktop, exist_ok=True)
        fallback = os.path.join(desktop, f'設備保養基準書_{int(time.time())}.xlsx')
        wb.save(fallback)
        return jsonify({'success': True,
                        'message': f'原始檔案無法寫入（{type(e).__name__}，可能被開啟中），'
                                   f'已改存 {len(rows)} 個項目到桌面：{fallback}'})


# ══════════════════════════════════════════════════════════
#  設備保養排程（P2：展開 → 待辦 → 回報）
# ══════════════════════════════════════════════════════════
#  展開規則見 docs/equipment-maintenance.md 3.4／3.5：
#  ① 固定錨點、不順延——晚做只標「逾期完成」，不把整條時間軸往後推。
#     年度計畫甘特表要事先畫得出來、稽核要能對照，浮動週期做不到這件事。
#  ② 日／週不預先展開（189 台 × 10 項 × 365 天 ≈ 69 萬筆空待辦，會把網芳上的
#     SQLite 拖垮），清單上以「虛擬列」即時算出來，真的回報了才寫進 mt_plan。

_MT_PLAN_STATUSES = ('待辦', '已完成', '跳過')
_MT_PLAN_CYCLES = ('month', 'quarter', 'half', 'year')   # 會預先展開的週期
_MT_LAZY_CYCLES = ('day', 'week')                        # 只在清單視窗內即時算
_MT_STEP_MONTHS = {'month': 1, 'quarter': 3, 'half': 6, 'year': 12}
_MT_STEP_DAYS = {'day': 1, 'week': 7}
_MT_ACTIVE_STATUS = ('使用中', '閒置')   # 報廢／已售出／轉賣不展開保養


def _mt_date(s):
    """把各種寫法的日期字串轉成 date（2015/3/1、2015-03-01…），看不懂回 None"""
    s = (s or '').strip().replace('/', '-')
    m = re.match(r'^(\d{4})-(\d{1,2})-(\d{1,2})', s)
    if not m:
        return None
    try:
        return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    except ValueError:
        return None


def _mt_add_months(d, n):
    """加 n 個月；落在該月沒有的日子取當月最後一天（1/31 加一個月 = 2/28）"""
    import calendar as _cal_mod
    total = d.year * 12 + (d.month - 1) + n
    y, m = divmod(total, 12)
    m += 1
    return date(y, m, min(d.day, _cal_mod.monthrange(y, m)[1]))


def _mt_period_key(kind, d):
    """週期代碼 + 落點日期 → 期別鍵（同一期只會有一筆待辦）"""
    if kind == 'year':
        return f'{d.year}'
    if kind == 'half':
        return f'{d.year}-H{1 if d.month <= 6 else 2}'
    if kind == 'quarter':
        return f'{d.year}-Q{(d.month - 1) // 3 + 1}'
    if kind == 'month':
        return f'{d.year}-{d.month:02d}'
    if kind == 'week':
        iso = d.isocalendar()
        return f'{iso[0]}-W{iso[1]:02d}'
    return d.strftime('%Y-%m-%d')


def _mt_occurrences(anchor, kind, n, start, end):
    """從 anchor 起每 n 個週期一次，回傳落在 [start, end] 之間的日期。

    固定錨點：不管實際什麼時候做完，落點都是從 anchor 算出來的那幾天。"""
    n = max(1, int(n or 1))
    out = []
    if kind in _MT_STEP_DAYS:
        step = _MT_STEP_DAYS[kind] * n
        k = max(0, -(-(start - anchor).days // step))       # 向上取整，跳過 start 之前的期數
        d = anchor + timedelta(days=k * step)
        while d <= end:
            if d >= start:
                out.append(d)
            d += timedelta(days=step)
        return out
    months = _MT_STEP_MONTHS.get(kind, 12) * n
    # 先估一個接近 start 的期數再往回退兩期，避免從 anchor 一步一步走（採購日可能是 20 年前）
    k = max(0, ((start.year * 12 + start.month) - (anchor.year * 12 + anchor.month)) // months - 2)
    while True:
        d = _mt_add_months(anchor, k * months)
        if d > end:
            break
        if d >= start:
            out.append(d)
        k += 1
    return out


def _mt_shift_workday(d, ov, limit=14):
    """落在非工作日就往後推到最近的工作日（ov = 預先載好的行事曆例外表）。

    limit 是保險絲：行事曆若被設成整段連假，最多推 14 天就放棄，不要無限迴圈。"""
    for i in range(limit + 1):
        x = d + timedelta(days=i)
        row = ov.get(x.strftime('%Y-%m-%d'))
        if (row['kind'] if row else _cal_default_kind(x)) == 'work':
            return x
    return d


def _mt_load_plan_context(conn):
    """一次把展開需要的東西全部載進記憶體（設備／基準書／項目／單機覆寫）。

    網芳上的 SQLite 最怕 N+1 查詢，189 台設備逐台查會慢到不能用。"""
    equips = [dict(r) for r in conn.execute(
        'SELECT code, group_code, type_code, status, buy_date FROM equipment '
        'WHERE status IN (%s)' % ','.join('?' * len(_MT_ACTIVE_STATUS)), _MT_ACTIVE_STATUS)]
    std_by_type, std_by_code = {}, {}
    for r in conn.execute('SELECT * FROM mt_standard'):
        d = dict(r)
        if d['scope'] == 'type':
            std_by_type[(d['group_code'], d['type_code'])] = d
        else:
            std_by_code[d['code']] = d
    items_by_std = {}
    for r in conn.execute('SELECT * FROM mt_item WHERE IFNULL(active,1)=1'):
        items_by_std.setdefault(r['std_id'], []).append(dict(r))
    ov_by_code = {}
    for r in conn.execute('SELECT * FROM mt_equip_item'):
        ov_by_code.setdefault(r['code'], {})[r['item_id']] = dict(r)
    return equips, std_by_type, std_by_code, items_by_std, ov_by_code


def _mt_equip_plan_items(eq, std_by_type, std_by_code, items_by_std, ov_by_code):
    """某台設備實際要做的保養項目（範本 + 單機專屬，套用覆寫、剔除被停用的）。

    回傳 [(item, cycle_kind, cycle_n, anchor_date)]；anchor 依序取
    單機覆寫 → 項目本身 → 設備採購日 → 今天。"""
    out = []
    ov = ov_by_code.get(eq['code'], {})
    srcs = [std_by_type.get((eq['group_code'], eq['type_code'])), std_by_code.get(eq['code'])]
    for std in srcs:
        if not std or (std.get('status') or '') == '停用':
            continue
        for it in items_by_std.get(std['id'], []):
            o = ov.get(it['id']) or {}
            if o.get('disabled'):
                continue
            kind = o.get('cycle_kind') or it['cycle_kind']
            n = o.get('cycle_n') or it['cycle_n'] or 1
            anchor = (_mt_date(o.get('anchor')) or _mt_date(it.get('anchor'))
                      or _mt_date(eq.get('buy_date')) or date.today())
            out.append((it, kind, n, anchor))
    return out


@app.route('/api/equipment_master/maint/plan/expand', methods=['POST'])
def maint_plan_expand():
    """把基準書展開成保養待辦（可重複執行，只補缺的）。

    - 只展開 `使用中`／`閒置` 的設備，`停用` 的基準書整份跳過
    - 只展開月以上的週期（日／週見上面的說明）
    - 視窗：今天往前 `back_days` 天（讓剛過期的也建得出來）到往後 `months` 個月。
      **刻意不從採購日開始補完整歷史**——那會一次生出好幾萬筆從來沒人做過的「逾期」，
      清單直接不能看，而且那些過去根本沒發生的保養也不該被記成漏做。
    - 已經存在的期別只更新 `待辦` 那些的 due_date（有人改了週期就跟著修正），
      已完成／跳過的一律不動
    - 已經不適用的 `待辦`（項目被刪、被某台停用、設備報廢）會被清掉，避免留下殭屍待辦
    """
    d = request.get_json(silent=True) or {}
    try:
        months = min(24, max(1, int(d.get('months') or 12)))
    except (TypeError, ValueError):
        months = 12
    back_days = 30
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        today = date.today()
        win_start, win_end = today - timedelta(days=back_days), _mt_add_months(today, months)
        ov_cal = _cal_overrides(win_start, win_end + timedelta(days=20))
        equips, std_by_type, std_by_code, items_by_std, ov_by_code = _mt_load_plan_context(conn)

        now = _mt_now()
        valid_pairs, rows, eq_hit = set(), [], set()
        for eq in equips:
            for it, kind, n, anchor in _mt_equip_plan_items(
                    eq, std_by_type, std_by_code, items_by_std, ov_by_code):
                valid_pairs.add((eq['code'], it['id']))
                if kind not in _MT_PLAN_CYCLES:
                    continue
                for raw in _mt_occurrences(anchor, kind, n, win_start, win_end):
                    due = _mt_shift_workday(raw, ov_cal)
                    rows.append((eq['code'], it['id'], due.strftime('%Y-%m-%d'),
                                 _mt_period_key(kind, raw), now, now))
                    eq_hit.add(eq['code'])

        before = conn.execute("SELECT COUNT(*) FROM mt_plan WHERE status='待辦'").fetchone()[0]
        conn.executemany(
            "INSERT INTO mt_plan (code, item_id, due_date, period_key, status, created_at, updated_at)"
            " VALUES (?,?,?,?,'待辦',?,?)"
            " ON CONFLICT(code, item_id, period_key) DO UPDATE SET due_date=excluded.due_date,"
            " updated_at=excluded.updated_at WHERE mt_plan.status='待辦'", rows)
        # 清掉不再適用的待辦（項目刪了／某台停用了／設備報廢了）
        removed = 0
        for r in conn.execute("SELECT id, code, item_id FROM mt_plan WHERE status='待辦'").fetchall():
            if (r['code'], r['item_id']) not in valid_pairs:
                conn.execute('DELETE FROM mt_plan WHERE id=?', (r['id'],))
                removed += 1
        conn.commit()
        after = conn.execute("SELECT COUNT(*) FROM mt_plan WHERE status='待辦'").fetchone()[0]
    finally:
        conn.close()
    return jsonify({'success': True, 'message':
                    f'展開完成：{len(eq_hit)} 台設備、待辦 {before} → {after} 筆'
                    + (f'（清掉 {removed} 筆已不適用）' if removed else ''),
                    'equips': len(eq_hit), 'todo': after, 'removed': removed,
                    'range': [win_start.strftime('%Y-%m-%d'), win_end.strftime('%Y-%m-%d')]})


def _mt_plan_window(args):
    """把前端的 scope 參數換算成查詢區間"""
    today = date.today()
    scope = (args.get('scope') or 'due').strip()
    if scope == 'range':
        s = _mt_date(args.get('start')) or today
        e = _mt_date(args.get('end')) or (s + timedelta(days=30))
        return scope, s, e
    if scope == 'month':
        ym = _mt_date((args.get('ym') or today.strftime('%Y-%m')) + '-01') or today
        return scope, date(ym.year, ym.month, 1), _mt_add_months(date(ym.year, ym.month, 1), 1) - timedelta(days=1)
    if scope == 'all':
        return scope, today - timedelta(days=3650), today + timedelta(days=3650)
    # 預設：逾期 + 未來 7 天（現場真正會問的是「現在該做什麼」）
    return 'due', today - timedelta(days=3650), today + timedelta(days=7)


@app.route('/api/equipment_master/maint/plan/list')
def maint_plan_list():
    """保養待辦清單。日／週週期的項目在視窗內以虛擬列即時算出來（還沒進資料庫），
    真的回報之後才變成 mt_plan 的一筆——虛擬列的 id 是 null，前端用 code+item_id+period_key 回報。"""
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    today = date.today()
    scope, win_s, win_e = _mt_plan_window(request.args)
    status_f = (request.args.get('status') or '').strip()
    # 手機版只要這一台（現場網路不好，不該把全廠的待辦都抓下來）
    only = (request.args.get('code') or '').strip().upper()
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT p.*, i.part, i.name AS item_name, i.method, i.criteria, i.cycle_kind, i.cycle_n,
                   i.need_stop, i.owner, i.tools, i.oil_code, i.oil, i.duration_min, i.attach,
                   e.group_code, e.type_code, e.old_code, e.location, e.status AS eq_status,
                   t.name AS type_name
              FROM mt_plan p
              JOIN mt_item i ON i.id = p.item_id
              JOIN equipment e ON e.code = p.code
              LEFT JOIN eq_type t ON t.group_code = e.group_code AND t.code = e.type_code
             WHERE p.due_date BETWEEN ? AND ? AND (?='' OR p.code=?)
             ORDER BY p.due_date, p.code""",
            (win_s.strftime('%Y-%m-%d'), win_e.strftime('%Y-%m-%d'), only, only))]

        # 日／週的虛擬列：視窗**獨立算**、而且一定要夾在 31 天內。
        # 不能直接用查詢區間——scope='due' 為了撈到所有逾期待辦會往回抓 10 年，
        # 拿那個區間去算日檢就是一天一筆生上千列（2026-08-11 實測踩到，當時虛擬列直接被跳過）。
        lazy_s = max(win_s, today - timedelta(days=14))
        lazy_e = min(win_e, today + timedelta(days=16))
        if scope == 'due':
            # 預設視圖回答的是「現在該做什麼」：日檢只看今天、週檢只看本週。
            # 前天沒做的日常點檢不會有人回頭補做，列出來只會把真正該處理的
            # 逾期月保養埋掉（實測 15 台 × 14 天 = 229 筆虛擬列洗版）。
            day_win = (today, today)
            week_win = (today - timedelta(days=today.weekday()),
                        today + timedelta(days=6 - today.weekday()))
        else:
            day_win = week_win = (lazy_s, lazy_e)
        if lazy_s <= lazy_e:
            have = {(r['code'], r['item_id'], r['period_key']) for r in rows}
            ov_cal = _cal_overrides(lazy_s, lazy_e + timedelta(days=20))
            equips, sbt, sbc, ibs, ovc = _mt_load_plan_context(conn)
            names = {r['code']: dict(r) for r in conn.execute("""
                SELECT e.code, e.group_code, e.type_code, e.old_code, e.location,
                       e.status AS eq_status, t.name AS type_name
                  FROM equipment e
                  LEFT JOIN eq_type t ON t.group_code = e.group_code AND t.code = e.type_code""")}
            for eq in equips:
                if only and eq['code'] != only:
                    continue
                for it, kind, n, anchor in _mt_equip_plan_items(eq, sbt, sbc, ibs, ovc):
                    if kind not in _MT_LAZY_CYCLES:
                        continue
                    w = day_win if kind == 'day' else week_win
                    for raw in _mt_occurrences(anchor, kind, n, w[0], w[1]):
                        key = _mt_period_key(kind, raw)
                        if (eq['code'], it['id'], key) in have:
                            continue
                        due = _mt_shift_workday(raw, ov_cal)
                        rows.append(dict(names.get(eq['code'], {}),
                                         id=None, code=eq['code'], item_id=it['id'],
                                         due_date=due.strftime('%Y-%m-%d'), period_key=key,
                                         status='待辦', done_date=None, done_by=None, result=None,
                                         minutes=None, remark=None,
                                         part=it['part'], item_name=it['name'], method=it['method'],
                                         criteria=it['criteria'], cycle_kind=kind, cycle_n=n,
                                         need_stop=it['need_stop'], owner=it['owner'],
                                         tools=it['tools'], oil_code=it['oil_code'], oil=it['oil'],
                                         duration_min=it['duration_min'], attach=it['attach']))
        _mt_fill_oil(rows)
    finally:
        conn.close()

    for r in rows:
        due = _mt_date(r['due_date'])
        r['overdue_days'] = (today - due).days if (due and r['status'] == '待辦' and due < today) else 0
        r['is_today'] = bool(due and due == today)
        r['cycle_label'] = _mt_cycle_label(r['cycle_kind'], r['cycle_n'])
        r['eq_name'] = ' '.join(x for x in [r.get('type_name') or '', r.get('old_code') or ''] if x)
    if scope == 'due':
        # 已完成的舊資料不要洗版，但**今天剛回報的一定要留著**——
        # 按了完成結果那列直接消失，會讓人以為沒存進去。
        ts = today.strftime('%Y-%m-%d')
        rows = [r for r in rows if r['status'] == '待辦' or r.get('done_date') == ts
                or (_mt_date(r['due_date']) or today) >= today - timedelta(days=7)]
    if status_f:
        rows = [r for r in rows if r['status'] == status_f]
    rows.sort(key=lambda r: (r['due_date'], r['code']))
    return jsonify({'success': True, 'data': rows,
                    'range': [win_s.strftime('%Y-%m-%d'), win_e.strftime('%Y-%m-%d')],
                    'today': today.strftime('%Y-%m-%d')})


@app.route('/api/equipment_master/maint/plan/stats')
def maint_plan_stats():
    """到期提醒用的數字：逾期／今天／本週／本月，另附最近一次展開到哪一天"""
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': True, 'overdue': 0, 'today': 0, 'week': 0, 'month': 0})
    today = date.today()
    week_end = today + timedelta(days=6 - today.weekday())     # 到本週日
    month_end = _mt_add_months(date(today.year, today.month, 1), 1) - timedelta(days=1)
    ts = today.strftime('%Y-%m-%d')
    try:
        def cnt(sql, params):
            return conn.execute('SELECT COUNT(*) FROM mt_plan WHERE ' + sql, params).fetchone()[0]
        out = {
            'overdue': cnt("status='待辦' AND due_date < ?", (ts,)),
            'today':   cnt("status='待辦' AND due_date = ?", (ts,)),
            'week':    cnt("status='待辦' AND due_date BETWEEN ? AND ?",
                           (ts, week_end.strftime('%Y-%m-%d'))),
            'month':   cnt("status='待辦' AND due_date BETWEEN ? AND ?",
                           (ts, month_end.strftime('%Y-%m-%d'))),
            'done_month': cnt("status='已完成' AND done_date BETWEEN ? AND ?",
                              (date(today.year, today.month, 1).strftime('%Y-%m-%d'),
                               month_end.strftime('%Y-%m-%d'))),
            'ng_month': cnt("result='NG' AND done_date BETWEEN ? AND ?",
                            (date(today.year, today.month, 1).strftime('%Y-%m-%d'),
                             month_end.strftime('%Y-%m-%d'))),
        }
        last = conn.execute("SELECT MAX(due_date) FROM mt_plan WHERE status='待辦'").fetchone()[0]
        out['expanded_to'] = last or ''
    finally:
        conn.close()
    out['success'] = True
    return jsonify(out)


@app.route('/api/equipment_master/maint/plan/done', methods=['POST'])
def maint_plan_done():
    """回報保養結果（OK／NG／跳過）。

    日／週的虛擬列沒有 id，帶 code+item_id+period_key 進來時在這裡才真的寫進 mt_plan
    （lazy 建立，見檔頭說明）。

    **只有月以上週期或 NG 才寫 eq_history**：日／週點檢每天 189 筆會把移轉／狀態變更
    這些真正重要的歷程洗掉（同 `_EQ_TRACKED_FIELDS` 刻意不追蹤備註的理由）。
    """
    d = request.get_json(silent=True) or {}
    pid = d.get('id')
    code = (d.get('code') or '').strip().upper()
    item_id = d.get('item_id')
    period = (d.get('period_key') or '').strip()
    status = (d.get('status') or '已完成').strip()
    if status not in _MT_PLAN_STATUSES:
        return jsonify({'success': False, 'error': f'不支援的狀態 {status}'}), 400
    result = (d.get('result') or '').strip().upper()
    if result and result not in ('OK', 'NG'):
        return jsonify({'success': False, 'error': '判定只能是 OK 或 NG'}), 400
    done_date = (d.get('done_date') or '').strip() or date.today().strftime('%Y-%m-%d')
    done_by = (d.get('done_by') or '').strip()
    remark = (d.get('remark') or '').strip()
    photo = (d.get('photo') or '').strip()      # 手機現場拍的照片檔名（多張用 | 分隔）
    try:
        minutes = float(d.get('minutes')) if str(d.get('minutes') or '').strip() else None
    except (TypeError, ValueError):
        minutes = None
    if status == '已完成' and not result:
        result = 'OK'

    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        now = _mt_now()
        if pid:
            row = conn.execute('SELECT * FROM mt_plan WHERE id=?', (pid,)).fetchone()
            if row is None:
                return jsonify({'success': False, 'error': '找不到這筆保養工作'}), 404
            code, item_id, period = row['code'], row['item_id'], row['period_key']
            due_date = row['due_date']
        else:
            if not (code and item_id and period):
                return jsonify({'success': False, 'error': '參數不足'}), 400
            due_date = (d.get('due_date') or done_date)
        conn.execute(
            'INSERT INTO mt_plan (code, item_id, due_date, period_key, status, done_date, done_by,'
            ' result, minutes, remark, photo, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)'
            ' ON CONFLICT(code, item_id, period_key) DO UPDATE SET status=excluded.status,'
            ' done_date=excluded.done_date, done_by=excluded.done_by, result=excluded.result,'
            ' minutes=excluded.minutes, remark=excluded.remark, photo=excluded.photo,'
            ' updated_at=excluded.updated_at',
            (code, item_id, due_date, period, status, done_date, done_by, result, minutes,
             remark, photo, now, now))
        it = conn.execute('SELECT name, part, cycle_kind, cycle_n FROM mt_item WHERE id=?',
                          (item_id,)).fetchone()
        logged = False
        if it is not None and status == '已完成' and (it['cycle_kind'] in _MT_PLAN_CYCLES or result == 'NG'):
            label = _mt_cycle_label(it['cycle_kind'], it['cycle_n'])
            detail = f'{label}保養：{it["name"]}'
            if it['part']:
                detail = f'{label}保養：[{it["part"]}] {it["name"]}'
            detail += f'（判定 {result}）' if result else ''
            if remark:
                detail += f' {remark}'
            conn.execute("INSERT INTO eq_history (code, date, action, detail, user)"
                         " VALUES (?,?,'保養',?,'system')", (code, done_date, detail))
            logged = True
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'logged_history': logged})


@app.route('/api/equipment_master/maint/plan/undo', methods=['POST'])
def maint_plan_undo():
    """把已回報的保養退回待辦（填錯時用）。

    連同那筆自動寫進 eq_history 的保養記錄一起刪掉——留著會變成「歷程說做了、
    排程說沒做」的矛盾資料。只刪 `user='system'` 且同一天同一台的那筆，
    不會動到人工補登的歷程。"""
    d = request.get_json(silent=True) or {}
    pid = d.get('id')
    if not pid:
        return jsonify({'success': False, 'error': '缺少 id'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        row = conn.execute('SELECT * FROM mt_plan WHERE id=?', (pid,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': '找不到這筆保養工作'}), 404
        it = conn.execute('SELECT name FROM mt_item WHERE id=?', (row['item_id'],)).fetchone()
        if it is not None and row['done_date']:
            conn.execute("DELETE FROM eq_history WHERE code=? AND date=? AND action='保養'"
                         " AND user='system' AND detail LIKE ?",
                         (row['code'], row['done_date'], f'%{it["name"]}%'))
        conn.execute("UPDATE mt_plan SET status='待辦', done_date=NULL, done_by=NULL, result=NULL,"
                     ' minutes=NULL, remark=NULL, updated_at=? WHERE id=?', (_mt_now(), pid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


# ══════════════════════════════════════════════════════════
#  手機掃碼回報（P3）  詳見 docs/equipment-maintenance.md 第六節
# ══════════════════════════════════════════════════════════
#  設備上貼 QR（內容是 http://<服務端>:5088/m/eq/<PDM編碼>），現場用手機相機掃了
#  直接開回報頁。要能用得先把 config.MOBILE_ACCESS 打開（預設關閉，見 _flask_host）。

@app.route('/m')
@app.route('/m/')
def mobile_home():
    """手機版首頁：搜尋設備（掃不到碼或標籤掉了的時候用）"""
    return render_template('m_equip.html', code='', app_version=APP_VERSION)


@app.route('/m/eq/<code>')
def mobile_equip(code):
    """手機版設備保養回報頁（QR 掃進來的落點）"""
    return render_template('m_equip.html', code=(code or '').strip().upper(),
                           app_version=APP_VERSION)


@app.route('/equipment_master/labels')
def equipment_labels_page():
    """設備 QR 標籤列印頁（貼在機台上給現場掃）"""
    return render_template('equipment_labels.html', app_version=APP_VERSION,
                           base_url=getattr(config, 'MOBILE_BASE_URL', '') or '')


@app.route('/api/equipment_master/maint/mobile')
def maint_mobile_detail():
    """手機版要的全部資料，一次給完（現場網路不好，來回越少越好）：
    設備基本資料 + 封面照片 + 這台現在該做的保養。

    **只回這一台**——手機不該把全廠 343 筆待辦抓下來。"""
    code = (request.args.get('code') or '').strip().upper()
    if not code:
        return jsonify({'success': False, 'error': '缺少設備編碼'}), 400
    conn = _mt_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '設備索引尚未建立'}), 500
    try:
        eq = conn.execute("""
            SELECT e.*, g.name AS group_name, t.name AS type_name, a.name AS attr_name
              FROM equipment e
              LEFT JOIN eq_group g ON g.code = e.group_code
              LEFT JOIN eq_type  t ON t.group_code = e.group_code AND t.code = e.type_code
              LEFT JOIN eq_attr  a ON a.group_code = e.group_code AND a.type_code = e.type_code
                                  AND a.code = e.attr_code
             WHERE e.code = ?""", (code,)).fetchone()
        if eq is None:
            return jsonify({'success': False, 'error': f'查無設備 {code}'}), 404
        photo = conn.execute(
            'SELECT relpath FROM eq_photo WHERE code=? ORDER BY is_cover DESC, sort LIMIT 1',
            (code,)).fetchone()
    finally:
        conn.close()
    out = dict(eq)
    out['photo'] = photo['relpath'] if photo else ''
    return jsonify({'success': True, 'data': out})


@app.route('/api/equipment_master/maint/mobile/search')
def maint_mobile_search():
    """手機版搜尋設備（編碼／舊編號／類型關鍵字），只回使用中與閒置的"""
    q = (request.args.get('q') or '').strip()
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': True, 'data': []})
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT e.code, e.old_code, e.location, t.name AS type_name
              FROM equipment e
              LEFT JOIN eq_type t ON t.group_code = e.group_code AND t.code = e.type_code
             WHERE e.status IN ('使用中','閒置') ORDER BY e.code""")]
    finally:
        conn.close()
    if q:
        keys = [k for k in q.upper().split() if k]
        rows = [r for r in rows if all(
            k in ' '.join(str(r.get(f) or '') for f in ('code', 'old_code', 'type_name')).upper()
            for k in keys)]
    return jsonify({'success': True, 'data': rows[:60]})


@app.route('/api/equipment_master/maint/plan/photo', methods=['POST'])
def maint_plan_photo():
    """保養回報時拍的照片，存到網芳 EQUIPMENT_MAINT_PHOTO_ROOT\\<設備編碼>\\。

    回報照片跟設備照片刻意分開放：設備照片是「這台長什麼樣」，回報照片是「那天現場的狀況」，
    混在一起會把設備相簿灌爆（`scan_photos()` 也會把它們當成設備照片收進 eq_photo）。"""
    code = (request.form.get('code') or '').strip().upper()
    files = request.files.getlist('files')
    if not code:
        return jsonify({'success': False, 'error': '缺少設備編碼'}), 400
    if not files or not any(f and f.filename for f in files):
        return jsonify({'success': False, 'error': '請選擇照片'}), 400
    root = getattr(config, 'EQUIPMENT_MAINT_PHOTO_ROOT', '')
    dest_dir = _eq_safe_path(root, code) if root else None
    if not dest_dir:
        return jsonify({'success': False, 'error': '保養回報照片資料夾未設定'}), 500
    try:
        os.makedirs(dest_dir, exist_ok=True)
    except OSError as e:
        return jsonify({'success': False, 'error': f'無法建立資料夾：{e}'}), 500
    saved, rejected = [], []
    for f in files:
        if not f or not f.filename:
            continue
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in _EQ_PHOTO_EXT:
            rejected.append(f.filename)
            continue
        filename = _cnc_safe_filename(f'{code}_{int(time.time() * 1000)}{ext}')
        try:
            f.save(os.path.join(dest_dir, filename))
        except OSError as e:
            rejected.append(f'{f.filename}（存檔失敗：{e}）')
            continue
        saved.append(filename)
    if not saved:
        return jsonify({'success': False,
                        'error': '沒有成功上傳的照片' + ('：' + '、'.join(rejected) if rejected else '')}), 400
    return jsonify({'success': True, 'saved': saved, 'folder': code})


@app.route('/api/equipment_master/maint/plan/photo/get')
def maint_plan_photo_get():
    """讀回報照片"""
    from flask import send_file
    full = _eq_safe_path(getattr(config, 'EQUIPMENT_MAINT_PHOTO_ROOT', ''),
                         os.path.join(request.args.get('code', ''), request.args.get('f', '')))
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404
    return send_file(full)


# ══════════════════════════════════════════════════════════
#  CNC 程式管理
# ══════════════════════════════════════════════════════════

def _cnc_safe_path(relpath):
    """將相對路徑轉成絕對路徑，並確保仍在 CNC_PROGRAM_ROOT_PATH 底下（防止路徑跳脫）"""
    root = os.path.normpath(config.CNC_PROGRAM_ROOT_PATH)
    full = os.path.normpath(os.path.join(root, relpath or ''))
    if full.lower() != root.lower() and not full.lower().startswith(root.lower() + os.sep):
        return None
    return full


def _cnc_safe_filename(name):
    """保留中文字元，只移除路徑分隔符與 Windows 不允許的字元"""
    name = os.path.basename(name or '').strip()
    name = re.sub(r'[\\/:*?"<>|]', '_', name)
    return name or 'upload.txt'


@app.route('/cnc_program')
def cnc_program_page():
    """CNC 程式管理頁面"""
    return render_template('cnc_program.html', app_version=APP_VERSION)


@app.route('/zumen')
def zumen_page():
    """ZUMEN 線上圖面管理頁面（內嵌本機 Node 工具）"""
    return render_template('zumen.html', app_version=APP_VERSION)


@app.route('/api/cnc_program/search')
def cnc_program_search():
    """搜尋 CNC 程式索引（空格=AND，-前綴=NOT，比對系列/品號/機台/檔名/備註）"""
    if not os.path.exists(CNC_DB_PATH):
        return jsonify({'success': False, 'error': '索引尚未建立，請先按「重建索引」'}), 500

    q = request.args.get('q', '').strip()
    conn = sqlite3.connect(CNC_DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        rows = [dict(r) for r in conn.execute(
            'SELECT * FROM cnc_program_index ORDER BY mtime DESC').fetchall()]
    finally:
        conn.close()

    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def haystack(r):
            return ' '.join([r.get('top_folder') or '', r.get('model') or '',
                              r.get('machine') or '', r.get('filename') or '',
                              r.get('remark') or '']).lower()

        rows = [r for r in rows if all(m in haystack(r) for m in must)
                and not any(m in haystack(r) for m in must_not)]

    truncated = len(rows) > 300
    return jsonify({'success': True, 'count': len(rows), 'truncated': truncated,
                    'data': rows[:300]})


@app.route('/api/cnc_program/tree')
def cnc_program_tree():
    """回傳系列→品號→機台的階層結構，供上傳/新增範本選單使用"""
    if not os.path.exists(CNC_DB_PATH):
        return jsonify({'success': True, 'tree': {}})

    conn = sqlite3.connect(CNC_DB_PATH)
    try:
        rows = conn.execute(
            'SELECT DISTINCT top_folder, model, machine FROM cnc_program_index').fetchall()
    finally:
        conn.close()

    tree = {}
    for top, model, machine in rows:
        top = top or ''
        model = model or ''
        machine = machine or ''
        bucket = tree.setdefault(top, {})
        if model:
            machines = bucket.setdefault(model, set())
            if machine:
                machines.add(machine)
    result = {t: {m: sorted(ms) for m, ms in models.items()} for t, models in tree.items()}
    return jsonify({'success': True, 'tree': result})


_cnc_reindex_state = {
    'running': False,
    'phase':   'idle',   # idle | scanning | done | error
    'count':   0,        # 目前已索引筆數（掃網路資料夾，事先不知道總數，無法算 %）
    'message': '',
    'error':   '',
}
_cnc_reindex_lock = _threading.Lock()


def _run_cnc_reindex():
    """背景執行 CNC 程式索引重建，解析 stdout 更新進度狀態。
    掃的是網路資料夾（\\\\192.168.1.99\\...），事先不知道檔案總數，
    只能顯示「目前已索引 N 筆」，不像治檢具/DCN 那樣能算出百分比。"""
    global _cnc_reindex_state
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_cnc_program_index.py')

    with _cnc_reindex_lock:
        _cnc_reindex_state.update(running=True, phase='scanning', count=0,
                                   message='啟動中...', error='')

    try:
        child_env = os.environ.copy()
        child_env['PYTHONIOENCODING'] = 'utf-8'
        # 子行程 stdout 接到管線（非終端機）時 Python 預設整批緩衝，進度輸出會卡到緩衝區滿
        # 或行程結束才一次送出，畫面看起來像卡住不動；設這個讓它逐行即時送出
        child_env['PYTHONUNBUFFERED'] = '1'
        proc = _subprocess.Popen(
            [sys.executable, script, '--deploy'],
            stdout=_subprocess.PIPE, stderr=_subprocess.STDOUT,
            text=True, encoding='utf-8', errors='replace',
            cwd=os.path.dirname(script), env=child_env
        )
        for line in proc.stdout:
            line = line.rstrip()
            if not line:
                continue
            with _cnc_reindex_lock:
                _cnc_reindex_state['message'] = line

            m = _re.search(r'已索引\s*(\d+)\s*筆', line)
            if m:
                with _cnc_reindex_lock:
                    _cnc_reindex_state['count'] = int(m.group(1))
                continue

            m = _re.search(r'完成！索引\s*(\d+)\s*筆', line)
            if m:
                with _cnc_reindex_lock:
                    _cnc_reindex_state['count'] = int(m.group(1))

        proc.wait()

        if proc.returncode == 0:
            with _cnc_reindex_lock:
                cnt = _cnc_reindex_state['count']
                _cnc_reindex_state.update(running=False, phase='done',
                                           message=f'更新完成，共 {cnt} 筆')
        else:
            with _cnc_reindex_lock:
                _cnc_reindex_state.update(running=False, phase='error',
                                           error='重建失敗，請查看伺服器日誌')
    except Exception as exc:
        with _cnc_reindex_lock:
            _cnc_reindex_state.update(running=False, phase='error', error=str(exc))


@app.route('/api/cnc_program/rebuild', methods=['POST'])
def cnc_program_rebuild():
    """啟動 CNC 程式索引重建（背景執行，前端輪詢 /api/cnc_program/rebuild/status 取得進度）"""
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_cnc_program_index.py')
    if not os.path.exists(script):
        return jsonify({'success': False, 'error': 'build_cnc_program_index.py 不存在'}), 500
    with _cnc_reindex_lock:
        if _cnc_reindex_state['running']:
            return jsonify({'success': False, 'error': '索引重建已在執行中，請稍候'}), 409
    t = _threading.Thread(target=_run_cnc_reindex, daemon=True)
    t.start()
    return jsonify({'success': True, 'message': '索引重建已啟動'})


@app.route('/api/cnc_program/rebuild/status', methods=['GET'])
def cnc_program_rebuild_status():
    """回傳 CNC 程式索引重建進度狀態"""
    with _cnc_reindex_lock:
        return jsonify(dict(_cnc_reindex_state))


@app.route('/api/cnc_program/view')
def cnc_program_view():
    """檢視程式內容（.txt 純文字 / 圖片轉 base64）"""
    full = _cnc_safe_path(request.args.get('relpath', ''))
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404

    ext = os.path.splitext(full)[1].lower()
    if ext in ('.jpg', '.jpeg', '.png', '.bmp'):
        import base64
        mime = 'image/jpeg' if ext in ('.jpg', '.jpeg') else f'image/{ext[1:]}'
        with open(full, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode('ascii')
        return jsonify({'success': True, 'type': 'image', 'data': f'data:{mime};base64,{b64}'})

    with open(full, 'rb') as f:
        raw = f.read()
    try:
        text = raw.decode('utf-8')
        enc = 'utf-8'
    except UnicodeDecodeError:
        text = raw.decode('cp950', errors='replace')
        enc = 'cp950'
    # encoding 一併回傳給前端，編輯存檔時要用同一種編碼寫回去，避免中文備註變亂碼
    return jsonify({'success': True, 'type': 'text', 'data': text, 'encoding': enc})


def _cnc_index_touch(relpath, full):
    """更新索引裡該檔案的 size/mtime（存檔後呼叫，索引不存在時忽略）"""
    try:
        conn = sqlite3.connect(CNC_DB_PATH)
        conn.execute('UPDATE cnc_program_index SET size=?, mtime=datetime("now","localtime") WHERE relpath=?',
                     (os.path.getsize(full), relpath))
        conn.commit()
        conn.close()
    except Exception:
        pass


@app.route('/api/cnc_program/save', methods=['POST'])
def cnc_program_save():
    """儲存編輯後的程式內容（覆寫原檔）。
    寫入前一定先把原檔備份到同層「已刪除」資料夾（時間戳記檔名），
    這是實際控制機台的 NC 程式，改壞了要能救回來，備份失敗就直接取消這次存檔。"""
    data = request.get_json(silent=True) or {}
    relpath = data.get('relpath', '')
    content = data.get('content', '')
    encoding = data.get('encoding') or 'utf-8'
    full = _cnc_safe_path(relpath)
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404

    import shutil
    parent = os.path.dirname(full)
    backup_dir = os.path.join(parent, '已刪除')
    try:
        os.makedirs(backup_dir, exist_ok=True)
        stem, ext = os.path.splitext(os.path.basename(full))
        backup_dest = os.path.join(backup_dir, f'{stem}_編輯前備份_{int(time.time())}{ext}')
        shutil.copy2(full, backup_dest)
    except OSError as e:
        return jsonify({'success': False, 'error': f'備份原檔失敗，已取消儲存（原檔未變動）：{e}'}), 500

    try:
        enc = 'cp950' if encoding == 'cp950' else 'utf-8'
        with open(full, 'w', encoding=enc, newline='') as f:
            f.write(content)
    except (OSError, UnicodeEncodeError) as e:
        return jsonify({'success': False, 'error': f'寫入失敗（原檔已備份於同層「已刪除」資料夾）：{e}'}), 500

    _cnc_index_touch(relpath, full)
    return jsonify({'success': True})


@app.route('/api/cnc_program/save_as', methods=['POST'])
def cnc_program_save_as():
    """另存新檔：內容存成新檔名，預設放在跟原檔同一個資料夾，分類資訊（系列/品號/機台）沿用原檔的索引。"""
    data = request.get_json(silent=True) or {}
    relpath = data.get('relpath', '')
    filename = (data.get('filename') or '').strip()
    content = data.get('content', '')
    encoding = data.get('encoding') or 'utf-8'
    full = _cnc_safe_path(relpath)
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '原始檔案不存在'}), 404
    if not filename:
        return jsonify({'success': False, 'error': '請輸入檔名'}), 400

    filename = _cnc_safe_filename(filename)
    folder = os.path.dirname(full)
    dest = os.path.join(folder, filename)
    if os.path.exists(dest):
        return jsonify({'success': False, 'error': f'檔案已存在：{filename}，請改檔名後再存'}), 409

    try:
        enc = 'cp950' if encoding == 'cp950' else 'utf-8'
        with open(dest, 'w', encoding=enc, newline='') as f:
            f.write(content)
    except (OSError, UnicodeEncodeError) as e:
        return jsonify({'success': False, 'error': f'寫入失敗：{e}'}), 500

    root = config.CNC_PROGRAM_ROOT_PATH
    new_relpath = os.path.relpath(dest, root).replace(os.sep, '/')

    try:
        conn = sqlite3.connect(CNC_DB_PATH)
        row = conn.execute(
            'SELECT top_folder, model, machine FROM cnc_program_index WHERE relpath=?', (relpath,)
        ).fetchone()
        top_folder, model, machine = row if row else ('', '', '')
        conn.execute(
            'INSERT OR REPLACE INTO cnc_program_index '
            '(top_folder, model, machine, filename, remark, relpath, ext, size, mtime) '
            'VALUES (?,?,?,?,?,?,?,?,datetime("now","localtime"))',
            (top_folder, model, machine, filename, '', new_relpath,
             os.path.splitext(filename)[1].lower(), os.path.getsize(dest))
        )
        conn.commit()
        conn.close()
    except Exception:
        pass

    return jsonify({'success': True, 'relpath': new_relpath, 'filename': filename})


@app.route('/api/cnc_program/download')
def cnc_program_download():
    """下載程式檔案"""
    from flask import send_file
    full = _cnc_safe_path(request.args.get('relpath', ''))
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404
    return send_file(full, as_attachment=True, download_name=os.path.basename(full))


@app.route('/api/cnc_program/upload', methods=['POST'])
def cnc_program_upload():
    """上傳程式檔到指定 系列/品號[/機台] 資料夾"""
    top_folder = request.form.get('top_folder', '').strip()
    model = request.form.get('model', '').strip()
    machine = request.form.get('machine', '').strip()
    f = request.files.get('file')

    if not f or not f.filename:
        return jsonify({'success': False, 'error': '請選擇檔案'}), 400
    if not top_folder or not model:
        return jsonify({'success': False, 'error': '請選擇系列與品號'}), 400

    root = config.CNC_PROGRAM_ROOT_PATH
    folder = os.path.join(root, top_folder, f'[{model}]')
    if machine:
        folder = os.path.join(folder, f'【{machine}】')

    try:
        os.makedirs(folder, exist_ok=True)
    except OSError as e:
        return jsonify({'success': False, 'error': f'無法建立資料夾：{e}'}), 500

    filename = _cnc_safe_filename(f.filename)
    dest = os.path.join(folder, filename)
    if os.path.exists(dest):
        return jsonify({'success': False,
                         'error': f'檔案已存在：{filename}，請改檔名後再上傳（例如加上日期）'}), 409

    f.save(dest)
    relpath = os.path.relpath(dest, root).replace(os.sep, '/')

    # 立即補進索引，不用等下次重建
    try:
        conn = sqlite3.connect(CNC_DB_PATH)
        conn.execute(
            'INSERT OR REPLACE INTO cnc_program_index '
            '(top_folder, model, machine, filename, remark, relpath, ext, size, mtime) '
            'VALUES (?,?,?,?,?,?,?,?,datetime("now","localtime"))',
            (top_folder, model, machine, filename, '', relpath,
             os.path.splitext(filename)[1].lower(), os.path.getsize(dest))
        )
        conn.commit()
        conn.close()
    except Exception:
        pass

    return jsonify({'success': True, 'relpath': relpath})


@app.route('/api/cnc_program/delete', methods=['POST'])
def cnc_program_delete():
    """軟刪除：移到同層的「已刪除」子資料夾，並從索引移除"""
    data = request.get_json(silent=True) or {}
    relpath = data.get('relpath', '')
    full = _cnc_safe_path(relpath)
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404

    import shutil
    parent = os.path.dirname(full)
    trash = os.path.join(parent, '已刪除')
    try:
        os.makedirs(trash, exist_ok=True)
        fn = os.path.basename(full)
        dest = os.path.join(trash, fn)
        if os.path.exists(dest):
            stem, ext = os.path.splitext(fn)
            dest = os.path.join(trash, f'{stem}_{int(time.time())}{ext}')
        shutil.move(full, dest)
    except OSError as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    try:
        conn = sqlite3.connect(CNC_DB_PATH)
        conn.execute('DELETE FROM cnc_program_index WHERE relpath=?', (relpath,))
        conn.commit()
        conn.close()
    except Exception:
        pass

    return jsonify({'success': True})


@app.route('/api/cnc_program/open_folder', methods=['POST'])
def cnc_program_open_folder():
    """用 Shell 開啟程式檔所在的資料夾"""
    data = request.get_json(silent=True) or {}
    relpath = data.get('relpath', '')
    full = _cnc_safe_path(relpath)
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404
    try:
        os.startfile(os.path.dirname(full))
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/cnc_program/new_model', methods=['POST'])
def cnc_program_new_model():
    """新增範本：在指定系列底下建立新品號資料夾，複製【空白範本】整套機台子資料夾"""
    data = request.get_json(silent=True) or {}
    top_folder = data.get('top_folder', '').strip()
    model = data.get('model', '').strip()
    if not top_folder or not model:
        return jsonify({'success': False, 'error': '請輸入系列與品號'}), 400

    root = config.CNC_PROGRAM_ROOT_PATH
    series_dir = os.path.join(root, top_folder)
    target = os.path.join(series_dir, f'[{model}]')
    if os.path.exists(target):
        return jsonify({'success': False, 'error': '此品號資料夾已存在'}), 409

    template = config.CNC_PROGRAM_TEMPLATE_PATH
    if not os.path.isdir(template):
        return jsonify({'success': False, 'error': '範本資料夾不存在：' + template}), 500

    import shutil
    try:
        os.makedirs(series_dir, exist_ok=True)
        shutil.copytree(template, target)
    except OSError as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    return jsonify({'success': True, 'path': target})


@app.route('/api/scrap/list')
def scrap_list():
    """讀取 Google 試算表『報廢統計』資料"""
    cache_key = 'scrap_list'
    if request.args.get('refresh'):
        cache_clear(cache_key)
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.SCRAP_SHEET_ID, gid=config.SCRAP_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    CATEGORY_MAP = {'A': 'CNC', 'B': '鑽床', 'C': '沖床', 'L': '車床'}
    scraps = []
    for r in rows:
        raw_amount = (r.get('報廢金額') or '').replace(',', '').strip()
        try:
            amount = float(raw_amount)
        except ValueError:
            amount = 0
        raw_qty = (r.get('出庫異動數量') or '').replace(',', '').strip()
        try:
            qty = float(raw_qty)
        except ValueError:
            qty = 0
        cat = (r.get('分類') or '').strip().upper()
        scraps.append({
            'date':     (r.get('異動日期') or '').strip(),
            'doc_no':   (r.get('單別-單號') or '').strip(),
            'part_no':  (r.get('品號') or '').strip(),
            'name':     (r.get('品名') or '').strip(),
            'spec':     (r.get('規格') or '').strip(),
            'dept':     (r.get('部門名稱') or '').strip(),
            'remark':   (r.get('備註') or '').strip(),
            'qty':      qty,
            'amount':   amount,
            'category': cat,
            'cat_name': CATEGORY_MAP.get(cat, cat),
        })

    def _sort_key(item):
        try:
            parts = item['date'].replace('-', '/').split('/')
            return tuple(int(p) for p in parts)
        except Exception:
            return (0, 0, 0)

    scraps.sort(key=_sort_key, reverse=True)

    result = {'success': True, 'count': len(scraps), 'scraps': scraps}
    cache_set(cache_key, result)
    return jsonify(result)


def _build_category_map():
    """讀取 K1_P2.ref 試算表，建立 {品號+製程: A/B/C} 對照表（含 Flask 快取）"""
    cache_key = 'category_map'
    cached = cache_get(cache_key)
    if cached is not None and cached.get('map'):
        return cached['map']
    try:
        cat_rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        cat_map = {}
        for r in cat_rows:
            part_no      = (r.get('品號')   or '').strip()
            process_code = (r.get('製程')   or '').strip()
            category     = (r.get('A/B/C分類') or r.get('A/B/C') or '').strip()
            if part_no and process_code and category:
                cat_map[part_no + process_code] = category
        # 空表視同失敗（試算表欄位變動或抓取異常），不快取，避免整批變未分類卡 TTL
        if not cat_map:
            return {}
        cache_set(cache_key, {'success': True, 'map': cat_map})
        return cat_map
    except Exception:
        return {}


@app.route('/api/prod_report/list')
def prod_report_list():
    """讀取 Google 試算表『生產日報表』資料（含 ABC 分類）"""
    cache_key = 'prod_report_list'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'prod_report_monthly_stats_v2', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, sheet_name=config.PROD_REPORT_SHEET_NAME)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    cat_map = _build_category_map()

    records = []
    for r in rows:
        raw_qty = (r.get('生產數') or '').replace(',', '').strip()
        try:    qty = float(raw_qty)
        except ValueError: qty = 0
        raw_sec = (r.get('秒數') or '').replace(',', '').strip()
        try:    seconds = float(raw_sec)
        except ValueError: seconds = 0
        part_no      = (r.get('品號') or '').strip()
        process_code = (r.get('製程代號') or '').strip()
        machine_name = (r.get('機台名稱') or '').strip()
        machine_code = (r.get('機台代號') or '').strip()

        # 分類判斷：① 試算表直接欄位 → ② 品號+製程對照表 → ③ 機台名稱/代號關鍵字
        cat = (r.get('A/B/C') or r.get('類別') or r.get('分類') or '').strip()
        if cat not in ('A', 'B', 'C', 'L'):
            cat = cat_map.get(part_no + process_code, '')
        if cat not in ('A', 'B', 'C', 'L'):
            kw = (machine_name + machine_code + process_code).upper()
            if 'CNC' in kw or '加工中心' in kw or '綜合加工' in kw:
                cat = 'A'
            elif '鑽' in kw:
                cat = 'B'
            elif '沖' in kw or '衝' in kw:
                cat = 'C'
            elif '車' in kw:
                cat = 'L'

        records.append({
            'std_time':     (r.get('標工') or '').strip(),
            'date':         (r.get('生產日期') or '').strip(),
            'operator':     (r.get('生產人員') or '').strip(),
            'work_order':   (r.get('製令') or '').strip(),
            'part_no':      part_no,
            'name':         (r.get('品名') or '').strip(),
            'process_code': process_code,
            'qty':          qty,
            'seconds':      seconds,
            'machine_code': machine_code,
            'machine_name': machine_name,
            'remark':       (r.get('備註') or '').strip(),
            'category':     cat,
        })

    def _sort_key(item):
        try:
            parts = item['date'].replace('-', '/').split('/')
            return tuple(int(p) for p in parts)
        except Exception:
            return (0, 0, 0)

    records.sort(key=_sort_key, reverse=True)
    result = {'success': True, 'count': len(records), 'records': records}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/sfcr06/list')
def sfcr06_list():
    """製令製程明細表（SFCR06）：直連 ERP SQL Server 查 SFCTA（製令途程交易檔），
    含投入/完成數量、標準人時、實際開完工日等 ERP 原生報表才有的明細，
    補上產品品號/品名/預計開工日/預計產量（沿用既有未完工製令清單，避免另外重查 ERP）。
    日期區間依「預計完工日」篩選（製令層級，但只看查詢線別自己的製程完工日取最大值，
    不含其他線別/外包——加工課的績效看加工課自己排到哪，不是整張單送到最後一關才算），
    細節見 docs/sfcr06.md「查詢區間依據」。"""
    date_from = request.args.get('date_from', '').replace('-', '')
    date_to = request.args.get('date_to', '').replace('-', '')
    keyword = request.args.get('keyword', '').strip()
    line = request.args.get('line', '').strip()

    if not date_from and not date_to and not keyword:
        today = date.today()
        date_from = today.replace(day=1).strftime('%Y%m%d')
        date_to = today.strftime('%Y%m%d')

    # 關鍵字比對「製令單號」還是「品號/品名」：SFCTA 本身沒有品號/品名欄位（那是後面
    # info_map 另外查出來補上的），SQL 階段只能比對製令單號。要連品號/品名一起搜，
    # 只能等補完品號/品名之後在 Python 端再篩一次——但這樣就得先把整段日期區間的資料
    # 都撈出來，沒有日期區間時（純關鍵字、搜全部時間）撈全表太貴，所以只有「有給日期區間」
    # 時才 defer 到後面比對品號/品名；沒給日期區間就維持原本 SQL 層只比對製令單號的做法。
    defer_keyword = bool(keyword) and bool(date_from or date_to)

    select_cols = """a.TA001, a.TA002, a.TA003, a.TA004,
               RTRIM(ISNULL(w.MW002,'')),
               RTRIM(ISNULL(a.TA005,'')), RTRIM(ISNULL(a.TA006,'')), RTRIM(ISNULL(a.TA007,'')),
               a.TA008, a.TA009, a.TA010, a.TA011,
               a.TA022, a.TA023, a.TA030, a.TA031, RTRIM(ISNULL(a.TA032,'')),
               RTRIM(ISNULL(a.TA024,''))"""
    from_join = "FROM SFCTA a LEFT JOIN CMSMW w ON RTRIM(w.MW001) = RTRIM(a.TA004)"

    try:
        conn = get_erp_conn()
        cur = conn.cursor()

        if line:
            # 有指定線別：預計完工日只看「這條線自己」的製程完工日取最大值，
            # 不看其他線別/外包（同一張單如果還要送別的線別加工，那是別條線的事）。
            # ①找候選製令號：只看這條線自己的 TA009 是否落在區間內——只要這條線
            #   自己的預計完工日落在區間，達到最大值那筆製程的 TA009 本身必定也
            #   落在區間內，這一步不會漏單。
            cand_sql = "SELECT DISTINCT a.TA001, a.TA002 FROM SFCTA a WHERE RTRIM(a.TA007) = ?"
            cand_params = [line]
            if date_from:
                cand_sql += ' AND a.TA009 >= ?'
                cand_params.append(date_from)
            if date_to:
                cand_sql += ' AND a.TA009 <= ?'
                cand_params.append(date_to)
            if keyword and not defer_keyword:
                cand_sql += " AND (a.TA001 + '-' + a.TA002 LIKE ?)"
                cand_params.append(f'%{keyword}%')
            cur.execute(cand_sql, cand_params)
            cand_oids = [(r[0], r[1]) for r in cur.fetchall()]

            # ②撈候選製令「這條線自己」的全部製程（同一張單同一條線可能有多道製程）
            rows = []
            CHUNK = 200
            for i in range(0, len(cand_oids), CHUNK):
                chunk = cand_oids[i:i + CHUNK]
                where = ' OR '.join('(a.TA001=? AND a.TA002=?)' for _ in chunk)
                params = [v for pair in chunk for v in pair] + [line]
                cur.execute(f"SELECT {select_cols} {from_join} WHERE ({where}) AND RTRIM(a.TA007)=?", params)
                rows.extend(cur.fetchall())

            order_max = {}
            for r in rows:
                oid, d9 = (r[0], r[1]), (r[9] or '')
                if d9:
                    order_max[oid] = max(order_max.get(oid, ''), d9)

            if date_from or date_to:
                def in_range(d9):
                    if date_from and d9 < date_from:
                        return False
                    if date_to and d9 > date_to:
                        return False
                    return True
                keep_oids = {oid for oid, d9 in order_max.items() if in_range(d9)}
            else:
                keep_oids = set(order_max.keys())  # 純關鍵字查詢：不篩日期，找到的候選單全留

            rows = [r for r in rows if (r[0], r[1]) in keep_oids]
            predicted_map = {f'{oid[0]}-{oid[1]}': order_max[oid] for oid in keep_oids}
        else:
            # 沒指定線別（全部線別）：無法定義「哪條線自己」的預計完工日，
            # 退回最單純的做法——直接用每筆製程自己的 TA009 篩選日期區間。
            sql = f"SELECT {select_cols} {from_join} WHERE 1=1"
            params = []
            if date_from:
                sql += ' AND a.TA009 >= ?'
                params.append(date_from)
            if date_to:
                sql += ' AND a.TA009 <= ?'
                params.append(date_to)
            if keyword and not defer_keyword:
                sql += " AND (a.TA001 + '-' + a.TA002 LIKE ?)"
                params.append(f'%{keyword}%')
            cur.execute(sql, params)
            rows = cur.fetchall()
            predicted_map = {}

        conn.close()
    except Exception as e:
        return jsonify({'success': False, 'error': f'ERP 資料庫查詢失敗：{e}'}), 500

    # 依 TA009 DESC, TA001/TA002/TA003 ASC 排序（等同原本的 SQL ORDER BY，
    # 用兩次 stable sort 達成：先排 tiebreak 欄位，再依 TA009 反向排序）
    rows.sort(key=lambda r: (r[0], r[1], r[2]))
    rows.sort(key=lambda r: r[9] or '', reverse=True)

    def fmt_date(s):
        s = (s or '').strip()
        return f'{s[:4]}/{s[4:6]}/{s[6:8]}' if len(s) == 8 else ''

    def fmt_hms(seconds):
        try:
            sec = int(float(seconds or 0))
        except (TypeError, ValueError):
            return ''
        if sec <= 0:
            return ''
        h, rem = divmod(sec, 3600)
        m, s = divmod(rem, 60)
        return f'{h}:{m:02d}:{s:02d}'

    data = []
    for r in rows:
        oid = f'{r[0]}-{r[1]}'
        data.append({
            '製令單號': oid,
            '加工順序': r[2],
            '製程代號': r[3],
            '製程名稱': r[4],
            '性質': {'1': '廠內製程', '2': '外包'}.get(r[5], r[5]),
            '線別/廠商代號': r[6],
            '線別/廠商名稱': r[7],
            '製程開工日': fmt_date(r[8]),
            '製程完工日': fmt_date(r[9]),
            '投入數量': float(r[10]) if r[10] else 0,
            '完成數量': float(r[11]) if r[11] else 0,
            '標準人時': fmt_hms(r[12]),
            '標準機時': fmt_hms(r[13]),
            '實際開工日': fmt_date(r[14]),
            '實際完工日': fmt_date(r[15]),
            # 完工碼統一轉大寫：ERP 原始資料裡有 19.5% 的已完工紀錄寫的是小寫 'y'
            # （2026-08-03 實測，加工課 13,225 筆 'Y' + 3,199 筆 'y'），前端十幾處判斷
            # 完工與否／延誤都用 === 'Y' 這種嚴格比對，不轉大寫會把這些單全部誤判成未完工
            '完工碼': (r[16] or '').upper(),
            '製程敘述': r[17],
        })

    # 補上產品品號/品名（先查未完工製令清單，已完工的舊製令改查 SFCTC 移轉單明細補齊）
    info_map = {}
    try:
        for u in fetch_all_unfinished('all'):
            oid = u.get('單別', '').strip()
            if oid and oid not in info_map:
                info_map[oid] = {
                    '產品品號': u.get('品號', '').strip(),
                    '產品品名': u.get('品名', '').strip(),
                    '預計開工日': u.get('製令預計開工日', '').strip(),
                    '預計產量': u.get('預計生產數', '').strip(),
                }
    except Exception:
        pass

    # SFCTC 分批 OR 查詢（每批 200 筆），移除 500 筆上限，避免全年查詢遺漏品名
    missing_oids = list({d['製令單號'] for d in data if d['製令單號'] not in info_map})
    if missing_oids:
        try:
            conn2 = get_erp_conn()
            cur2 = conn2.cursor()
            CHUNK = 200
            for i in range(0, len(missing_oids), CHUNK):
                chunk = missing_oids[i:i+CHUNK]
                pairs = [oid.split('-', 1) for oid in chunk]
                where = ' OR '.join('(TC004=? AND TC005=?)' for _ in pairs)
                params2 = [v for p in pairs for v in p]
                cur2.execute(
                    f"SELECT TC004, TC005, RTRIM(ISNULL(TC047,'')), RTRIM(ISNULL(TC048,'')) "
                    f"FROM SFCTC WHERE {where}", params2
                )
                for row in cur2.fetchall():
                    oid = f'{row[0]}-{row[1]}'
                    if oid not in info_map and row[2]:
                        info_map[oid] = {
                            '產品品號': row[2],
                            '產品品名': row[3].rstrip('|').strip(),
                            '預計開工日': '', '預計產量': '',
                        }
            conn2.close()
        except Exception:
            pass

    cat_map = _build_category_map()

    for d in data:
        info = info_map.get(d['製令單號'], {})
        d['產品品號'] = info.get('產品品號', '')
        d['產品品名'] = info.get('產品品名', '')
        d['預計開工日'] = info.get('預計開工日', '')
        d['預計產量'] = info.get('預計產量', '')
        # predicted_map：該線別自己的製程完工日取最大值（沒指定線別時查無此概念，留空）
        d['預計完工日'] = fmt_date(predicted_map.get(d['製令單號'], ''))
        # 分類：① 品號+製程對照表 → ② 製程名稱/代號關鍵字備援（同生產日報表慣例）
        cat = cat_map.get(d['產品品號'] + d['製程代號'], '')
        if cat not in ('A', 'B', 'C', 'L'):
            kw = (d['製程名稱'] + d['製程代號']).upper()
            if 'CNC' in kw or '加工中心' in kw or '綜合加工' in kw:
                cat = 'A'
            elif '鑽' in kw:
                cat = 'B'
            elif '沖' in kw or '衝' in kw:
                cat = 'C'
            elif '車' in kw:
                cat = 'L'
            else:
                cat = ''
        d['分類'] = cat

    if defer_keyword:
        # SQL 階段沒篩關鍵字（見上面 defer_keyword 說明），這裡補上；
        # 品號/品名要等這裡才有值，所以連同製令單號一起比對，關鍵字才搜得到「FTB150」這種品號
        kw = keyword.lower()
        data = [d for d in data if kw in d['製令單號'].lower()
                or kw in d['產品品號'].lower() or kw in d['產品品名'].lower()]

    return jsonify({'success': True, 'count': len(data), 'data': data})


@app.route('/api/prod_report/monthly_stats')
def prod_report_monthly_stats():
    """生產日報表圖表資料：讀『P5.3生產日報表data_ref』分頁，
    依 A/B/C/L 分類 × 年月 彙整生產數與實際產出秒數（秒數欄=該筆總秒數，
    資料來源與『生產報工統計P2』的出站數量不同）"""
    cache_key = 'prod_report_monthly_stats_v2'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'prod_report_list', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.PROD_REPORT_SHEET_ID, sheet_name=config.PROD_REPORT_SHEET_NAME)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表，請確認網路狀態'}), 502

    cat_map = _build_category_map()
    from collections import defaultdict
    month_data = defaultdict(lambda: {
        'A': 0, 'B': 0, 'C': 0, 'L': 0, 'other': 0,
        'sec_A': 0.0, 'sec_B': 0.0, 'sec_C': 0.0, 'sec_L': 0.0, 'sec_other': 0.0,
        'records': 0, 'uncat': 0,
    })

    for r in rows:
        date_str = (r.get('生產日期') or '').strip()
        if not date_str:
            continue
        clean = date_str.split(' ')[0].replace('-', '/')
        parts = clean.split('/')
        if len(parts) < 2:
            continue
        try:
            ym = f'{int(parts[0])}-{int(parts[1]):02d}'
        except Exception:
            continue

        raw_qty = (r.get('生產數') or '').replace(',', '').strip()
        try:
            qty = int(float(raw_qty)) if raw_qty else 0
        except ValueError:
            qty = 0

        raw_sec = (r.get('秒數') or '').replace(',', '').strip()
        try:
            sec = float(raw_sec) if raw_sec else 0.0
        except ValueError:
            sec = 0.0

        part_no      = (r.get('品號') or '').strip()
        process_code = (r.get('製程代號') or '').strip()
        machine_code = (r.get('機台代號') or '').strip()
        machine_name = (r.get('機台名稱') or '').strip()

        # 分類：① 品號+製程對照表 → ② 機台名稱/代號/製程關鍵字
        cat = cat_map.get(part_no + process_code, '')
        if cat not in ('A', 'B', 'C', 'L'):
            kw = (machine_name + machine_code + process_code).upper()
            if 'CNC' in kw or '加工中心' in kw or '綜合加工' in kw:
                cat = 'A'
            elif '鑽' in kw:
                cat = 'B'
            elif '沖' in kw or '衝' in kw:
                cat = 'C'
            elif '車' in kw:
                cat = 'L'

        d = month_data[ym]
        if cat in ('A', 'B', 'C', 'L'):
            if qty > 0:
                d[cat]       += qty
                d['records'] += 1
            if sec > 0:
                d[f'sec_{cat}'] += sec
        else:
            d['uncat'] += 1
            if qty > 0:
                d['other'] += qty
            if sec > 0:
                d['sec_other'] += sec

    months = []
    for ym in sorted(month_data.keys()):
        d = month_data[ym]
        classified = d['A'] + d['B'] + d['C'] + d['L']
        months.append({
            'month': ym,
            'A': d['A'], 'B': d['B'], 'C': d['C'], 'L': d['L'],
            'other': d['other'],
            'total': classified + d['other'],
            'sec_A': round(d['sec_A']), 'sec_B': round(d['sec_B']),
            'sec_C': round(d['sec_C']), 'sec_L': round(d['sec_L']),
            'sec_other': round(d['sec_other']),
            'records': d['records'],
            'uncat': d['uncat'],
        })

    result = {'success': True, 'months': months}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/category/map')
def category_map():
    """讀取 Google 試算表，回傳 {品號+製程代號: A/B/C分類} 對照表"""
    cache_key = 'category_map'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    try:
        rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
    except Exception:
        return jsonify({'success': False, 'error': '無法連線至 Google 試算表'}), 502

    cat_map = {}
    for r in rows:
        # 欄位依序：製令(A)、預計數量(B)、品號(C)、品名(D)...製程代號(F)...A/B/C(Q)
        # fetch_google_sheet_csv 已用第一列 header 做 key，直接用 index 取值更安全
        vals = list(r.values())
        if len(vals) < 17:
            continue
        part_no      = vals[2].strip()   # C欄：品號
        process_code = vals[5].strip()   # F欄：製程代號
        category     = vals[16].strip()  # Q欄：A/B/C 分類
        if part_no and process_code and category:
            cat_map[part_no + process_code] = category

    result = {'success': True, 'map': cat_map}
    cache_set(cache_key, result)
    return jsonify(result)


@app.route('/api/k1p2/list')
def k1p2_list():
    """讀取 K1_P2.ref 全部明細記錄"""
    cache_key = 'k1p2_list'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'k1p2_monthly_stats_v3', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)
    try:
        rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        records = []
        for r in rows:
            out_str = (r.get('欄位格式化') or r.get('出站時間') or '').strip()
            ym = ''
            if out_str:
                parts = out_str.replace('-', '/').split('/')
                if len(parts) >= 2:
                    try: ym = f'{int(parts[0])}-{int(parts[1]):02d}'
                    except: pass
            records.append({
                'wo':           (r.get('製令')           or '').strip(),
                'plan_qty':     (r.get('預計產量')        or '').strip(),
                'part_no':      (r.get('品號')            or '').strip(),
                'name':         (r.get('品名')            or '').strip(),
                'seq':          (r.get('加工順序')        or '').strip(),
                'process':      (r.get('製程')            or '').strip(),
                'process_name': (r.get('製程名稱')        or '').strip(),
                'qty':          (r.get('出站數量')        or '').strip(),
                'seconds':      (r.get('出站總工時(秒)')  or '').strip(),
                'in_time':      (r.get('進站時間')        or '').strip(),
                'out_time':     out_str,
                'operator':     (r.get('報工人員')        or '').strip(),
                'category':     (r.get('A/B/C分類') or r.get('A/B/C') or '').strip(),
                'ym':           ym,
            })
        records.sort(key=lambda x: x.get('out_time', ''), reverse=True)
        result = {'success': True, 'count': len(records), 'records': records}
        cache_set(cache_key, result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/k1p2/std_time', methods=['POST'])
def k1p2_std_time():
    """P2 明細表用：批次查標準工時，讓「標工」欄可以跟「出站總秒數/出站數量」算出的
    實際效率對照。沿用 P5.3 生產日報表同一份 fetch_std_time（含 TTL 快取），不重新造輪子。
    由前端傳目前畫面上實際會用到的製令號碼（已經篩選過），不是整張 K1_P2 表全撈，
    避免對 SSRS 打出大量沒必要的查詢。上限 300 張：完全不篩選（全部月份/全部分類）
    時 P2 明細常有數千筆、對應數百張不重複的製令，超過上限的部分標工欄會是空白——
    這是設計上的節流，不是 bug，篩月份/分類後通常就在上限內。"""
    body = request.get_json(silent=True) or {}
    orders = body.get('orders') or []
    orders = [o for o in orders if isinstance(o, str) and '-' in o][:300]

    result = {}
    to_fetch = []
    for oid in orders:
        parts = oid.split('-', 1)
        if len(parts) != 2:
            continue
        otype, onum = parts
        cache_key = f'std_{otype}-{onum}'
        cached = cache_get(cache_key)
        if cached is not None:
            result[oid] = cached
        else:
            to_fetch.append((oid, otype, onum))

    if to_fetch:
        with ThreadPoolExecutor(max_workers=min(8, len(to_fetch))) as pool:
            futures = {pool.submit(fetch_std_time, otype, onum): oid for oid, otype, onum in to_fetch}
            for fut in as_completed(futures):
                oid = futures[fut]
                try:
                    result[oid] = fut.result()
                except Exception:
                    result[oid] = {}

    return jsonify({'success': True, 'data': result})


@app.route('/api/k1p2/monthly_stats')
def k1p2_monthly_stats():
    """K1_P2.ref 月份分類統計：出站數量 × A/B/C/L × 年月"""
    cache_key = 'k1p2_monthly_stats_v3'
    if request.args.get('refresh'):
        cache_clear(cache_key, 'k1p2_list', 'category_map')
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)
    try:
        rows = fetch_google_sheet_csv(config.CATEGORY_SHEET_ID, gid=config.CATEGORY_SHEET_GID)
        from collections import defaultdict
        month_data = defaultdict(lambda: {
            'A': 0, 'B': 0, 'C': 0, 'L': 0, 'other': 0,
            'sec_A': 0, 'sec_B': 0, 'sec_C': 0, 'sec_L': 0,
            'records': 0, 'uncat': 0, 'seconds': 0,
        })
        for r in rows:
            # 優先用欄位格式化（YYYY/MM/DD），備援出站時間（含中文上午/下午）
            date_str = (r.get('欄位格式化') or r.get('出站時間') or '').strip()
            qty_str  = (r.get('出站數量') or '').strip()
            # 相容半型 () 和全型（）括號
            sec_str  = (r.get('出站總工時(秒)') or r.get('出站總工時（秒）') or '').strip()
            cat      = (r.get('A/B/C') or r.get('A/B/C分類') or '').strip()
            try:
                qty = int(float(qty_str.replace(',', ''))) if qty_str else 0
            except Exception:
                qty = 0
            try:
                sec = int(float(sec_str.replace(',', ''))) if sec_str else 0
            except Exception:
                sec = 0
            if not date_str:
                continue
            # 日期解析 → 年月（先取日期部份再分割，避免時間字串干擾）
            ym = ''
            clean_date = date_str.split(' ')[0]
            if '/' in clean_date:
                parts = clean_date.split('/')
                if len(parts) >= 2:
                    try:
                        ym = f'{int(parts[0])}-{int(parts[1]):02d}'
                    except Exception:
                        pass
            elif '-' in clean_date:
                parts = clean_date.split('-')
                if len(parts) >= 2:
                    try:
                        ym = f'{parts[0]}-{int(parts[1]):02d}'
                    except Exception:
                        pass
            if not ym:
                continue
            month_data[ym]['seconds'] += sec
            if cat in ('A', 'B', 'C', 'L'):
                month_data[ym][f'sec_{cat}'] += sec   # 秒數：不論 qty 是否為 0 都累加
                if qty > 0:
                    month_data[ym][cat]       += qty
                    month_data[ym]['records'] += 1
            else:
                month_data[ym]['uncat'] += 1
                if qty > 0:
                    month_data[ym]['other'] += qty

        months = []
        for ym in sorted(month_data.keys()):
            d = month_data[ym]
            classified = d['A'] + d['B'] + d['C'] + d['L']
            months.append({
                'month': ym,
                'A': d['A'], 'B': d['B'], 'C': d['C'], 'L': d['L'],
                'sec_A': d['sec_A'], 'sec_B': d['sec_B'],
                'sec_C': d['sec_C'], 'sec_L': d['sec_L'],
                'other': d['other'],
                'total': classified + d['other'],
                'records': d['records'],
                'uncat': d['uncat'],
                'seconds': d['seconds'],
            })

        result = {'success': True, 'months': months}
        cache_set(cache_key, result)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/jig/list')
def jig_list():
    """讀取治檢具清單索引（PDM 資料夾資料卡，build_jig_index.py 建立）"""
    conn = get_pdm_db()
    if not conn:
        return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

    try:
        cur = conn.execute(
            'SELECT folder_name, folder_path, product_model, item_name, '
            '       handler, submitter, status, apply_date, unit, remarks '
            'FROM jig_index ORDER BY folder_name DESC'
        )
        jigs = [dict(r) for r in cur.fetchall()]
        row = conn.execute('SELECT MAX(indexed_at) FROM jig_index').fetchone()
        last_updated = row[0] if row and row[0] else None
    except sqlite3.OperationalError:
        jigs = []
        last_updated = None
    finally:
        conn.close()

    return jsonify({'success': True, 'count': len(jigs), 'jigs': jigs,
                    'last_updated': last_updated})


_jig_reindex_state = {
    'running': False,
    'phase':   'idle',   # idle | scanning | indexing | done | error
    'idx':     0,
    'total':   0,
    'count':   0,
    'message': '',
    'error':   '',
}
_jig_reindex_lock = _threading.Lock()


def _run_jig_reindex():
    """背景執行治檢具索引重建，解析 stdout 更新進度狀態（比照 PDM 圖面 _run_reindex 的做法）"""
    global _jig_reindex_state
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_jig_index.py')

    with _jig_reindex_lock:
        _jig_reindex_state.update(running=True, phase='scanning', idx=0, total=0,
                                   count=0, message='啟動中...', error='')

    try:
        # 子行程預設用主控台編碼（機台常見為 cp950）輸出中文，跟這裡宣告的 utf-8 解碼對不上，
        # 中文字全部變亂碼、下面的正規表達式永遠比對不到 → 進度卡在 0、count 也讀不到。
        # 強制子行程改用 UTF-8 輸出即可解掉
        child_env = os.environ.copy()
        child_env['PYTHONIOENCODING'] = 'utf-8'
        # 子行程 stdout 接到管線（非終端機）時 Python 預設整批緩衝，進度輸出會卡到緩衝區滿
        # 或行程結束才一次送出，畫面看起來像卡住不動；設這個讓它逐行即時送出
        child_env['PYTHONUNBUFFERED'] = '1'
        proc = _subprocess.Popen(
            [sys.executable, script, '--deploy'],
            stdout=_subprocess.PIPE, stderr=_subprocess.STDOUT,
            text=True, encoding='utf-8', errors='replace',
            cwd=os.path.dirname(script), env=child_env
        )
        for line in proc.stdout:
            line = line.rstrip()
            if not line:
                continue
            with _jig_reindex_lock:
                _jig_reindex_state['message'] = line

            m = _re.search(r'找到\s*(\d+)\s*個子資料夾', line)
            if m:
                with _jig_reindex_lock:
                    _jig_reindex_state['total'] = int(m.group(1))
                continue

            m = _re.search(r'\[(\d+)/(\d+)\]', line)
            if m:
                with _jig_reindex_lock:
                    _jig_reindex_state['phase'] = 'indexing'
                    _jig_reindex_state['idx']   = int(m.group(1))
                    _jig_reindex_state['total'] = int(m.group(2))
                continue

            m = _re.search(r'完成！寫入：(\d+)\s*筆', line)
            if m:
                with _jig_reindex_lock:
                    _jig_reindex_state['count'] = int(m.group(1))

        proc.wait()

        if proc.returncode == 0:
            with _jig_reindex_lock:
                cnt = _jig_reindex_state['count']
                _jig_reindex_state.update(running=False, phase='done',
                                           message=f'更新完成，共 {cnt} 筆')
        else:
            with _jig_reindex_lock:
                _jig_reindex_state.update(running=False, phase='error',
                                           error='重建失敗，請查看伺服器日誌')
    except Exception as exc:
        with _jig_reindex_lock:
            _jig_reindex_state.update(running=False, phase='error', error=str(exc))


@app.route('/api/jig/rebuild', methods=['POST'])
def jig_rebuild():
    """啟動治檢具索引重建（背景執行，前端輪詢 /api/jig/rebuild/status 取得進度）"""
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_jig_index.py')
    if not os.path.exists(script):
        return jsonify({'success': False, 'error': 'build_jig_index.py 不存在'}), 500
    with _jig_reindex_lock:
        if _jig_reindex_state['running']:
            return jsonify({'success': False, 'error': '索引重建已在執行中，請稍候'}), 409
    t = _threading.Thread(target=_run_jig_reindex, daemon=True)
    t.start()
    return jsonify({'success': True, 'message': '索引重建已啟動'})


@app.route('/api/jig/rebuild/status', methods=['GET'])
def jig_rebuild_status():
    """回傳治檢具索引重建進度狀態"""
    with _jig_reindex_lock:
        return jsonify(dict(_jig_reindex_state))


@app.route('/api/jig/open-folder', methods=['POST'])
def jig_open_folder():
    """用 Shell 開啟治檢具資料夾"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400

    jig_root = os.path.normpath(config.JIG_VAULT_PATH)
    norm_path = os.path.normpath(path)
    if not norm_path.lower().startswith(jig_root.lower()):
        return jsonify({'ok': False, 'error': '無效的路徑'}), 400
    if not os.path.isdir(norm_path):
        return jsonify({'ok': False, 'error': '資料夾不存在'}), 404

    try:
        os.startfile(norm_path)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/jig/open-file', methods=['POST'])
def jig_open_file():
    """開啟治檢具資料夾內的申請單 Excel（優先找『申請單』xlsm，否則第一個 xlsm）"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400

    jig_root = os.path.normpath(config.JIG_VAULT_PATH)
    norm_path = os.path.normpath(path)
    if not norm_path.lower().startswith(jig_root.lower()):
        return jsonify({'ok': False, 'error': '無效的路徑'}), 400
    if not os.path.isdir(norm_path):
        return jsonify({'ok': False, 'error': '資料夾不存在'}), 404

    try:
        xlsms = [f for f in os.listdir(norm_path) if f.lower().endswith('.xlsm')]
        if not xlsms:
            # 本機尚未快取：透過 PDM API 列出檔案並下載
            import pythoncom
            pythoncom.CoInitialize()
            try:
                import win32com.client
                vault = _pdm_vault_login()
                folder = vault.GetFolderFromPath(norm_path)
                f5 = win32com.client.CastTo(folder, 'IEdmFolder5')
                pos = f5.GetFirstFilePosition()
                while not pos.IsNull:
                    f = f5.GetNextFile(pos)
                    if f.Name.lower().endswith('.xlsm'):
                        win32com.client.CastTo(f, 'IEdmFile5').GetFileCopy(0)
                        xlsms.append(f.Name)
            finally:
                pythoncom.CoUninitialize()
        if not xlsms:
            return jsonify({'ok': False, 'error': '此資料夾沒有 Excel 申請單'}), 404
        target = next((f for f in xlsms if '申請單' in f), xlsms[0])
        os.startfile(os.path.join(norm_path, target))
        return jsonify({'ok': True, 'file': target})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ══════════════════════════════════════════════════════════
#  治檢具申請單：直接寫入 PDM（建立 PT 專案資料夾 + 資料卡）
# ══════════════════════════════════════════════════════════

JIG_APPLY_SERNO = 'PT_PT專案序號(2026後)'   # PDM 序號產生器名稱（與檔案總管範本同一個，確保不撞號）
_jig_apply_opts_cache = {'ts': 0, 'data': None}


def _pdm_vault_login():
    """登入 PDM Vault（LoginAuto：沿用本機已登入的 PDM 工作階段，不需帳密）。
    gen_py 快取損毀時的自動清除重試邏輯同 build_pdm_index.py connect_vault()，
    這是多台電腦共用同一份 dist_embed 時的已知問題，見該函式註解。"""
    import win32com.client
    from win32com.client import gencache
    try:
        v = gencache.EnsureDispatch('ConisioLib.EdmVault')
    except AttributeError as e:
        # gen_py 快取損毀的變體不只一種（CLSIDToPackageMap／CLSIDToClassMap 等），
        # 共同特徵是訊息含 win32com.gen_py，見 build_pdm_index.py connect_vault() 註解
        if 'win32com.gen_py' not in str(e):
            raise
        try:
            import shutil
            gen_path = win32com.__gen_path__
            if gen_path and os.path.isdir(gen_path):
                shutil.rmtree(gen_path, ignore_errors=True)
            v = gencache.EnsureDispatch('ConisioLib.EdmVault')
        except Exception:
            v = win32com.client.Dispatch('ConisioLib.EdmVault')
    except Exception:
        v = win32com.client.Dispatch('ConisioLib.EdmVault')
    v.LoginAuto('MAXCLAW', 0)
    return v


# 範本建立時要一併產生的標準子資料夾
JIG_APPLY_SUBFOLDERS = (
    '01-模具圖檔(機型+品名+類別-00流水號)',
    '02-模具照片(模具編碼-機型+品名-00流水號)',
    '03-單據掃描檔備存(專案代號+供應商-00流水號)',
    '04-其他(日期_檔案內容)',
)


def _fill_xlsm_custom_vars(src_path, dst_path, vals):
    """複製 xlsm 並改寫 docProps/custom.xml 內的 PDM 卡片變數值。
    PDM 2021 的 Excel 卡片變數存在 custom.xml（attribute mapping），
    檔案加入 vault 時 PDM 會直接從這裡讀出顯示在資料卡上。"""
    import zipfile
    from xml.sax.saxutils import escape
    with zipfile.ZipFile(src_path, 'r') as zin:
        xml = zin.read('docProps/custom.xml').decode('utf-8')
        for k, v in vals.items():
            xml = re.sub(
                r'(name="%s"[^>]*>\s*<vt:lpwstr>)[^<]*(</vt:lpwstr>)' % re.escape(k),
                lambda m, _v=v: m.group(1) + escape(str(_v)) + m.group(2),
                xml, count=1)
        with zipfile.ZipFile(dst_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'docProps/custom.xml':
                    zout.writestr(item, xml.encode('utf-8'))
                else:
                    zout.writestr(item, zin.read(item.filename))


def _write_xlsm_defined_names(xlsm_path, vals):
    """依 Excel Defined Names 把值寫進工作表儲存格（Excel 開啟/列印時直接看得到，
    治檢具索引也讀這些儲存格）。openpyxl 重寫會動到 custom.xml 的 PDM 卡片格式，
    所以先備份、寫完再還原。"""
    import zipfile
    with zipfile.ZipFile(xlsm_path) as z:
        custom_bak = z.read('docProps/custom.xml')
    wb = openpyxl.load_workbook(xlsm_path, keep_vba=True)
    for name, v in vals.items():
        if name not in wb.defined_names:
            continue
        for sn, coord in wb.defined_names[name].destinations:
            try:
                wb[sn][coord] = v
            except Exception:
                pass
    wb.save(xlsm_path)
    tmp_zip = xlsm_path + '.tmpz'
    with zipfile.ZipFile(xlsm_path, 'r') as zin, \
         zipfile.ZipFile(tmp_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            zout.writestr(item, custom_bak if item.filename == 'docProps/custom.xml' else zin.read(item.filename))
    os.replace(tmp_zip, xlsm_path)


@app.route('/api/jig/apply/defaults')
def jig_apply_defaults():
    """新增治檢具申請單：取得預設值（下一個 PT 編號預覽、申請人、模具類型選項）"""
    import pythoncom
    pythoncom.CoInitialize()
    try:
        import win32com.client
        vault = _pdm_vault_login()

        # 下一號預覽：先配號再立即 Rollback 還回去，不佔號
        sg = win32com.client.CastTo(vault, 'IEdmSerNoGen7')
        sv = sg.AllocSerNoValue(JIG_APPLY_SERNO)
        next_no = sv.Value
        sv.Rollback()

        # 申請人 = PDM 登入者名稱（FullName 在 IEdmUser6 介面上）
        applicant = ''
        try:
            umgr = win32com.client.CastTo(vault, 'IEdmUserMgr5')
            u = umgr.GetLoggedInUser()
            try:
                u6 = win32com.client.CastTo(u, 'IEdmUser6')
                applicant = (u6.FullName or '').strip()
            except Exception:
                pass
            if not applicant:
                applicant = (u.Name or '').strip()
        except Exception:
            pass

        # 模具類型選項：掃現有 PT 資料夾的 PT選單 值（快取 10 分鐘）
        now = time.time()
        if not _jig_apply_opts_cache['data'] or now - _jig_apply_opts_cache['ts'] > 600:
            opts = set()
            root = vault.GetFolderFromPath(config.JIG_VAULT_PATH)
            f5 = win32com.client.CastTo(root, 'IEdmFolder5')
            pos = f5.GetFirstSubFolderPosition()
            while not pos.IsNull:
                sub = f5.GetNextSubFolder(pos)
                try:
                    ev = win32com.client.CastTo(sub, 'IEdmEnumeratorVariable5')
                    ok, val = ev.GetVar('PT選單', '')
                    if ok and val:
                        opts.add(str(val).strip())
                except Exception:
                    pass
            _jig_apply_opts_cache.update(ts=now, data=sorted(opts))

        return jsonify({'success': True, 'next_no': next_no, 'applicant': applicant,
                        'today': datetime.now().strftime('%Y/%m/%d'),
                        'mold_types': _jig_apply_opts_cache['data']})
    except Exception as e:
        return jsonify({'success': False, 'error': f'PDM 連線失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


# 模具特殊需求選項文字（照空白範本原始值；勾選=■ 前綴、未勾=□ 前綴）
JIG_CARD_OPTIONS = {
    1: ('□ 嚴禁夾傷  ',            '■嚴禁夾傷'),
    2: ('□ 允許輕微夾痕',           '■允許輕微夾痕'),
    3: ('□ 需快速換模機構 ',        '■需快速換模機構'),
    4: ('□ 具備防呆機制 (避免反裝)', '■具備防呆機制 (避免反裝)'),
    5: ('□ 易耗損零件(需備品) ',    '■易耗損零件(需備品)'),
    6: ('□ 需易於清理排屑',         '■需易於清理排屑'),
}


@app.route('/api/jig/apply/card_save', methods=['POST'])
def jig_apply_card_save():
    """填寫申請單資料卡：對取出中的申請單 xlsm 寫入主要欄位（等同在 PDM 資料卡按儲存）"""
    d = request.get_json(silent=True) or {}
    folder_path = (d.get('folder') or '').strip()
    jig_root = os.path.normpath(config.JIG_VAULT_PATH)
    norm = os.path.normpath(folder_path)
    if not folder_path or not norm.lower().startswith(jig_root.lower()):
        return jsonify({'success': False, 'error': '無效的資料夾路徑'}), 400

    import pythoncom
    pythoncom.CoInitialize()
    try:
        import win32com.client
        vault = _pdm_vault_login()

        # 找資料夾內的申請單 xlsm
        folder = vault.GetFolderFromPath(norm)
        f5folder = win32com.client.CastTo(folder, 'IEdmFolder5')
        target = None
        pos = f5folder.GetFirstFilePosition()
        while not pos.IsNull:
            f = f5folder.GetNextFile(pos)
            if f.Name.lower().endswith('.xlsm') and '申請單' in f.Name:
                target = f
                break
        if target is None:
            return jsonify({'success': False, 'error': '找不到申請單 Excel'}), 404

        f5 = win32com.client.CastTo(target, 'IEdmFile5')
        # 已簽入（未取出）的單子：不能改卡片內容，但仍可單獨「提出申請」
        # （補救先前存回成功、送審失敗的情況，不用進 PDM 重新取出）
        skip_card_write = False
        if not f5.IsLocked:
            if d.get('submit_flow'):
                skip_card_write = True
            else:
                return jsonify({'success': False, 'error': '申請單未取出，請先在 PDM 取出檔案再填寫'}), 409

        # 本機副本若殘留唯讀屬性，後面 CloseFile(True)/openpyxl 都會失敗，先清掉
        _ro_path = os.path.join(norm, target.Name)
        try:
            if os.path.isfile(_ro_path):
                os.chmod(_ro_path, 0o666)
        except Exception:
            pass

        vals = {}
        if d.get('urgency'):     vals['YC_緊急程度'] = d['urgency'].strip()
        if d.get('nature'):      vals['YC_性質']     = d['nature'].strip()
        if d.get('part_no'):     vals['PT_品號']     = d['part_no'].strip()
        if d.get('proc_code'):   vals['YC_製程代號'] = d['proc_code'].strip()
        if d.get('unit'):        vals['單位']        = d['unit'].strip()
        if d.get('expect_date'): vals['YC_期望完成日'] = d['expect_date'].strip().replace('-', '/')
        checked = set(d.get('options') or [])
        for n, (off, on) in JIG_CARD_OPTIONS.items():
            vals[f'YC_選項{n}'] = on if n in checked else off

        # 取出中的檔案，資料卡顯示的是本機檔案內的屬性值：
        # Flush 只寫資料庫，必須用 IEdmEnumeratorVariable8.CloseFile(True)
        # 把值同步寫進檔案本體（custom.xml），卡片才會顯示、簽入時也不會被檔案舊值蓋回。
        # 組態名稱用 ''（實測手動建立的申請單變數都在空名組態；用 '@' 會多出一個空白組態分頁）
        if not skip_card_write:
            ev = win32com.client.CastTo(f5.GetEnumeratorVariable(), 'IEdmEnumeratorVariable8')
            for k, v in vals.items():
                try:
                    ev.SetVar(k, '', v)
                except Exception:
                    pass
            ev.CloseFile(True)

        # 工作表內容：卡片值同步寫進儲存格（Defined Names）＋示意圖（E7）＋其他需求文字（D15）
        # 直接用 openpyxl 改取出中的本機 xlsm
        img_b64    = d.get('image_b64') or ''
        other_text = (d.get('other_text') or '').strip()
        if not skip_card_write:
            if not _OPENPYXL_OK:
                return jsonify({'success': False, 'error': '伺服器未安裝 openpyxl'}), 500
            import base64, tempfile, zipfile
            xlsm_path = os.path.join(norm, target.Name)
            try:
                # 先備份 custom.xml（openpyxl 重寫時會丟失 PDM 卡片連結的 linkTarget 屬性）
                with zipfile.ZipFile(xlsm_path) as z:
                    custom_bak = z.read('docProps/custom.xml')

                wb = openpyxl.load_workbook(xlsm_path, keep_vba=True)
                ws = wb.worksheets[0]
                # 卡片欄位值依 Defined Names 寫進儲存格（Excel 開啟/列印直接看得到，索引也讀這裡）
                for name, v in vals.items():
                    if name not in wb.defined_names:
                        continue
                    for sn, coord in wb.defined_names[name].destinations:
                        try:
                            wb[sn][coord] = v
                        except Exception:
                            pass
                if other_text:
                    from openpyxl.styles import Alignment
                    cell = ws['D15']
                    cell.value = other_text
                    cell.alignment = Alignment(wrap_text=True, vertical='center',
                                               horizontal=cell.alignment.horizontal)
                    # 依行數放大列高（每行約 16pt，至少保留原高度）
                    lines = other_text.count('\n') + 1
                    cur_h = ws.row_dimensions[15].height or 16
                    ws.row_dimensions[15].height = max(cur_h, lines * 16 + 4)
                img_tmp = None
                if img_b64:
                    from openpyxl.drawing.image import Image as XLImage
                    raw = base64.b64decode(img_b64.split(',')[-1])
                    img_tmp = os.path.join(tempfile.gettempdir(), f'jig_sketch_{int(time.time())}.png')
                    with open(img_tmp, 'wb') as fimg:
                        fimg.write(raw)
                    img = XLImage(img_tmp)
                    # 等比縮放到示意圖區域（E7:H11 約 420x150）
                    ratio = min(420 / img.width, 150 / img.height)
                    img.width  = int(img.width * ratio)
                    img.height = int(img.height * ratio)
                    ws.add_image(img, 'E7')
                wb.save(xlsm_path)
                if img_tmp:
                    try:
                        os.remove(img_tmp)
                    except Exception:
                        pass

                # 還原 custom.xml，保住 PDM 卡片變數與 linkTarget
                tmp_zip = xlsm_path + '.tmp'
                with zipfile.ZipFile(xlsm_path, 'r') as zin, \
                     zipfile.ZipFile(tmp_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        if item.filename == 'docProps/custom.xml':
                            zout.writestr(item, custom_bak)
                        else:
                            zout.writestr(item, zin.read(item.filename))
                os.replace(tmp_zip, xlsm_path)
            except PermissionError:
                return jsonify({'success': False,
                                'error': '卡片已儲存，但 Excel 檔案被開啟中，示意圖/文字無法寫入，請先關閉 Excel 再試一次'}), 409
            except Exception as e:
                return jsonify({'success': False, 'error': f'卡片已儲存，但寫入示意圖/文字失敗：{e}'}), 502

        # 自動存回（簽入）：把填好的申請單簽入 PDM
        if d.get('checkin'):
            try:
                f5.Refresh()
                if f5.IsLocked:
                    f5.UnlockFile(0, '申請單填寫完成')
            except Exception as e:
                return jsonify({'success': True,
                                'warning': f'卡片已儲存，但存回失敗（請在 PDM 手動存回）：{e}'})

        # 提出申請：走 workflow 轉換「00-提出申請」到「單位主管審核」。
        # 注意（2026-07-07 實測）：
        # 1. Python IDispatch 呼叫 IEdmFile5.ChangeState 一律回 DISP_E_MEMBERNOTFOUND，
        #    必須走 .NET Interop 的 vtable（pdm_change_state.ps1 helper）。
        # 2. 「00-提出申請」轉換設了身分驗證，要用 ChangeState3 帶使用者的 PDM 密碼。
        if d.get('submit_flow'):
            try:
                comment = (d.get('comment') or '').strip() or '提出申請'
                password = d.get('flow_password') or ''
                if not password:
                    return jsonify({'success': True,
                                    'warning': '已存回，但提出申請需要輸入 PDM 登入密碼（此轉換有身分驗證），'
                                               '請重新按「續填」勾選提出申請並輸入密碼'})
                helper = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pdm_change_state.ps1')
                env = os.environ.copy()
                env.update({
                    'PDM_CS_FILE':       os.path.join(norm, target.Name),
                    'PDM_CS_STATE':      '單位主管審核',
                    'PDM_CS_TRANSITION': '00-提出申請',
                    'PDM_CS_COMMENT':    comment,
                    'PDM_CS_PASSWORD':   password,
                })
                import subprocess
                r = subprocess.run(
                    ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-File', helper],
                    env=env, capture_output=True, timeout=90)
                out = (r.stdout or b'').decode('utf-8', errors='replace').strip().splitlines()
                result = out[-1] if out else 'ERROR: 無輸出'
                if not result.startswith('OK'):
                    return jsonify({'success': True,
                                    'warning': f'已存回，但提出申請未成功：{result}。'
                                               '（密碼錯誤或欄位不齊；可在 PDM 手動變更狀態）'})
            except Exception as e:
                return jsonify({'success': True,
                                'warning': f'已存回，但提出申請失敗（請在 PDM 手動變更狀態）：{e}'})

        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': f'儲存卡片失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


@app.route('/api/jig/apply', methods=['POST'])
def jig_apply():
    """新增治檢具申請單：配一個 PT 序號，在 PDM 建立資料夾並填寫資料卡"""
    d = request.get_json(silent=True) or {}
    machine    = (d.get('machine') or '').strip()
    item       = (d.get('item') or '').strip()
    mold       = (d.get('mold_type') or '').strip()
    rd_no      = (d.get('rd_no') or '').strip() or '--'
    applicant  = (d.get('applicant') or '').strip()
    apply_date = (d.get('date') or '').strip() or datetime.now().strftime('%Y/%m/%d')
    if not machine or not item or not mold:
        return jsonify({'success': False, 'error': '機型、品名、模具類型為必填'}), 400
    if not applicant:
        return jsonify({'success': False, 'error': '申請人為必填'}), 400

    import pythoncom
    pythoncom.CoInitialize()
    sv = None
    new_folder = None
    parent = None
    try:
        import win32com.client
        vault = _pdm_vault_login()

        # 1. 從 PDM 序號產生器配號（跟檔案總管範本同一個計數器）
        sg = win32com.client.CastTo(vault, 'IEdmSerNoGen7')
        sv = sg.AllocSerNoValue(JIG_APPLY_SERNO)
        pt_no = str(sv.Value)

        # 2. 建立 PT 專案資料夾（poData 必須明確傳 None，gen_py 預設值 0 會轉型失敗）
        root = vault.GetFolderFromPath(config.JIG_VAULT_PATH)
        parent = win32com.client.CastTo(root, 'IEdmFolder5')
        new_folder = parent.AddFolder(0, pt_no, None)

        # 3. 寫入資料夾資料卡變數（欄位對應實測自現有 PT 資料夾）
        ev = win32com.client.CastTo(new_folder, 'IEdmEnumeratorVariable5')
        for k, v in {
            'PT_專案代號': pt_no,
            'PT選單':      mold,
            'YC_專案代號': rd_no,
            'YC_機型':     machine,
            'YC_品名':     item,
            'YC_提出人員': applicant,
            'YC_日期':     apply_date,
            'YC_狀態':     'PT設計發包',
        }.items():
            ev.SetVar(k, '', v)
        ev.Flush()

        # 4. 建立 4 個標準子資料夾（照檔案總管範本的結構）
        nf5 = win32com.client.CastTo(new_folder, 'IEdmFolder5')
        warnings = []
        for sub in JIG_APPLY_SUBFOLDERS:
            try:
                nf5.AddFolder(0, sub, None)
            except Exception as e:
                warnings.append(f'子資料夾 {sub} 建立失敗：{e}')

        # 5. 複製空白申請單範本 → 改寫卡片變數 → 加入 PDM → 入庫
        xlsm_name = f'{pt_no}機器模檢治具申請單.xlsm'
        try:
            # 5a. 確保範本檔本機快取是最新版（第二參數必須明確傳 None，gen_py 預設值會轉型失敗）
            ret = vault.GetFileFromPath(config.JIG_APPLY_TEMPLATE_XLSM, None)
            tpl_file = ret[0] if isinstance(ret, tuple) else ret
            try:
                win32com.client.CastTo(tpl_file, 'IEdmFile5').GetFileCopy(0)
            except Exception:
                pass  # 快取已存在時可略過
            if not os.path.isfile(config.JIG_APPLY_TEMPLATE_XLSM):
                raise RuntimeError(f'找不到範本檔：{config.JIG_APPLY_TEMPLATE_XLSM}')

            # 5b. 在暫存目錄產生已填值的 xlsm
            import tempfile
            fill_vals = {
                'PT_專案代號':  pt_no,
                'YC_機型':      machine,
                'YC_品名':      item,
                'YC_提出人員':  applicant,
                'YC_日期':      apply_date,
                'YC_專案代號':  rd_no,
                'PT選單':       mold,
                'YC_部門':      '加工部',
                '單位':         '加工部',
                '版別':         '--',
                '00文件狀態':   '申請單建立',
                '00文件分類':   '機器模檢治具申請單(2026後)',
            }
            tmp_xlsm = os.path.join(tempfile.gettempdir(), xlsm_name)
            _fill_xlsm_custom_vars(config.JIG_APPLY_TEMPLATE_XLSM, tmp_xlsm, fill_vals)
            # 同步寫進工作表儲存格（Excel 開啟直接看得到，治檢具索引也讀儲存格）
            try:
                _write_xlsm_defined_names(tmp_xlsm, fill_vals)
            except Exception as e:
                warnings.append(f'工作表儲存格寫入失敗（卡片資料不受影響）：{e}')

            # 5c. 加入 PDM → 變數寫入資料庫 → 簽入 → 自動分派後再取出給使用者。
            # 工作流程的自動轉換靠「00文件分類」等 DB 變數判斷分派；只寫在 custom.xml 裡的值
            # 在簽入當下還沒進資料庫，條件不成立會掉進「其他文件歸檔」死路（實測），
            # 所以必須先 SetVar 寫 DB 再簽入，分派到「申請單建立」後再 LockFile 取出讓使用者續填。
            file_id = nf5.AddFile(0, tmp_xlsm, xlsm_name)
            try:
                os.remove(tmp_xlsm)
            except Exception:
                pass
            fobj = vault.GetObject(1, file_id)  # 1 = EdmObject_File
            f5 = win32com.client.CastTo(fobj, 'IEdmFile5')
            # 組態名稱用 ''（跟手動建立一致；用 '@' 會在卡片上多出一個空白組態分頁）
            ev2 = f5.GetEnumeratorVariable()
            for k, v in fill_vals.items():
                try:
                    ev2.SetVar(k, '', v)
                except Exception:
                    pass
            ev2.Flush()
            f5.UnlockFile(0, '系統建立申請單')

            # 等待自動分派完成（實測約 2 秒）
            final_state = ''
            for _ in range(5):
                time.sleep(2)
                f5.Refresh()
                final_state = f5.CurrentState.Name
                if final_state == '申請單建立':
                    break
            if final_state == '申請單建立':
                try:
                    f5.LockFile(new_folder.ID, 0)
                    # 取出後本機副本可能殘留唯讀屬性（PDM 簽入時設的），
                    # 不清掉的話資料卡儲存/openpyxl 寫入都會失敗（2026-07-07 實際發生）
                    local_xlsm = os.path.join(config.JIG_VAULT_PATH, pt_no, xlsm_name)
                    if os.path.isfile(local_xlsm):
                        os.chmod(local_xlsm, 0o666)
                except Exception as e:
                    warnings.append(f'申請單已建立但自動取出失敗（可手動取出）：{e}')
            else:
                warnings.append(f'申請單流程狀態異常（目前：{final_state or "未知"}），請聯絡 PDM 管理員')
        except Exception as e:
            warnings.append(f'申請單 Excel 建立失敗：{e}')

        sv = None  # 序號已正式使用，不再 rollback
        return jsonify({'success': True, 'pt_no': pt_no,
                        'folder': os.path.join(config.JIG_VAULT_PATH, pt_no),
                        'warning': '；'.join(warnings) if warnings else None})
    except Exception as e:
        # 失敗善後：刪除半成品資料夾、把序號還回去
        try:
            if new_folder is not None and parent is not None:
                parent.DeleteFolder(0, new_folder.ID, True)
        except Exception:
            pass
        try:
            if sv is not None:
                sv.Rollback()
        except Exception:
            pass
        return jsonify({'success': False, 'error': f'建立失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


# ══════════════════════════════════════════════════════════
#  設計變更申請單（DCN/ECR）PDM 直寫新增功能
#  比照治檢具申請單的做法（docs/pdm-jig-application-sop.md），四項前置確認見該文件；
#  空白範本檔案在 PDM Admin 登記、查不到實際路徑，改用既有申請單當結構範本（唯讀複製）
# ══════════════════════════════════════════════════════════
DCN_APPLY_SERNO = 'RD_RR設計變更申請單'   # PDM 序號產生器名稱（跟檔案總管新增精靈同一個，不會撞號）

# 申請設變原因勾選框：{key: (文字變數名, 未勾選文字, 勾選文字, 勾選旗標變數名)}
# 這張卡片的每個原因其實有「兩個」變數，兩個都要寫（2026-08-18 實測 RR2608016 才發現）：
#   1. `..._tasky`  → ■/□ 文字，對應 xlsm 的 docProps 屬性＋工作表儲存格，只管 Excel 印出來的樣子
#   2. `...F_tasky` → '1'/'0'，**PDM 資料卡上那個勾選框真正綁的變數**，而且它是純資料庫變數
#      （xlsm 的 custom.xml 跟 defined names 都查無此名），所以只能 SetVar + Flush 寫進 DB，
#      不會出現在檔案裡，也不能拿 _dcn_read_local_var 讀回比對，要用 GetVar 驗。
# 只寫 1 不寫 2 的後果：Excel 印出來有 ■，但 PDM 資料卡七個框全空（使用者看到的就是「沒打勾」）。
DCN_REASON_OPTIONS = {
    'customer':    ('PP_R_004_客戶要求_tasky',   '□客戶要求',   '■客戶要求',   'PP_R_004_客戶要求F_tasky'),
    'design_err':  ('PP_R_004_設計錯誤_tasky',   '□設計錯誤',   '■設計錯誤',   'PP_R_004_設計錯誤F_tasky'),
    'material':    ('PP_R_004_物料需求_tasky',   '□物料需求',   '■物料需求',   'PP_R_004_物料需求F_tasky'),
    'manufacture': ('PP_R_004_製造問題_tasky',   '□製造問題',   '■製造問題',   'PP_R_004_製造問題F_tasky'),
    'improve':     ('PP_R_004_功能改善_tasky',   '□功能改善',   '■功能改善',   'PP_R_004_功能改善F_tasky'),
    'vendor':      ('PP_R_004_變更供應商_tasky', '□變更供應商', '■變更供應商', 'PP_R_004_變更供應商F_tasky'),
    'other':       ('PP_R_004_其他_tasky',       '□其他',       '■其他',       'PP_R_004_其他F_tasky'),
}
DCN_FLAG_ON, DCN_FLAG_OFF = '1', '0'   # 勾選旗標變數的值（實測自人工建立的既有申請單）


def _dcn_crlf(text):
    """PDM 資料卡的多行文字框只認得 CRLF：只寫 \\n 的話卡片上會擠成一整段（RR2608016 實測），
    Excel 儲存格則兩種都吃。統一正規化成 \\r\\n。"""
    return (text or '').replace('\r\n', '\n').replace('\r', '\n').replace('\n', '\r\n')


def _dcn_same_text(a, b):
    """比對卡片值：忽略換行符差異。
    （SetVar 寫入 \\r\\n，但讀回 docProps/custom.xml 時 XML 解析器會把 CRLF 正規化成 LF，
      直接字串比對會永遠不相等）"""
    if a is None or b is None:
        return a == b
    return str(a).replace('\r\n', '\n').replace('\r', '\n') == str(b).replace('\r\n', '\n').replace('\r', '\n')


@app.route('/api/dcn/apply/defaults')
def dcn_apply_defaults():
    """新增設計變更申請單：取得預設值（下一個 RR 編號預覽、申請人、申請原因選項）"""
    import pythoncom
    pythoncom.CoInitialize()
    try:
        import win32com.client
        vault = _pdm_vault_login()

        # 下一號預覽：先配號再立即 Rollback 還回去，不佔號
        sg = win32com.client.CastTo(vault, 'IEdmSerNoGen7')
        sv = sg.AllocSerNoValue(DCN_APPLY_SERNO)
        next_no = sv.Value
        sv.Rollback()

        applicant = ''
        try:
            umgr = win32com.client.CastTo(vault, 'IEdmUserMgr5')
            u = umgr.GetLoggedInUser()
            try:
                u6 = win32com.client.CastTo(u, 'IEdmUser6')
                applicant = (u6.FullName or '').strip()
            except Exception:
                pass
            if not applicant:
                applicant = (u.Name or '').strip()
        except Exception:
            pass

        # 機型代號選項：PDM 新增精靈的下拉選單是 Admin 登記的固定清單，查不到來源，
        # 改從既有 DCN 索引（pdm_search.db）的機型欄位取開頭英文字母代號去重，當自動完成建議
        # （不是強制選單，使用者仍可自行輸入清單沒有的代號）
        model_prefixes = []
        try:
            conn = get_pdm_db()
            if conn:
                rows = conn.execute(
                    "SELECT DISTINCT product_model FROM dcn_index WHERE product_model IS NOT NULL AND product_model != ''"
                ).fetchall()
                conn.close()
                prefixes = set()
                for r in rows:
                    m = re.match(r'^([A-Za-z]+)', (r['product_model'] or '').strip())
                    if m:
                        prefixes.add(m.group(1).upper())
                model_prefixes = sorted(prefixes)
        except Exception:
            model_prefixes = []

        return jsonify({
            'success': True, 'next_no': next_no, 'applicant': applicant,
            'today': datetime.now().strftime('%Y/%m/%d'),
            'model_prefixes': model_prefixes,
            'reason_options': [{'key': k, 'label': v[1].lstrip('□')} for k, v in DCN_REASON_OPTIONS.items()],
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'PDM 連線失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


@app.route('/api/dcn/apply', methods=['POST'])
def dcn_apply():
    """新增設計變更申請單：配一個 RR 序號，在 PDM 建立資料夾並填寫資料卡"""
    d = request.get_json(silent=True) or {}
    model      = (d.get('model') or '').strip()
    applicant  = (d.get('applicant') or '').strip()
    apply_date = (d.get('date') or '').strip() or datetime.now().strftime('%Y/%m/%d')
    if not model:
        return jsonify({'success': False, 'error': '機型為必填'}), 400
    if not applicant:
        return jsonify({'success': False, 'error': '申請人為必填'}), 400

    import pythoncom
    pythoncom.CoInitialize()
    sv = None
    new_folder = None
    parent = None
    try:
        import win32com.client
        vault = _pdm_vault_login()

        # 1. 配號（跟檔案總管新增精靈同一個計數器，不會撞號）
        sg = win32com.client.CastTo(vault, 'IEdmSerNoGen7')
        sv = sg.AllocSerNoValue(DCN_APPLY_SERNO)
        rr_no = str(sv.Value)

        # 2. 建立資料夾（跟現有申請單同一種「每筆一個同名資料夾」結構，poData 明確傳 None）
        root = vault.GetFolderFromPath(config.DCN_VAULT_PATH)
        parent = win32com.client.CastTo(root, 'IEdmFolder5')
        new_folder = parent.AddFolder(0, rr_no, None)

        # 3. 寫入資料夾卡片變數（欄位對應實測自現有 RR 資料夾）
        # 注意：PDM 檔案總管在資料夾層級顯示的「資料卡」是讀資料夾層的變數，不是檔案層的——
        # 「申請變更」欄位就是綁資料夾層的 YC_設變申請編號，漏寫這個會讓 PDS 建立的申請單
        # 在資料夾卡片上「申請變更」顯示空白（檔案層的 YC_設變申請編號 有正確寫入，但那是另一份卡片）
        ev = win32com.client.CastTo(new_folder, 'IEdmEnumeratorVariable5')
        for k, v in {
            'YC_設變申請編號': rr_no,
            'YC_日期':     apply_date,
            'YC_機型':     model,
            'YC_提出人員': applicant,
            'YC_經辦':     '--',
        }.items():
            ev.SetVar(k, '', v)
        ev.Flush()

        # 4. 複製既有申請單當結構範本 → 改寫卡片變數（申請專屬欄位全部覆寫成這次的值）→
        #    加入 PDM → 變數寫進資料庫 → 簽入 → 等自動分派到「ECR建立」→ 取出給使用者續填
        xlsm_name = f'{rr_no}設計變更申請單.xlsm'
        warnings = []
        try:
            ret = vault.GetFileFromPath(config.DCN_APPLY_TEMPLATE_SOURCE, None)
            tpl_file = ret[0] if isinstance(ret, tuple) else ret
            try:
                win32com.client.CastTo(tpl_file, 'IEdmFile5').GetFileCopy(0)
            except Exception:
                pass  # 本機快取已存在時可略過
            if not os.path.isfile(config.DCN_APPLY_TEMPLATE_SOURCE):
                raise RuntimeError(f'找不到範本來源檔：{config.DCN_APPLY_TEMPLATE_SOURCE}')

            import tempfile
            fill_vals = {
                'YC_設變申請編號': rr_no,
                'YC_機型':         model,
                'YC_提出人員':     applicant,
                'YC_日期':         apply_date,
                'YC_日期01':       '--',
                'YC_經辦':         '--',
                'YC_經辦01':       '--',
                'YC_審核':         '--',
                'YC_核決':         '--',
                'YC_實際完成日':   '--',
                'YC_期望完成日01': '--',
                '00文件狀態':      'ECR建立',
                'PP_R_004_變更規格敘述_tasky': '',
                '資材部庫存回報':  '□資材部庫存回報',
                '鼎新品號編修':    '□鼎新品號編修',
            }
            # 新申請預設全部未勾選，交給 Step2 卡片畫面勾選。旗標變數（F）是純資料庫變數，
            # 只能靠下面的 SetVar/Flush 寫，寫進 fill_vals 給 _fill_xlsm_custom_vars 也是空轉
            flag_vals = {}
            for _key, (var_name, off_text, _on_text, flag_var) in DCN_REASON_OPTIONS.items():
                fill_vals[var_name] = off_text
                flag_vals[flag_var] = DCN_FLAG_OFF

            tmp_xlsm = os.path.join(tempfile.gettempdir(), xlsm_name)
            _fill_xlsm_custom_vars(config.DCN_APPLY_TEMPLATE_SOURCE, tmp_xlsm, fill_vals)
            # PDM 的「資料卡」畫面實際渲染的是工作表儲存格內容（依 Excel Defined Names 對應），
            # 跟 docProps/custom.xml 的卡片變數是兩個完全獨立的儲存位置——只寫 custom.xml
            # 的話，SetVar/GetVar 讀得到正確值，但 PDM Explorer 開卡片看到的還是範本原始的
            # 空白/舊值（實測 RR2607050 才發現這個問題）。比照治檢具的做法兩處都要寫。
            try:
                _write_xlsm_defined_names(tmp_xlsm, fill_vals)
            except Exception as e:
                warnings.append(f'工作表儲存格寫入失敗（卡片資料不受影響）：{e}')

            nf5 = win32com.client.CastTo(new_folder, 'IEdmFolder5')
            file_id = nf5.AddFile(0, tmp_xlsm, xlsm_name)
            try:
                os.remove(tmp_xlsm)
            except Exception:
                pass
            fobj = vault.GetObject(1, file_id)  # 1 = EdmObject_File
            f5 = win32com.client.CastTo(fobj, 'IEdmFile5')
            ev2 = f5.GetEnumeratorVariable()
            for k, v in list(fill_vals.items()) + list(flag_vals.items()):
                try:
                    ev2.SetVar(k, '', v)
                except Exception:
                    pass
            ev2.Flush()
            f5.UnlockFile(0, '系統建立設計變更申請單')

            # 等待自動分派完成（治檢具實測約 2 秒）
            final_state = ''
            for _ in range(5):
                time.sleep(2)
                f5.Refresh()
                final_state = f5.CurrentState.Name
                if final_state == 'ECR建立':
                    break
            if final_state == 'ECR建立':
                try:
                    f5.LockFile(new_folder.ID, 0)
                    local_xlsm = os.path.join(config.DCN_VAULT_PATH, rr_no, xlsm_name)
                    if os.path.isfile(local_xlsm):
                        os.chmod(local_xlsm, 0o666)
                except Exception as e:
                    warnings.append(f'申請單已建立但自動取出失敗（可手動取出）：{e}')
            else:
                warnings.append(f'申請單流程狀態異常（目前：{final_state or "未知"}），請聯絡 PDM 管理員')
        except Exception as e:
            warnings.append(f'申請單 Excel 建立失敗：{e}')

        sv = None  # 序號已正式使用，不再 rollback
        return jsonify({'success': True, 'rr_no': rr_no,
                        'folder': os.path.join(config.DCN_VAULT_PATH, rr_no),
                        'warning': '；'.join(warnings) if warnings else None})
    except Exception as e:
        try:
            if new_folder is not None and parent is not None:
                parent.DeleteFolder(0, new_folder.ID, True)
        except Exception:
            pass
        try:
            if sv is not None:
                sv.Rollback()
        except Exception:
            pass
        return jsonify({'success': False, 'error': f'建立失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


@app.route('/api/dcn/apply/card_save', methods=['POST'])
def dcn_apply_card_save():
    """填寫設計變更申請單資料卡：對取出中的申請單 xlsm 寫入申請原因/變更內容說明
    （等同在 PDM 資料卡按儲存），可選存回（簽入）、可選提出申請（00提出申請 轉換到 ECR單位主管審核）"""
    d = request.get_json(silent=True) or {}
    folder_path = (d.get('folder') or '').strip()
    dcn_root = os.path.normpath(config.DCN_VAULT_PATH)
    norm = os.path.normpath(folder_path)
    if not folder_path or not norm.lower().startswith(dcn_root.lower()):
        return jsonify({'success': False, 'error': '無效的資料夾路徑'}), 400

    import pythoncom
    pythoncom.CoInitialize()
    try:
        import win32com.client
        vault = _pdm_vault_login()

        folder = vault.GetFolderFromPath(norm)
        f5folder = win32com.client.CastTo(folder, 'IEdmFolder5')
        target = None
        pos = f5folder.GetFirstFilePosition()
        while not pos.IsNull:
            f = f5folder.GetNextFile(pos)
            if f.Name.lower().endswith('.xlsm') and '申請單' in f.Name:
                target = f
                break
        if target is None:
            return jsonify({'success': False, 'error': '找不到申請單 Excel'}), 404

        f5 = win32com.client.CastTo(target, 'IEdmFile5')
        # 已簽入（未取出）的單子：不能改卡片內容，但仍可單獨「提出申請」
        # （補救先前存回成功、送審失敗的情況，不用進 PDM 重新取出）
        skip_card_write = False
        if not f5.IsLocked:
            if d.get('submit_flow'):
                skip_card_write = True
            else:
                return jsonify({'success': False, 'error': '申請單未取出，請先在 PDM 取出檔案再填寫'}), 409

        # 本機副本若殘留唯讀屬性，CloseFile(True) 會失敗，先清掉
        # （實測發現：使用者在 PDM 檔案總管開著這份檔案的資料卡面板時，本機快取會被鎖回唯讀，
        #  導致 SetVar/CloseFile 全部失敗——但拋出的是「發生一個不明確的錯誤」這種看不出原因的
        #  com_error，之前又被 except Exception: pass 整組吞掉，使用者只會看到假的「成功」，
        #  資料其實完全沒寫進去。現在改成：失敗就清一次唯讀屬性重試一次，還是失敗就老實回報錯誤，
        #  不再偽裝成功）
        _ro_path = os.path.join(norm, target.Name)
        def _dcn_clear_readonly():
            try:
                if os.path.isfile(_ro_path):
                    os.chmod(_ro_path, 0o666)
            except Exception:
                pass
        _dcn_clear_readonly()

        vals = {}        # 檔案屬性變數（docProps + 工作表儲存格）
        flag_vals = {}   # 純資料庫變數（資料卡勾選框綁的那組），只能 SetVar + Flush
        checked = set(d.get('reasons') or [])
        for key, (var_name, off_text, on_text, flag_var) in DCN_REASON_OPTIONS.items():
            vals[var_name] = on_text if key in checked else off_text
            flag_vals[flag_var] = DCN_FLAG_ON if key in checked else DCN_FLAG_OFF
        if 'description' in d:
            vals['PP_R_004_變更規格敘述_tasky'] = _dcn_crlf((d.get('description') or '').strip())

        # 取出中的檔案，資料卡顯示的是本機檔案內的屬性值：Flush 只寫資料庫，
        # 必須用 IEdmEnumeratorVariable8.CloseFile(True) 同步寫進檔案本體，卡片才會顯示。
        # 注意（RR2607050 實測才發現）：PDM Explorer 的「資料卡」畫面實際渲染的是工作表儲存格
        # （依 Excel Defined Names 對應，例如 PP_R_004_製造問題_tasky -> 工作表1!$E$8），
        # 跟 docProps/custom.xml 是兩個完全獨立的儲存位置——SetVar/CloseFile 只會寫 docProps，
        # GetVar 讀得到正確值、索引（build_dcn_index.py）也讀得到，但 PDM 卡片畫面還是顯示範本
        # 原始的空白/舊值。跟治檢具一樣，兩處都要寫，見下面 _write_xlsm_defined_names 呼叫。
        if not skip_card_write:
            def _dcn_write_vals():
                failed = []
                ev = win32com.client.CastTo(f5.GetEnumeratorVariable(), 'IEdmEnumeratorVariable8')
                for k, v in vals.items():
                    try:
                        ev.SetVar(k, '', v)
                    except Exception as e:
                        failed.append(f'{k}：{e}')
                ev.CloseFile(True)
                return failed

            def _dcn_read_local_var(name):
                """直接讀本機檔案的 docProps/custom.xml（不透過 PDM API，排除 DB／檔案兩處不同步的疑慮）"""
                import zipfile
                import xml.etree.ElementTree as ET
                ns_custom = 'http://schemas.openxmlformats.org/officeDocument/2006/custom-properties'
                ns_vt = 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'
                with zipfile.ZipFile(_ro_path, 'r') as zf:
                    root = ET.fromstring(zf.read('docProps/custom.xml'))
                for prop in root.findall(f'{{{ns_custom}}}property'):
                    if prop.get('name') == name:
                        elem = prop.find(f'{{{ns_vt}}}lpwstr')
                        return elem.text if elem is not None else ''
                return None

            CONFLICT_HINT = ('。最常見原因是這份申請單目前在 PDM 資料卡面板或 Excel 裡被開啟中，'
                              '請先關閉該視窗（或在 PDM 檔案總管切到別的檔案），再按一次「儲存卡片」重試')
            # PDM 剛簽出的檔案偶爾會「回成功但沒真的寫入」（尤其是這批剛 AddFile 沒多久的新檔案），
            # 不能只看 SetVar/CloseFile 有沒有拋例外就回報成功，要讀回本機檔案比對，不一致就重試
            # （docs/pdm-jig-application-sop.md 第六節已記載這是通則，不是治檢具專屬）
            # 注意：之前只挑第一個欄位當代表比對，結果那個欄位剛好是使用者沒勾選的原因（值本來就
            # 等於未勾選的預設值），導致整份寫入其實完全失敗，比對卻誤判成功（RR2607045 實測踩到）。
            # 現在改成每個欄位都讀回比對，不能再抄捷徑。
            ok = False
            last_err = None
            for attempt in range(3):
                try:
                    write_failed = _dcn_write_vals()
                except Exception as e:
                    last_err = str(e)
                    _dcn_clear_readonly()
                    time.sleep(1.5)
                    continue
                if write_failed:
                    last_err = '；'.join(write_failed)
                    time.sleep(1.5)
                    continue
                mismatches = []
                for k, v in vals.items():
                    actual = _dcn_read_local_var(k)
                    if not _dcn_same_text(actual, v):
                        mismatches.append(f'{k} 預期 {v!r}，實際讀到 {actual!r}')
                if not mismatches:
                    ok = True
                    break
                last_err = '讀回比對不一致（' + '；'.join(mismatches) + '）'
                time.sleep(1.5)
            if not ok:
                return jsonify({'success': False, 'error': f'寫入卡片失敗：{last_err}{CONFLICT_HINT}'}), 502

            # docProps 已確認寫入成功，接著同步寫進工作表儲存格（Excel 列印出來的內容）。
            # 儲存格裡的換行用 LF（Excel 慣例，人工建立的既有申請單也是 LF），CRLF 只給 PDM 變數用
            cell_vals = {k: (v.replace('\r\n', '\n') if isinstance(v, str) else v) for k, v in vals.items()}
            try:
                _write_xlsm_defined_names(_ro_path, cell_vals)
            except Exception as e:
                return jsonify({'success': False, 'error': f'卡片資料已存到資料庫，但工作表儲存格寫入失敗：{e}{CONFLICT_HINT}'}), 502

            # 資料卡勾選框綁的旗標變數（純資料庫變數，檔案裡沒有對應屬性）：
            # 只能 SetVar + Flush 寫 DB，驗證也只能用 GetVar 讀 DB，不能讀檔案比對
            flag_err = None
            for _attempt in range(3):
                try:
                    evf = f5.GetEnumeratorVariable()
                    for k, v in flag_vals.items():
                        evf.SetVar(k, '', v)
                    evf.Flush()
                except Exception as e:
                    flag_err = str(e)
                    time.sleep(1.5)
                    continue
                bad = []
                evr = f5.GetEnumeratorVariable()
                for k, v in flag_vals.items():
                    try:
                        r = evr.GetVar(k, '')
                        actual = r[1] if isinstance(r, tuple) else r
                    except Exception as e:
                        actual = f'<讀取失敗 {e}>'
                    if str(actual) != v:
                        bad.append(f'{k} 預期 {v!r}，實際讀到 {actual!r}')
                if not bad:
                    flag_err = None
                    break
                flag_err = '；'.join(bad)
                time.sleep(1.5)
            if flag_err:
                return jsonify({'success': False,
                                'error': f'申請設變原因的勾選狀態寫入失敗：{flag_err}{CONFLICT_HINT}'}), 502

        # 自動存回（簽入）
        if d.get('checkin'):
            try:
                f5.Refresh()
                if f5.IsLocked:
                    f5.UnlockFile(0, '申請單填寫完成')
            except Exception as e:
                return jsonify({'success': True,
                                'warning': f'卡片已儲存，但存回失敗（請在 PDM 手動存回）：{e}'})

        # 提出申請：走 workflow 轉換「00提出申請」到「ECR單位主管審核」
        # （轉換名稱／目的狀態實測自現有 workflow，見 docs/dcn-index.md；沿用治檢具同一套
        # ChangeState3 + .NET Interop helper，若這個轉換也設了身分驗證會需要 PDM 密碼）
        if d.get('submit_flow'):
            try:
                comment = (d.get('comment') or '').strip() or '提出申請'
                password = d.get('flow_password') or ''
                if not password:
                    return jsonify({'success': True,
                                    'warning': '已存回，但提出申請需要輸入 PDM 登入密碼（此轉換可能有身分驗證），'
                                               '請重新按「續填」勾選提出申請並輸入密碼'})
                helper = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pdm_change_state.ps1')
                env = os.environ.copy()
                env.update({
                    'PDM_CS_FILE':       os.path.join(norm, target.Name),
                    'PDM_CS_STATE':      'ECR單位主管審核',
                    'PDM_CS_TRANSITION': '00提出申請',
                    'PDM_CS_COMMENT':    comment,
                    'PDM_CS_PASSWORD':   password,
                })
                import subprocess
                r = subprocess.run(
                    ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-File', helper],
                    env=env, capture_output=True, timeout=90)
                out = (r.stdout or b'').decode('utf-8', errors='replace').strip().splitlines()
                result = out[-1] if out else 'ERROR: 無輸出'
                if not result.startswith('OK'):
                    return jsonify({'success': True,
                                    'warning': f'已存回，但提出申請未成功：{result}。'
                                               '（密碼錯誤或欄位不齊；可在 PDM 手動變更狀態）'})
            except Exception as e:
                return jsonify({'success': True,
                                'warning': f'已存回，但提出申請失敗（請在 PDM 手動變更狀態）：{e}'})

        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': f'儲存卡片失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


def _dcn_paste_as_reference(vault, root_file_id, ref_paths):
    """把 ref_paths 這些檔案掛成 root_file_id 的 PDM「參考」（等同檔案總管的複製→貼上為參考）。

    2026-08-18 實測打通（RR2608017）。要點：
    - 用 EdmUtil_AddCustomRefs（=8）→ AddReferencesPath → CreateTree → CreateReferences。
    - **不能用 gen_py 包裝或 dynamic dispatch 呼叫**：gen_py 產生的 IEdmAddCustomRefs 類別
      根本沒有這些方法（只有 CLSID），dynamic dispatch 則會回「類型不符」——因為
      AddReferencesPath 的第二個參數是「SAFEARRAY(BSTR) 的指標」，pywin32 推不出來。
      要用 InvokeTypes 明確指定 VT_BYREF|VT_ARRAY|VT_BSTR (0x4000|0x2000|8 = 24584)。
    - memid：AddReferencesPath=2、CreateTree=4、CreateReferences=6（唯讀取自型別庫）。
    - 人工建立的申請單都是這種結構（既有 RR2608015／RR2607050 實測，附件都是 xlsm 的參考）。
    """
    import pythoncom as _pc
    VT_I4, VT_VOID, VT_BOOL = 3, 24, 11
    VT_SAFEARRAY_BSTR_BYREF = 0x4000 | 0x2000 | 8
    ob = vault.CreateUtility(8)   # EdmUtil_AddCustomRefs
    ob = ob._oleobj_ if hasattr(ob, '_oleobj_') else ob
    ob.InvokeTypes(2, 0, _pc.DISPATCH_METHOD, (VT_VOID, 0),
                   ((VT_I4, 1), (VT_SAFEARRAY_BSTR_BYREF, 1)), root_file_id, list(ref_paths))
    ob.InvokeTypes(4, 0, _pc.DISPATCH_METHOD, (VT_BOOL, 0), ((VT_I4, 1),), 0)   # CreateTree
    ob.InvokeTypes(6, 0, _pc.DISPATCH_METHOD, (VT_BOOL, 0), ())                 # CreateReferences


@app.route('/api/dcn/attachment/upload', methods=['POST'])
def dcn_attachment_upload():
    """上傳附件到指定的 DCN 資料夾，並把它掛成申請單 xlsm 的 PDM「參考」
    （等同人工做的複製→在 Excel 上貼上參考，見 _dcn_paste_as_reference）"""
    folder_path = (request.form.get('folder') or '').strip()
    dcn_root = os.path.normpath(config.DCN_VAULT_PATH)
    norm = os.path.normpath(folder_path)
    if not folder_path or not norm.lower().startswith(dcn_root.lower()):
        return jsonify({'success': False, 'error': '無效的資料夾路徑'}), 400
    f = request.files.get('file')
    if not f or not f.filename:
        return jsonify({'success': False, 'error': '請選擇檔案'}), 400

    import pythoncom
    pythoncom.CoInitialize()
    try:
        import win32com.client
        vault = _pdm_vault_login()
        folder = vault.GetFolderFromPath(norm)
        f5folder = win32com.client.CastTo(folder, 'IEdmFolder5')

        import tempfile
        safe_name = os.path.basename(f.filename)
        tmp_path = os.path.join(tempfile.gettempdir(), f'dcn_upload_{int(time.time())}_{safe_name}')
        f.save(tmp_path)
        warning = None
        try:
            file_id = f5folder.AddFile(0, tmp_path, safe_name)
            fobj = vault.GetObject(1, file_id)  # 1 = EdmObject_File
            f5file = win32com.client.CastTo(fobj, 'IEdmFile5')
            try:
                f5file.UnlockFile(0, '附件上傳')   # 簽入，讓它跟其他附件一樣落到預設歸檔狀態
            except Exception:
                pass
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

        # 掛成申請單 xlsm 的參考（失敗只警告，附件本身已經進 vault 了，不該讓整個上傳算失敗）
        try:
            xlsm = None
            pos = f5folder.GetFirstFilePosition()
            while not pos.IsNull:
                ff = f5folder.GetNextFile(pos)
                if ff.Name.lower().endswith('.xlsm') and '申請單' in ff.Name:
                    xlsm = ff
                    break
            if xlsm is None:
                warning = '附件已上傳，但這個資料夾找不到申請單 Excel，無法建立參考關聯'
            else:
                _dcn_paste_as_reference(vault, xlsm.ID, [os.path.join(norm, safe_name)])
        except Exception as e:
            warning = f'附件已上傳，但建立參考關聯失敗（可在 PDM 手動複製→貼上參考）：{e}'
        return jsonify({'success': True, 'filename': safe_name, 'warning': warning})
    except Exception as e:
        return jsonify({'success': False, 'error': f'上傳失敗：{e}'}), 502
    finally:
        pythoncom.CoUninitialize()


@app.route('/api/dcn/list')
def dcn_list():
    """讀取設計變更通知單索引（PDM 資料夾，build_dcn_index.py 建立）"""
    conn = get_pdm_db()
    if not conn:
        return jsonify({'success': False, 'error': 'PDM 索引資料庫不存在'}), 500

    try:
        cur = conn.execute(
            'SELECT folder_name, folder_path, issue_date, product_model, '
            '       submitter, handler, reviewer, approver, reason, description, status '
            'FROM dcn_index ORDER BY folder_name DESC'
        )
        rows = [dict(r) for r in cur.fetchall()]
        row = conn.execute('SELECT MAX(indexed_at) FROM dcn_index').fetchone()
        last_updated = row[0] if row and row[0] else None
    except sqlite3.OperationalError:
        rows = []
        last_updated = None
    finally:
        conn.close()

    return jsonify({'success': True, 'count': len(rows), 'dcns': rows,
                    'last_updated': last_updated})


_dcn_reindex_state = {
    'running': False,
    'phase':   'idle',   # idle | scanning | indexing | done | error
    'idx':     0,
    'total':   0,
    'count':   0,
    'message': '',
    'error':   '',
}
_dcn_reindex_lock = _threading.Lock()


def _run_dcn_reindex():
    """背景執行 DCN 索引重建，解析 stdout 更新進度狀態（比照治檢具 _run_jig_reindex 的做法）"""
    global _dcn_reindex_state
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_dcn_index.py')

    with _dcn_reindex_lock:
        _dcn_reindex_state.update(running=True, phase='scanning', idx=0, total=0,
                                   count=0, message='啟動中...', error='')

    try:
        child_env = os.environ.copy()
        child_env['PYTHONIOENCODING'] = 'utf-8'
        # 子行程 stdout 接到管線（非終端機）時 Python 預設整批緩衝，進度輸出會卡到緩衝區滿
        # 或行程結束才一次送出，畫面看起來像卡住不動；設這個讓它逐行即時送出
        child_env['PYTHONUNBUFFERED'] = '1'
        proc = _subprocess.Popen(
            [sys.executable, script, '--deploy'],
            stdout=_subprocess.PIPE, stderr=_subprocess.STDOUT,
            text=True, encoding='utf-8', errors='replace',
            cwd=os.path.dirname(script), env=child_env
        )
        for line in proc.stdout:
            line = line.rstrip()
            if not line:
                continue
            with _dcn_reindex_lock:
                _dcn_reindex_state['message'] = line

            m = _re.search(r'找到\s*(\d+)\s*個子資料夾', line)
            if m:
                with _dcn_reindex_lock:
                    _dcn_reindex_state['total'] = int(m.group(1))
                continue

            m = _re.search(r'\[(\d+)/(\d+)\]', line)
            if m:
                with _dcn_reindex_lock:
                    _dcn_reindex_state['phase'] = 'indexing'
                    _dcn_reindex_state['idx']   = int(m.group(1))
                    _dcn_reindex_state['total'] = int(m.group(2))
                continue

            m = _re.search(r'索引共\s*(\d+)\s*筆', line)
            if m:
                with _dcn_reindex_lock:
                    _dcn_reindex_state['count'] = int(m.group(1))

        proc.wait()

        if proc.returncode == 0:
            with _dcn_reindex_lock:
                cnt = _dcn_reindex_state['count']
                _dcn_reindex_state.update(running=False, phase='done',
                                           message=f'更新完成，共 {cnt} 筆')
        else:
            with _dcn_reindex_lock:
                _dcn_reindex_state.update(running=False, phase='error',
                                           error='重建失敗，請查看伺服器日誌')
    except Exception as exc:
        with _dcn_reindex_lock:
            _dcn_reindex_state.update(running=False, phase='error', error=str(exc))


@app.route('/api/dcn/rebuild', methods=['POST'])
def dcn_rebuild():
    """啟動設計變更通知單索引重建（背景執行，前端輪詢 /api/dcn/rebuild/status 取得進度）"""
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'build_dcn_index.py')
    if not os.path.exists(script):
        return jsonify({'success': False, 'error': 'build_dcn_index.py 不存在'}), 500
    with _dcn_reindex_lock:
        if _dcn_reindex_state['running']:
            return jsonify({'success': False, 'error': '索引重建已在執行中，請稍候'}), 409
    t = _threading.Thread(target=_run_dcn_reindex, daemon=True)
    t.start()
    return jsonify({'success': True, 'message': '索引重建已啟動'})


@app.route('/api/dcn/rebuild/status', methods=['GET'])
def dcn_rebuild_status():
    """回傳 DCN 索引重建進度狀態"""
    with _dcn_reindex_lock:
        return jsonify(dict(_dcn_reindex_state))


@app.route('/api/dcn/open-folder', methods=['POST'])
def dcn_open_folder():
    """用 Shell 開啟設計變更通知單資料夾"""
    data = request.get_json(silent=True) or {}
    path = data.get('path', '')
    if not path:
        return jsonify({'ok': False, 'error': '缺少 path 參數'}), 400

    dcn_root = os.path.normpath(config.DCN_VAULT_PATH)
    norm_path = os.path.normpath(path)
    if not norm_path.lower().startswith(dcn_root.lower()):
        return jsonify({'ok': False, 'error': '無效的路徑'}), 400
    if not os.path.isdir(norm_path):
        return jsonify({'ok': False, 'error': '資料夾不存在'}), 404

    try:
        os.startfile(norm_path)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ══════════════════════════════════════════════════════════
#  油品管理（油品主檔／MSDS／使用單位／更換記錄）  詳見 docs/oil-management.md
# ══════════════════════════════════════════════════════════
#  正本是網芳上的 oil.db（config.OIL_DB_PATH），MSDS/簡介的實體 PDF 留在網芳資料夾，
#  這裡只存索引與對應關係。schema 定義在 build_oil_index.py，本檔不另外抄一份。

_OIL_STATUSES = ('使用中', '停用')
_OIL_ROOTS = {'msds': 'OIL_MSDS_ROOT', 'doc': 'OIL_DOC_ROOT'}
# 附件允許的副檔名：MSDS 是 PDF，簡介/注意事項常是 Word，另外開放 Excel——
# 使用者要求「平台可以放 EXCEL 資料存放」，油品的用量表/濃度記錄表就是丟這裡
_OIL_UPLOAD_EXT = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.xlsm', '.csv',
                   '.ppt', '.pptx', '.txt', '.jpg', '.jpeg', '.png'}
# 編輯時會寫進異動歷程的欄位（跟設備主檔一樣，備註/說明這種隨手改的刻意不追蹤）
_OIL_TRACKED_FIELDS = {
    'status':   '狀態變更',
    'supplier': '資料修改',
    'category': '資料修改',
    'spec':     '資料修改',
    'pack':     '資料修改',
    'replaced_by': '資料修改',
}
_OIL_FIELD_LABEL = {'status': '狀態', 'supplier': '供應商', 'category': '分類',
                    'spec': '規格', 'pack': '包裝', 'replaced_by': '替代油品'}


# 清單排序用的分類順序：切削類 → 潤滑類 → 其他用途，跟新增表單的分類下拉選單同一個順序。
# 不能直接 ORDER BY category——中文字是照 Unicode 碼位排（主軸油會排在切削油前面），
# 對現場來說毫無意義，同類的油品也不會排在一起
_OIL_CATEGORY_ORDER = ['水性切削液', '切削油', '切削油膏', '液壓油', '滑道油', '主軸油',
                       '機油', '齒輪油', '潤滑脂', '防銹油', '清潔劑', '燃料', '其他']
_OIL_CAT_SORT = {name: i for i, name in enumerate(_OIL_CATEGORY_ORDER)}


def _oil_sort_key(r):
    """分類優先 → 油品代號（分類沒在名單裡的排最後，代號大小寫不影響）

    代號前面掛的中文說明（`「全合成切削液」CS-1010`）排序時先拿掉，照真正的代號排——
    不然中文字的碼位比英數大，所有加了說明的油品都會被擠到該分類的最後面。"""
    code = re.sub(r'^「[^」]*」\s*', '', (r.get('code') or '')).upper()
    return (_OIL_CAT_SORT.get((r.get('category') or '').strip(), len(_OIL_CATEGORY_ORDER)), code)


def _oil_root(root):
    """'msds'/'doc' → 網芳實體資料夾路徑"""
    key = _OIL_ROOTS.get(root)
    return getattr(config, key, '') if key else ''


def _oil_conn():
    """開啟油品主檔資料庫（不存在時自動建立空的 schema，讓沒跑過匯入的電腦也能用）。

    跟 equipment.db 同樣是網芳上的共用檔案：timeout=15 讓撞到別人寫入時等待重試，
    刻意不開 WAL（Windows SMB 對 WAL 需要的 shared memory 支援不穩定）。"""
    try:
        from build_oil_index import SCHEMA as _OIL_SCHEMA
    except Exception:
        _OIL_SCHEMA = None
    if not os.path.exists(OIL_DB_PATH):
        if _OIL_SCHEMA is None:
            return None
        try:
            os.makedirs(os.path.dirname(OIL_DB_PATH), exist_ok=True)
        except OSError:
            return None
    conn = sqlite3.connect(OIL_DB_PATH, timeout=15)
    conn.row_factory = sqlite3.Row
    if _OIL_SCHEMA:
        conn.executescript(_OIL_SCHEMA)     # 全部 IF NOT EXISTS，已有資料不受影響
        conn.commit()
    return conn


def _oil_now():
    return datetime.now().strftime('%Y-%m-%d %H:%M:%S')


def _oil_log(cur, code, action, detail, user='system'):
    cur.execute('INSERT INTO oil_history (code, date, action, detail, user) VALUES (?,?,?,?,?)',
                (code, datetime.now().strftime('%Y-%m-%d'), action, detail, user))


def _oil_log_changes(cur, code, before, after):
    """比對被追蹤欄位，有變動就一個欄位記一筆。原值必須在 UPDATE **之前**撈好。"""
    for field, action in _OIL_TRACKED_FIELDS.items():
        old = (before.get(field) or '').strip()
        new = (after.get(field) or '').strip()
        if old != new:
            label = _OIL_FIELD_LABEL.get(field, field)
            _oil_log(cur, code, action, f'{label}：{old or "（空白）"} → {new or "（空白）"}')


@app.route('/oil')
def oil_page():
    """油品管理頁面（油品主檔／MSDS／使用單位／更換記錄）"""
    return render_template('oil.html', app_version=APP_VERSION)


@app.route('/api/oil/search')
def oil_search():
    """油品清單查詢（空格=AND、-前綴=NOT，比對代號/品名/品牌/供應商/分類/規格/說明/使用單位）"""
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟，請確認網芳路徑'}), 500
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT o.*,
                   (SELECT COUNT(*) FROM oil_file f
                     WHERE f.code = o.code AND f.root='msds' AND f.obsolete=0) AS msds_count,
                   (SELECT COUNT(*) FROM oil_file f
                     WHERE f.code = o.code AND f.root='doc')  AS doc_count,
                   (SELECT MAX(f.file_date) FROM oil_file f
                     WHERE f.code = o.code AND f.root='msds' AND f.obsolete=0) AS msds_date,
                   (SELECT COUNT(*) FROM oil_change c WHERE c.code = o.code) AS change_count,
                   (SELECT MAX(c.date)  FROM oil_change c WHERE c.code = o.code) AS last_change
              FROM oil o""").fetchall()]
        units = {}
        for r in conn.execute('SELECT code, unit, equip_code FROM oil_unit ORDER BY code, sort'):
            units.setdefault(r['code'], []).append(
                r['unit'] + (f"({r['equip_code']})" if r['equip_code'] else ''))
    finally:
        conn.close()
    for r in rows:
        r['units'] = units.get(r['code'], [])
    rows.sort(key=_oil_sort_key)

    q = request.args.get('q', '').strip()
    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def haystack(r):
            return ' '.join([str(r.get(k) or '') for k in
                             ('code', 'name', 'brand', 'supplier', 'category', 'spec',
                              'pack', 'status', 'usage_note', 'remark', 'replaced_by')]
                            + r['units']).lower()

        rows = [r for r in rows if all(m in haystack(r) for m in must)
                and not any(m in haystack(r) for m in must_not)]
    return jsonify({'success': True, 'count': len(rows), 'data': rows})


@app.route('/api/oil/stats')
def oil_stats():
    """分類／供應商／狀態的支數統計，供篩選列顯示數字（徽章篩選慣例）"""
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': True, 'categories': [], 'suppliers': [], 'statuses': []})
    try:
        # GROUP BY 一定要寫完整運算式，不能寫 `GROUP BY name`——oil 表本身就有一個
        # name 欄位（品名），別名會被同名欄位蓋掉，變成每支油品各自一組（實測踩過）
        cats = [dict(r) for r in conn.execute(
            "SELECT IFNULL(NULLIF(category,''),'未分類') AS name, COUNT(*) AS cnt FROM oil "
            "GROUP BY IFNULL(NULLIF(category,''),'未分類') ORDER BY cnt DESC")]
        sups = [dict(r) for r in conn.execute(
            "SELECT IFNULL(NULLIF(supplier,''),'未填') AS name, COUNT(*) AS cnt FROM oil "
            "GROUP BY IFNULL(NULLIF(supplier,''),'未填') ORDER BY cnt DESC")]
        sts = [dict(r) for r in conn.execute(
            "SELECT IFNULL(NULLIF(status,''),'未填') AS name, COUNT(*) AS cnt FROM oil "
            "GROUP BY IFNULL(NULLIF(status,''),'未填') ORDER BY cnt DESC")]
        # 「使用中卻沒有現行 MSDS」是這個模組最該被看到的風險，放在工具列紅字提示
        no_msds = conn.execute(
            "SELECT COUNT(*) FROM oil o WHERE o.status='使用中' AND NOT EXISTS "
            "(SELECT 1 FROM oil_file f WHERE f.code=o.code AND f.root='msds' AND f.obsolete=0)"
        ).fetchone()[0]
        orphan = conn.execute('SELECT COUNT(*) FROM oil_file WHERE code IS NULL').fetchone()[0]
    finally:
        conn.close()
    return jsonify({'success': True, 'categories': cats, 'suppliers': sups,
                    'statuses': sts, 'no_msds': no_msds, 'orphan_files': orphan})


@app.route('/api/oil/detail')
def oil_detail():
    """單一油品詳情：MSDS/文件清單、使用單位、更換記錄、異動歷程"""
    code = request.args.get('code', '').strip()
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        row = conn.execute('SELECT * FROM oil WHERE code=?', (code,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
        d = dict(row)
        d['files'] = [dict(r) for r in conn.execute(
            'SELECT rowid AS id, relpath, root, filename, ext, size, mtime, file_date, obsolete '
            'FROM oil_file WHERE code=? ORDER BY obsolete, root, file_date DESC, filename', (code,))]
        d['units'] = [dict(r) for r in conn.execute(
            'SELECT rowid AS id, unit, equip_code, note FROM oil_unit WHERE code=? ORDER BY sort, rowid',
            (code,))]
        d['changes'] = [dict(r) for r in conn.execute(
            'SELECT id, equip_code, equip_name, date, qty, operator, note, user '
            'FROM oil_change WHERE code=? ORDER BY date DESC, id DESC', (code,))]
        d['history'] = [dict(r) for r in conn.execute(
            'SELECT rowid AS id, date, action, detail, user FROM oil_history WHERE code=? '
            'ORDER BY date DESC, rowid DESC', (code,))]
    finally:
        conn.close()
    return jsonify({'success': True, 'data': d})


@app.route('/api/oil/save', methods=['POST'])
def oil_save():
    """新增或編輯油品。改代號時所有子表（檔案/使用單位/更換記錄/歷程）一起搬過去。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip()
    orig = (d.get('orig_code') or '').strip()
    if not code:
        return jsonify({'success': False, 'error': '油品代號必填'}), 400
    if d.get('status') and d['status'] not in _OIL_STATUSES:
        return jsonify({'success': False, 'error': '狀態只能是使用中或停用'}), 400

    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        cur = conn.cursor()
        fields = ('name', 'brand', 'supplier', 'category', 'spec', 'pack',
                  'status', 'replaced_by', 'usage_note', 'remark')
        vals = {f: (d.get(f) or '').strip() for f in fields}
        vals['status'] = vals['status'] or '使用中'

        exists = cur.execute('SELECT * FROM oil WHERE code=?', (orig or code,)).fetchone()
        if orig and orig != code:
            if cur.execute('SELECT 1 FROM oil WHERE code=?', (code,)).fetchone():
                return jsonify({'success': False, 'error': f'代號 {code} 已存在'}), 400

        if exists is None:
            if cur.execute('SELECT 1 FROM oil WHERE code=?', (code,)).fetchone():
                return jsonify({'success': False, 'error': f'代號 {code} 已存在'}), 400
            cur.execute("""INSERT INTO oil (code, name, brand, supplier, category, spec, pack,
                                            status, replaced_by, usage_note, remark,
                                            source, origin, created_at, updated_at)
                           VALUES (?,?,?,?,?,?,?,?,?,?,?,'manual','manual',?,?)""",
                        (code, vals['name'], vals['brand'], vals['supplier'], vals['category'],
                         vals['spec'], vals['pack'], vals['status'], vals['replaced_by'],
                         vals['usage_note'], vals['remark'], _oil_now(), _oil_now()))
            _oil_log(cur, code, '新增', f"在系統內新增油品：{vals['name'] or code}")
        else:
            before = dict(exists)
            if orig and orig != code:
                for tbl in ('oil_file', 'oil_unit', 'oil_change', 'oil_history'):
                    cur.execute(f'UPDATE {tbl} SET code=? WHERE code=?', (code, orig))
                cur.execute('UPDATE oil SET code=? WHERE code=?', (code, orig))
                _oil_log(cur, code, '重新編碼', f'{orig} → {code}')
            cur.execute("""UPDATE oil SET name=?, brand=?, supplier=?, category=?, spec=?,
                                          pack=?, status=?, replaced_by=?, usage_note=?,
                                          remark=?, source='manual', updated_at=?
                            WHERE code=?""",
                        (vals['name'], vals['brand'], vals['supplier'], vals['category'],
                         vals['spec'], vals['pack'], vals['status'], vals['replaced_by'],
                         vals['usage_note'], vals['remark'], _oil_now(), code))
            _oil_log_changes(cur, code, before, vals)

        # 使用單位整組覆寫（前端送完整清單，比逐筆增刪簡單也不會漏）
        if isinstance(d.get('units'), list):
            cur.execute('DELETE FROM oil_unit WHERE code=?', (code,))
            for i, u in enumerate(d['units']):
                unit = (u.get('unit') or '').strip()
                eq = (u.get('equip_code') or '').strip().upper()
                if not unit and not eq:
                    continue
                cur.execute('INSERT INTO oil_unit (code, unit, equip_code, note, sort) '
                            'VALUES (?,?,?,?,?)', (code, unit, eq, (u.get('note') or '').strip(), i))
        conn.commit()
    finally:
        conn.close()
    # 改代號的話，設備保養基準書引用到的代號要一起改（見 _oil_sync_maint_code）
    if orig and orig != code:
        _oil_sync_maint_code(orig, code)
    return jsonify({'success': True, 'code': code})


# 可以在清單表格裡直接改的欄位（白名單，不接受任意欄位名）
_OIL_INLINE_FIELDS = {'name': '品名'}   # 一般欄位（直接 UPDATE）；'code' 走下面的改代號分支


def _oil_sync_maint_code(old, new):
    """油品改代號時，把設備保養基準書引用到的代號一起改過去。

    保養項目的 `mt_item.oil_code` 存的是代號字串，而 oil.db 與 equipment.db 是兩個
    不同的資料庫檔案，沒有外鍵可以連動——不主動同步的話，改完代號基準書上就會顯示成
    紅色「（主檔查無）」。清單頁的快速編輯讓改代號變得很容易按，這條同步就更不能省。
    回傳更新筆數；設備資料庫還沒建立時回 0，不影響油品這邊的改名。"""
    conn = _eq_conn()
    if conn is None:
        return 0
    try:
        cols = {r['name'] for r in conn.execute('PRAGMA table_info(mt_item)')}
        if 'oil_code' not in cols:      # 還沒用過保養基準書的資料庫
            return 0
        n = conn.execute('UPDATE mt_item SET oil_code=? WHERE oil_code=?', (new, old)).rowcount
        conn.commit()
        return n
    except sqlite3.Error:
        return 0
    finally:
        conn.close()


@app.route('/api/oil/inline_save', methods=['POST'])
def oil_inline_save():
    """清單表格內直接編輯單一欄位（品名、油品代號）。

    一律把 source 標成 manual——這些欄位正是重新掃描時會從 MSDS 檔名覆蓋回去的，
    使用者親手改過就不該再被檔名蓋掉。品名刻意不寫進異動歷程（跟 _OIL_TRACKED_FIELDS
    的取捨一致：隨手修正的欄位記進去只會把歷程洗版），但**改代號會記**——那是會牽動
    一整串子表的異動，事後一定會想知道什麼時候改的。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip()
    field = (d.get('field') or '').strip()
    value = (d.get('value') or '').strip()
    if field != 'code' and field not in _OIL_INLINE_FIELDS:
        return jsonify({'success': False, 'error': f'不支援編輯欄位 {field}'}), 400
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        cur = conn.cursor()
        if field == 'code':
            if not value:
                return jsonify({'success': False, 'error': '油品代號不可以空白'}), 400
            if value == code:
                return jsonify({'success': True, 'code': code, 'field': field, 'value': value})
            if cur.execute('SELECT 1 FROM oil WHERE code=?', (value,)).fetchone():
                return jsonify({'success': False, 'error': f'代號 {value} 已經有人用了'}), 400
            if cur.execute('SELECT 1 FROM oil WHERE code=?', (code,)).fetchone() is None:
                return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
            # 子表一起搬（跟 /api/oil/save 改代號同一套做法）
            for tbl in ('oil_file', 'oil_unit', 'oil_change', 'oil_history'):
                cur.execute(f'UPDATE {tbl} SET code=? WHERE code=?', (value, code))
            cur.execute("UPDATE oil SET code=?, source='manual', updated_at=? WHERE code=?",
                        (value, _oil_now(), code))
            _oil_log(cur, value, '重新編碼', f'{code} → {value}（清單快速編輯）')
            conn.commit()
            n = _oil_sync_maint_code(code, value)
            msg = f'代號已改為 {value}'
            if n:
                msg += f'，並同步更新 {n} 筆保養基準書的引用'
            return jsonify({'success': True, 'code': value, 'old_code': code,
                            'field': field, 'value': value, 'message': msg})

        cur.execute(f"UPDATE oil SET {field}=?, source='manual', updated_at=? WHERE code=?",
                    (value, _oil_now(), code))
        if cur.rowcount == 0:
            return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'code': code, 'field': field, 'value': value})


@app.route('/api/oil/delete', methods=['POST'])
def oil_delete():
    """預設軟刪除（狀態改停用）；hard=true 只允許刪系統內新增的（origin='manual'）。

    判斷用 origin 不是 source——source 只要在系統內編輯過就會變 manual，
    拿它當門檻會讓掃描來的油品也變成可硬刪（設備主檔踩過這個雷）。"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip()
    hard = bool(d.get('hard'))
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        row = conn.execute('SELECT * FROM oil WHERE code=?', (code,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
        cur = conn.cursor()
        if hard:
            if (row['origin'] or 'msds') != 'manual':
                return jsonify({'success': False,
                                'error': '這支油品來自 MSDS 掃描，只能停用不能刪除'}), 400
            for tbl in ('oil_unit', 'oil_change', 'oil_history'):
                cur.execute(f'DELETE FROM {tbl} WHERE code=?', (code,))
            cur.execute('UPDATE oil_file SET code=NULL WHERE code=?', (code,))
            cur.execute('DELETE FROM oil WHERE code=?', (code,))
            msg = f'已刪除 {code}'
        else:
            cur.execute("UPDATE oil SET status='停用', source='manual', updated_at=? WHERE code=?",
                        (_oil_now(), code))
            _oil_log(cur, code, '狀態變更', f"狀態：{row['status']} → 停用")
            msg = f'{code} 已改為停用'
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': msg})


@app.route('/api/oil/file')
def oil_file():
    """輸出 MSDS／文件（PDF 直接內嵌預覽，其餘下載）"""
    from flask import send_file
    root = request.args.get('root', 'msds')
    base = _oil_root(root)
    full = _eq_safe_path(base, request.args.get('relpath', '')) if base else None
    if not full or not os.path.isfile(full):
        return jsonify({'success': False, 'error': '檔案不存在'}), 404
    dl = request.args.get('dl') == '1' or os.path.splitext(full)[1].lower() != '.pdf'
    return send_file(full, as_attachment=dl, download_name=os.path.basename(full))


@app.route('/api/oil/file/upload', methods=['POST'])
def oil_file_upload():
    """上傳 MSDS 或其他文件（含 Excel）到網芳，並立即補進索引。

    MSDS 存進 OIL_MSDS_ROOT 根目錄（沿用既有的扁平結構，檔名照 [供應商][代號] 慣例
    自動組出來，重新掃描時才認得出是哪支油品）；其他文件存進 OIL_DOC_ROOT\\<代號>\\。
    兩者都標 claimed=1，重新掃描不會被搶走對應。"""
    code = (request.form.get('code') or '').strip()
    kind = (request.form.get('kind') or 'doc').strip()
    files = request.files.getlist('files')
    if not code:
        return jsonify({'success': False, 'error': '缺少油品代號'}), 400
    if kind not in _OIL_ROOTS:
        return jsonify({'success': False, 'error': '檔案類型只能是 msds 或 doc'}), 400
    if not files or not any(f and f.filename for f in files):
        return jsonify({'success': False, 'error': '請選擇檔案'}), 400

    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        row = conn.execute('SELECT supplier FROM oil WHERE code=?', (code,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
        base = _oil_root(kind)
        dest_dir = base if kind == 'msds' else os.path.join(base, code)
        try:
            os.makedirs(dest_dir, exist_ok=True)
        except OSError as e:
            return jsonify({'success': False, 'error': f'無法建立資料夾：{e}'}), 500

        saved, rejected = [], []
        for f in files:
            if not f or not f.filename:
                continue
            ext = os.path.splitext(f.filename)[1].lower()
            if ext not in _OIL_UPLOAD_EXT:
                rejected.append(f.filename)
                continue
            filename = _cnc_safe_filename(f.filename)
            if kind == 'msds' and not filename.startswith('['):
                # 照既有慣例組檔名，下次重新掃描才對得回這支油品
                sup = (row['supplier'] or '').strip()
                prefix = (f'[{sup}]' if sup else '') + f'[{code}]'
                filename = prefix + filename
            dest = os.path.join(dest_dir, filename)
            if os.path.exists(dest):
                stem, ex = os.path.splitext(filename)
                filename = f'{stem}_{int(time.time())}{ex}'
                dest = os.path.join(dest_dir, filename)
            try:
                f.save(dest)
            except OSError as e:
                rejected.append(f'{f.filename}（存檔失敗：{e}）')
                continue
            # relpath 是 oil_file 的主鍵，一定要跟 build_oil_index.py 掃描時寫入的格式一致
            # （os.path.relpath 的原生分隔符，Windows 上是反斜線）。這裡若改寫成正斜線，
            # 下次重新掃描會把同一個檔案當成另一筆插進去，清單上就會出現兩份（實測踩過）
            relpath = os.path.relpath(dest, base)
            conn.execute("""INSERT OR REPLACE INTO oil_file
                (relpath, root, code, folder, filename, ext, size, mtime, file_date, obsolete, claimed)
                VALUES (?,?,?,?,?,?,?,datetime('now','localtime'),'',0,1)""",
                (relpath, kind, code, '' if kind == 'msds' else code, filename, ext,
                 os.path.getsize(dest)))
            saved.append(filename)
        conn.commit()
    finally:
        conn.close()

    if not saved:
        err = '沒有成功上傳的檔案'
        if rejected:
            err += '：' + '、'.join(rejected) + '（僅支援 ' + '/'.join(sorted(_OIL_UPLOAD_EXT)) + '）'
        return jsonify({'success': False, 'error': err}), 400
    msg = f'已上傳 {len(saved)} 個檔案'
    if rejected:
        msg += f'，{len(rejected)} 個格式不支援已略過'
    return jsonify({'success': True, 'message': msg, 'saved': saved})


@app.route('/api/oil/file/claim', methods=['POST'])
def oil_file_claim():
    """把掃描時歸不了位的檔案人工指定給某支油品（claimed=1，重新掃描保留）"""
    d = request.get_json(silent=True) or {}
    relpath = (d.get('relpath') or '').strip()
    code = (d.get('code') or '').strip()
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        if code and not conn.execute('SELECT 1 FROM oil WHERE code=?', (code,)).fetchone():
            return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
        cur = conn.execute('UPDATE oil_file SET code=?, claimed=? WHERE relpath=?',
                           (code or None, 1 if code else 0, relpath))
        if cur.rowcount == 0:
            return jsonify({'success': False, 'error': '查無這個檔案'}), 404
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/oil/file/list')
def oil_file_list():
    """MSDS 子頁：跨油品列出所有檔案（含歸不了位的），支援同一套搜尋語法"""
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT f.rowid AS id, f.relpath, f.root, f.code, f.folder, f.filename, f.ext,
                   f.size, f.mtime, f.file_date, f.obsolete, f.claimed,
                   IFNULL(o.name,'') AS oil_name, IFNULL(o.supplier,'') AS supplier,
                   IFNULL(o.category,'') AS category, IFNULL(o.status,'') AS oil_status
              FROM oil_file f LEFT JOIN oil o ON o.code = f.code
             ORDER BY f.obsolete, f.root, IFNULL(f.code,'zzz'), f.filename""").fetchall()]
    finally:
        conn.close()

    q = request.args.get('q', '').strip()
    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def hay(r):
            return ' '.join(str(r.get(k) or '') for k in
                            ('code', 'filename', 'oil_name', 'supplier', 'category', 'folder')).lower()
        rows = [r for r in rows if all(m in hay(r) for m in must)
                and not any(m in hay(r) for m in must_not)]
    return jsonify({'success': True, 'count': len(rows), 'data': rows})


@app.route('/api/oil/change/list')
def oil_change_list():
    """更換記錄清單（跨油品），搜尋語法同其他清單"""
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT c.*, IFNULL(o.name,'') AS oil_name, IFNULL(o.category,'') AS category,
                   IFNULL(o.supplier,'') AS supplier
              FROM oil_change c LEFT JOIN oil o ON o.code = c.code
             ORDER BY c.date DESC, c.id DESC""").fetchall()]
    finally:
        conn.close()
    q = request.args.get('q', '').strip()
    if q:
        must, must_not = [], []
        for tok in q.split():
            if tok.startswith('-') and len(tok) > 1:
                must_not.append(tok[1:].lower())
            else:
                must.append(tok.lower())

        def hay(r):
            return ' '.join(str(r.get(k) or '') for k in
                            ('code', 'oil_name', 'equip_code', 'equip_name', 'date',
                             'operator', 'note', 'category', 'supplier')).lower()
        rows = [r for r in rows if all(m in hay(r) for m in must)
                and not any(m in hay(r) for m in must_not)]
    return jsonify({'success': True, 'count': len(rows), 'data': rows})


@app.route('/api/oil/change/save', methods=['POST'])
def oil_change_save():
    """新增或編輯一筆更換記錄（帶 id = 編輯）"""
    d = request.get_json(silent=True) or {}
    code = (d.get('code') or '').strip()
    date_s = (d.get('date') or '').strip()
    if not code or not date_s:
        return jsonify({'success': False, 'error': '油品與更換日期必填'}), 400
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        if not conn.execute('SELECT 1 FROM oil WHERE code=?', (code,)).fetchone():
            return jsonify({'success': False, 'error': f'查無油品 {code}'}), 404
        vals = (code, (d.get('equip_code') or '').strip().upper(),
                (d.get('equip_name') or '').strip(), date_s,
                (d.get('qty') or '').strip(), (d.get('operator') or '').strip(),
                (d.get('note') or '').strip())
        cur = conn.cursor()
        if d.get('id'):
            cur.execute("""UPDATE oil_change SET code=?, equip_code=?, equip_name=?, date=?,
                                                 qty=?, operator=?, note=? WHERE id=?""",
                        vals + (int(d['id']),))
        else:
            cur.execute("""INSERT INTO oil_change
                (code, equip_code, equip_name, date, qty, operator, note, user)
                VALUES (?,?,?,?,?,?,?, 'user')""", vals)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/oil/change/delete', methods=['POST'])
def oil_change_delete():
    """刪除更換記錄。只能刪系統內登錄的（user='user'）——Excel 匯入的是舊資料軌跡，
    要清掉請重跑 build_oil_index.py（它會整批重建 user='excel' 的部分）。"""
    d = request.get_json(silent=True) or {}
    rid = d.get('id')
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        row = conn.execute('SELECT user FROM oil_change WHERE id=?', (rid,)).fetchone()
        if row is None:
            return jsonify({'success': False, 'error': '查無這筆記錄'}), 404
        if row['user'] != 'user':
            return jsonify({'success': False, 'error': 'Excel 匯入的記錄不能刪除'}), 400
        conn.execute('DELETE FROM oil_change WHERE id=?', (rid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/oil/equipment_options')
def oil_equipment_options():
    """使用單位／更換記錄選設備用的下拉選項（來自設備主檔，只列使用中與閒置的）"""
    conn = _eq_conn()
    if conn is None:
        return jsonify({'success': True, 'data': []})
    try:
        rows = [dict(r) for r in conn.execute("""
            SELECT e.code, IFNULL(e.old_code,'') AS old_code, IFNULL(e.location,'') AS location,
                   CASE WHEN e.needs_fix = 1 AND IFNULL(e.type_name_raw,'') <> ''
                        THEN e.type_name_raw ELSE IFNULL(t.name,'') END AS type_name
              FROM equipment e
              LEFT JOIN eq_type t ON t.group_code = e.group_code AND t.code = e.type_code
             WHERE e.status IN ('使用中','閒置') ORDER BY e.code""")]
    finally:
        conn.close()
    return jsonify({'success': True, 'data': rows})


@app.route('/api/oil/rebuild', methods=['POST'])
def oil_rebuild():
    """重新掃描 MSDS／簡介資料夾（直接呼叫 build_oil_index 的函式，不另外開 process）。

    掃描是合併式的：只更新 source='msds' 的資料，系統內編輯過的（manual）不受影響。"""
    try:
        import build_oil_index as B
    except Exception as e:
        return jsonify({'success': False, 'error': f'找不到 build_oil_index.py：{e}'}), 500
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        new_oil, upd_oil = B.scan_msds(conn)
        ndoc = B.scan_docs(conn)
        msg = f'掃描完成：新增 {new_oil} 支油品、更新 {upd_oil} 支、文件索引 {ndoc} 個檔案'
        if (request.get_json(silent=True) or {}).get('with_change'):
            msg += f'；更換記錄匯入 {B.import_change_xlsx(conn)} 筆'
    except Exception as e:
        return jsonify({'success': False, 'error': f'掃描失敗：{e}'}), 500
    finally:
        conn.close()
    return jsonify({'success': True, 'message': msg})


_OIL_EXPORT_HEADERS = ['油品代號', '品名', '品牌', '供應商', '分類', '規格', '包裝', '狀態',
                       '替代油品', '使用單位', 'MSDS份數', 'MSDS日期', '使用說明', '備註',
                       '資料來源', '最後更新']


@app.route('/api/oil/export', methods=['POST'])
def oil_export():
    """把油品主檔／使用單位／更換記錄匯出成獨立的 Excel（config.OIL_EXPORT_XLSX）。

    每次都用 openpyxl.Workbook() 開一本全新的活頁簿整份重建，**不會**去開既有的
    「油品更換記錄表.xlsx」——用 openpyxl 開檔存檔會把公式的快取值洗掉
    （設備主檔實測踩過，見 docs/equipment-master.md）。檔案被開啟中時退存到桌面。"""
    if not _OPENPYXL_OK:
        return jsonify({'success': False, 'error': '未安裝 openpyxl'}), 500
    conn = _oil_conn()
    if conn is None:
        return jsonify({'success': False, 'error': '油品資料庫無法開啟'}), 500
    try:
        oils = [dict(r) for r in conn.execute("""
            SELECT o.*,
                   (SELECT COUNT(*) FROM oil_file f WHERE f.code=o.code AND f.root='msds' AND f.obsolete=0) AS msds_count,
                   (SELECT MAX(f.file_date) FROM oil_file f WHERE f.code=o.code AND f.root='msds' AND f.obsolete=0) AS msds_date
              FROM oil o""")]
        oils.sort(key=_oil_sort_key)
        units = {}
        for r in conn.execute('SELECT code, unit, equip_code FROM oil_unit ORDER BY code, sort'):
            units.setdefault(r['code'], []).append(
                r['unit'] + (f"({r['equip_code']})" if r['equip_code'] else ''))
        changes = [dict(r) for r in conn.execute("""
            SELECT c.date, c.code, IFNULL(o.name,'') AS oil_name, c.equip_code, c.equip_name,
                   c.qty, c.operator, c.note, c.user
              FROM oil_change c LEFT JOIN oil o ON o.code=c.code
             ORDER BY c.date DESC, c.id DESC""")]
    finally:
        conn.close()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '油品主檔'
    ws.append(_OIL_EXPORT_HEADERS)
    for o in oils:
        ws.append([o['code'], o['name'], o['brand'], o['supplier'], o['category'], o['spec'],
                   o['pack'], o['status'], o['replaced_by'], '、'.join(units.get(o['code'], [])),
                   o['msds_count'], o['msds_date'] or '', o['usage_note'], o['remark'],
                   'MSDS掃描' if (o['origin'] or '') != 'manual' else '系統內新增', o['updated_at']])
    ws2 = wb.create_sheet('更換記錄')
    ws2.append(['更換日期', '油品代號', '品名', '設備編碼', '設備名稱', '數量', '登錄人', '備註', '來源'])
    for c in changes:
        ws2.append([c['date'], c['code'], c['oil_name'], c['equip_code'], c['equip_name'],
                    c['qty'], c['operator'], c['note'],
                    'Excel匯入' if c['user'] == 'excel' else '系統登錄'])
    for sheet, cols, widths in (
            (ws,  'ABCDEFGHIJKLMNOP', (12, 34, 12, 12, 14, 18, 14, 10, 12, 30, 10, 12, 40, 30, 12, 20)),
            (ws2, 'ABCDEFGHI',        (12, 12, 30, 12, 14, 10, 12, 24, 12))):
        for col, w in zip(cols, widths):
            sheet.column_dimensions[col].width = w
        sheet.freeze_panes = 'A2'

    path = getattr(config, 'OIL_EXPORT_XLSX', '') or os.path.join(
        os.path.expanduser('~'), 'Desktop', '油品主檔.xlsx')
    try:
        wb.save(path)
    except Exception:
        path = os.path.join(os.path.expanduser('~'), 'Desktop', '油品主檔.xlsx')
        try:
            wb.save(path)
        except Exception as e:
            return jsonify({'success': False, 'error': f'存檔失敗：{e}'}), 500
    return jsonify({'success': True, 'path': path,
                    'message': f'已匯出 {len(oils)} 支油品、{len(changes)} 筆更換記錄'})


@app.route('/api/oil/open_folder', methods=['POST'])
def oil_open_folder():
    """用檔案總管開啟油品資料夾（放 Excel／其他資料用）"""
    root = (request.get_json(silent=True) or {}).get('root', '')
    path = _oil_root(root) if root in _OIL_ROOTS else getattr(config, 'OIL_ROOT_PATH', '')
    if not path or not os.path.isdir(path):
        return jsonify({'success': False, 'error': '資料夾不存在'}), 404
    try:
        os.startfile(path)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    return jsonify({'success': True, 'path': path})


def _preload_cache():
    """背景預載 SSRS 資料到快取（加速首次搜尋）"""
    try:
        with ThreadPoolExecutor(max_workers=2) as pool:
            pool.submit(fetch_all_unfinished)
            pool.submit(fetch_efficiency_data)
    except Exception:
        pass


def _flask_host():
    """要綁哪個網卡。

    預設維持 config.FLASK_HOST（127.0.0.1，只有本機連得到，跟以前完全一樣）；
    只有把 config.MOBILE_ACCESS 打開時才改綁 0.0.0.0，讓同一個內網的手機掃 QR 進來
    回報保養（見 docs/equipment-maintenance.md P3）。這個開關預設關閉是刻意的——
    系統沒有登入機制，開了就等於內網任何人都能用。"""
    if getattr(config, 'MOBILE_ACCESS', False):
        return '0.0.0.0'
    return config.FLASK_HOST


if __name__ == '__main__':
    # 獨立執行模式（除錯用）— GUI 版由 main.py 啟動
    threading.Thread(target=_preload_cache, daemon=True).start()
    app.run(host=_flask_host(), port=config.FLASK_PORT, debug=config.FLASK_DEBUG)
